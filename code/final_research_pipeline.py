#!/usr/bin/env python3
from __future__ import annotations

import html
import json
import os
import re
import shutil
import warnings
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple
from zipfile import ZIP_DEFLATED, ZipFile

THESIS_DIR = Path(__file__).resolve().parents[2]
PROJECT_DIR = Path(__file__).resolve().parents[1]
MATERIALS_DIR = PROJECT_DIR / "materials"
INPUTS_DIR = MATERIALS_DIR / "inputs"
REFERENCES_DIR = MATERIALS_DIR / "references"
FULL_UAH = INPUTS_DIR / "full_uah_final.xlsx"
OUTPUT_DIR = PROJECT_DIR / "outputs"
DATA_DIR = PROJECT_DIR / "data"
FIG_DIR = PROJECT_DIR / "figures"
DOC_DIR = PROJECT_DIR / "documents"
LOG_DIR = PROJECT_DIR / "logs"
CODE_DIR = PROJECT_DIR / "code"
MODEL_OUTPUT_DIR = OUTPUT_DIR / "single_model_tables"
MODEL_DIAG_DIR = OUTPUT_DIR / "single_model_diagnostics"
MODEL_NOTE_DIR = OUTPUT_DIR / "single_model_notes"
VECM_DETAIL_DIR = OUTPUT_DIR / "vecm_detail"
FIG_SEQ_DIR = FIG_DIR / "sequence"
FIG_CH5_DIR = FIG_DIR / "chapter5_data"
FIG_CH6_DIR = FIG_DIR / "chapter6_results"
FIG_MODEL_DIR = FIG_DIR / "model_specific"
FIG_RELIABILITY_DIR = FIG_DIR / "reliability"
SOURCE_DRAFT = REFERENCES_DIR / "Charniuk_Maksym_MScThesis_Draft_correctedformat.docx"
REFERENCE_CHAPTER = REFERENCES_DIR / "data_estiamtion_updated_conclusion_fullversion.md"
ADDED_DATASET_PATH = INPUTS_DIR / "full_uah_final_whatadded_matched_smoothed.xlsx"

os.environ.setdefault("MPLCONFIGDIR", str(PROJECT_DIR / "_mplconfig"))

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import statsmodels.api as sm
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from statsmodels.stats.diagnostic import acorr_ljungbox, het_arch, het_breuschpagan, het_white
from statsmodels.stats.stattools import jarque_bera
from statsmodels.tsa.ardl import UECM, ardl_select_order
from statsmodels.tsa.stattools import adfuller, coint, kpss
from statsmodels.tsa.vector_ar.vecm import VECM, select_coint_rank, select_order


warnings.filterwarnings("ignore", category=RuntimeWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

PRODUCT_LABELS = {
    "milk": "Drinking milk / fermented milk",
    "butter": "Butter",
    "cheese": "CHEESE",
    "sour_cream": "Sour cream",
    "cream": "Cream",
    "yogurt_dessert": "Yogurt / dessert",
    "condensed_milk": "Condensed milk",
    "milk_powder": "Milk powder",
    "other": "Other",
    "hard_cheese": "CHEESE",
    "cottage_cheese": "CHEESE",
}

LITERAL_LABELS = {
    "milk": "Milk",
    "kefir": "Kefir",
    "ryazhanka": "Ryazhanka",
    "ayran": "Ayran / fermented milk drink",
    "yogurt": "Yogurt",
    "dessert": "Dairy dessert / glazed snack",
    "cottage_cheese": "Cheese / tvorog",
    "hard_cheese": "Cheese",
    "sour_cream": "Sour cream",
    "cream": "Cream",
    "butter": "Butter",
    "condensed_milk": "Condensed milk",
    "milk_powder": "Milk powder",
    "other_dairy": "Other",
}

PRODUCT_ORDER = [
    "milk",
    "butter",
    "cheese",
    "sour_cream",
    "cream",
    "yogurt_dessert",
    "condensed_milk",
    "milk_powder",
    "other",
]

RETAIL_STOPWORDS = {
    "молоко",
    "сир",
    "сметана",
    "вершки",
    "йогурт",
    "кефір",
    "ряжанка",
    "масло",
    "сирок",
    "кисломолочний",
    "питний",
    "солодковершкове",
    "ультрапастеризоване",
    "пастеризоване",
    "стерилізовані",
    "стерилізоване",
    "домашній",
    "ванночка",
    "стакан",
    "пляшка",
    "пакет",
    "пастоподібний",
    "білий",
    "натуральний",
    "термостатна",
    "термостатний",
    "екстра",
    "селянське",
    "селянська",
    "родинне",
    "особливе",
    "т",
    "б",
    "п",
    "е",
    "pet",
    "тетрапак",
    "шт",
}

GENERIC_BRANDS = {
    "",
    "nan",
    "продукт",
    "молочний",
    "напій",
    "сир",
    "яйця",
    "тістечко",
    "milk",
    "butter",
    "cream",
    "yogurt",
    "dessert",
    "cottage_cheese",
    "sour_cream",
}

PLANT_BASED_BRANDS = {
    "alpro",
    "oatly",
    "sojasun",
    "joya",
    "wanted",
    "wanted vegan",
    "green smile",
    "vega",
    "vega milk",
    "feels good",
}

HORIZONS = [0, 1, 3, 7, 14, 21]
MIN_LP_OBS = 32
MIN_MARGIN_OBS = 32
MIN_DISCOUNT_OBS = 30
MAX_HAC_LAG = 7
MIN_CHAIN_OBS = 24


def ensure_dirs() -> None:
    for path in [
        OUTPUT_DIR,
        DATA_DIR,
        FIG_DIR,
        DOC_DIR,
        LOG_DIR,
        MODEL_OUTPUT_DIR,
        MODEL_DIAG_DIR,
        MODEL_NOTE_DIR,
        VECM_DETAIL_DIR,
        FIG_SEQ_DIR,
        FIG_CH5_DIR,
        FIG_CH6_DIR,
        FIG_MODEL_DIR,
        FIG_RELIABILITY_DIR,
        Path(os.environ["MPLCONFIGDIR"]),
    ]:
        path.mkdir(parents=True, exist_ok=True)


def clear_generated_outputs() -> None:
    stale_dirs = [
        MODEL_OUTPUT_DIR,
        MODEL_DIAG_DIR,
        MODEL_NOTE_DIR,
        VECM_DETAIL_DIR,
        OUTPUT_DIR / "chapter_tables",
        FIG_SEQ_DIR,
        FIG_CH5_DIR,
        FIG_CH6_DIR,
        FIG_MODEL_DIR,
        FIG_RELIABILITY_DIR,
    ]
    for path in stale_dirs:
        if path.exists():
            shutil.rmtree(path)
        path.mkdir(parents=True, exist_ok=True)


def ntext(value: object) -> str:
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return ""
    text = str(value).strip().lower()
    text = text.replace("’", "'").replace("`", "'")
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_retail_text(value: object) -> str:
    text = ntext(value)
    text = text.replace("м’який", "мякий").replace("’", "'")
    text = re.sub(r"[%]", " % ", text)
    text = re.sub(r"[^\w\s%]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def normalize_brand_name(brand: object, entity: object = "") -> str:
    b = normalize_retail_text(brand)
    e = normalize_retail_text(entity)
    b = "" if b in GENERIC_BRANDS else b
    e = "" if e in GENERIC_BRANDS else e
    return b or e


def retail_full_text(*fields: object) -> str:
    return normalize_retail_text(" ".join(ntext(v) for v in fields if ntext(v)))


def harmonize_product_code(value: object) -> str:
    product = ntext(value)
    mapping = {
        "hard_cheese": "cheese",
        "cottage_cheese": "cheese",
        "other_dairy": "other",
        "інше": "other",
        "other": "other",
    }
    return mapping.get(product, product or "other")


def is_non_dairy_retail(text: str, brand_norm: str) -> bool:
    if brand_norm in PLANT_BASED_BRANDS:
        return True
    strong_patterns = [
        r"\bтофу\b",
        r"\bvegan\b",
        r"\bвеган",
        r"\bрослин",
        r"\bсоєв",
        r"\bsoy\b",
        r"\bsoya\b",
        r"\bвівся",
        r"\boat\b",
        r"\bмигдал",
        r"\balmond\b",
        r"\bbarista\b",
        r"сирно[\s-]?рослин",
        r"рослинно[\s-]?вершков",
    ]
    context_patterns = [
        r"snail.+кокосовому крем",
        r"кокосов.+рис",
        r"рис.+кокос",
        r"напій.+кокосов",
        r"напій.+рисов",
        r"напій.+мигдал",
    ]
    for pattern in strong_patterns + context_patterns:
        if re.search(pattern, text):
            return True
    return False


def literal_retail_type(*fields: object) -> str:
    text = retail_full_text(*fields)
    if not text:
        return "other_dairy"
    if any(token in text for token in ["яйц", "egg", "eggs"]):
        return "eggs"
    if any(token in text for token in ["молоко сух", "сухе молоко", "сухі вершки"]):
        return "milk_powder"
    if "згущ" in text:
        return "condensed_milk"
    if "масло" in text:
        return "butter"
    if "сметан" in text:
        return "sour_cream"
    if "вершк" in text:
        return "cream"
    if any(token in text for token in ["сирок", "десерт", "пудинг"]):
        return "dessert"
    if any(token in text for token in ["сир кисломолоч", "творог", "творож", "cottage cheese", "творожна"]):
        return "cottage_cheese"
    if any(token in text for token in ["йогурт", "скір", "skyr"]):
        return "yogurt"
    if "айран" in text:
        return "ayran"
    if "ряжан" in text:
        return "ryazhanka"
    if "кеф" in text:
        return "kefir"
    cheese_tokens = ["моцарел", "гауда", "gouda", "edam", "едам", "чеддер", "cheddar", "маасдам", "сулугун", "камамбер", "брі", "пармезан"]
    if "сир" in text or any(token in text for token in cheese_tokens):
        return "hard_cheese"
    if any(token in text for token in ["молоко", "молочний", "коктейль", "закваска"]):
        return "milk"
    return "other_dairy"


def retail_product_from_literal(literal_type: str) -> str:
    if literal_type in {"milk", "kefir", "ryazhanka", "ayran"}:
        return "milk"
    if literal_type in {"yogurt", "dessert"}:
        return "yogurt_dessert"
    if literal_type in {"cottage_cheese", "hard_cheese"}:
        return "cheese"
    if literal_type in {"sour_cream"}:
        return "sour_cream"
    if literal_type in {"cream"}:
        return "cream"
    if literal_type in {"butter"}:
        return "butter"
    if literal_type in {"condensed_milk"}:
        return "condensed_milk"
    if literal_type in {"milk_powder"}:
        return "milk_powder"
    if literal_type in {"other_dairy"}:
        return "other"
    return "other"


def extract_pack_std(row: pd.Series) -> Tuple[float, str]:
    qty = safe_num(pd.Series([row.get("qty_std", np.nan)])).iloc[0]
    unit = normalize_retail_text(row.get("unit_std", ""))
    if pd.notna(qty) and qty > 0 and unit in {"kg", "l", "liter"}:
        return float(qty), "kg" if unit == "kg" else "l"
    qty2 = safe_num(pd.Series([row.get("pack_qty_final", np.nan)])).iloc[0]
    unit2 = normalize_retail_text(row.get("pack_unit_final", ""))
    if pd.isna(qty2) or qty2 <= 0:
        return np.nan, ""
    if unit2 in {"г", "g", "гр"}:
        return float(qty2) / 1000.0, "kg"
    if unit2 in {"кг", "kg"}:
        return float(qty2), "kg"
    if unit2 in {"мл", "ml"}:
        return float(qty2) / 1000.0, "l"
    if unit2 in {"л", "l", "liter"}:
        return float(qty2), "l"
    return np.nan, ""


def canonical_item_name(product_name: object, product_title: object, brand_norm: str, product_type: str) -> str:
    text = normalize_retail_text(product_name or product_title)
    if brand_norm:
        for token in brand_norm.split():
            if len(token) >= 3:
                text = re.sub(rf"\b{re.escape(token)}\b", " ", text)
    text = re.sub(r"\b\d+[,.]?\d*\s*%\b", " ", text)
    text = re.sub(r"\b\d+[,.]?\d*\b", " ", text)
    tokens = []
    for tok in text.split():
        if tok in RETAIL_STOPWORDS:
            continue
        if len(tok) <= 1:
            continue
        tokens.append(tok)
    if not tokens:
        fallback = [t for t in normalize_retail_text(product_name).split() if t not in RETAIL_STOPWORDS]
        tokens = fallback[:5] if fallback else [product_type]
    return " ".join(tokens[:8]).strip()


def item_key(product: str, brand_norm: str, canonical_name: str) -> str:
    return "|".join([product or "na", brand_norm or "na", canonical_name or "na"])


def item_key_strict(product: str, brand_norm: str, canonical_name: str, fat_pct: object, pack_std: object, pack_unit: str) -> str:
    fat = safe_num(pd.Series([fat_pct])).iloc[0]
    fat_key = "na" if pd.isna(fat) else f"{float(fat):.1f}"
    pack_val = safe_num(pd.Series([pack_std])).iloc[0]
    pack_key = "na" if pd.isna(pack_val) else f"{float(pack_val):.3f}"
    return "|".join([product or "na", brand_norm or "na", canonical_name or "na", fat_key, pack_key, pack_unit or "na"])


def standardize_product(*fields: object) -> str:
    text = " | ".join(ntext(v) for v in fields if ntext(v))
    if not text:
        return "other"
    if "яйц" in text:
        return "eggs"
    if "сух" in text and "молок" in text:
        return "milk_powder"
    if "згущ" in text:
        return "condensed_milk"
    if "кисломолоч" in text or "творог" in text or "cottage" in text:
        return "cheese"
    if "масло" in text or "butter" in text:
        return "butter"
    if "сметан" in text or "sour cream" in text:
        return "sour_cream"
    if "вершк" in text or "cream" in text:
        return "cream"
    if "йогурт" in text or "десерт" in text or "yogurt" in text:
        return "yogurt_dessert"
    if "сир" in text or "cheese" in text or "gouda" in text or "edam" in text or "cheddar" in text:
        return "cheese"
    if "кеф" in text or "ряжан" in text or "молок" in text or "milk" in text:
        return "milk"
    return "other"


def as_date(series: pd.Series) -> pd.Series:
    return pd.to_datetime(series, errors="coerce").dt.normalize()


def safe_num(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def safe_log(series: pd.Series) -> pd.Series:
    values = safe_num(series)
    return np.log(values.where(values > 0))


def winsorize_by_group(df: pd.DataFrame, value_col: str, group_cols: Sequence[str], lo: float = 0.01, hi: float = 0.99) -> pd.Series:
    out = safe_num(df[value_col]).copy()
    for _, idx in df.groupby(list(group_cols), dropna=False).groups.items():
        vals = out.loc[idx].dropna()
        if len(vals) < 8:
            continue
        lower = vals.quantile(lo)
        upper = vals.quantile(hi)
        out.loc[idx] = out.loc[idx].clip(lower, upper)
    return out


def prioritized_products(products: Sequence[str], extras: Sequence[str], limit: int) -> List[str]:
    ordered: List[str] = []
    seen = set()
    for product in list(products) + list(extras):
        product = harmonize_product_code(product)
        if not product or product in seen:
            continue
        ordered.append(product)
        seen.add(product)
        if len(ordered) >= limit:
            break
    return ordered


def focus_products_from_counts(frame: pd.DataFrame, product_col: str = "product", limit: int = 6) -> List[str]:
    if frame.empty or product_col not in frame.columns:
        return []
    counts = frame[product_col].value_counts(dropna=True)
    main = [p for p in ["milk", "butter", "cheese"] if p in counts.index]
    extras = [p for p in counts.index.tolist() if p not in main]
    return prioritized_products(main, extras, limit)


def interpolate_panel(panel: pd.DataFrame, cols: Sequence[str], limits: Dict[str, int]) -> pd.DataFrame:
    out = panel.sort_values(["product", "date"]).copy()
    for col in cols:
        limit = limits.get(col, 0)
        model_col = f"{col}_model"
        if limit <= 0:
            out[model_col] = out[col]
            continue
        pieces = []
        for _, g in out[["product", "date", col]].groupby("product", dropna=False):
            s = g.set_index("date")[col].sort_index()
            s_interp = s.interpolate(method="time", limit=limit, limit_area="inside")
            pieces.append(s_interp.rename(model_col).reset_index().assign(product=g["product"].iloc[0]))
        interp = pd.concat(pieces, ignore_index=True) if pieces else pd.DataFrame(columns=["date", "product", model_col])
        out = out.drop(columns=[model_col], errors="ignore").merge(interp, on=["product", "date"], how="left")
    return out


def load_farm_gate() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="FarmGate_UA").copy()
    df["date"] = as_date(df["date"])
    df["price_linear"] = safe_num(df["price_linear"])
    df["price_pchip"] = safe_num(df["price_pchip"])
    df["region"] = df.get("region", "Україна").astype(str).str.strip()
    df["farm_type"] = df.get("farm_type", "Україна").astype(str).str.strip()
    df = df.dropna(subset=["date"])
    df = df[df["price_linear"].notna() | df["price_pchip"].notna()].copy()

    farm_map = {
        "Підприємства": "enterprise",
        "Господарства населення": "household",
    }
    df["farm_type_key"] = df["farm_type"].map(farm_map).fillna("other")
    national = df[df["region"].eq("Україна")].copy()
    if national.empty:
        national = (
            df.groupby(["date", "farm_type_key"], as_index=False)[["price_linear", "price_pchip"]]
            .median()
            .assign(region="regional_median")
        )
    else:
        national = national.groupby(["date", "farm_type_key"], as_index=False)[["price_linear", "price_pchip"]].median()

    wide = national.pivot_table(index="date", columns="farm_type_key", values=["price_linear", "price_pchip"], aggfunc="median")
    wide.columns = [f"farmgate_{farm_type}_{variant}" for variant, farm_type in wide.columns]
    wide = wide.reset_index().sort_values("date")
    rename_map = {}
    for col in wide.columns:
        if isinstance(col, str) and col.endswith("_price_linear"):
            rename_map[col] = col.replace("_price_linear", "_linear")
        elif isinstance(col, str) and col.endswith("_price_pchip"):
            rename_map[col] = col.replace("_price_pchip", "_pchip")
    if rename_map:
        wide = wide.rename(columns=rename_map)
    for variant in ["linear", "pchip"]:
        enterprise_col = f"farmgate_enterprise_{variant}"
        household_col = f"farmgate_household_{variant}"
        if enterprise_col in wide.columns and household_col in wide.columns:
            wide[f"farmgate_combined_{variant}"] = wide[[enterprise_col, household_col]].median(axis=1, skipna=True)
        elif enterprise_col in wide.columns:
            wide[f"farmgate_combined_{variant}"] = wide[enterprise_col]
        elif household_col in wide.columns:
            wide[f"farmgate_combined_{variant}"] = wide[household_col]
        else:
            wide[f"farmgate_combined_{variant}"] = np.nan
    return wide


def load_producer() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="Producer_UA")
    df["date"] = as_date(df["date"])
    df["product"] = df["ua_product"].apply(standardize_product).apply(harmonize_product_code)
    df["producer_linear"] = safe_num(df["price_linear"])
    df["producer_pchip"] = safe_num(df["price_pchip"])
    df = df[df["product"].ne("eggs")]
    out = (
        df.groupby(["date", "product"], as_index=False)[["producer_linear", "producer_pchip"]]
        .median()
        .sort_values(["product", "date"])
    )
    return out


def load_prozorro() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="Prozorro")
    col_map = {}
    for target, options in {
        "date": ["Дата", "date"],
        "raw_product": ["Product", "product_ua", "product", "title_ua", "Профіль", "Товар"],
        "unit": ["Одиниця виміру", "unit_ua", "unit_en", "unit"],
        "unit_price": ["Ціна за одиницю", "unit_price", "unit price"],
        "region": ["Регіон організатора", "organizer_region_ua", "organizer_region_en", "region"],
        "expected": ["Очікувана вартість", "expected_value", "expected"],
        "sum_initial": ["Сума договорів початкова", "contract_value_initial", "sum_initial"],
        "sum_current": ["Сума договорів поточна", "contract_value_current", "sum_current"],
    }.items():
        for opt in options:
            if opt in df.columns:
                col_map[opt] = target
                break
    df = df.rename(columns=col_map)
    if "date" not in df.columns:
        df["date"] = df.get("Дата", df.get("date", pd.NaT))
    if "raw_product" not in df.columns:
        raw_product = df.get("product_ua")
        if raw_product is None:
            raw_product = df.get("title_ua")
        if raw_product is None:
            raw_product = df.get("product")
        if raw_product is None:
            product_like = [c for c in df.columns if "product" in str(c).lower() or "товар" in str(c).lower()]
            raw_product = df[product_like[0]] if product_like else pd.Series("", index=df.index)
        df["raw_product"] = raw_product
    if "unit" not in df.columns:
        df["unit"] = df.get("unit_ua", df.get("unit_en", pd.Series("", index=df.index)))
    if "unit_price" not in df.columns:
        df["unit_price"] = df.get("Ціна за одиницю", df.get("unit price", np.nan))
    if "region" not in df.columns:
        df["region"] = df.get("organizer_region_ua", df.get("organizer_region_en", pd.Series("", index=df.index)))
    if "expected" not in df.columns:
        df["expected"] = df.get("expected_value", np.nan)
    if "sum_initial" not in df.columns:
        df["sum_initial"] = df.get("contract_value_initial", np.nan)
    if "sum_current" not in df.columns:
        df["sum_current"] = df.get("contract_value_current", np.nan)
    df["date"] = as_date(df["date"])
    df["product"] = df["raw_product"].apply(standardize_product).apply(harmonize_product_code)
    df["unit_norm"] = df["unit"].astype(str).str.lower().str.strip()
    df["price"] = safe_num(df["unit_price"])
    df["expected"] = safe_num(df.get("expected", pd.Series(np.nan, index=df.index)))
    df["sum_initial"] = safe_num(df.get("sum_initial", pd.Series(np.nan, index=df.index)))
    df["sum_current"] = safe_num(df.get("sum_current", pd.Series(np.nan, index=df.index)))
    unit_ok = df["unit_norm"].str.contains("кілограм|кг|kg|літр|л", regex=True, na=False)
    df = df[unit_ok & df["product"].ne("eggs") & df["price"].gt(0)].copy()
    df["price"] = winsorize_by_group(df, "price", ["product"])
    out = (
        df.groupby(["date", "product"], as_index=False)
        .agg(
            prozorro=("price", "median"),
            prozorro_n=("price", "size"),
            prozorro_regions=("region", pd.Series.nunique),
            prozorro_expected=("expected", "sum"),
            prozorro_sum_initial=("sum_initial", "sum"),
            prozorro_sum_current=("sum_current", "sum"),
        )
        .sort_values(["product", "date"])
    )
    return out


def load_prozorro_region_profile() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="Prozorro")
    rename_map = {}
    for target, options in {
        "date": ["Дата", "date"],
        "raw_product": ["Product", "product_ua", "product", "title_ua", "Профіль", "Товар"],
        "region": ["Регіон організатора", "organizer_region_ua", "organizer_region_en", "region"],
        "unit": ["Одиниця виміру", "unit_ua", "unit_en", "unit"],
        "unit_price": ["Ціна за одиницю", "unit_price", "unit price"],
    }.items():
        for opt in options:
            if opt in df.columns:
                rename_map[opt] = target
                break
    df = df.rename(columns=rename_map)
    if "raw_product" not in df.columns:
        df["raw_product"] = df.get("product_ua", df.get("title_ua", df.get("product", "")))
    if "region" not in df.columns:
        df["region"] = df.get("organizer_region_ua", df.get("organizer_region_en", ""))
    if "unit" not in df.columns:
        df["unit"] = df.get("unit_ua", df.get("unit_en", ""))
    if "unit_price" not in df.columns:
        df["unit_price"] = df.get("Ціна за одиницю", df.get("unit price", np.nan))
    df["date"] = as_date(df.get("date", df.get("Дата", pd.NaT)))
    df["product"] = df["raw_product"].apply(standardize_product).apply(harmonize_product_code)
    df["region"] = df["region"].astype(str).str.strip()
    df["unit_norm"] = df["unit"].astype(str).str.lower().str.strip()
    df["unit_price"] = safe_num(df["unit_price"])
    unit_ok = df["unit_norm"].str.contains("кілограм|кг|kg|літр|л", regex=True, na=False)
    df = df[unit_ok & df["unit_price"].gt(0) & df["product"].ne("eggs")].copy()
    if df.empty:
        return df
    df["unit_price"] = winsorize_by_group(df, "unit_price", ["product", "region"])
    out = (
        df.groupby(["product", "region"], as_index=False)
        .agg(
            n_obs=("unit_price", "size"),
            n_dates=("date", pd.Series.nunique),
            median_price=("unit_price", "median"),
            mean_price=("unit_price", "mean"),
            sd_price=("unit_price", "std"),
            date_min=("date", "min"),
            date_max=("date", "max"),
        )
        .sort_values(["product", "n_obs", "region"], ascending=[True, False, True])
    )
    out["product_label"] = out["product"].map(PRODUCT_LABELS).fillna(out["product"])
    return out


def load_consumer() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="Consumer_UA")
    df["date"] = as_date(df["date"])
    df["product"] = df["ua_product"].apply(standardize_product).apply(harmonize_product_code)
    df["consumer_linear"] = safe_num(df["price_linear"])
    df["consumer_pchip"] = safe_num(df["price_pchip"])
    df = df[df["product"].ne("eggs")].copy()
    out = (
        df.groupby(["date", "product"], as_index=False)[["consumer_linear", "consumer_pchip"]]
        .median()
        .sort_values(["product", "date"])
    )
    return out


def standardize_eu_product(value: object) -> str:
    text = ntext(value)
    if any(token in text for token in ["drinking milk", "whole milk", "skimmed milk", "milk"]):
        return "milk"
    if "butter" in text:
        return "butter"
    if any(token in text for token in ["emmental", "edam", "gouda", "cheddar", "cheese"]):
        return "hard_cheese"
    if "cream" in text:
        return "cream"
    if any(token in text for token in ["smp", "wmp", "milk powder"]):
        return "milk_powder"
    return "other"


def load_cme_benchmark() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="CME III").copy()
    date_col = "Date" if "Date" in df.columns else df.columns[0]
    value_col = "CME III UAH" if "CME III UAH" in df.columns else df.columns[1]
    df["date"] = as_date(df[date_col])
    df["cme_class3_uah"] = safe_num(df[value_col])
    df = df[["date", "cme_class3_uah"]].dropna().sort_values("date")
    return df


def load_europe_benchmark() -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name="Europe").copy()
    df["date"] = as_date(df["date"])
    df["product"] = df["Product"].apply(standardize_eu_product)
    df["country"] = df["Country"].astype(str).str.strip()
    df["eu_uah_kg"] = safe_num(df["Price (UAH/kg)"])
    df = df[df["product"].ne("other") & df["eu_uah_kg"].gt(0)].copy()
    out = (
        df.groupby(["date", "product"], as_index=False)
        .agg(
            eu_price_uah=("eu_uah_kg", "median"),
            eu_country_n=("country", pd.Series.nunique),
        )
        .sort_values(["product", "date"])
    )
    return out


def load_retail_items(sheet: str, product_col: str, retailer: str) -> pd.DataFrame:
    df = pd.read_excel(FULL_UAH, sheet_name=sheet).copy()
    if product_col not in df.columns:
        for fallback in ["product", "product_ua", "Product", "product_ua.1"]:
            if fallback in df.columns:
                product_col = fallback
                break
    df["date"] = as_date(df["date"])
    df["retailer"] = retailer
    df["brand_norm"] = [normalize_brand_name(b, e) for b, e in zip(df.get("brand", ""), df.get("entity", ""))]
    df["raw_sheet_product"] = df[product_col].apply(ntext)
    df["retail_text"] = [
        retail_full_text(sheet_prod, pn, pt, b, e)
        for sheet_prod, pn, pt, b, e in zip(
            df[product_col],
            df.get("product_name", ""),
            df.get("product_title", ""),
            df.get("brand", ""),
            df.get("entity", ""),
        )
    ]
    literals = []
    for sheet_prod, pn, pt, b, e in zip(
        df[product_col],
        df.get("product_name", ""),
        df.get("product_title", ""),
        df.get("brand", ""),
        df.get("entity", ""),
    ):
        literal = literal_retail_type(pn, pt, b, e)
        if literal == "other_dairy":
            literal = literal_retail_type(sheet_prod, pn, pt, b, e)
        literals.append(literal)
    df["product_literal"] = literals
    df["product"] = df["product_literal"].map(retail_product_from_literal).apply(harmonize_product_code)
    df["non_dairy_flag"] = [int(is_non_dairy_retail(text, brand)) for text, brand in zip(df["retail_text"], df["brand_norm"])]
    df = df[df["non_dairy_flag"].eq(0) & df["product"].ne("eggs")].copy()
    df["effective_price"] = safe_num(df.get("price_current", pd.Series(np.nan, index=df.index)))
    raw_discount_amount = safe_num(df.get("discount_value", pd.Series(np.nan, index=df.index)))
    df["discount_dummy_discount"] = safe_num(df.get("discount_dummy_discount", pd.Series(0, index=df.index))).fillna(0).astype(int)
    df["discount_dummy_bulk"] = safe_num(df.get("discount_dummy_bulk", pd.Series(0, index=df.index))).fillna(0).astype(int)
    df["discount_dummy_regular"] = safe_num(df.get("discount_dummy_regular", pd.Series(0, index=df.index))).fillna(0).astype(int)
    explicit_discount = safe_num(df.get("discount_%", pd.Series(np.nan, index=df.index)))
    valid_discount_pct = explicit_discount.where(explicit_discount.between(0, 95))
    df["discount_present"] = (
        df["discount_dummy_discount"].eq(1)
        | df["discount_dummy_bulk"].eq(1)
        | raw_discount_amount.fillna(0).gt(0)
        | valid_discount_pct.fillna(0).gt(0)
    ).astype(int)
    df["discount_type"] = np.select(
        [df["discount_dummy_bulk"].eq(1), df["discount_dummy_discount"].eq(1), df["discount_present"].eq(0)],
        ["bulk", "markdown", "regular"],
        default="unknown",
    )
    inferred_baseline = pd.Series(
        np.where(
            valid_discount_pct.notna() & df["effective_price"].gt(0),
            df["effective_price"] / (1.0 - valid_discount_pct / 100.0),
            np.nan,
        ),
        index=df.index,
    )
    df["discount_amount"] = raw_discount_amount.copy()
    need_amount_from_pct = df["discount_amount"].fillna(0).le(0) & inferred_baseline.notna()
    df.loc[need_amount_from_pct, "discount_amount"] = inferred_baseline.loc[need_amount_from_pct] - df.loc[need_amount_from_pct, "effective_price"]
    df["discount_amount"] = df["discount_amount"].fillna(0)
    df["baseline_price"] = np.where(df["discount_amount"].gt(0), df["effective_price"] + df["discount_amount"], df["effective_price"])
    df["markdown_rate"] = np.where(df["baseline_price"].gt(0), (df["baseline_price"] - df["effective_price"]) / df["baseline_price"], np.nan)
    df["discount_percent"] = valid_discount_pct.where(valid_discount_pct.notna(), df["markdown_rate"] * 100.0)
    packs = df.apply(extract_pack_std, axis=1, result_type="expand")
    df["pack_std"] = packs[0]
    df["pack_std_unit"] = packs[1]
    df["unit_effective_price"] = safe_num(df.get("unit_price", pd.Series(np.nan, index=df.index)))
    need_fallback = df["unit_effective_price"].isna() & df["pack_std"].gt(0)
    df.loc[need_fallback, "unit_effective_price"] = df.loc[need_fallback, "effective_price"] / df.loc[need_fallback, "pack_std"]
    df["unit_baseline_price"] = np.where(df["pack_std"].gt(0), df["baseline_price"] / df["pack_std"], np.nan)
    df["fat_pct_clean"] = safe_num(df.get("fat_pct", pd.Series(np.nan, index=df.index)))
    df["canonical_name"] = [
        canonical_item_name(pn, pt, bn, prod)
        for pn, pt, bn, prod in zip(df.get("product_name", ""), df.get("product_title", ""), df["brand_norm"], df["product"])
    ]
    df["item_key"] = [item_key(prod, bn, cn) for prod, bn, cn in zip(df["product"], df["brand_norm"], df["canonical_name"])]
    df["item_key_strict"] = [
        item_key_strict(prod, bn, cn, fat, pack, unit)
        for prod, bn, cn, fat, pack, unit in zip(
            df["product"], df["brand_norm"], df["canonical_name"], df["fat_pct_clean"], df["pack_std"], df["pack_std_unit"]
        )
    ]
    keep = [
        "retailer",
        "date",
        "product",
        "product_literal",
        "brand_norm",
        "canonical_name",
        "item_key",
        "item_key_strict",
        "raw_sheet_product",
        "retail_text",
        "product_title",
        "product_name",
        "brand",
        "entity",
        "fat_pct_clean",
        "pack_std",
        "pack_std_unit",
        "effective_price",
        "baseline_price",
        "unit_effective_price",
        "unit_baseline_price",
        "discount_amount",
        "discount_percent",
        "discount_present",
        "discount_type",
        "markdown_rate",
        "discount_dummy_discount",
        "discount_dummy_bulk",
        "discount_dummy_regular",
    ]
    df = df[keep].copy()
    df = df[df["effective_price"].gt(0)].copy()
    df["effective_price"] = winsorize_by_group(df, "effective_price", ["product", "retailer"])
    df["baseline_price"] = winsorize_by_group(df, "baseline_price", ["product", "retailer"])
    df["unit_effective_price"] = df["unit_effective_price"].where(df["unit_effective_price"].between(0.5, 5000))
    df["unit_baseline_price"] = df["unit_baseline_price"].where(df["unit_baseline_price"].between(0.5, 5000))
    return df.sort_values(["retailer", "date", "product", "item_key"])


def build_retail_catalog_and_matches(silpo_items: pd.DataFrame, novus_items: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    all_items = pd.concat([silpo_items, novus_items], ignore_index=True)
    catalog = (
        all_items.groupby(["retailer", "item_key"], as_index=False)
        .agg(
            product=("product", "first"),
            brand_norm=("brand_norm", "first"),
            canonical_name=("canonical_name", "first"),
            n_strict_keys=("item_key_strict", pd.Series.nunique),
            fat_pct=("fat_pct_clean", "median"),
            pack_std=("pack_std", "median"),
            pack_std_unit=("pack_std_unit", "first"),
            first_title=("product_title", "first"),
            first_name=("product_name", "first"),
            n_rows=("effective_price", "size"),
            date_min=("date", "min"),
            date_max=("date", "max"),
        )
        .sort_values(["retailer", "product", "brand_norm", "canonical_name"])
    )
    sil = catalog[catalog["retailer"].eq("Silpo")].drop(columns=["retailer"]).rename(columns=lambda c: f"silpo_{c}" if c != "item_key" else c)
    nov = catalog[catalog["retailer"].eq("Novus")].drop(columns=["retailer"]).rename(columns=lambda c: f"novus_{c}" if c != "item_key" else c)
    audit = sil.merge(nov, on="item_key", how="outer", indicator=True)
    audit["match_status"] = audit["_merge"].map({"both": "matched_both_shops", "left_only": "silpo_only", "right_only": "novus_only"})
    audit["strict_alignment_flag"] = np.where(
        audit["match_status"].eq("matched_both_shops")
        & audit["silpo_n_strict_keys"].fillna(0).eq(1)
        & audit["novus_n_strict_keys"].fillna(0).eq(1),
        1,
        0,
    )
    audit = audit.drop(columns=["_merge"])
    return catalog, audit


def aggregate_retail_items(items: pd.DataFrame, prefix: str = "") -> pd.DataFrame:
    grouped = (
        items.groupby(["date", "product"], as_index=False)
        .agg(
            observed=("effective_price", "median"),
            baseline=("baseline_price", "median"),
            discount_share=("discount_present", "mean"),
            discount_depth=("markdown_rate", "median"),
            discount_percent=("discount_percent", "median"),
            discount_amount_median=("discount_amount", "median"),
            n_items=("effective_price", "size"),
            n_item_keys=("item_key", pd.Series.nunique),
            n_brands=("brand_norm", pd.Series.nunique),
            n_literal_types=("product_literal", pd.Series.nunique),
        )
        .sort_values(["product", "date"])
    )
    grouped["discount_depth"] = grouped["discount_depth"].fillna(0)
    grouped["category_observed"] = grouped["observed"]
    grouped["category_baseline"] = grouped["baseline"]
    if prefix:
        rename = {c: f"{prefix}_{c}" for c in grouped.columns if c not in {"date", "product"}}
        grouped = grouped.rename(columns=rename)
    return grouped


def build_consumer_linked_retail(retail: pd.DataFrame, consumer: pd.DataFrame) -> pd.DataFrame:
    merged = retail.merge(consumer, on=["date", "product"], how="left")
    merged["consumer_support_flag"] = merged.get("consumer_linear", pd.Series(np.nan, index=merged.index)).notna().astype(int)
    merged["retail_vs_consumer_gap"] = safe_log(merged.get("retail_category_observed", pd.Series(np.nan, index=merged.index))) - safe_log(
        merged.get("consumer_linear", pd.Series(np.nan, index=merged.index))
    )
    keep_cols = list(retail.columns) + ["consumer_support_flag", "retail_vs_consumer_gap"]
    existing = [c for c in keep_cols if c in merged.columns]
    return merged[existing]


def build_retail_brand_panels(items: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    brand_daily = (
        items.groupby(["date", "product", "product_literal", "retailer", "brand_norm"], as_index=False)
        .agg(
            observed=("effective_price", "median"),
            baseline=("baseline_price", "median"),
            discount_share=("discount_present", "mean"),
            markdown_rate=("markdown_rate", "median"),
            n_items=("effective_price", "size"),
            n_item_keys=("item_key", pd.Series.nunique),
        )
        .sort_values(["product", "retailer", "brand_norm", "date"])
    )
    brand_support = (
        items.groupby(["product", "product_literal", "retailer", "brand_norm"], as_index=False)
        .agg(
            n_rows=("effective_price", "size"),
            n_dates=("date", pd.Series.nunique),
            n_item_keys=("item_key", pd.Series.nunique),
            matched_share=("matched_cross_shop", "mean"),
            mean_discount_share=("discount_present", "mean"),
            median_price=("effective_price", "median"),
        )
        .sort_values(["product", "retailer", "n_rows", "n_dates"], ascending=[True, True, False, False])
    )
    return brand_daily, brand_support


def build_retail_literal_summary(items: pd.DataFrame) -> pd.DataFrame:
    return (
        items.groupby(["product", "product_literal", "retailer"], as_index=False)
        .agg(
            n_rows=("effective_price", "size"),
            n_dates=("date", pd.Series.nunique),
            n_item_keys=("item_key", pd.Series.nunique),
            n_brands=("brand_norm", pd.Series.nunique),
            matched_share=("matched_cross_shop", "mean"),
            mean_discount_share=("discount_present", "mean"),
            median_price=("effective_price", "median"),
        )
        .sort_values(["product", "retailer", "n_item_keys"], ascending=[True, True, False])
    )


def build_reconciliation_examples(audit: pd.DataFrame) -> pd.DataFrame:
    if audit.empty:
        return pd.DataFrame()
    matched = audit[audit["match_status"].eq("matched_both_shops")].copy()
    if matched.empty:
        return matched
    matched["combined_rows"] = matched["silpo_n_rows"].fillna(0) + matched["novus_n_rows"].fillna(0)
    cols = [
        "item_key",
        "silpo_product",
        "silpo_brand_norm",
        "silpo_canonical_name",
        "silpo_first_name",
        "novus_brand_norm",
        "novus_canonical_name",
        "novus_first_name",
        "silpo_n_rows",
        "novus_n_rows",
        "strict_alignment_flag",
        "combined_rows",
    ]
    return matched[cols].sort_values(["strict_alignment_flag", "combined_rows"], ascending=[False, False]).head(250)


def retail_candidate_specs(include_optimal: bool = False) -> List[Dict[str, str]]:
    specs = [
        {
            "candidate": "retail_merged",
            "label": "Retail merged full list",
            "observed_col": "retail_category_observed_model",
            "baseline_col": "retail_category_baseline_model",
            "discount_col": "retail_discount_share",
            "support_col": "retail_n_item_keys",
            "link_forward": "ProZorro -> Retail merged",
            "link_reverse": "Retail merged -> ProZorro",
        },
        {
            "candidate": "retail_matched",
            "label": "Retail matched cross-shop",
            "observed_col": "retail_matched_category_observed_model",
            "baseline_col": "retail_matched_category_baseline_model",
            "discount_col": "retail_matched_discount_share",
            "support_col": "retail_matched_n_item_keys",
            "link_forward": "ProZorro -> Retail matched",
            "link_reverse": "Retail matched -> ProZorro",
        },
        {
            "candidate": "silpo",
            "label": "Silpo only",
            "observed_col": "silpo_category_observed_model",
            "baseline_col": "silpo_category_baseline_model",
            "discount_col": "silpo_discount_share",
            "support_col": "silpo_n_item_keys",
            "link_forward": "ProZorro -> Silpo",
            "link_reverse": "Silpo -> ProZorro",
        },
        {
            "candidate": "novus",
            "label": "Novus only",
            "observed_col": "novus_category_observed_model",
            "baseline_col": "novus_category_baseline_model",
            "discount_col": "novus_discount_share",
            "support_col": "novus_n_item_keys",
            "link_forward": "ProZorro -> Novus",
            "link_reverse": "Novus -> ProZorro",
        },
    ]
    if include_optimal:
        specs.append(
            {
                "candidate": "retail_optimal",
                "label": "Optimal retail level",
                "observed_col": "retail_optimal_observed_model",
                "baseline_col": "retail_optimal_baseline_model",
                "discount_col": "retail_optimal_discount_share",
                "support_col": "retail_optimal_n_item_keys",
                "link_forward": "ProZorro -> Retail optimal",
                "link_reverse": "Retail optimal -> ProZorro",
            }
        )
    return specs


def best_abs_lag_corr(x: pd.Series, y: pd.Series, max_lag: int = 28) -> Tuple[float, float, int]:
    best_corr = np.nan
    best_lag = np.nan
    best_n = 0
    lx = safe_log(x)
    ly = safe_log(y)
    for lag in range(0, max_lag + 1):
        d = pd.concat([lx.shift(lag).rename("x"), ly.rename("y")], axis=1).dropna()
        if len(d) < 20:
            continue
        corr = d["x"].corr(d["y"])
        if pd.isna(corr):
            continue
        if pd.isna(best_corr) or abs(corr) > abs(best_corr):
            best_corr = float(corr)
            best_lag = float(lag)
            best_n = int(len(d))
    return best_corr, best_lag, best_n


def select_optimal_retail_levels(panel: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    records: List[Dict[str, object]] = []
    for product, g in panel.groupby("product"):
        g = g.sort_values("date").copy()
        for spec in retail_candidate_specs(include_optimal=False):
            obs_col = spec["observed_col"]
            support_col = spec["support_col"]
            disc_col = spec["discount_col"]
            if obs_col not in g.columns:
                continue
            n_obs = int(g[obs_col].notna().sum())
            if n_obs < MIN_LP_OBS:
                continue
            prozorro_series = g["prozorro_model"] if "prozorro_model" in g.columns else pd.Series(np.nan, index=g.index)
            consumer_series = g["consumer_linear_model"] if "consumer_linear_model" in g.columns else pd.Series(np.nan, index=g.index)
            corr_p, lag_p, n_p = best_abs_lag_corr(prozorro_series, g[obs_col], max_lag=28)
            corr_c, lag_c, n_c = best_abs_lag_corr(consumer_series, g[obs_col], max_lag=14)
            discount_std = float(safe_num(g.get(disc_col, pd.Series(np.nan, index=g.index))).std(skipna=True))
            median_support = float(safe_num(g.get(support_col, pd.Series(np.nan, index=g.index))).median(skipna=True))
            records.append(
                {
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "candidate": spec["candidate"],
                    "candidate_label": spec["label"],
                    "observed_col": obs_col,
                    "baseline_col": spec["baseline_col"],
                    "discount_col": disc_col,
                    "support_col": support_col,
                    "n_obs": n_obs,
                    "coverage_share": float(g[obs_col].notna().mean()),
                    "best_corr_prozorro": corr_p,
                    "best_lag_prozorro": lag_p,
                    "best_corr_consumer": corr_c,
                    "best_lag_consumer": lag_c,
                    "discount_std": discount_std,
                    "median_item_support": median_support,
                    "overlap_prozorro_n": n_p,
                    "overlap_consumer_n": n_c,
                }
            )
    score_df = pd.DataFrame(records)
    if score_df.empty:
        return score_df, score_df

    for col in ["coverage_share", "best_corr_prozorro", "best_corr_consumer", "discount_std", "median_item_support"]:
        base = score_df[col].abs() if "corr" in col else score_df[col]
        denom = base.max()
        score_df[f"{col}_scaled"] = 0.0 if pd.isna(denom) or denom <= 0 else (base / denom).fillna(0)

    score_df["selection_score"] = (
        0.30 * score_df["coverage_share_scaled"]
        + 0.30 * score_df["best_corr_prozorro_scaled"]
        + 0.15 * score_df["best_corr_consumer_scaled"]
        + 0.15 * score_df["median_item_support_scaled"]
        + 0.10 * score_df["discount_std_scaled"]
    )
    chosen = (
        score_df.sort_values(
            ["product", "selection_score", "best_corr_prozorro_scaled", "coverage_share_scaled"],
            ascending=[True, False, False, False],
        )
        .groupby("product", as_index=False)
        .head(1)
        .reset_index(drop=True)
    )
    chosen["rank_note"] = "Highest composite score across coverage, procurement alignment, consumer alignment, item support, and discount variation."
    return score_df.sort_values(["product", "selection_score"], ascending=[True, False]), chosen


def build_panel() -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, pd.DataFrame]]:
    farm = load_farm_gate()
    producer = load_producer()
    prozorro = load_prozorro()
    prozorro_region_profile = load_prozorro_region_profile()
    consumer = load_consumer()
    europe = load_europe_benchmark()
    cme = load_cme_benchmark()
    silpo_items = load_retail_items("Silpo", "Product", "Silpo")
    novus_items = load_retail_items("Novus", "product", "Novus")
    retail_catalog, retail_match_audit = build_retail_catalog_and_matches(silpo_items, novus_items)
    matched_keys = set(retail_match_audit.loc[retail_match_audit["match_status"].eq("matched_both_shops"), "item_key"])
    silpo_items["matched_cross_shop"] = silpo_items["item_key"].isin(matched_keys).astype(int)
    novus_items["matched_cross_shop"] = novus_items["item_key"].isin(matched_keys).astype(int)
    retail_items_full = pd.concat([silpo_items, novus_items], ignore_index=True).sort_values(["date", "product", "retailer", "item_key"])
    retail_brand_daily, retail_brand_support = build_retail_brand_panels(retail_items_full)
    retail_literal_summary = build_retail_literal_summary(retail_items_full)
    reconciliation_examples = build_reconciliation_examples(retail_match_audit)
    silpo = aggregate_retail_items(silpo_items, prefix="silpo")
    novus = aggregate_retail_items(novus_items, prefix="novus")
    retail = aggregate_retail_items(retail_items_full, prefix="retail")
    matched_only = aggregate_retail_items(retail_items_full[retail_items_full["matched_cross_shop"].eq(1)].copy(), prefix="retail_matched")
    retail = retail.merge(matched_only, on=["date", "product"], how="left")
    retail = retail.merge(silpo, on=["date", "product"], how="left")
    retail = retail.merge(novus, on=["date", "product"], how="left")
    retail = build_consumer_linked_retail(retail, consumer)

    products = sorted(
        set(producer["product"]).union(prozorro["product"]).union(retail["product"]).union(consumer["product"]) - {"eggs"}
    )
    min_date = min(producer["date"].min(), farm["date"].min(), prozorro["date"].min(), retail["date"].min(), consumer["date"].min())
    max_date = max(producer["date"].max(), farm["date"].max(), prozorro["date"].max(), retail["date"].max(), consumer["date"].max())
    grid = pd.MultiIndex.from_product(
        [pd.date_range(min_date, max_date, freq="D"), products],
        names=["date", "product"],
    ).to_frame(index=False)

    panel = grid.merge(farm, on="date", how="left")
    panel = panel.merge(producer, on=["date", "product"], how="left")
    panel = panel.merge(prozorro, on=["date", "product"], how="left")
    panel = panel.merge(consumer, on=["date", "product"], how="left")
    panel = panel.merge(retail, on=["date", "product"], how="left")
    panel["product_label"] = panel["product"].map(PRODUCT_LABELS).fillna(panel["product"])
    farm_aliases = {
        "farmgate_initial_linear": "farmgate_enterprise_linear",
        "farmgate_initial_pchip": "farmgate_enterprise_pchip",
        "farmgate_filled_linear": "farmgate_household_linear",
        "farmgate_filled_pchip": "farmgate_household_pchip",
    }
    for alias, source_col in farm_aliases.items():
        if source_col in panel.columns:
            panel[alias] = panel[source_col]
        else:
            panel[alias] = np.nan
    if "farmgate_combined_linear" not in panel.columns:
        panel["farmgate_combined_linear"] = panel[["farmgate_initial_linear", "farmgate_filled_linear"]].median(axis=1, skipna=True)
    if "farmgate_combined_pchip" not in panel.columns:
        panel["farmgate_combined_pchip"] = panel[["farmgate_initial_pchip", "farmgate_filled_pchip"]].median(axis=1, skipna=True)

    model_limits = {
        "farmgate_combined_linear": 7,
        "farmgate_combined_pchip": 7,
        "prozorro": 14,
        "prozorro_expected": 14,
        "prozorro_sum_initial": 14,
        "prozorro_sum_current": 14,
        "consumer_linear": 7,
        "consumer_pchip": 7,
        "retail_observed": 3,
        "retail_baseline": 3,
        "retail_category_observed": 3,
        "retail_category_baseline": 3,
        "retail_matched_observed": 3,
        "retail_matched_baseline": 3,
        "retail_matched_category_observed": 3,
        "retail_matched_category_baseline": 3,
        "silpo_observed": 3,
        "silpo_baseline": 3,
        "silpo_category_observed": 3,
        "silpo_category_baseline": 3,
        "novus_observed": 3,
        "novus_category_observed": 3,
        "novus_category_baseline": 3,
    }
    price_cols = [
        "farmgate_initial_linear",
        "farmgate_initial_pchip",
        "farmgate_filled_linear",
        "farmgate_filled_pchip",
        "farmgate_combined_linear",
        "farmgate_combined_pchip",
        "producer_linear",
        "producer_pchip",
        "consumer_linear",
        "consumer_pchip",
        "prozorro",
        "prozorro_expected",
        "prozorro_sum_initial",
        "prozorro_sum_current",
        "retail_observed",
        "retail_baseline",
        "retail_category_observed",
        "retail_category_baseline",
        "retail_matched_observed",
        "retail_matched_baseline",
        "retail_matched_category_observed",
        "retail_matched_category_baseline",
        "silpo_observed",
        "silpo_baseline",
        "silpo_category_observed",
        "silpo_category_baseline",
        "novus_observed",
        "novus_category_observed",
        "novus_category_baseline",
    ]
    panel = interpolate_panel(panel, [c for c in price_cols if c in panel.columns], model_limits)

    level_scores, level_chosen = select_optimal_retail_levels(panel)
    if not level_chosen.empty:
        panel = panel.merge(
            level_chosen[["product", "candidate", "observed_col", "baseline_col", "discount_col", "support_col"]],
            on="product",
            how="left",
        )
        panel["retail_optimal_observed"] = np.nan
        panel["retail_optimal_baseline"] = np.nan
        panel["retail_optimal_observed_model"] = np.nan
        panel["retail_optimal_baseline_model"] = np.nan
        panel["retail_optimal_discount_share"] = np.nan
        panel["retail_optimal_n_item_keys"] = np.nan
        for row in level_chosen.itertuples():
            mask = panel["product"].eq(row.product)
            raw_observed_col = row.observed_col.replace("_model", "")
            raw_baseline_col = row.baseline_col.replace("_model", "")
            if raw_observed_col in panel.columns:
                panel.loc[mask, "retail_optimal_observed"] = panel.loc[mask, raw_observed_col]
            if raw_baseline_col in panel.columns:
                panel.loc[mask, "retail_optimal_baseline"] = panel.loc[mask, raw_baseline_col]
            panel.loc[mask, "retail_optimal_observed_model"] = panel.loc[mask, row.observed_col]
            panel.loc[mask, "retail_optimal_baseline_model"] = panel.loc[mask, row.baseline_col]
            panel.loc[mask, "retail_optimal_discount_share"] = panel.loc[mask, row.discount_col] if row.discount_col in panel.columns else np.nan
            panel.loc[mask, "retail_optimal_n_item_keys"] = panel.loc[mask, row.support_col] if row.support_col in panel.columns else np.nan
        panel = panel.drop(columns=["candidate", "observed_col", "baseline_col", "discount_col", "support_col"], errors="ignore")

    source_frames = {
        "farm_gate": farm,
        "producer": producer,
        "consumer": consumer,
        "europe_benchmark": europe,
        "cme_benchmark": cme,
        "prozorro": prozorro,
        "prozorro_region_profile": prozorro_region_profile,
        "silpo_daily": silpo,
        "novus_daily": novus,
        "retail_combined": retail,
        "retail_items_full": retail_items_full,
        "retail_catalog": retail_catalog,
        "retail_match_audit": retail_match_audit,
        "retail_brand_daily": retail_brand_daily,
        "retail_brand_support": retail_brand_support,
        "retail_literal_summary": retail_literal_summary,
        "retail_name_reconciliation": reconciliation_examples,
        "retail_level_scores": level_scores,
        "retail_level_selection": level_chosen,
    }
    inventory_rows = []
    for name, df in source_frames.items():
        inventory_rows.append(
            {
                "source": name,
                "rows": len(df),
                "products": df["product"].nunique() if "product" in df.columns else np.nan,
                "date_min": df["date"].min() if "date" in df.columns else pd.NaT,
                "date_max": df["date"].max() if "date" in df.columns else pd.NaT,
                "columns": ", ".join(map(str, df.columns[:20])),
            }
        )
    inventory = pd.DataFrame(inventory_rows)
    return panel, inventory, source_frames


def coverage_table(panel: pd.DataFrame) -> pd.DataFrame:
    cols = [
        "farmgate_initial_linear",
        "farmgate_filled_linear",
        "producer_linear",
        "consumer_linear",
        "prozorro",
        "prozorro_model",
        "retail_observed",
        "retail_observed_model",
        "retail_baseline",
        "retail_baseline_model",
        "retail_matched_observed",
        "retail_matched_observed_model",
        "retail_optimal_observed",
        "retail_optimal_observed_model",
        "retail_optimal_baseline",
        "retail_optimal_baseline_model",
        "retail_discount_share",
        "silpo_observed",
        "silpo_observed_model",
        "novus_observed",
        "novus_observed_model",
    ]
    rows = []
    for product, g in panel.groupby("product"):
        row = {
            "product": product,
            "product_label": PRODUCT_LABELS.get(product, product),
            "date_min": g["date"].min(),
            "date_max": g["date"].max(),
            "n_days": len(g),
        }
        for col in cols:
            if col in g.columns:
                row[f"{col}_n"] = int(g[col].notna().sum())
                row[f"{col}_share"] = float(g[col].notna().mean())
        rows.append(row)
    return pd.DataFrame(rows).sort_values("product")


def rolling_group_median(series: pd.Series, window: int) -> pd.Series:
    if window <= 1:
        return series.copy()
    return series.rolling(window=window, min_periods=1, center=True).median()


def build_weekly_panels(panel: pd.DataFrame) -> pd.DataFrame:
    value_cols = [
        c
        for c in [
            "farmgate_enterprise_linear",
            "farmgate_enterprise_pchip",
            "farmgate_household_linear",
            "farmgate_household_pchip",
            "farmgate_combined_linear",
            "farmgate_combined_pchip",
            "farmgate_initial_linear",
            "farmgate_initial_pchip",
            "farmgate_filled_linear",
            "farmgate_filled_pchip",
            "producer_linear",
            "producer_pchip",
            "consumer_linear",
            "consumer_pchip",
            "prozorro",
            "prozorro_model",
            "prozorro_expected",
            "prozorro_sum_initial",
            "prozorro_sum_current",
            "retail_observed",
            "retail_observed_model",
            "retail_baseline",
            "retail_baseline_model",
            "retail_category_observed",
            "retail_category_observed_model",
            "retail_category_baseline",
            "retail_category_baseline_model",
            "retail_matched_observed",
            "retail_matched_observed_model",
            "retail_matched_baseline",
            "retail_matched_baseline_model",
            "retail_matched_category_observed",
            "retail_matched_category_observed_model",
            "retail_matched_category_baseline",
            "retail_matched_category_baseline_model",
            "retail_optimal_observed",
            "retail_optimal_observed_model",
            "retail_optimal_baseline",
            "retail_optimal_baseline_model",
            "retail_discount_share",
            "retail_discount_depth",
            "silpo_observed",
            "silpo_observed_model",
            "silpo_baseline",
            "silpo_baseline_model",
            "silpo_category_observed",
            "silpo_category_observed_model",
            "silpo_category_baseline",
            "silpo_category_baseline_model",
            "silpo_discount_share",
            "silpo_discount_depth",
            "novus_observed",
            "novus_observed_model",
            "novus_category_observed",
            "novus_category_observed_model",
            "novus_category_baseline",
            "novus_category_baseline_model",
            "novus_discount_share",
        ]
        if c in panel.columns
    ]
    sum_cols = [c for c in ["prozorro_expected", "prozorro_sum_initial", "prozorro_sum_current"] if c in value_cols]
    median_cols = [c for c in value_cols if c not in sum_cols]
    pieces: List[pd.DataFrame] = []
    for product, g in panel.groupby("product", dropna=False):
        work = g.copy()
        work["week"] = pd.to_datetime(work["date"], errors="coerce").dt.to_period("W-SUN").dt.start_time
        weekly_med = work.groupby("week", as_index=False)[median_cols].median() if median_cols else pd.DataFrame(columns=["week"])
        weekly_sum = work.groupby("week", as_index=False)[sum_cols].sum(min_count=1) if sum_cols else pd.DataFrame(columns=["week"])
        if weekly_med.empty and weekly_sum.empty:
            continue
        wk = weekly_med.merge(weekly_sum, on="week", how="outer") if (not weekly_med.empty and not weekly_sum.empty) else (weekly_med if not weekly_med.empty else weekly_sum)
        wk["product"] = product
        pieces.append(wk)
    weekly = pd.concat(pieces, ignore_index=True) if pieces else pd.DataFrame(columns=["week", "product"])
    if weekly.empty:
        return weekly
    weekly["product_label"] = weekly["product"].map(PRODUCT_LABELS).fillna(weekly["product"])
    rename = {c: f"{c}_weekly" for c in value_cols}
    weekly = weekly.rename(columns=rename).sort_values(["product", "week"])

    smoothing_windows = {
        "prozorro_weekly": 3,
        "prozorro_model_weekly": 3,
        "prozorro_expected_weekly": 3,
        "prozorro_sum_initial_weekly": 3,
        "prozorro_sum_current_weekly": 3,
        "retail_observed_model_weekly": 2,
        "retail_baseline_model_weekly": 2,
        "retail_category_observed_model_weekly": 2,
        "retail_category_baseline_model_weekly": 2,
        "retail_matched_category_observed_model_weekly": 2,
        "retail_matched_category_baseline_model_weekly": 2,
        "retail_optimal_observed_model_weekly": 2,
        "retail_optimal_baseline_model_weekly": 2,
        "silpo_observed_model_weekly": 2,
        "silpo_baseline_model_weekly": 2,
        "silpo_category_observed_model_weekly": 2,
        "silpo_category_baseline_model_weekly": 2,
        "novus_observed_model_weekly": 2,
        "novus_category_observed_model_weekly": 2,
        "novus_category_baseline_model_weekly": 2,
        "producer_linear_weekly": 3,
        "producer_pchip_weekly": 3,
        "consumer_linear_weekly": 3,
        "consumer_pchip_weekly": 3,
        "farmgate_enterprise_linear_weekly": 3,
        "farmgate_enterprise_pchip_weekly": 3,
        "farmgate_household_linear_weekly": 3,
        "farmgate_household_pchip_weekly": 3,
        "farmgate_combined_linear_weekly": 3,
        "farmgate_combined_pchip_weekly": 3,
    }
    for col, window in smoothing_windows.items():
        if col not in weekly.columns:
            continue
        weekly[f"{col}_smooth"] = (
            weekly.groupby("product", dropna=False)[col]
            .transform(lambda s: rolling_group_median(s, window))
        )
    return weekly


def build_aggregate_chain_indices(panel: pd.DataFrame, source_frames: Dict[str, pd.DataFrame]) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if panel.empty:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    coverage = (
        panel.groupby("product", as_index=False)
        .agg(
            producer_obs=("producer_linear", lambda s: int(s.notna().sum())),
            procurement_obs=("prozorro", lambda s: int(s.notna().sum())),
            retail_obs=("retail_category_baseline", lambda s: int(s.notna().sum())),
            consumer_obs=("consumer_linear", lambda s: int(s.notna().sum())),
        )
        .sort_values("product")
    )
    common = coverage[
        coverage["producer_obs"].ge(30)
        & coverage["procurement_obs"].ge(20)
        & coverage["retail_obs"].ge(20)
        & coverage["consumer_obs"].ge(30)
    ]["product"].tolist()
    if not common:
        common = coverage["product"].tolist()
    if not common:
        return pd.DataFrame(), pd.DataFrame(), coverage

    retail_items = source_frames.get("retail_items_full", pd.DataFrame())
    retail_share = (
        retail_items.groupby("product")["item_key"].nunique().reindex(common).fillna(0).astype(float)
        if not retail_items.empty and "item_key" in retail_items.columns
        else pd.Series(0.0, index=common)
    )
    proc_share = (
        panel.groupby("product")["prozorro_n"].sum(min_count=1).reindex(common).fillna(0).astype(float)
        if "prozorro_n" in panel.columns
        else panel.groupby("product")["prozorro"].apply(lambda s: float(s.notna().sum())).reindex(common).fillna(0).astype(float)
    )
    equal_share = pd.Series(1.0 / len(common), index=common)
    retail_share = retail_share / retail_share.sum() if retail_share.sum() > 0 else equal_share.copy()
    proc_share = proc_share / proc_share.sum() if proc_share.sum() > 0 else equal_share.copy()
    weights = (
        pd.DataFrame(
            {
                "product": common,
                "product_label": [PRODUCT_LABELS.get(p, p) for p in common],
                "procurement_weight_proxy": proc_share.values,
                "retail_weight_proxy": retail_share.values,
                "equal_weight_anchor": equal_share.values,
            }
        )
        .assign(
            weight=lambda d: 0.45 * d["procurement_weight_proxy"] + 0.45 * d["retail_weight_proxy"] + 0.10 * d["equal_weight_anchor"]
        )
        .sort_values("product")
        .reset_index(drop=True)
    )
    weights["weight"] = weights["weight"] / weights["weight"].sum()

    chain_specs = [
        ("producer_linear", "producer_index"),
        ("prozorro", "procurement_index"),
        ("retail_category_baseline", "retail_index"),
        ("consumer_linear", "consumer_index"),
    ]
    base = (
        panel[panel["product"].isin(common)][["date", "product", "farmgate_combined_linear"] + [c for c, _ in chain_specs if c in panel.columns]]
        .copy()
        .sort_values(["date", "product"])
    )
    weight_map = weights.set_index("product")["weight"].to_dict()
    rows: List[Dict[str, object]] = []
    for date, g in base.groupby("date", dropna=False):
        row: Dict[str, object] = {"date": date}
        row["farmgate_index"] = float(g["farmgate_combined_linear"].median()) if "farmgate_combined_linear" in g.columns else np.nan
        row["farmgate_weight_share"] = 1.0 if pd.notna(row["farmgate_index"]) else 0.0
        row["farmgate_n_products"] = 1 if pd.notna(row["farmgate_index"]) else 0
        for src_col, out_col in chain_specs:
            if src_col not in g.columns:
                row[out_col] = np.nan
                row[f"{out_col}_weight_share"] = 0.0
                row[f"{out_col}_n_products"] = 0
                continue
            work = g[["product", src_col]].dropna().copy()
            work = work[work[src_col].gt(0)]
            if work.empty:
                row[out_col] = np.nan
                row[f"{out_col}_weight_share"] = 0.0
                row[f"{out_col}_n_products"] = 0
                continue
            work["weight"] = work["product"].map(weight_map).fillna(0)
            raw_weight = float(work["weight"].sum())
            row[f"{out_col}_weight_share"] = raw_weight
            row[f"{out_col}_n_products"] = int(len(work))
            if raw_weight < 0.60:
                row[out_col] = np.nan
                continue
            work["weight_norm"] = work["weight"] / raw_weight
            row[out_col] = float(np.exp((work["weight_norm"] * np.log(work[src_col])).sum()))
        rows.append(row)
    daily = pd.DataFrame(rows).sort_values("date")
    if daily.empty:
        return daily, pd.DataFrame(), weights
    daily["week"] = pd.to_datetime(daily["date"], errors="coerce").dt.to_period("W-SUN").dt.start_time
    weekly = (
        daily.groupby("week", as_index=False)
        .agg(
            farmgate_index_weekly=("farmgate_index", "median"),
            producer_index_weekly=("producer_index", "median"),
            procurement_index_weekly=("procurement_index", "median"),
            retail_index_weekly=("retail_index", "median"),
            consumer_index_weekly=("consumer_index", "median"),
            producer_weight_share=("producer_index_weight_share", "mean"),
            procurement_weight_share=("procurement_index_weight_share", "mean"),
            retail_weight_share=("retail_index_weight_share", "mean"),
            consumer_weight_share=("consumer_index_weight_share", "mean"),
        )
        .sort_values("week")
    )
    overlap_mask = weekly["retail_index_weekly"].notna() & weekly["consumer_index_weekly"].notna()
    if overlap_mask.sum() >= 6:
        retail_consumer_gap = float(
            np.nanmedian(
                safe_log(weekly.loc[overlap_mask, "retail_index_weekly"]) - safe_log(weekly.loc[overlap_mask, "consumer_index_weekly"])
            )
        )
    elif weekly["consumer_index_weekly"].notna().sum() >= 6:
        retail_consumer_gap = 0.0
    else:
        retail_consumer_gap = np.nan
    if pd.notna(retail_consumer_gap):
        consumer_scaled_weekly = np.exp(safe_log(weekly["consumer_index_weekly"]) + retail_consumer_gap)
        weekly["downstream_extension_weekly"] = weekly["retail_index_weekly"].where(weekly["retail_index_weekly"].notna(), consumer_scaled_weekly)
        weekly["downstream_extension_source"] = np.where(weekly["retail_index_weekly"].notna(), "retail_observed", "consumer_scaled_extension")
    else:
        weekly["downstream_extension_weekly"] = weekly["retail_index_weekly"]
        weekly["downstream_extension_source"] = np.where(weekly["retail_index_weekly"].notna(), "retail_observed", "")
    for col, window in {
        "producer_index_weekly": 3,
        "procurement_index_weekly": 3,
        "retail_index_weekly": 2,
        "consumer_index_weekly": 3,
        "farmgate_index_weekly": 3,
        "downstream_extension_weekly": 3,
    }.items():
        if col in weekly.columns:
            weekly[f"{col}_smooth"] = rolling_group_median(weekly[col], window)
    if pd.notna(retail_consumer_gap):
        consumer_scaled_daily = np.exp(safe_log(daily["consumer_index"]) + retail_consumer_gap)
        daily["downstream_extension"] = daily["retail_index"].where(daily["retail_index"].notna(), consumer_scaled_daily)
    else:
        daily["downstream_extension"] = daily.get("retail_index", np.nan)
    return daily, weekly, weights


def run_aggregate_index_models(index_weekly: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, pd.DataFrame]]:
    if index_weekly.empty:
        return pd.DataFrame(), pd.DataFrame(), {}
    links = [
        ("farmgate_index_weekly", "producer_index_weekly", "Index FarmGate -> Producer"),
        ("producer_index_weekly", "procurement_index_weekly", "Index Producer -> Procurement"),
        ("procurement_index_weekly", "retail_index_weekly", "Index Procurement -> Retail"),
        ("retail_index_weekly", "consumer_index_weekly", "Index Retail -> Consumer"),
    ]
    if "downstream_extension_weekly" in index_weekly.columns:
        links.extend(
            [
                ("farmgate_index_weekly", "downstream_extension_weekly", "Index FarmGate -> Downstream extension"),
                ("producer_index_weekly", "downstream_extension_weekly", "Index Producer -> Downstream extension"),
                ("procurement_index_weekly", "downstream_extension_weekly", "Index Procurement -> Downstream extension"),
            ]
        )
    rows: List[Dict[str, object]] = []
    series_rows: List[Dict[str, object]] = []
    details: Dict[str, pd.DataFrame] = {}
    for series_col in [c for pair in links for c in pair[:2]]:
        if series_col in index_weekly.columns:
            series_rows.append({"series_col": series_col, **series_stationarity(safe_log(index_weekly[series_col]))})
    for x_col, y_col, label in links:
        if x_col not in index_weekly.columns or y_col not in index_weekly.columns:
            continue
        x = safe_log(index_weekly[x_col])
        y = safe_log(index_weekly[y_col])
        for x_use, y_use, variant_name in [
            (x_col, y_col, "weekly_raw"),
            (f"{x_col}_smooth" if f"{x_col}_smooth" in index_weekly.columns else x_col, f"{y_col}_smooth" if f"{y_col}_smooth" in index_weekly.columns else y_col, "weekly_smoothed"),
        ]:
            if x_use not in index_weekly.columns or y_use not in index_weekly.columns:
                continue
            x_var = safe_log(index_weekly[x_use])
            y_var = safe_log(index_weekly[y_use])
            for fit_fn in [_fit_ardl_pair, _fit_ecm_pair, _fit_nardl_pair]:
                res = fit_fn(y_var, x_var)
                if res is None:
                    continue
                if res["model_family"] == "ECM":
                    reliability = "reliable" if res.get("ect_coef", np.nan) < 0 and res.get("ect_pvalue", 1.0) < 0.10 and res.get("ljungbox_p", 1.0) >= 0.05 else "conditionally_usable"
                elif res["model_family"] == "NARDL":
                    reliability = "reliable" if res.get("ect_coef", np.nan) < 0 and res.get("asymmetry_pvalue", 1.0) < 0.10 else "conditionally_usable"
                else:
                    reliability = "reliable" if res.get("cointegration_p", 1.0) < 0.10 else "conditionally_usable"
                rows.append(
                    {
                        "link": label,
                        "x_col": x_use,
                        "y_col": y_use,
                        "data_variant": variant_name,
                        "model_reliability": reliability,
                        **res,
                    }
                )

    vecm_summary_rows: List[Dict[str, object]] = []
    vecm_specs = [
        ("aggregate_chain_strict", ["farmgate_index_weekly", "producer_index_weekly", "procurement_index_weekly", "retail_index_weekly", "consumer_index_weekly"]),
        ("aggregate_chain_extended", ["farmgate_index_weekly", "producer_index_weekly", "procurement_index_weekly", "downstream_extension_weekly"]),
        ("aggregate_midstream_extended", ["producer_index_weekly", "procurement_index_weekly", "downstream_extension_weekly"]),
    ]
    for system_name, raw_cols in vecm_specs:
        vecm_cols = [c for c in raw_cols if c in index_weekly.columns]
        if len(vecm_cols) < 3:
            continue
        panel_df = index_weekly[["week"] + vecm_cols].dropna().copy()
        key_prefix = f"{system_name}_vecm"
        details[f"{key_prefix}_table1_series"] = panel_df
        details[f"{key_prefix}_table2_stationarity"] = pd.DataFrame([{"series_col": c, **series_stationarity(safe_log(panel_df[c]))} for c in vecm_cols])
        if len(panel_df) < MIN_CHAIN_OBS + 8:
            vecm_summary_rows.append({"system_name": system_name, "status": "infeasible", "reason": "too_short", "n_obs": int(len(panel_df))})
            continue
        try:
            log_levels = panel_df[vecm_cols].apply(safe_log)
            lag_sel = select_order(log_levels, maxlags=4, deterministic="ci")
            details[f"{key_prefix}_table3_lag_selection"] = pd.DataFrame([{"aic": lag_sel.aic, "bic": lag_sel.bic, "hqic": lag_sel.hqic, "fpe": lag_sel.fpe}])
            k_diff = max(1, int(lag_sel.aic if lag_sel.aic is not None else 2) - 1)
            rank_res = select_coint_rank(log_levels, det_order=0, k_ar_diff=k_diff, signif=0.05)
            details[f"{key_prefix}_table4_johansen"] = pd.DataFrame([{"selected_rank": int(rank_res.rank), "k_ar_diff": k_diff}])
            if int(rank_res.rank) < 1:
                vecm_summary_rows.append({"system_name": system_name, "status": "infeasible", "reason": "rank_zero", "n_obs": int(len(panel_df))})
                continue
            fit = VECM(log_levels, k_ar_diff=k_diff, coint_rank=int(rank_res.rank), deterministic="ci").fit()
            alpha = pd.DataFrame(fit.alpha, index=vecm_cols, columns=[f"relation_{i+1}" for i in range(fit.alpha.shape[1])]).reset_index().rename(columns={"index": "series"})
            beta = pd.DataFrame(fit.beta, index=vecm_cols, columns=[f"relation_{i+1}" for i in range(fit.beta.shape[1])]).reset_index().rename(columns={"index": "series"})
            details[f"{key_prefix}_table5_alpha"] = alpha
            details[f"{key_prefix}_table5_beta"] = beta
            details[f"{key_prefix}_table6_speed"] = alpha.copy()
            irf = fit.irf(8)
            irf_rows = []
            irfs = np.asarray(irf.irfs)
            for h in range(irfs.shape[0]):
                for i, target in enumerate(vecm_cols):
                    for j, impulse in enumerate(vecm_cols):
                        irf_rows.append({"horizon": h, "target": target, "impulse": impulse, "response": float(irfs[h, i, j])})
            details[f"{key_prefix}_table7_irf"] = pd.DataFrame(irf_rows)
            details[f"{key_prefix}_table8_fevd"] = compute_fevd_from_irf(irf, 8, vecm_cols)
            vecm_summary_rows.append(
                {
                    "system_name": system_name,
                    "status": "ok",
                    "reason": "",
                    "n_obs": int(len(panel_df)),
                    "vecm_rank": int(rank_res.rank),
                    "k_ar_diff": int(k_diff),
                    **{f"alpha_{name}": float(fit.alpha[idx, 0]) for idx, name in enumerate(vecm_cols)},
                }
            )
        except Exception as exc:
            vecm_summary_rows.append({"system_name": system_name, "status": "failed", "reason": str(exc), "n_obs": int(len(panel_df))})
    return (
        pd.DataFrame(series_rows),
        pd.DataFrame(rows),
        {"summary": pd.DataFrame(vecm_summary_rows), **details},
    )


def build_link21_summary(intersections: pd.DataFrame, corr_scan: pd.DataFrame, chain_models: pd.DataFrame) -> pd.DataFrame:
    target_links = [
        "FarmGate -> Producer",
        "Producer -> Procurement",
        "Procurement -> Retail (Novus)",
        "Procurement -> Retail (Silpo)",
        "Procurement -> Retail combined",
        "Producer -> Retail (Novus)",
        "Producer -> Retail (Silpo)",
        "Producer -> Retail combined",
        "FarmGate -> Procurement",
        "FarmGate -> Retail (Novus)",
        "FarmGate -> Retail (Silpo)",
        "FarmGate -> Retail combined",
        "Retail (Novus) -> Procurement",
        "Retail (Silpo) -> Procurement",
        "Retail combined -> Procurement",
        "Retail (Novus) -> Producer",
        "Retail (Silpo) -> Producer",
        "Retail combined -> Producer",
        "Retail (Novus) -> FarmGate",
        "Retail (Silpo) -> FarmGate",
        "Retail combined -> FarmGate",
    ]
    status_rank = {"strong": 4, "acceptable": 3, "weak_extension": 2, "unusable": 1}
    reliability_rank = {"reliable": 3, "conditionally_usable": 2}
    rows: List[Dict[str, object]] = []
    for link in target_links:
        i_sub = intersections[intersections["link"].eq(link)].copy() if not intersections.empty else pd.DataFrame()
        c_sub = corr_scan[corr_scan["link"].eq(link)].copy() if not corr_scan.empty else pd.DataFrame()
        m_sub = chain_models[chain_models["link"].eq(link)].copy() if not chain_models.empty else pd.DataFrame()
        best_status = "unusable"
        overlap = np.nan
        continuity = np.nan
        mapping = ""
        if not i_sub.empty:
            i_sub["status_rank"] = i_sub["admissibility_status"].map(status_rank).fillna(0)
            best_i = i_sub.sort_values(["status_rank", "overlap_weeks", "continuity_share"], ascending=[False, False, False]).iloc[0]
            best_status = best_i["admissibility_status"]
            overlap = best_i["overlap_weeks"]
            continuity = best_i["continuity_share"]
            mapping = best_i["mapping_type"]
        best_corr = np.nan
        best_lag = np.nan
        corr_product = ""
        if not c_sub.empty:
            c_sub["abs_corr"] = c_sub["best_abs_corr"].abs()
            best_c = c_sub.sort_values(["abs_corr", "n_obs"], ascending=[False, False]).iloc[0]
            best_corr = best_c["best_abs_corr"]
            best_lag = best_c["best_lag_weeks"]
            corr_product = best_c["product_label"]
        best_family = ""
        best_coef = np.nan
        best_ect = np.nan
        best_reliability = ""
        best_product = ""
        if not m_sub.empty:
            m_sub["reliability_rank"] = m_sub["model_reliability"].map(reliability_rank).fillna(0)
            m_sub["display_coef"] = m_sub["lr_coef"].where(m_sub["lr_coef"].notna(), m_sub["sr_coef"])
            m_sub["abs_coef"] = m_sub["display_coef"].abs()
            best_m = m_sub.sort_values(["reliability_rank", "abs_coef", "n_obs"], ascending=[False, False, False]).iloc[0]
            best_family = best_m["model_family"]
            best_coef = best_m["display_coef"]
            best_ect = best_m["ect_coef"]
            best_reliability = best_m["model_reliability"]
            best_product = best_m["product_label"]
        if best_reliability == "reliable":
            reading = "Retained weekly evidence."
        elif best_reliability == "conditionally_usable":
            reading = "Directional evidence, but diagnostics or overlap remain weaker."
        elif best_status in {"strong", "acceptable", "weak_extension"}:
            reading = "Overlap exists, but no retained weekly model passed the reporting rule."
        else:
            reading = "Weekly overlap too thin for retained estimation."
        rows.append(
            {
                "link": link,
                "best_admissibility": best_status,
                "median_overlap_weeks": overlap,
                "continuity_share": continuity,
                "mapping_type": mapping,
                "best_corr": best_corr,
                "best_lag_weeks": best_lag,
                "corr_product": corr_product,
                "best_model_family": best_family,
                "best_model_product": best_product,
                "display_coef": best_coef,
                "ect_coef": best_ect,
                "model_reliability": best_reliability,
                "reading": reading,
            }
        )
    return pd.DataFrame(rows)


def build_model_reliability_overview(
    chain_models: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    index_models: pd.DataFrame,
    index_vecm_summary: pd.DataFrame,
) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    if not chain_models.empty:
        for fam, g in chain_models.groupby("model_family", dropna=False):
            rows.append(
                {
                    "block": "Product-level weekly models",
                    "family": fam,
                    "rows": int(len(g)),
                    "reliable_rows": int(g["model_reliability"].eq("reliable").sum()),
                    "conditionally_usable_rows": int(g["model_reliability"].eq("conditionally_usable").sum()),
                }
            )
    if not index_models.empty:
        for fam, g in index_models.groupby("model_family", dropna=False):
            rows.append(
                {
                    "block": "Aggregate index models",
                    "family": fam,
                    "rows": int(len(g)),
                    "reliable_rows": int(g["model_reliability"].eq("reliable").sum()),
                    "conditionally_usable_rows": int(g["model_reliability"].eq("conditionally_usable").sum()),
                }
            )
    if not discounts.empty:
        rows.append({"block": "Discount models", "family": "OLS-HAC", "rows": int(len(discounts)), "reliable_rows": int(discounts["discount_strategy_signal"].sum()), "conditionally_usable_rows": int(len(discounts) - int(discounts["discount_strategy_signal"].sum()))})
    if not scale_models.empty:
        rows.append({"block": "Procurement-scale models", "family": "OLS-HAC", "rows": int(len(scale_models)), "reliable_rows": int(scale_models["scale_signal_flag"].sum()), "conditionally_usable_rows": int(len(scale_models) - int(scale_models["scale_signal_flag"].sum()))})
    if not vecm_results.empty:
        rows.append({"block": "Product-level VECM", "family": "VECM", "rows": int(len(vecm_results)), "reliable_rows": int(vecm_results["status"].eq("ok").sum()), "conditionally_usable_rows": int(vecm_results["status"].eq("infeasible").sum())})
    if not index_vecm_summary.empty:
        rows.append({"block": "Aggregate index VECM", "family": "VECM", "rows": int(len(index_vecm_summary)), "reliable_rows": int(index_vecm_summary["status"].eq("ok").sum()), "conditionally_usable_rows": int(index_vecm_summary["status"].eq("infeasible").sum())})
    return pd.DataFrame(rows)


def continuity_share(mask: pd.Series) -> float:
    flags = mask.fillna(False).astype(bool).tolist()
    if not flags:
        return np.nan
    total = sum(flags)
    if total == 0:
        return 0.0
    best = 0
    cur = 0
    for flag in flags:
        if flag:
            cur += 1
            best = max(best, cur)
        else:
            cur = 0
    return float(best / total)


def mapping_type_for_link(link: str, product: str) -> str:
    link_l = link.lower()
    if "farmgate" in link_l and product != "milk":
        return "inferred"
    if "retail matched" in link_l or "retail optimal" in link_l:
        return "widened"
    if "consumer" in link_l and product == "cheese":
        return "approximate"
    return "direct"


def raw_support_share(weekly: pd.DataFrame, model_col: str) -> float:
    base = model_col.replace("_smooth", "")
    if base.endswith("_weekly"):
        raw_guess = base.replace("_model_weekly", "_weekly")
    else:
        raw_guess = base.replace("_model", "")
    if raw_guess not in weekly.columns:
        return np.nan
    return float(weekly[raw_guess].notna().mean())


def pairwise_weekly_specs() -> List[Tuple[str, str, str]]:
    return [
        ("farmgate_combined_linear_weekly", "producer_linear_weekly", "FarmGate -> Producer"),
        ("producer_linear_weekly", "prozorro_weekly", "Producer -> Procurement"),
        ("prozorro_weekly", "novus_category_baseline_weekly", "Procurement -> Retail (Novus)"),
        ("prozorro_weekly", "silpo_category_baseline_weekly", "Procurement -> Retail (Silpo)"),
        ("prozorro_weekly", "retail_category_baseline_weekly", "Procurement -> Retail combined"),
        ("producer_linear_weekly", "novus_category_baseline_weekly", "Producer -> Retail (Novus)"),
        ("producer_linear_weekly", "silpo_category_baseline_weekly", "Producer -> Retail (Silpo)"),
        ("producer_linear_weekly", "retail_category_baseline_weekly", "Producer -> Retail combined"),
        ("farmgate_combined_linear_weekly", "prozorro_weekly", "FarmGate -> Procurement"),
        ("farmgate_combined_linear_weekly", "novus_category_baseline_weekly", "FarmGate -> Retail (Novus)"),
        ("farmgate_combined_linear_weekly", "silpo_category_baseline_weekly", "FarmGate -> Retail (Silpo)"),
        ("farmgate_combined_linear_weekly", "retail_category_baseline_weekly", "FarmGate -> Retail combined"),
        ("novus_category_baseline_weekly", "prozorro_weekly", "Retail (Novus) -> Procurement"),
        ("silpo_category_baseline_weekly", "prozorro_weekly", "Retail (Silpo) -> Procurement"),
        ("retail_category_baseline_weekly", "prozorro_weekly", "Retail combined -> Procurement"),
        ("novus_category_baseline_weekly", "producer_linear_weekly", "Retail (Novus) -> Producer"),
        ("silpo_category_baseline_weekly", "producer_linear_weekly", "Retail (Silpo) -> Producer"),
        ("retail_category_baseline_weekly", "producer_linear_weekly", "Retail combined -> Producer"),
        ("novus_category_baseline_weekly", "farmgate_combined_linear_weekly", "Retail (Novus) -> FarmGate"),
        ("silpo_category_baseline_weekly", "farmgate_combined_linear_weekly", "Retail (Silpo) -> FarmGate"),
        ("retail_category_baseline_weekly", "farmgate_combined_linear_weekly", "Retail combined -> FarmGate"),
    ]


def build_intersection_table(weekly: pd.DataFrame) -> pd.DataFrame:
    specs = pairwise_weekly_specs()
    rows: List[Dict[str, object]] = []
    for product, g in weekly.groupby("product", dropna=False):
        g = g.sort_values("week").copy()
        for x_col, y_col, link in specs:
            if x_col not in g.columns or y_col not in g.columns:
                continue
            mask = g[x_col].notna() & g[y_col].notna()
            overlap = int(mask.sum())
            if overlap == 0:
                continue
            mapping_type = mapping_type_for_link(link, str(product))
            x_support = raw_support_share(g.loc[mask], x_col)
            y_support = raw_support_share(g.loc[mask], y_col)
            cont = continuity_share(mask)
            min_support = float(np.nanmin([x_support, y_support])) if not (pd.isna(x_support) and pd.isna(y_support)) else np.nan
            if overlap >= 100 and cont >= 0.80 and (pd.isna(min_support) or min_support >= 0.50) and mapping_type in {"direct", "approximate"}:
                status = "strong"
            elif overlap >= 60 and cont >= 0.65 and (pd.isna(min_support) or min_support >= 0.25):
                status = "acceptable"
            elif overlap >= 24 and cont >= 0.35:
                status = "weak_extension"
            else:
                status = "unusable"
            rows.append(
                {
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "link": link,
                    "x_col": x_col,
                    "y_col": y_col,
                    "week_min": g.loc[mask, "week"].min(),
                    "week_max": g.loc[mask, "week"].max(),
                    "overlap_weeks": overlap,
                    "continuity_share": cont,
                    "x_raw_support_share": x_support,
                    "y_raw_support_share": y_support,
                    "mapping_type": mapping_type,
                    "interpolation_dominant_flag": int(pd.notna(min_support) and min_support < 0.25),
                    "admissibility_status": status,
                }
            )
    return pd.DataFrame(rows).sort_values(["product", "link"]) if rows else pd.DataFrame()


def build_product_definition_audit() -> Tuple[pd.DataFrame, pd.DataFrame]:
    audit_rows: List[Dict[str, object]] = []

    farm = pd.read_excel(FULL_UAH, sheet_name="FarmGate_UA")
    if not farm.empty:
        raw_col = "ua_product" if "ua_product" in farm.columns else farm.columns[0]
        for raw_product, g in farm.groupby(raw_col, dropna=False):
            audit_rows.append(
                {
                    "dataset": "FarmGate_UA",
                    "raw_product": str(raw_product),
                    "standardized_product": "milk",
                    "mapping_basis": "official farm-gate product label",
                    "mapping_type": "direct",
                    "ambiguity_flag": 0,
                    "n_rows": int(len(g)),
                }
            )

    for sheet_name, dataset_label in [("Producer_UA", "Producer_UA"), ("Consumer_UA", "Consumer_UA")]:
        df = pd.read_excel(FULL_UAH, sheet_name=sheet_name)
        if df.empty or "ua_product" not in df.columns:
            continue
        for raw_product, g in df.groupby("ua_product", dropna=False):
            std = harmonize_product_code(standardize_product(raw_product))
            audit_rows.append(
                {
                    "dataset": dataset_label,
                    "raw_product": str(raw_product),
                    "standardized_product": std,
                    "mapping_basis": "government row label",
                    "mapping_type": "direct" if std != "other" else "approximate",
                    "ambiguity_flag": int(std == "other"),
                    "n_rows": int(len(g)),
                }
            )

    pz = pd.read_excel(FULL_UAH, sheet_name="Prozorro")
    if not pz.empty:
        raw_col = "Product" if "Product" in pz.columns else pz.columns[0]
        title_col = "Товар" if "Товар" in pz.columns else raw_col
        for raw_product, g in pz.groupby(raw_col, dropna=False):
            std = harmonize_product_code(standardize_product(raw_product))
            examples = g[title_col].astype(str).head(3).tolist()
            audit_rows.append(
                {
                    "dataset": "Prozorro",
                    "raw_product": str(raw_product),
                    "standardized_product": std,
                    "mapping_basis": "procurement profile plus tender title",
                    "mapping_type": "direct" if std != "other" else "approximate",
                    "ambiguity_flag": int(std == "other" or len(set(examples)) > 10),
                    "n_rows": int(len(g)),
                }
            )

    silpo_items = load_retail_items("Silpo", "Product", "Silpo")
    novus_items = load_retail_items("Novus", "product", "Novus")
    for dataset_label, df in [("Silpo", silpo_items), ("Novus", novus_items)]:
        if df.empty:
            continue
        for raw_product, g in df.groupby("raw_sheet_product", dropna=False):
            std_modes = g["product"].dropna().astype(str).map(harmonize_product_code).value_counts()
            std = std_modes.index[0] if not std_modes.empty else "other"
            audit_rows.append(
                {
                    "dataset": dataset_label,
                    "raw_product": str(raw_product),
                    "standardized_product": std,
                    "mapping_basis": "retail title, product title, literal dairy typing, and brand-normalized item text",
                    "mapping_type": "inferred",
                    "ambiguity_flag": int(g["product"].nunique(dropna=True) > 1 or std == "other"),
                    "n_rows": int(len(g)),
                }
            )

    audit_long = pd.DataFrame(audit_rows).sort_values(["dataset", "standardized_product", "raw_product"]) if audit_rows else pd.DataFrame()
    if audit_long.empty:
        return audit_long, audit_long
    dictionary = (
        audit_long.groupby(["dataset", "standardized_product"], as_index=False)
        .agg(
            raw_variants=("raw_product", pd.Series.nunique),
            example_raw_products=("raw_product", lambda s: "; ".join(pd.Series(s).astype(str).head(5))),
            mapping_types=("mapping_type", lambda s: "; ".join(sorted(set(pd.Series(s).astype(str))))),
            ambiguity_rows=("ambiguity_flag", "sum"),
            n_rows=("n_rows", "sum"),
        )
        .sort_values(["dataset", "standardized_product"])
    )
    return audit_long, dictionary


def series_stationarity(series: pd.Series) -> Dict[str, float]:
    s = pd.Series(series).replace([np.inf, -np.inf], np.nan).dropna()
    if len(s) < MIN_CHAIN_OBS:
        return {
            "n_obs": int(len(s)),
            "adf_p": np.nan,
            "kpss_p": np.nan,
            "integration_class": "too_short",
            "stability_flag": 1,
        }
    try:
        adf_p = float(adfuller(s, autolag="AIC")[1])
    except Exception:
        adf_p = np.nan
    try:
        kpss_p = float(kpss(s, regression="c", nlags="auto")[1])
    except Exception:
        kpss_p = np.nan
    ds = s.diff().dropna()
    try:
        adf_d1 = float(adfuller(ds, autolag="AIC")[1]) if len(ds) >= MIN_CHAIN_OBS else np.nan
    except Exception:
        adf_d1 = np.nan
    try:
        kpss_d1 = float(kpss(ds, regression="c", nlags="auto")[1]) if len(ds) >= MIN_CHAIN_OBS else np.nan
    except Exception:
        kpss_d1 = np.nan
    if pd.notna(adf_p) and pd.notna(kpss_p) and adf_p < 0.05 and kpss_p > 0.05:
        cls = "I(0)"
    elif pd.notna(adf_d1) and pd.notna(kpss_d1) and adf_d1 < 0.05 and kpss_d1 > 0.05:
        cls = "I(1)"
    else:
        cls = "ambiguous"
    mid = len(s) // 2
    denom = float(s.std(ddof=1)) if len(s) > 1 else np.nan
    stability = int(pd.notna(denom) and denom > 0 and abs(float(s.iloc[:mid].mean() - s.iloc[mid:].mean())) / denom > 0.75)
    return {
        "n_obs": int(len(s)),
        "adf_p": adf_p,
        "kpss_p": kpss_p,
        "adf_diff1_p": adf_d1,
        "kpss_diff1_p": kpss_d1,
        "integration_class": cls,
        "stability_flag": stability,
    }


def model_residual_diagnostics(resid: pd.Series, exog: Optional[pd.DataFrame] = None) -> Dict[str, float]:
    r = pd.Series(resid).replace([np.inf, -np.inf], np.nan).dropna()
    if len(r) < 12:
        return {
            "ljungbox_p": np.nan,
            "bp_p": np.nan,
            "white_p": np.nan,
            "jb_p": np.nan,
            "stability_flag": 1,
        }
    lag = min(8, max(2, len(r) // 5))
    try:
        lb = float(acorr_ljungbox(r, lags=[lag], return_df=True)["lb_pvalue"].iloc[0])
    except Exception:
        lb = np.nan
    try:
        jb = float(jarque_bera(r)[1])
    except Exception:
        jb = np.nan
    bp_p = np.nan
    white_p = np.nan
    if exog is not None:
        x = pd.DataFrame(exog).replace([np.inf, -np.inf], np.nan).dropna()
        if not x.empty:
            try:
                x = sm.add_constant(x, has_constant="add")
                aligned = pd.concat([r.rename("resid"), x], axis=1).dropna()
                if len(aligned) >= 12:
                    bp_p = float(het_breuschpagan(aligned["resid"], aligned.drop(columns=["resid"]))[1])
                    white_p = float(het_white(aligned["resid"], aligned.drop(columns=["resid"]))[1])
            except Exception:
                bp_p = np.nan
                white_p = np.nan
    mid = len(r) // 2
    denom = float(r.std(ddof=1)) if len(r) > 1 else np.nan
    stability = int(pd.notna(denom) and denom > 0 and abs(float(r.iloc[:mid].mean() - r.iloc[mid:].mean())) / denom > 0.75)
    return {
        "ljungbox_p": lb,
        "bp_p": bp_p,
        "white_p": white_p,
        "jb_p": jb,
        "stability_flag": stability,
    }


def best_abs_lag_corr_weeks(x: pd.Series, y: pd.Series, max_lag: int = 8) -> Tuple[float, float, int]:
    best_corr = np.nan
    best_lag = np.nan
    best_n = 0
    lx = safe_log(x)
    ly = safe_log(y)
    for lag in range(0, max_lag + 1):
        d = pd.concat([lx.shift(lag).rename("x"), ly.rename("y")], axis=1).dropna()
        if len(d) < MIN_CHAIN_OBS:
            continue
        corr = d["x"].corr(d["y"])
        if pd.isna(corr):
            continue
        if pd.isna(best_corr) or abs(corr) > abs(best_corr):
            best_corr = float(corr)
            best_lag = float(lag)
            best_n = int(len(d))
    return best_corr, best_lag, best_n


def _fit_ardl_pair(log_y: pd.Series, log_x: pd.Series, max_lag: int = 6) -> Optional[Dict[str, object]]:
    d = pd.concat([log_y.rename("y"), log_x.rename("x")], axis=1).dropna()
    if len(d) < MIN_CHAIN_OBS:
        return None
    for lag_try in [max_lag, min(4, max_lag), min(3, max_lag), 2]:
        try:
            sel = ardl_select_order(d["y"], lag_try, d[["x"]], lag_try, ic="aic", trend="c")
            fit = sel.model.fit(cov_type="HAC", cov_kwds={"maxlags": 2})
            phi = float(np.nansum([v for k, v in fit.params.items() if str(k).startswith("y.L")]))
            beta = float(np.nansum([v for k, v in fit.params.items() if str(k).startswith("x.L")]))
            lr_coef = beta / (1.0 - phi) if abs(1.0 - phi) > 1e-8 else np.nan
            uecm_res = UECM.from_ardl(sel.model).fit(cov_type="HAC", cov_kwds={"maxlags": 2})
            bt = uecm_res.bounds_test(case=3)
            diag = model_residual_diagnostics(fit.resid, pd.DataFrame(fit.model.exog, columns=fit.model.exog_names))
            return {
                "model_family": "ARDL",
                "n_obs": int(len(d)),
                "sr_coef": float(fit.params.get("x.L0", np.nan)),
                "lr_coef": float(lr_coef) if pd.notna(lr_coef) else np.nan,
                "ect_coef": np.nan,
                "ect_pvalue": np.nan,
                "cointegration_p": float(bt.p_values.get("upper", np.nan)),
                "lag_order": int(lag_try),
                "notes": f"bounds_upper_p={float(bt.p_values.get('upper', np.nan)):.4g}",
                **diag,
            }
        except Exception:
            continue
    return None


def _fit_ecm_pair(log_y: pd.Series, log_x: pd.Series, max_lag: int = 4) -> Optional[Dict[str, object]]:
    d = pd.concat([log_y.rename("y"), log_x.rename("x")], axis=1).dropna()
    if len(d) < MIN_CHAIN_OBS + 8:
        return None
    try:
        coint_p = float(coint(d["y"], d["x"])[1])
    except Exception:
        return None
    if not (pd.notna(coint_p) and coint_p < 0.10):
        return None
    lr = sm.OLS(d["y"], sm.add_constant(d[["x"]], has_constant="add")).fit()
    d["ect_l1"] = lr.resid.shift(1)
    d["dy"] = d["y"].diff()
    d["dx"] = d["x"].diff()
    best_fit = None
    best_table = None
    best_aic = None
    best_lag = None
    for p in range(1, max_lag + 1):
        t = d[["dy", "dx", "ect_l1"]].copy()
        for i in range(1, p + 1):
            t[f"dy_l{i}"] = t["dy"].shift(i)
            t[f"dx_l{i}"] = t["dx"].shift(i)
        t = t.dropna()
        if len(t) < MIN_CHAIN_OBS:
            continue
        exog = sm.add_constant(t.drop(columns=["dy"]), has_constant="add")
        fit = sm.OLS(t["dy"], exog).fit(cov_type="HAC", cov_kwds={"maxlags": 2})
        if best_aic is None or fit.aic < best_aic:
            best_aic = float(fit.aic)
            best_fit = fit
            best_table = t
            best_lag = p
    if best_fit is None or best_table is None:
        return None
    diag = model_residual_diagnostics(best_fit.resid, best_table.drop(columns=["dy"]))
    return {
        "model_family": "ECM",
        "n_obs": int(len(best_table)),
        "sr_coef": float(best_fit.params.get("dx", np.nan)),
        "lr_coef": float(lr.params.get("x", np.nan)),
        "ect_coef": float(best_fit.params.get("ect_l1", np.nan)),
        "ect_pvalue": float(best_fit.pvalues.get("ect_l1", np.nan)),
        "cointegration_p": coint_p,
        "lag_order": int(best_lag) if best_lag is not None else np.nan,
        "notes": f"engle_granger_p={coint_p:.4g}",
        **diag,
    }


def _fit_nardl_pair(log_y: pd.Series, log_x: pd.Series) -> Optional[Dict[str, object]]:
    d = pd.concat([log_y.rename("y"), log_x.rename("x")], axis=1).dropna()
    if len(d) < MIN_CHAIN_OBS + 10:
        return None
    d["dx"] = d["x"].diff()
    d["x_pos"] = d["dx"].clip(lower=0).cumsum()
    d["x_neg"] = d["dx"].clip(upper=0).abs().cumsum()
    d["dy"] = d["y"].diff()
    d["dy_l1"] = d["dy"].shift(1)
    d["dx_pos"] = d["dx"].clip(lower=0)
    d["dx_neg_abs"] = d["dx"].clip(upper=0).abs()
    d["y_l1"] = d["y"].shift(1)
    d["x_pos_l1"] = d["x_pos"].shift(1)
    d["x_neg_l1"] = d["x_neg"].shift(1)
    t = d[["dy", "dy_l1", "dx_pos", "dx_neg_abs", "y_l1", "x_pos_l1", "x_neg_l1"]].dropna()
    if len(t) < MIN_CHAIN_OBS:
        return None
    fit = sm.OLS(t["dy"], sm.add_constant(t.drop(columns=["dy"]), has_constant="add")).fit(cov_type="HAC", cov_kwds={"maxlags": 2})
    lam = float(fit.params.get("y_l1", np.nan))
    psi_p = float(fit.params.get("x_pos_l1", np.nan))
    psi_n = float(fit.params.get("x_neg_l1", np.nan))
    lr_pos = -psi_p / lam if pd.notna(lam) and abs(lam) > 1e-8 else np.nan
    lr_neg = -psi_n / lam if pd.notna(lam) and abs(lam) > 1e-8 else np.nan
    try:
        asym_p = float(fit.f_test("x_pos_l1 = x_neg_l1").pvalue)
    except Exception:
        asym_p = np.nan
    diag = model_residual_diagnostics(fit.resid, t.drop(columns=["dy"]))
    return {
        "model_family": "NARDL",
        "n_obs": int(len(t)),
        "sr_coef": float(fit.params.get("dx_pos", np.nan) - fit.params.get("dx_neg_abs", np.nan)),
        "lr_coef": float(lr_pos - lr_neg) if pd.notna(lr_pos) and pd.notna(lr_neg) else np.nan,
        "ect_coef": lam,
        "ect_pvalue": float(fit.pvalues.get("y_l1", np.nan)),
        "cointegration_p": np.nan,
        "lag_order": 1,
        "asymmetry_pvalue": asym_p,
        "notes": "Positive and negative weekly shock decomposition.",
        **diag,
    }


def _fit_vecm_system(panel_df: pd.DataFrame, columns: Sequence[str], target_col: str, max_lag: int = 4) -> Dict[str, object]:
    d = panel_df[list(columns)].dropna()
    if len(d) < MIN_CHAIN_OBS + 12:
        return {"status": "infeasible", "reason": "too_short", "n_obs": int(len(d))}
    try:
        sel = select_order(d, maxlags=max_lag, deterministic="ci")
        aic_lag = sel.aic if sel.aic is not None else 2
        k_diff = max(1, int(aic_lag) - 1)
        rank = int(select_coint_rank(d, det_order=0, k_ar_diff=k_diff, signif=0.05).rank)
        if rank < 1:
            return {"status": "infeasible", "reason": "rank_zero", "n_obs": int(len(d))}
        fit = VECM(d, k_ar_diff=k_diff, coint_rank=rank, deterministic="ci").fit()
        resid = pd.DataFrame(fit.resid, columns=d.columns)
        alpha_map = {f"alpha_{name}": float(fit.alpha[idx, 0]) for idx, name in enumerate(d.columns)}
        target_resid = resid[target_col] if target_col in resid.columns else resid.iloc[:, -1]
        diag = model_residual_diagnostics(target_resid)
        return {
            "status": "ok",
            "reason": "",
            "n_obs": int(len(d)),
            "vecm_rank": rank,
            "k_ar_diff": k_diff,
            **alpha_map,
            **diag,
        }
    except Exception as exc:
        return {"status": "failed", "reason": str(exc), "n_obs": int(len(d))}


def run_weekly_correlation_scan(weekly: pd.DataFrame, intersections: pd.DataFrame) -> pd.DataFrame:
    if weekly.empty:
        return pd.DataFrame()
    rows: List[Dict[str, object]] = []
    for row in intersections.itertuples():
        if row.admissibility_status == "unusable":
            continue
        g = weekly[weekly["product"].eq(row.product)].sort_values("week")
        if row.x_col not in g.columns or row.y_col not in g.columns:
            continue
        corr, lag, n_obs = best_abs_lag_corr_weeks(g[row.x_col], g[row.y_col], max_lag=8)
        rows.append(
            {
                "product": row.product,
                "product_label": row.product_label,
                "link": row.link,
                "x_col": row.x_col,
                "y_col": row.y_col,
                "best_abs_corr": corr,
                "best_lag_weeks": lag,
                "n_obs": n_obs,
                "admissibility_status": row.admissibility_status,
                "mapping_type": row.mapping_type,
            }
        )
    return pd.DataFrame(rows).sort_values(["product", "link"]) if rows else pd.DataFrame()


def run_core_chain_models(weekly: pd.DataFrame, intersections: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if weekly.empty:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    rows: List[Dict[str, object]] = []
    series_rows: List[Dict[str, object]] = []
    vecm_rows: List[Dict[str, object]] = []
    seen_series: set[tuple[str, str]] = set()

    for row in intersections.itertuples():
        if row.admissibility_status == "unusable":
            continue
        product_weekly = weekly[weekly["product"].eq(row.product)].sort_values("week").copy()
        x_col = row.x_col.replace("_weekly", "_weekly_smooth") if f"{row.x_col}_smooth".endswith("_smooth") else row.x_col
        y_col = row.y_col
        variants = [(row.x_col, row.y_col, "weekly_raw")]
        x_smooth = f"{row.x_col}_smooth" if f"{row.x_col}_smooth" in product_weekly.columns else None
        y_smooth = f"{row.y_col}_smooth" if f"{row.y_col}_smooth" in product_weekly.columns else None
        if x_smooth or y_smooth:
            variants.append((x_smooth or row.x_col, y_smooth or row.y_col, "weekly_smoothed"))

        for series_col in {row.x_col, row.y_col}:
            key = (row.product, series_col)
            if key not in seen_series and series_col in product_weekly.columns:
                tests = series_stationarity(safe_log(product_weekly[series_col]))
                series_rows.append(
                    {
                        "product": row.product,
                        "product_label": row.product_label,
                        "series_col": series_col,
                        **tests,
                    }
                )
                seen_series.add(key)

        for x_use, y_use, variant_name in variants:
            if x_use is None or y_use is None or x_use not in product_weekly.columns or y_use not in product_weekly.columns:
                continue
            x = safe_log(product_weekly[x_use])
            y = safe_log(product_weekly[y_use])
            for fit_fn in [_fit_ardl_pair, _fit_ecm_pair, _fit_nardl_pair]:
                result = fit_fn(y, x)
                if result is None:
                    continue
                if result["model_family"] == "ECM":
                    reliability = "reliable" if result.get("ect_coef", np.nan) < 0 and result.get("ect_pvalue", 1.0) < 0.10 and result.get("ljungbox_p", 1.0) >= 0.05 else "conditionally_usable"
                elif result["model_family"] == "NARDL":
                    reliability = "reliable" if result.get("ect_coef", np.nan) < 0 and result.get("asymmetry_pvalue", 1.0) < 0.10 else "conditionally_usable"
                else:
                    reliability = "reliable" if result.get("cointegration_p", 1.0) < 0.10 else "conditionally_usable"
                rows.append(
                    {
                        "product": row.product,
                        "product_label": row.product_label,
                        "link": row.link,
                        "x_col": x_use,
                        "y_col": y_use,
                        "data_variant": variant_name,
                        "admissibility_status": row.admissibility_status,
                        "mapping_type": row.mapping_type,
                        "model_reliability": reliability,
                        **result,
                    }
                )

    vecm_specs = [
        ("staged_chain", ["producer_linear_weekly", "prozorro_weekly", "retail_category_baseline_weekly"], "retail_category_baseline_weekly"),
        ("staged_chain_smoothed", ["producer_linear_weekly_smooth", "prozorro_weekly_smooth", "retail_category_baseline_weekly_smooth"], "retail_category_baseline_weekly_smooth"),
        ("full_chain", ["farmgate_combined_linear_weekly", "producer_linear_weekly", "prozorro_weekly", "retail_category_baseline_weekly", "consumer_linear_weekly"], "consumer_linear_weekly"),
        ("full_chain_smoothed", ["farmgate_combined_linear_weekly_smooth", "producer_linear_weekly_smooth", "prozorro_weekly_smooth", "retail_category_baseline_weekly_smooth", "consumer_linear_weekly_smooth"], "consumer_linear_weekly_smooth"),
    ]
    for product, g in weekly.groupby("product", dropna=False):
        for system_name, cols, target in vecm_specs:
            usable_cols = [c for c in cols if c in g.columns]
            if len(usable_cols) < 3 or target not in usable_cols:
                continue
            vecm = _fit_vecm_system(g.sort_values("week"), usable_cols, target)
            vecm_rows.append(
                {
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "system_name": system_name,
                    "target_series": target,
                    "included_series": "; ".join(usable_cols),
                    **vecm,
                }
            )
    return (
        pd.DataFrame(series_rows).sort_values(["product", "series_col"]) if series_rows else pd.DataFrame(),
        pd.DataFrame(rows).sort_values(["product", "link", "model_family", "data_variant"]) if rows else pd.DataFrame(),
        pd.DataFrame(vecm_rows).sort_values(["product", "system_name"]) if vecm_rows else pd.DataFrame(),
    )


def compute_fevd_from_irf(irf_obj, steps: int, names: Sequence[str]) -> pd.DataFrame:
    if not hasattr(irf_obj, "orth_irfs"):
        return pd.DataFrame()
    orth = np.asarray(irf_obj.orth_irfs)
    rows: List[Dict[str, object]] = []
    for h in range(1, min(steps + 1, orth.shape[0])):
        contrib = (orth[: h + 1] ** 2).sum(axis=0)
        denom = contrib.sum(axis=1, keepdims=True)
        share = np.divide(contrib, denom, out=np.full_like(contrib, np.nan), where=denom > 0)
        for i, target in enumerate(names):
            for j, shock in enumerate(names):
                rows.append(
                    {
                        "horizon": h,
                        "target": target,
                        "shock": shock,
                        "fevd_share": float(share[i, j]),
                    }
                )
    return pd.DataFrame(rows)


def run_vecm_detailed_tables(weekly: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame]]:
    summary_rows: List[Dict[str, object]] = []
    details: Dict[str, pd.DataFrame] = {}
    vecm_specs = [
        ("full_chain", ["farmgate_combined_linear_weekly", "producer_linear_weekly", "prozorro_weekly", "retail_category_baseline_weekly", "consumer_linear_weekly"]),
        ("staged_chain", ["producer_linear_weekly", "prozorro_weekly", "retail_category_baseline_weekly"]),
    ]
    for product, g in weekly.groupby("product", dropna=False):
        g = g.sort_values("week").copy()
        for system_name, cols in vecm_specs:
            usable_cols = [c for c in cols if c in g.columns]
            if len(usable_cols) < 3:
                continue
            panel_df = g[["week"] + usable_cols].dropna().copy()
            key_prefix = f"{product}_{system_name}"
            details[f"{key_prefix}_table1_series"] = panel_df
            stat_rows = []
            for col in usable_cols:
                tests = series_stationarity(safe_log(panel_df[col]))
                stat_rows.append({"series_col": col, **tests})
            details[f"{key_prefix}_table2_stationarity"] = pd.DataFrame(stat_rows)
            if len(panel_df) < MIN_CHAIN_OBS + 12:
                summary_rows.append(
                    {
                        "product": product,
                        "product_label": PRODUCT_LABELS.get(product, product),
                        "system_name": system_name,
                        "status": "infeasible",
                        "reason": "too_short",
                        "n_obs": int(len(panel_df)),
                    }
                )
                continue
            log_levels = panel_df[usable_cols].apply(safe_log)
            try:
                lag_sel = select_order(log_levels, maxlags=4, deterministic="ci")
                lag_table = pd.DataFrame(
                    [
                        {
                            "aic": lag_sel.aic,
                            "bic": lag_sel.bic,
                            "hqic": lag_sel.hqic,
                            "fpe": lag_sel.fpe,
                        }
                    ]
                )
                details[f"{key_prefix}_table3_lag_selection"] = lag_table
                k_diff = max(1, int(lag_sel.aic if lag_sel.aic is not None else 2) - 1)
                rank_res = select_coint_rank(log_levels, det_order=0, k_ar_diff=k_diff, signif=0.05)
                details[f"{key_prefix}_table4_johansen"] = pd.DataFrame(
                    [{"selected_rank": int(rank_res.rank), "trace_stat_5pct_test": "see selected rank output"}]
                )
                if int(rank_res.rank) < 1:
                    summary_rows.append(
                        {
                            "product": product,
                            "product_label": PRODUCT_LABELS.get(product, product),
                            "system_name": system_name,
                            "status": "infeasible",
                            "reason": "rank_zero",
                            "n_obs": int(len(panel_df)),
                        }
                    )
                    continue
                fit = VECM(log_levels, k_ar_diff=k_diff, coint_rank=int(rank_res.rank), deterministic="ci").fit()
                alpha = pd.DataFrame(fit.alpha, index=usable_cols, columns=[f"cointegration_relation_{i+1}" for i in range(fit.alpha.shape[1])]).reset_index().rename(columns={"index": "series"})
                beta = pd.DataFrame(fit.beta, index=usable_cols, columns=[f"cointegration_relation_{i+1}" for i in range(fit.beta.shape[1])]).reset_index().rename(columns={"index": "series"})
                gamma = pd.DataFrame(fit.gamma)
                details[f"{key_prefix}_table5_vecm_estimation_alpha"] = alpha
                details[f"{key_prefix}_table5_vecm_estimation_beta"] = beta
                details[f"{key_prefix}_table5_vecm_short_run_gamma"] = gamma
                details[f"{key_prefix}_table6_speed_adjustment"] = alpha.copy()
                irf = fit.irf(8)
                irf_rows: List[Dict[str, object]] = []
                irfs = np.asarray(irf.irfs)
                for h in range(irfs.shape[0]):
                    for i, target in enumerate(usable_cols):
                        for j, impulse in enumerate(usable_cols):
                            irf_rows.append(
                                {
                                    "horizon": h,
                                    "target": target,
                                    "impulse": impulse,
                                    "irf_response": float(irfs[h, i, j]),
                                }
                            )
                details[f"{key_prefix}_table7_irf"] = pd.DataFrame(irf_rows)
                details[f"{key_prefix}_table8_fevd"] = compute_fevd_from_irf(irf, 8, usable_cols)
                summary_rows.append(
                    {
                        "product": product,
                        "product_label": PRODUCT_LABELS.get(product, product),
                        "system_name": system_name,
                        "status": "ok",
                        "reason": "",
                        "n_obs": int(len(panel_df)),
                        "vecm_rank": int(rank_res.rank),
                        "k_ar_diff": int(k_diff),
                    }
                )
            except Exception as exc:
                summary_rows.append(
                    {
                        "product": product,
                        "product_label": PRODUCT_LABELS.get(product, product),
                        "system_name": system_name,
                        "status": "failed",
                        "reason": str(exc),
                        "n_obs": int(len(panel_df)),
                    }
                )
    return pd.DataFrame(summary_rows), details


def fit_hac_ols(y: pd.Series, x: pd.DataFrame, hac_lag: int) -> Optional[sm.regression.linear_model.RegressionResultsWrapper]:
    d = pd.concat([y.rename("_y"), x], axis=1).replace([np.inf, -np.inf], np.nan).dropna()
    if len(d) < 20:
        return None
    y2 = d["_y"]
    x2 = d.drop(columns=["_y"])
    if x2.empty:
        return None
    x2 = sm.add_constant(x2, has_constant="add")
    if x2.drop(columns=["const"], errors="ignore").std(numeric_only=True).fillna(0).sum() <= 0:
        return None
    try:
        return sm.OLS(y2, x2).fit(cov_type="HAC", cov_kwds={"maxlags": int(max(1, min(MAX_HAC_LAG, hac_lag)))})
    except Exception:
        return None


def residual_diagnostics(resid: pd.Series) -> Dict[str, float]:
    r = pd.Series(resid).replace([np.inf, -np.inf], np.nan).dropna()
    if len(r) < 20:
        return {"ljungbox_p": np.nan, "arch_p": np.nan, "jb_p": np.nan}
    lag = min(10, max(2, len(r) // 5))
    try:
        lb = float(acorr_ljungbox(r, lags=[lag], return_df=True)["lb_pvalue"].iloc[0])
    except Exception:
        lb = np.nan
    try:
        arch = float(het_arch(r)[1])
    except Exception:
        arch = np.nan
    try:
        jb = float(jarque_bera(r)[1])
    except Exception:
        jb = np.nan
    return {"ljungbox_p": lb, "arch_p": arch, "jb_p": jb}


def run_local_projection_for_pair(
    g: pd.DataFrame,
    product: str,
    y_col: str,
    x_col: str,
    link: str,
    link_direction: str,
    price_variant: str,
    farm_gate_source: str,
    reconstruction_variant: str,
) -> List[Dict[str, object]]:
    rows: List[Dict[str, object]] = []
    data = g[["date", y_col, x_col, "retail_discount_share", "retail_discount_depth"]].copy()
    data["ly"] = safe_log(data[y_col])
    data["lx"] = safe_log(data[x_col])
    data["dy"] = data["ly"].diff()
    data["dx"] = data["lx"].diff()
    data["dy_l1"] = data["dy"].shift(1)
    data["dx_l1"] = data["dx"].shift(1)
    data["discount_share"] = safe_num(data["retail_discount_share"]).fillna(0)
    data["discount_depth"] = safe_num(data["retail_discount_depth"]).fillna(0)
    if data[[y_col, x_col]].dropna().shape[0] < MIN_LP_OBS:
        return rows

    for h in HORIZONS:
        data[f"response_h{h}"] = data["ly"].shift(-h) - data["ly"].shift(1)
        base = data[["date", f"response_h{h}", "dx", "dy_l1", "dx_l1", "discount_share", "discount_depth"]].copy()
        control_cols = ["dy_l1", "dx_l1"]
        if base["discount_share"].std(skipna=True) > 1e-8:
            control_cols.append("discount_share")
        if base["discount_depth"].std(skipna=True) > 1e-8:
            control_cols.append("discount_depth")

        fit = fit_hac_ols(base[f"response_h{h}"], base[["dx"] + control_cols], max(1, h))
        if fit is not None:
            d_used = pd.concat([base[f"response_h{h}"], base[["dx"] + control_cols]], axis=1).dropna()
            diag = residual_diagnostics(fit.resid)
            rows.append(
                {
                    "model": "LP_linear",
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "link": link,
                    "link_direction": link_direction,
                    "price_variant": price_variant,
                    "farm_gate_source": farm_gate_source,
                    "reconstruction_variant": reconstruction_variant,
                    "horizon_days": h,
                    "n_obs": int(len(d_used)),
                    "coef": float(fit.params.get("dx", np.nan)),
                    "pvalue": float(fit.pvalues.get("dx", np.nan)),
                    "stderr": float(fit.bse.get("dx", np.nan)),
                    "r2": float(getattr(fit, "rsquared", np.nan)),
                    "core_signal": int((len(d_used) >= MIN_LP_OBS) and (float(fit.pvalues.get("dx", 1.0)) < 0.10)),
                    **diag,
                }
            )

        asym = base.copy()
        asym["dx_pos"] = asym["dx"].clip(lower=0)
        asym["dx_neg_abs"] = (-asym["dx"].clip(upper=0)).clip(lower=0)
        asym_controls = ["dx_pos", "dx_neg_abs"] + control_cols
        if asym[["dx_pos", "dx_neg_abs"]].std(skipna=True).fillna(0).sum() > 1e-8:
            fit_a = fit_hac_ols(asym[f"response_h{h}"], asym[asym_controls], max(1, h))
            if fit_a is not None:
                d_used = pd.concat([asym[f"response_h{h}"], asym[asym_controls]], axis=1).dropna()
                try:
                    wald = fit_a.wald_test("dx_pos = dx_neg_abs", scalar=True)
                    asym_p = float(wald.pvalue)
                except Exception:
                    asym_p = np.nan
                diag = residual_diagnostics(fit_a.resid)
                rows.append(
                    {
                        "model": "LP_asymmetric",
                        "product": product,
                        "product_label": PRODUCT_LABELS.get(product, product),
                        "link": link,
                        "link_direction": link_direction,
                        "price_variant": price_variant,
                        "farm_gate_source": farm_gate_source,
                        "reconstruction_variant": reconstruction_variant,
                        "horizon_days": h,
                        "n_obs": int(len(d_used)),
                        "coef_positive_shock": float(fit_a.params.get("dx_pos", np.nan)),
                        "p_positive_shock": float(fit_a.pvalues.get("dx_pos", np.nan)),
                        "coef_negative_shock_abs": float(fit_a.params.get("dx_neg_abs", np.nan)),
                        "p_negative_shock_abs": float(fit_a.pvalues.get("dx_neg_abs", np.nan)),
                        "asymmetry_pvalue": asym_p,
                        "asymmetry_gap": float(fit_a.params.get("dx_pos", np.nan) - fit_a.params.get("dx_neg_abs", np.nan)),
                        "r2": float(getattr(fit_a, "rsquared", np.nan)),
                        "core_signal": int(
                            (len(d_used) >= MIN_LP_OBS)
                            and (
                                float(fit_a.pvalues.get("dx_pos", 1.0)) < 0.10
                                or float(fit_a.pvalues.get("dx_neg_abs", 1.0)) < 0.10
                            )
                        ),
                        **diag,
                    }
                )
    return rows


def lp_link_specs() -> List[Dict[str, str]]:
    specs: List[Dict[str, str]] = []
    for variant in ["linear", "pchip"]:
        producer = f"producer_{variant}_model"
        for fg_source in ["initial", "filled"]:
            farm = f"farmgate_{fg_source}_{variant}_model"
            specs.extend(
                [
                    {
                        "x": farm,
                        "y": producer,
                        "link": "FarmGateUA -> ProducerUA",
                        "direction": "forward",
                        "price_variant": "processor_price",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": farm,
                        "y": "prozorro_model",
                        "link": "FarmGateUA -> ProZorro",
                        "direction": "forward",
                        "price_variant": "procurement_price",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": farm,
                        "y": "retail_observed_model",
                        "link": "FarmGateUA -> Retail",
                        "direction": "forward",
                        "price_variant": "retail_observed",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": "retail_observed_model",
                        "y": farm,
                        "link": "Retail -> FarmGateUA",
                        "direction": "reverse",
                        "price_variant": "retail_observed",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": producer,
                        "y": farm,
                        "link": "ProducerUA -> FarmGateUA",
                        "direction": "reverse",
                        "price_variant": "processor_price",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": farm,
                        "y": "retail_optimal_observed_model",
                        "link": "FarmGateUA -> Retail optimal",
                        "direction": "forward",
                        "price_variant": "retail_optimal_observed",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "x": "retail_optimal_observed_model",
                        "y": farm,
                        "link": "Retail optimal -> FarmGateUA",
                        "direction": "reverse",
                        "price_variant": "retail_optimal_observed",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                ]
            )
        specs.extend(
            [
                {
                    "x": producer,
                    "y": "prozorro_model",
                    "link": "ProducerUA -> ProZorro",
                    "direction": "forward",
                    "price_variant": "procurement_price",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "retail_observed_model",
                    "link": "ProZorro -> Retail",
                    "direction": "forward",
                    "price_variant": "retail_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "retail_baseline_model",
                    "link": "ProZorro -> Retail",
                    "direction": "forward",
                    "price_variant": "retail_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_observed_model",
                    "y": "prozorro_model",
                    "link": "Retail -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_baseline_model",
                    "y": "prozorro_model",
                    "link": "Retail -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": producer,
                    "link": "ProZorro -> ProducerUA",
                    "direction": "reverse",
                    "price_variant": "processor_price",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
            ]
        )
        specs.extend(
            [
                {
                    "x": "prozorro_model",
                    "y": "retail_matched_observed_model",
                    "link": "ProZorro -> Retail matched",
                    "direction": "forward",
                    "price_variant": "retail_matched_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "retail_matched_baseline_model",
                    "link": "ProZorro -> Retail matched",
                    "direction": "forward",
                    "price_variant": "retail_matched_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_matched_observed_model",
                    "y": "prozorro_model",
                    "link": "Retail matched -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_matched_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_matched_baseline_model",
                    "y": "prozorro_model",
                    "link": "Retail matched -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_matched_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "silpo_observed_model",
                    "link": "ProZorro -> Silpo",
                    "direction": "forward",
                    "price_variant": "silpo_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "silpo_baseline_model",
                    "link": "ProZorro -> Silpo",
                    "direction": "forward",
                    "price_variant": "silpo_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "silpo_observed_model",
                    "y": "prozorro_model",
                    "link": "Silpo -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "silpo_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "silpo_baseline_model",
                    "y": "prozorro_model",
                    "link": "Silpo -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "silpo_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "novus_observed_model",
                    "link": "ProZorro -> Novus",
                    "direction": "forward",
                    "price_variant": "novus_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "novus_observed_model",
                    "y": "prozorro_model",
                    "link": "Novus -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "novus_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "retail_optimal_observed_model",
                    "link": "ProZorro -> Retail optimal",
                    "direction": "forward",
                    "price_variant": "retail_optimal_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "prozorro_model",
                    "y": "retail_optimal_baseline_model",
                    "link": "ProZorro -> Retail optimal",
                    "direction": "forward",
                    "price_variant": "retail_optimal_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_optimal_observed_model",
                    "y": "prozorro_model",
                    "link": "Retail optimal -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_optimal_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "x": "retail_optimal_baseline_model",
                    "y": "prozorro_model",
                    "link": "Retail optimal -> ProZorro",
                    "direction": "reverse",
                    "price_variant": "retail_optimal_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
            ]
        )
    return specs


def run_local_projections(panel: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    specs = lp_link_specs()
    available = set(panel.columns)
    for product, g in panel.groupby("product"):
        g = g.sort_values("date")
        for spec in specs:
            if spec["x"] not in available or spec["y"] not in available:
                continue
            rows.extend(
                run_local_projection_for_pair(
                    g,
                    product=product,
                    y_col=spec["y"],
                    x_col=spec["x"],
                    link=spec["link"],
                    link_direction=spec["direction"],
                    price_variant=spec["price_variant"],
                    farm_gate_source=spec["farm_gate_source"],
                    reconstruction_variant=spec["reconstruction_variant"],
                )
            )
    return pd.DataFrame(rows)


def summarize_lp(lp: pd.DataFrame) -> pd.DataFrame:
    if lp.empty:
        return pd.DataFrame()
    linear = lp[lp["model"].eq("LP_linear")].copy()
    if linear.empty:
        return pd.DataFrame()
    summary = (
        linear.groupby(["link", "link_direction", "price_variant", "horizon_days"], as_index=False)
        .agg(
            models=("coef", "size"),
            products=("product", pd.Series.nunique),
            variants=("reconstruction_variant", pd.Series.nunique),
            mean_coef=("coef", "mean"),
            median_coef=("coef", "median"),
            mean_abs_coef=("coef", lambda s: float(np.nanmean(np.abs(s)))),
            sig_share=("pvalue", lambda s: float(np.nanmean(s < 0.10))),
            core_share=("core_signal", "mean"),
            median_n_obs=("n_obs", "median"),
        )
        .sort_values(["horizon_days", "link"])
    )
    return summary


def margin_specs() -> List[Dict[str, str]]:
    specs = []
    for variant in ["linear", "pchip"]:
        producer = f"producer_{variant}_model"
        for fg_source in ["initial", "filled"]:
            farm = f"farmgate_{fg_source}_{variant}_model"
            specs.extend(
                [
                    {
                        "upstream": farm,
                        "downstream": producer,
                        "spread": "ProducerUA / FarmGateUA",
                        "stage": "producer_farmgate",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                {
                    "upstream": farm,
                    "downstream": "retail_observed_model",
                    "spread": "Retail observed / FarmGateUA",
                        "stage": "retail_farmgate_observed",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                    {
                        "upstream": farm,
                        "downstream": "retail_baseline_model",
                        "spread": "Retail baseline / FarmGateUA",
                        "stage": "retail_farmgate_baseline",
                        "farm_gate_source": fg_source,
                        "reconstruction_variant": variant,
                    },
                ]
            )
        specs.extend(
            [
                {
                    "upstream": producer,
                    "downstream": "prozorro_model",
                    "spread": "ProZorro / ProducerUA",
                    "stage": "procurement_producer",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_observed_model",
                    "spread": "Retail observed / ProZorro",
                    "stage": "retail_procurement_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_baseline_model",
                    "spread": "Retail baseline / ProZorro",
                    "stage": "retail_procurement_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_matched_observed_model",
                    "spread": "Retail matched observed / ProZorro",
                    "stage": "retail_matched_procurement_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_matched_baseline_model",
                    "spread": "Retail matched baseline / ProZorro",
                    "stage": "retail_matched_procurement_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "silpo_observed_model",
                    "spread": "Silpo observed / ProZorro",
                    "stage": "silpo_procurement_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "silpo_baseline_model",
                    "spread": "Silpo baseline / ProZorro",
                    "stage": "silpo_procurement_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "novus_observed_model",
                    "spread": "Novus observed / ProZorro",
                    "stage": "novus_procurement_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_optimal_observed_model",
                    "spread": "Retail optimal observed / ProZorro",
                    "stage": "retail_optimal_procurement_observed",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
                {
                    "upstream": "prozorro_model",
                    "downstream": "retail_optimal_baseline_model",
                    "spread": "Retail optimal baseline / ProZorro",
                    "stage": "retail_optimal_procurement_baseline",
                    "farm_gate_source": "not_applicable",
                    "reconstruction_variant": variant,
                },
            ]
        )
    return specs


def run_margin_models(panel: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    rows: List[Dict[str, object]] = []
    spread_rows: List[Dict[str, object]] = []
    specs = margin_specs()
    for product, g in panel.groupby("product"):
        g = g.sort_values("date").copy()
        for spec in specs:
            if spec["upstream"] not in g.columns or spec["downstream"] not in g.columns:
                continue
            d = g[["date", spec["upstream"], spec["downstream"], "retail_discount_share", "retail_discount_depth"]].copy()
            d["lu"] = safe_log(d[spec["upstream"]])
            d["ld"] = safe_log(d[spec["downstream"]])
            d["spread"] = d["ld"] - d["lu"]
            d["dspread"] = d["spread"].diff()
            d["lag_spread"] = d["spread"].shift(1)
            d["du"] = d["lu"].diff()
            d["upstream_pos"] = d["du"].clip(lower=0)
            d["upstream_neg_abs"] = (-d["du"].clip(upper=0)).clip(lower=0)
            d["discount_share"] = safe_num(d["retail_discount_share"]).fillna(0)
            d["discount_depth"] = safe_num(d["retail_discount_depth"]).fillna(0)
            usable = d[["spread", "dspread", "lag_spread", "upstream_pos", "upstream_neg_abs"]].dropna()
            if len(usable) < MIN_MARGIN_OBS:
                continue
            exog_cols = ["lag_spread", "upstream_pos", "upstream_neg_abs"]
            if d["discount_share"].std(skipna=True) > 1e-8:
                exog_cols.append("discount_share")
            if d["discount_depth"].std(skipna=True) > 1e-8:
                exog_cols.append("discount_depth")
            fit = fit_hac_ols(d["dspread"], d[exog_cols], hac_lag=7)
            spread_clean = d["spread"].dropna()
            spread_rows.append(
                {
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "stage": spec["stage"],
                    "spread": spec["spread"],
                    "farm_gate_source": spec["farm_gate_source"],
                    "reconstruction_variant": spec["reconstruction_variant"],
                    "n_obs": int(len(spread_clean)),
                    "mean_log_spread": float(spread_clean.mean()),
                    "median_log_spread": float(spread_clean.median()),
                    "sd_log_spread": float(spread_clean.std(ddof=1)),
                    "mean_price_ratio": float(np.exp(spread_clean.mean())),
                    "date_min": d.loc[d["spread"].notna(), "date"].min(),
                    "date_max": d.loc[d["spread"].notna(), "date"].max(),
                }
            )
            if fit is None:
                continue
            used = pd.concat([d["dspread"], d[exog_cols]], axis=1).dropna()
            try:
                wald = fit.wald_test("upstream_pos = upstream_neg_abs", scalar=True)
                asym_p = float(wald.pvalue)
            except Exception:
                asym_p = np.nan
            diag = residual_diagnostics(fit.resid)
            rows.append(
                {
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "stage": spec["stage"],
                    "spread": spec["spread"],
                    "farm_gate_source": spec["farm_gate_source"],
                    "reconstruction_variant": spec["reconstruction_variant"],
                    "n_obs": int(len(used)),
                    "lag_spread_coef": float(fit.params.get("lag_spread", np.nan)),
                    "lag_spread_p": float(fit.pvalues.get("lag_spread", np.nan)),
                    "upstream_pos_coef": float(fit.params.get("upstream_pos", np.nan)),
                    "upstream_pos_p": float(fit.pvalues.get("upstream_pos", np.nan)),
                    "upstream_neg_abs_coef": float(fit.params.get("upstream_neg_abs", np.nan)),
                    "upstream_neg_abs_p": float(fit.pvalues.get("upstream_neg_abs", np.nan)),
                    "asymmetry_pvalue": asym_p,
                    "discount_share_coef": float(fit.params.get("discount_share", np.nan)),
                    "discount_share_p": float(fit.pvalues.get("discount_share", np.nan)),
                    "discount_depth_coef": float(fit.params.get("discount_depth", np.nan)),
                    "discount_depth_p": float(fit.pvalues.get("discount_depth", np.nan)),
                    "r2": float(getattr(fit, "rsquared", np.nan)),
                    "persistent_margin_flag": int(
                        not (
                            float(fit.params.get("lag_spread", 0)) < 0
                            and float(fit.pvalues.get("lag_spread", 1.0)) < 0.10
                        )
                    ),
                    "asymmetric_margin_flag": int(pd.notna(asym_p) and asym_p < 0.10),
                    **diag,
                }
            )
    return pd.DataFrame(rows), pd.DataFrame(spread_rows)


def run_discount_models(panel: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for product, g in panel.groupby("product"):
        g = g.sort_values("date").copy()
        cols = [
            "retail_discount_share",
            "retail_discount_depth",
            "retail_discount_percent",
            "retail_category_observed_model",
            "retail_category_baseline_model",
            "consumer_linear_model",
            "prozorro_model",
            "producer_linear_model",
            "farmgate_combined_linear_model",
            "silpo_category_observed_model",
            "novus_category_observed_model",
            "prozorro_expected_model",
            "prozorro_sum_initial_model",
            "prozorro_sum_current_model",
        ]
        if any(c not in g.columns for c in cols):
            cols = [c for c in cols if c in g.columns]
            if "retail_discount_share" not in cols or "retail_category_observed_model" not in cols or "prozorro_model" not in cols:
                continue
        d = g[["date"] + cols].copy()
        d["discount_share"] = safe_num(d["retail_discount_share"]).fillna(0)
        d["discount_percent"] = safe_num(d.get("retail_discount_percent", pd.Series(np.nan, index=d.index))).fillna(d["discount_share"] * 100.0)
        if d["discount_share"].std(skipna=True) <= 1e-8:
            continue
        d["retail_gap"] = safe_log(d["retail_category_baseline_model"]) - safe_log(d["retail_category_observed_model"])
        d["retail_procurement_spread"] = safe_log(d["retail_category_baseline_model"]) - safe_log(d["prozorro_model"])
        d["retail_vs_consumer_gap"] = safe_log(d["retail_category_observed_model"]) - safe_log(d.get("consumer_linear_model", pd.Series(np.nan, index=d.index)))
        d["silpo_novus_gap"] = safe_log(d.get("silpo_category_observed_model", pd.Series(np.nan, index=d.index))) - safe_log(d.get("novus_category_observed_model", pd.Series(np.nan, index=d.index)))
        d["d_producer"] = safe_log(d["producer_linear_model"]).diff()
        d["d_prozorro"] = safe_log(d["prozorro_model"]).diff()
        d["d_farmgate"] = safe_log(d["farmgate_combined_linear_model"]).diff()
        d["d_expected"] = np.log1p(safe_num(d.get("prozorro_expected_model", pd.Series(np.nan, index=d.index)))).diff()
        d["d_sum_initial"] = np.log1p(safe_num(d.get("prozorro_sum_initial_model", pd.Series(np.nan, index=d.index)))).diff()
        d["d_sum_current"] = np.log1p(safe_num(d.get("prozorro_sum_current_model", pd.Series(np.nan, index=d.index)))).diff()
        d["abs_d_producer"] = d["d_producer"].abs()
        d["abs_d_prozorro"] = d["d_prozorro"].abs()
        d["abs_d_farmgate"] = d["d_farmgate"].abs()
        d["abs_d_expected"] = d["d_expected"].abs()
        d["abs_d_sum_initial"] = d["d_sum_initial"].abs()
        d["abs_d_sum_current"] = d["d_sum_current"].abs()
        d["lag_discount"] = d["discount_share"].shift(1)
        d["lag_spread"] = d["retail_procurement_spread"].shift(1)
        exog_cols = ["lag_discount", "lag_spread", "abs_d_producer", "abs_d_prozorro", "abs_d_farmgate", "discount_percent", "retail_vs_consumer_gap", "silpo_novus_gap", "abs_d_expected", "abs_d_sum_initial", "abs_d_sum_current"]
        exog_cols = [c for c in exog_cols if c in d.columns and d[c].std(skipna=True) > 1e-8]
        used = d[["discount_share"] + exog_cols].replace([np.inf, -np.inf], np.nan).dropna()
        if len(used) < MIN_DISCOUNT_OBS:
            continue
        fit = fit_hac_ols(d["discount_share"], d[exog_cols], hac_lag=7)
        if fit is None:
            continue
        rows.append(
            {
                "product": product,
                "product_label": PRODUCT_LABELS.get(product, product),
                "n_obs": int(len(used)),
                "mean_discount_share": float(d["discount_share"].mean()),
                "lag_discount_coef": float(fit.params.get("lag_discount", np.nan)),
                "lag_discount_p": float(fit.pvalues.get("lag_discount", np.nan)),
                "lag_spread_coef": float(fit.params.get("lag_spread", np.nan)),
                "lag_spread_p": float(fit.pvalues.get("lag_spread", np.nan)),
                "abs_d_producer_coef": float(fit.params.get("abs_d_producer", np.nan)),
                "abs_d_producer_p": float(fit.pvalues.get("abs_d_producer", np.nan)),
                "abs_d_prozorro_coef": float(fit.params.get("abs_d_prozorro", np.nan)),
                "abs_d_prozorro_p": float(fit.pvalues.get("abs_d_prozorro", np.nan)),
                "abs_d_farmgate_coef": float(fit.params.get("abs_d_farmgate", np.nan)),
                "abs_d_farmgate_p": float(fit.pvalues.get("abs_d_farmgate", np.nan)),
                "discount_percent_coef": float(fit.params.get("discount_percent", np.nan)),
                "discount_percent_p": float(fit.pvalues.get("discount_percent", np.nan)),
                "retail_vs_consumer_gap_coef": float(fit.params.get("retail_vs_consumer_gap", np.nan)),
                "retail_vs_consumer_gap_p": float(fit.pvalues.get("retail_vs_consumer_gap", np.nan)),
                "silpo_novus_gap_coef": float(fit.params.get("silpo_novus_gap", np.nan)),
                "silpo_novus_gap_p": float(fit.pvalues.get("silpo_novus_gap", np.nan)),
                "abs_d_expected_coef": float(fit.params.get("abs_d_expected", np.nan)),
                "abs_d_expected_p": float(fit.pvalues.get("abs_d_expected", np.nan)),
                "abs_d_sum_initial_coef": float(fit.params.get("abs_d_sum_initial", np.nan)),
                "abs_d_sum_initial_p": float(fit.pvalues.get("abs_d_sum_initial", np.nan)),
                "abs_d_sum_current_coef": float(fit.params.get("abs_d_sum_current", np.nan)),
                "abs_d_sum_current_p": float(fit.pvalues.get("abs_d_sum_current", np.nan)),
                "r2": float(getattr(fit, "rsquared", np.nan)),
                "discount_strategy_signal": int(
                    float(fit.pvalues.get("lag_spread", 1.0)) < 0.10
                    or float(fit.pvalues.get("abs_d_producer", 1.0)) < 0.10
                    or float(fit.pvalues.get("abs_d_prozorro", 1.0)) < 0.10
                    or float(fit.pvalues.get("silpo_novus_gap", 1.0)) < 0.10
                    or float(fit.pvalues.get("abs_d_sum_current", 1.0)) < 0.10
                ),
            }
        )
    return pd.DataFrame(rows)


def run_procurement_scale_models(weekly: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    needed = ["prozorro_weekly", "producer_linear_weekly", "prozorro_expected_weekly", "prozorro_sum_initial_weekly", "prozorro_sum_current_weekly"]
    for product, g in weekly.groupby("product", dropna=False):
        if any(c not in g.columns for c in needed):
            continue
        d = g.sort_values("week").copy()
        d["d_prozorro"] = safe_log(d["prozorro_weekly"]).diff()
        d["d_producer"] = safe_log(d["producer_linear_weekly"]).diff()
        d["d_farmgate"] = safe_log(d.get("farmgate_combined_linear_weekly", pd.Series(np.nan, index=d.index))).diff()
        d["d_expected"] = np.log1p(safe_num(d["prozorro_expected_weekly"])).diff()
        d["d_sum_initial"] = np.log1p(safe_num(d["prozorro_sum_initial_weekly"])).diff()
        d["d_sum_current"] = np.log1p(safe_num(d["prozorro_sum_current_weekly"])).diff()
        d["lag_price"] = safe_log(d["prozorro_weekly"]).shift(1)
        exog_cols = ["lag_price", "d_producer", "d_farmgate", "d_expected", "d_sum_initial", "d_sum_current"]
        exog_cols = [c for c in exog_cols if d[c].std(skipna=True) > 1e-8]
        used = d[["d_prozorro"] + exog_cols].replace([np.inf, -np.inf], np.nan).dropna()
        if len(used) < MIN_CHAIN_OBS:
            continue
        fit = fit_hac_ols(d["d_prozorro"], d[exog_cols], hac_lag=4)
        if fit is None:
            continue
        rows.append(
            {
                "product": product,
                "product_label": PRODUCT_LABELS.get(product, product),
                "n_obs": int(len(used)),
                "lag_price_coef": float(fit.params.get("lag_price", np.nan)),
                "lag_price_p": float(fit.pvalues.get("lag_price", np.nan)),
                "d_producer_coef": float(fit.params.get("d_producer", np.nan)),
                "d_producer_p": float(fit.pvalues.get("d_producer", np.nan)),
                "d_farmgate_coef": float(fit.params.get("d_farmgate", np.nan)),
                "d_farmgate_p": float(fit.pvalues.get("d_farmgate", np.nan)),
                "d_expected_coef": float(fit.params.get("d_expected", np.nan)),
                "d_expected_p": float(fit.pvalues.get("d_expected", np.nan)),
                "d_sum_initial_coef": float(fit.params.get("d_sum_initial", np.nan)),
                "d_sum_initial_p": float(fit.pvalues.get("d_sum_initial", np.nan)),
                "d_sum_current_coef": float(fit.params.get("d_sum_current", np.nan)),
                "d_sum_current_p": float(fit.pvalues.get("d_sum_current", np.nan)),
                "r2": float(getattr(fit, "rsquared", np.nan)),
                "scale_signal_flag": int(
                    float(fit.pvalues.get("d_expected", 1.0)) < 0.10
                    or float(fit.pvalues.get("d_sum_initial", 1.0)) < 0.10
                    or float(fit.pvalues.get("d_sum_current", 1.0)) < 0.10
                ),
            }
        )
    return pd.DataFrame(rows)


def robust_findings(lp: pd.DataFrame, margins: pd.DataFrame, discounts: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    if not lp.empty:
        linear = lp[(lp["model"].eq("LP_linear")) & (lp["horizon_days"].isin([7, 14]))].copy()
        linear["sign"] = np.sign(linear["coef"])
        grouped = linear.groupby(["link", "product", "price_variant"], dropna=False)
        for (link, product, price_variant), g in grouped:
            core = g[g["core_signal"].eq(1)]
            if core.empty:
                continue
            signs = core["sign"].dropna().unique()
            rows.append(
                {
                    "finding_type": "local_projection",
                    "link_or_stage": link,
                    "product": product,
                    "product_label": PRODUCT_LABELS.get(product, product),
                    "price_variant": price_variant,
                    "n_core_rows": int(len(core)),
                    "median_coef": float(core["coef"].median()),
                    "directionally_stable": int(len(signs) == 1),
                    "interpretation_flag": "stronger" if len(core) >= 3 and len(signs) == 1 else "selective",
                }
            )
    if not margins.empty:
        for _, row in margins[margins.get("asymmetric_margin_flag", 0).eq(1) | margins.get("persistent_margin_flag", 0).eq(1)].iterrows():
            rows.append(
                {
                    "finding_type": "margin_power_proxy",
                    "link_or_stage": row["stage"],
                    "product": row["product"],
                    "product_label": row["product_label"],
                    "price_variant": row["spread"],
                    "n_core_rows": 1,
                    "median_coef": row.get("lag_spread_coef", np.nan),
                    "directionally_stable": np.nan,
                    "interpretation_flag": "proxy_only",
                }
            )
    if not discounts.empty:
        for _, row in discounts[discounts.get("discount_strategy_signal", 0).eq(1)].iterrows():
            rows.append(
                {
                    "finding_type": "discount_strategy",
                    "link_or_stage": "Silpo/retail discount incidence",
                    "product": row["product"],
                    "product_label": row["product_label"],
                    "price_variant": "discount_share",
                    "n_core_rows": 1,
                    "median_coef": row.get("lag_spread_coef", np.nan),
                    "directionally_stable": np.nan,
                    "interpretation_flag": "retail_price_management_signal",
                }
            )
    return pd.DataFrame(rows)


def save_plot(fig_path: Path, extra_paths: Optional[Sequence[Path]] = None) -> None:
    fig_path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(fig_path, dpi=180, bbox_inches="tight")
    for extra in extra_paths or []:
        extra.parent.mkdir(parents=True, exist_ok=True)
        plt.savefig(extra, dpi=180, bbox_inches="tight")
    plt.close()


def plot_outputs(panel: pd.DataFrame, coverage: pd.DataFrame, lp_summary: pd.DataFrame, margins_summary: pd.DataFrame, discounts: pd.DataFrame, source_frames: Dict[str, pd.DataFrame]) -> None:
    gov_frames = []
    for key, col in [("farm_gate", "farmgate_combined_linear"), ("producer", "producer_linear"), ("consumer", "consumer_linear")]:
        df = source_frames.get(key, pd.DataFrame()).copy()
        if not df.empty and col in df.columns:
            work = df[["date", col]].dropna().copy()
            if not work.empty:
                work["series"] = key
                work = work.rename(columns={col: "value"})
                gov_frames.append(work)
    if gov_frames:
        chart = pd.concat(gov_frames, ignore_index=True)
        plt.figure(figsize=(11, 5))
        for series, g in chart.groupby("series"):
            daily = g.groupby("date", as_index=False)["value"].median()
            plt.plot(daily["date"], daily["value"], label=series, linewidth=1.5)
        plt.title("Raw corrected governmental series used in the core chain")
        plt.ylabel("UAH")
        plt.legend(frameon=False)
        save_plot(FIG_SEQ_DIR / "01_raw_government_layers.png", [FIG_CH5_DIR / "01_raw_government_layers.png"])

    retail_items = source_frames.get("retail_items_full", pd.DataFrame())
    if not retail_items.empty:
        retail_focus_products = focus_products_from_counts(retail_items, limit=6)
        focus = retail_items[retail_items["product"].isin(retail_focus_products)].copy()
        if not focus.empty:
            chart = (
                focus.groupby(["date", "retailer", "product"], as_index=False)["effective_price"]
                .median()
                .dropna()
            )
            n_focus = max(1, len(retail_focus_products))
            ncols = 2
            nrows = int(np.ceil(n_focus / ncols))
            fig, axes = plt.subplots(nrows, ncols, figsize=(12, 3.8 * nrows), sharex=False)
            axes_arr = np.array(axes).reshape(-1)
            for ax, product in zip(axes_arr, retail_focus_products):
                g = chart[chart["product"].eq(product)]
                if g.empty:
                    continue
                for retailer, r in g.groupby("retailer"):
                    ax.plot(r["date"], r["effective_price"], label=retailer, linewidth=1.4)
                ax.set_title(PRODUCT_LABELS.get(product, product))
                ax.legend(frameon=False, fontsize=8)
                ax.set_ylabel("UAH per observed package")
                valid_dates = pd.to_datetime(g["date"], errors="coerce").dropna()
                if not valid_dates.empty:
                    ax.set_xlim(valid_dates.min(), valid_dates.max())
            for ax in axes_arr[len(retail_focus_products):]:
                ax.axis("off")
            fig.suptitle("Raw retail observed-price series before transformation")
            save_plot(FIG_SEQ_DIR / "02_raw_retail_observed_series.png", [FIG_CH5_DIR / "02_raw_retail_observed_series.png"])

    europe = source_frames.get("europe_benchmark", pd.DataFrame())
    cme = source_frames.get("cme_benchmark", pd.DataFrame())
    if not europe.empty or not cme.empty:
        fig, axes = plt.subplots(2, 1, figsize=(11, 8))
        if not europe.empty:
            e = europe[europe["product"].isin(["milk", "butter", "hard_cheese", "milk_powder"])].copy()
            for product, g in e.groupby("product"):
                axes[0].plot(g["date"], g["eu_price_uah"], label=PRODUCT_LABELS.get(product, product), linewidth=1.4)
            axes[0].set_title("European benchmark series in UAH/kg")
            axes[0].legend(frameon=False, fontsize=8)
        if not cme.empty:
            axes[1].plot(cme["date"], cme["cme_class3_uah"], color="#6a51a3", linewidth=1.5)
            axes[1].set_title("CME Class III benchmark")
        fig.suptitle("External benchmark block")
        save_plot(FIG_SEQ_DIR / "03_raw_external_benchmarks.png", [FIG_CH5_DIR / "03_raw_external_benchmarks.png"])

    index_daily = source_frames.get("aggregate_index_daily", pd.DataFrame())
    if not index_daily.empty:
        fig, axes = plt.subplots(2, 2, figsize=(13, 9), sharex=False)
        dataset_specs = [
            ("producer", source_frames.get("producer", pd.DataFrame()), "producer_linear", "producer_index", "ProducerUA product lines and aggregate index"),
            ("prozorro", source_frames.get("prozorro", pd.DataFrame()), "prozorro", "procurement_index", "Procurement product lines and aggregate index"),
            ("retail", panel[["date", "product", "retail_category_baseline"]].dropna() if "retail_category_baseline" in panel.columns else pd.DataFrame(), "retail_category_baseline", "retail_index", "Retail combined product lines and aggregate index"),
            ("consumer", source_frames.get("consumer", pd.DataFrame()), "consumer_linear", "consumer_index", "ConsumerUA product lines and aggregate index"),
        ]
        focus_products = focus_products_from_counts(panel, limit=6)
        for ax, (_, df, value_col, index_col, title) in zip(axes.flat, dataset_specs):
            if df.empty or value_col not in df.columns:
                ax.axis("off")
                continue
            for product in focus_products:
                g = df[df["product"].eq(product)][["date", value_col]].dropna()
                if len(g) == 0:
                    continue
                ax.plot(g["date"], g[value_col], linewidth=0.8, alpha=0.55, label=PRODUCT_LABELS.get(product, product))
            if index_col in index_daily.columns:
                idx_sub = index_daily[["date", index_col]].dropna()
                ax.plot(idx_sub["date"], idx_sub[index_col], color="black", linewidth=2.0, label="Aggregate index")
            ax.set_title(title)
            ax.set_ylabel("UAH per kg/l")
            dates = []
            if not df.empty and "date" in df.columns:
                dates.extend(pd.to_datetime(df["date"], errors="coerce").dropna().tolist())
            if index_col in index_daily.columns:
                dates.extend(pd.to_datetime(index_daily.loc[index_daily[index_col].notna(), "date"], errors="coerce").dropna().tolist())
            if dates:
                ax.set_xlim(min(dates), max(dates))
        handles, labels = axes.flat[0].get_legend_handles_labels()
        if handles:
            fig.legend(handles, labels, loc="lower center", ncol=4, frameon=False, fontsize=8)
        fig.suptitle("Product-level price paths and aggregate dairy-chain indices")
        save_plot(FIG_CH5_DIR / "04_dataset_product_lines_and_indices.png", [FIG_DIR / "18_dataset_product_lines_and_indices.png"])

        plt.figure(figsize=(11, 5.5))
        for col, label, color in [
            ("farmgate_index", "Farm-gate index", "#636363"),
            ("producer_index", "Producer index", "#3182bd"),
            ("procurement_index", "Procurement index", "#756bb1"),
            ("retail_index", "Retail index", "#e6550d"),
            ("consumer_index", "Consumer index", "#31a354"),
        ]:
            if col in index_daily.columns and index_daily[col].notna().sum() > 0:
                plt.plot(index_daily["date"], index_daily[col], label=label, linewidth=1.7, color=color)
        plt.title("Aggregate dairy price indices by chain level")
        plt.ylabel("Index level, UAH-equivalent")
        plt.legend(frameon=False, ncol=3, fontsize=8)
        if "date" in index_daily.columns and not index_daily.empty:
            plt.xlim(index_daily["date"].min(), index_daily["date"].max())
        save_plot(FIG_CH5_DIR / "05_aggregate_chain_indices.png", [FIG_DIR / "19_aggregate_chain_indices.png"])

    cov_cols = [c for c in coverage.columns if c.endswith("_n") and any(k in c for k in ["producer", "prozorro", "retail_observed", "retail_matched", "retail_optimal", "farmgate_filled", "consumer", "silpo", "novus"])]
    if cov_cols:
        data = coverage.set_index("product_label")[cov_cols].fillna(0)
        plt.figure(figsize=(11, max(4, 0.45 * len(data))))
        plt.imshow(data.values, aspect="auto", cmap="Blues")
        plt.colorbar(label="non-missing product-days")
        plt.yticks(range(len(data.index)), data.index)
        plt.xticks(range(len(data.columns)), [c.replace("_n", "") for c in data.columns], rotation=45, ha="right")
        plt.title("Panel coverage by product and source")
        save_plot(FIG_DIR / "01_panel_coverage.png", [FIG_CH5_DIR / "06_panel_coverage.png"])

    if not lp_summary.empty:
        focus = lp_summary[
            lp_summary["link"].isin(
                [
                    "FarmGateUA -> ProZorro",
                    "ProducerUA -> ProZorro",
                    "ProZorro -> Retail",
                    "Retail -> ProZorro",
                    "Retail -> FarmGateUA",
                    "ProZorro -> Silpo",
                    "ProZorro -> Novus",
                    "Silpo -> ProZorro",
                    "Novus -> ProZorro",
                ]
            )
            & lp_summary["price_variant"].isin(["retail_observed", "procurement_price", "processor_price", "silpo_observed", "novus_observed"])
        ].copy()
        if not focus.empty:
            plt.figure(figsize=(11, 6))
            for label, g in focus.groupby("link"):
                h = g.groupby("horizon_days", as_index=False)["mean_coef"].mean()
                plt.plot(h["horizon_days"], h["mean_coef"], marker="o", label=label)
            plt.axhline(0, color="black", linewidth=0.8)
            plt.xlabel("Horizon, days")
            plt.ylabel("Mean local-projection coefficient")
            plt.title("Local-projection pass-through by horizon")
            plt.legend(fontsize=8)
            save_plot(FIG_DIR / "02_lp_pass_through_horizons.png", [FIG_CH6_DIR / "06_lp_pass_through_horizons.png"])

        h7 = lp_summary[lp_summary["horizon_days"].isin([7, 14])].copy()
        if not h7.empty:
            chart = h7.groupby(["link_direction", "link"], as_index=False)["core_share"].mean().sort_values("core_share")
            plt.figure(figsize=(11, max(4, 0.35 * len(chart))))
            colors = chart["link_direction"].map({"forward": "#2c7fb8", "reverse": "#f03b20"}).fillna("#969696")
            plt.barh(chart["link"], chart["core_share"], color=colors)
            plt.xlabel("Core-signal share at 7/14-day horizons")
            plt.title("Forward versus reverse second-stage evidence")
            save_plot(FIG_DIR / "03_forward_reverse_core_share.png", [FIG_CH6_DIR / "07_forward_reverse_core_share.png"])

    if not margins_summary.empty:
        chart = margins_summary.groupby("spread", as_index=False)["mean_price_ratio"].median().sort_values("mean_price_ratio")
        plt.figure(figsize=(10, max(4, 0.35 * len(chart))))
        plt.barh(chart["spread"], chart["mean_price_ratio"], color="#756bb1")
        plt.axvline(1, color="black", linewidth=0.8)
        plt.xlabel("Median exp(mean log spread)")
        plt.title("Vertical spread / market-power proxy by chain segment")
        save_plot(FIG_DIR / "04_vertical_spread_proxy.png", [FIG_CH6_DIR / "10_vertical_spread_proxy.png"])

    if not discounts.empty:
        chart = discounts.sort_values("mean_discount_share")
        plt.figure(figsize=(9, max(4, 0.45 * len(chart))))
        plt.barh(chart["product_label"], chart["mean_discount_share"], color="#31a354")
        plt.xlabel("Mean discount share")
        plt.title("Retail discount incidence by product")
        save_plot(FIG_DIR / "05_discount_incidence.png", [FIG_CH6_DIR / "11_discount_incidence.png"])

    if "retail_match_audit" in source_frames:
        audit = source_frames["retail_match_audit"].copy()
        if not audit.empty and "match_status" in audit.columns:
            chart = audit["match_status"].value_counts().rename_axis("match_status").reset_index(name="n")
            plt.figure(figsize=(8, 5))
            plt.bar(chart["match_status"], chart["n"], color=["#238b45", "#9ecae1", "#fc9272"][: len(chart)])
            plt.ylabel("Unique harmonized item keys")
            plt.title("Cross-shop retail item harmonization status")
            plt.xticks(rotation=15, ha="right")
            save_plot(FIG_DIR / "06_cross_shop_match_status.png", [FIG_CH5_DIR / "07_cross_shop_match_status.png"])

    literal_summary = source_frames.get("retail_literal_summary", pd.DataFrame())
    if not literal_summary.empty:
        chart = (
            literal_summary.groupby(["product", "retailer"], as_index=False)["n_item_keys"]
            .sum()
            .assign(product_label=lambda d: d["product"].map(PRODUCT_LABELS).fillna(d["product"]))
        )
        pivot = chart.pivot_table(index="product_label", columns="retailer", values="n_item_keys", aggfunc="sum", fill_value=0)
        if not pivot.empty:
            colors = {"Silpo": "#2c7fb8", "Novus": "#f03b20"}
            plt.figure(figsize=(11, max(4, 0.45 * len(pivot))))
            left = np.zeros(len(pivot))
            for retailer in pivot.columns:
                vals = pivot[retailer].to_numpy(dtype=float)
                plt.barh(pivot.index, vals, left=left, color=colors.get(retailer, "#969696"), label=retailer)
                left += vals
            plt.xlabel("Unique harmonized item keys")
            plt.title("Retail product distribution by retailer")
            plt.legend(frameon=False)
            save_plot(FIG_DIR / "07_retail_literal_mix.png", [FIG_CH5_DIR / "07_retail_product_distribution.png"])

    brand_support = source_frames.get("retail_brand_support", pd.DataFrame())
    if not brand_support.empty:
        chart = (
            brand_support.groupby(["product", "retailer"], as_index=False)
            .agg(n_brands=("brand_norm", lambda s: int(pd.Series(s).fillna("").replace("", np.nan).dropna().nunique())))
            .assign(product_label=lambda d: d["product"].map(PRODUCT_LABELS).fillna(d["product"]))
        )
        pivot = chart.pivot_table(index="product_label", columns="retailer", values="n_brands", aggfunc="sum", fill_value=0)
        if not pivot.empty:
            colors = {"Silpo": "#2c7fb8", "Novus": "#f03b20"}
            plt.figure(figsize=(11, max(4, 0.45 * len(pivot))))
            x = np.arange(len(pivot.index))
            width = 0.38
            for idx, retailer in enumerate(pivot.columns):
                plt.bar(
                    x + (idx - (len(pivot.columns) - 1) / 2) * width,
                    pivot[retailer].to_numpy(dtype=float),
                    width=width,
                    color=colors.get(retailer, "#969696"),
                    label=retailer,
                )
            plt.xticks(x, pivot.index, rotation=35, ha="right")
            plt.ylabel("Distinct normalized brands")
            plt.title("Retail brand distribution by retailer")
            plt.legend(frameon=False)
            save_plot(FIG_DIR / "08_dominant_brand_support.png", [FIG_CH5_DIR / "08_retail_brand_distribution.png"])

    region_profile = source_frames.get("prozorro_region_profile", pd.DataFrame())
    if not region_profile.empty:
        top_regions = region_profile.groupby("region", as_index=False)["n_obs"].sum().sort_values("n_obs", ascending=False).head(10)["region"]
        region_chart = region_profile[region_profile["region"].isin(top_regions)].copy()
        pivot = region_chart.pivot_table(index="product_label", columns="region", values="median_price", aggfunc="median")
        if not pivot.empty:
            plt.figure(figsize=(12, max(4, 0.45 * len(pivot))))
            plt.imshow(pivot.fillna(0).values, aspect="auto", cmap="YlOrBr")
            plt.colorbar(label="Median ProZorro unit price")
            plt.yticks(range(len(pivot.index)), pivot.index)
            plt.xticks(range(len(pivot.columns)), pivot.columns, rotation=45, ha="right")
            plt.title("Regional procurement-price profile for leading ProZorro regions")
            save_plot(FIG_CH5_DIR / "10_prozorro_region_profile.png", [FIG_DIR / "20_prozorro_region_profile.png"])

    level_scores = source_frames.get("retail_level_scores", pd.DataFrame())
    if not level_scores.empty:
        pivot = level_scores.pivot_table(index="product_label", columns="candidate_label", values="selection_score", aggfunc="max", fill_value=0)
        if not pivot.empty:
            plt.figure(figsize=(11, max(4, 0.45 * len(pivot))))
            plt.imshow(pivot.values, aspect="auto", cmap="PuBu")
            plt.colorbar(label="Composite selection score")
            plt.yticks(range(len(pivot.index)), pivot.index)
            plt.xticks(range(len(pivot.columns)), pivot.columns, rotation=45, ha="right")
            plt.title("Retail-level candidate scores by product")
            save_plot(FIG_DIR / "09_retail_level_scores.png", [FIG_CH5_DIR / "13_retail_level_scores.png"])

    level_selection = source_frames.get("retail_level_selection", pd.DataFrame())
    if not level_selection.empty:
        chosen = level_selection.sort_values("selection_score")
        codes = {label: idx for idx, label in enumerate(sorted(chosen["candidate_label"].unique()))}
        colors = [plt.cm.Set2(codes[label] / max(1, len(codes) - 1)) for label in chosen["candidate_label"]]
        plt.figure(figsize=(11, max(4, 0.45 * len(chosen))))
        plt.barh(chosen["product_label"], chosen["selection_score"], color=colors)
        plt.xlabel("Composite selection score")
        plt.title("Chosen optimal retail level by product")
        save_plot(FIG_DIR / "10_optimal_retail_level.png", [FIG_CH5_DIR / "14_optimal_retail_level.png"])

    if not lp_summary.empty:
        candidate_links = [
            "ProZorro -> Retail",
            "ProZorro -> Retail matched",
            "ProZorro -> Silpo",
            "ProZorro -> Novus",
            "ProZorro -> Retail optimal",
        ]
        chart = (
            lp_summary[
                lp_summary["link"].isin(candidate_links)
                & lp_summary["horizon_days"].isin([7, 14])
                & lp_summary["price_variant"].isin(
                    [
                        "retail_observed",
                        "retail_matched_observed",
                        "silpo_observed",
                        "novus_observed",
                        "retail_optimal_observed",
                    ]
                )
            ]
            .groupby("link", as_index=False)["core_share"]
            .mean()
            .sort_values("core_share")
        )
        if not chart.empty:
            plt.figure(figsize=(10, max(4, 0.45 * len(chart))))
            plt.barh(chart["link"], chart["core_share"], color="#dd3497")
            plt.xlabel("Mean 7/14-day core share")
            plt.title("Candidate downstream levels in procurement-to-retail LP tests")
            save_plot(FIG_DIR / "11_candidate_downstream_core_share.png", [FIG_CH6_DIR / "08_candidate_downstream_core_share.png"])

    items = source_frames.get("retail_items_full", pd.DataFrame())
    if not items.empty:
        silpo_source = items[items["retailer"].eq("Silpo")].copy()
        silpo_chart = (
            silpo_source.groupby("product", as_index=False)
            .agg(
                mean_discount_share=("discount_present", "mean"),
                median_markdown=(
                    "markdown_rate",
                    lambda s: float(s[s > 0].median()) if pd.Series(s).gt(0).any() else 0.0,
                ),
            )
            .assign(product_label=lambda d: d["product"].map(PRODUCT_LABELS).fillna(d["product"]))
            .sort_values("mean_discount_share")
        )
        if not silpo_chart.empty:
            fig, axes = plt.subplots(1, 2, figsize=(11, max(4, 0.45 * len(silpo_chart))), sharey=True)
            axes[0].barh(silpo_chart["product_label"], silpo_chart["mean_discount_share"], color="#31a354")
            axes[0].set_title("Silpo mean discount share")
            axes[0].set_xlabel("Share of discounted observations")
            axes[1].barh(silpo_chart["product_label"], silpo_chart["median_markdown"], color="#ef6548")
            axes[1].set_title("Silpo median markdown depth")
            axes[1].set_xlabel("Markdown depth")
            fig.suptitle("Silpo discount environment by dairy product")
            save_plot(FIG_DIR / "12_discount_environment.png", [FIG_CH5_DIR / "11_silpo_discount_environment.png"])


def plot_final_outputs(
    weekly: pd.DataFrame,
    corr_scan: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    scale_models: pd.DataFrame,
    index_models: pd.DataFrame,
    link21_summary: pd.DataFrame,
    reliability_overview: pd.DataFrame,
    discounts: pd.DataFrame,
    margins_summary: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
) -> None:
    if not weekly.empty:
        focus_products = prioritized_products(
            [p for p in ["milk", "butter", "cheese"] if p in set(weekly["product"])],
            [p for p in focus_products_from_counts(weekly, limit=6) if p not in {"milk", "butter", "cheese"}],
            6,
        )
        if focus_products:
            fig, axes = plt.subplots(len(focus_products), 1, figsize=(12, 3.5 * len(focus_products)), sharex=False)
            if len(focus_products) == 1:
                axes = [axes]
            for ax, product in zip(axes, focus_products):
                g = weekly[weekly["product"].eq(product)].sort_values("week")
                for col, label, color in [
                    ("farmgate_combined_linear_weekly", "FarmGate Ukraine average", "#636363"),
                    ("producer_linear_weekly", "Producer", "#3182bd"),
                    ("prozorro_weekly", "ProZorro unit price", "#756bb1"),
                    ("retail_category_baseline_weekly", "Retail combined", "#e6550d"),
                    ("consumer_linear_weekly", "ConsumerUA", "#31a354"),
                ]:
                    if col in g.columns and g[col].notna().sum() > 0:
                        sub = g.loc[g[col].notna(), ["week", col]].copy()
                        ax.plot(sub["week"], sub[col], label=label, linewidth=1.6, color=color)
                ax.set_title(PRODUCT_LABELS.get(product, product))
                valid_weeks = g[
                    g[
                        [
                            c
                            for c in [
                                "farmgate_combined_linear_weekly",
                                "producer_linear_weekly",
                                "prozorro_weekly",
                                "retail_category_baseline_weekly",
                                "consumer_linear_weekly",
                            ]
                            if c in g.columns
                        ]
                    ].notna().any(axis=1)
                ]["week"]
                if not valid_weeks.empty:
                    ax.set_xlim(valid_weeks.min(), valid_weeks.max())
                ax.set_ylabel("UAH per kg/l")
                ax.legend(frameon=False, ncol=3, fontsize=8)
            fig.suptitle("Weekly median chain overlays on corrected data", y=0.995)
            save_plot(FIG_DIR / "13_weekly_chain_overlay.png", [FIG_SEQ_DIR / "05_transformed_weekly_chain_overlay.png", FIG_CH5_DIR / "12_weekly_chain_overlay.png"])

    if not corr_scan.empty:
        chart = corr_scan.dropna(subset=["best_abs_corr"]).copy()
        chart["abs_corr"] = chart["best_abs_corr"].abs()
        chart = chart.sort_values("abs_corr", ascending=False).head(18)
        if not chart.empty:
            plt.figure(figsize=(12, max(4, 0.35 * len(chart))))
            labels = chart["product_label"] + " | " + chart["link"]
            plt.barh(labels, chart["abs_corr"], color="#6baed6")
            plt.gca().invert_yaxis()
            plt.xlabel("Best absolute weekly lag correlation")
            plt.title("Top weekly lag-correlation signals across admissible links")
            save_plot(FIG_DIR / "14_weekly_corr_scan.png", [FIG_SEQ_DIR / "06_weekly_lag_correlation_scan.png", FIG_CH6_DIR / "01_weekly_corr_scan.png"])
        fg_chart = corr_scan[corr_scan["link"].astype(str).str.contains("FarmGate")].dropna(subset=["best_abs_corr"]).copy()
        if not fg_chart.empty:
            fg_heat = fg_chart.pivot_table(index="product_label", columns="link", values="best_abs_corr", aggfunc="mean")
            if not fg_heat.empty:
                plt.figure(figsize=(12, max(4, 0.55 * len(fg_heat))))
                plt.imshow(fg_heat.fillna(0).values, aspect="auto", cmap="RdBu", vmin=-1, vmax=1)
                plt.colorbar(label="Best weekly lag correlation")
                plt.yticks(range(len(fg_heat.index)), fg_heat.index)
                plt.xticks(range(len(fg_heat.columns)), fg_heat.columns, rotation=35, ha="right")
                plt.title("Farm-gate transmission lag map across the chain")
                save_plot(FIG_CH6_DIR / "16_farmgate_lag_map.png", [FIG_DIR / "31_farmgate_lag_map.png"])

    if not link21_summary.empty:
        chart = link21_summary.copy()
        score_map = {"unusable": 0, "weak_extension": 1, "acceptable": 2, "strong": 3}
        chart["status_score"] = chart["best_admissibility"].map(score_map).fillna(0)
        chart["model_score"] = chart["model_reliability"].map({"": 0, "conditionally_usable": 1, "reliable": 2}).fillna(0)
        plot_df = chart[["status_score", "model_score"]].to_numpy(dtype=float)
        fig, axes = plt.subplots(1, 2, figsize=(12, max(6, 0.32 * len(chart))), sharey=True)
        im1 = axes[0].imshow(plot_df[:, [0]], aspect="auto", cmap="YlGnBu", vmin=0, vmax=3)
        axes[0].set_title("Admissibility")
        axes[0].set_xticks([0])
        axes[0].set_xticklabels(["status"])
        axes[0].set_yticks(range(len(chart)))
        axes[0].set_yticklabels(chart["link"])
        fig.colorbar(im1, ax=axes[0], fraction=0.046, pad=0.04)
        im2 = axes[1].imshow(plot_df[:, [1]], aspect="auto", cmap="OrRd", vmin=0, vmax=2)
        axes[1].set_title("Retained model reliability")
        axes[1].set_xticks([0])
        axes[1].set_xticklabels(["best model"])
        fig.colorbar(im2, ax=axes[1], fraction=0.046, pad=0.04)
        fig.suptitle("Summary of the 21 core directional chain links")
        save_plot(FIG_CH6_DIR / "02_link21_status_matrix.png", [FIG_DIR / "21_link21_status_matrix.png"])

    if not chain_models.empty:
        chart = chain_models[chain_models["model_reliability"].isin(["reliable", "conditionally_usable"])].copy()
        chart["coef_display"] = chart["lr_coef"].where(chart["lr_coef"].notna(), chart["sr_coef"])
        chart = chart.dropna(subset=["coef_display"]).sort_values("coef_display").head(20)
        if not chart.empty:
            plt.figure(figsize=(12, max(4, 0.35 * len(chart))))
            labels = chart["product_label"] + " | " + chart["model_family"] + " | " + chart["link"]
            colors = chart["model_reliability"].map({"reliable": "#238b45", "conditionally_usable": "#fd8d3c"}).fillna("#969696")
            plt.barh(labels, chart["coef_display"], color=colors)
            plt.axvline(0, color="black", linewidth=0.8)
            plt.xlabel("Displayed coefficient (long-run when available)")
            plt.title("Core weekly model coefficients on corrected data")
            save_plot(FIG_DIR / "15_core_model_coefficients.png", [FIG_SEQ_DIR / "08_core_weekly_model_coefficients.png", FIG_CH6_DIR / "03_core_model_coefficients.png"])

        ecm = chain_models[chain_models["model_family"].eq("ECM")].copy()
        if not ecm.empty:
            ecm = ecm.sort_values("ect_coef")
            plt.figure(figsize=(11, max(4, 0.35 * len(ecm))))
            labels = ecm["product_label"] + " | " + ecm["link"] + " | " + ecm["data_variant"]
            colors = ecm["model_reliability"].map({"reliable": "#2ca25f", "conditionally_usable": "#feb24c"}).fillna("#969696")
            plt.barh(labels, ecm["ect_coef"], color=colors)
            plt.axvline(0, color="black", linewidth=0.8)
            plt.xlabel("Error-correction coefficient")
            plt.title("ECM speed of adjustment across retained weekly links")
            save_plot(FIG_CH6_DIR / "04_ecm_speed_of_adjustment.png", [FIG_DIR / "22_ecm_speed_of_adjustment.png"])

        nardl = chain_models[chain_models["model_family"].eq("NARDL")].copy()
        if not nardl.empty:
            nardl = nardl.sort_values("asymmetry_pvalue")
            plt.figure(figsize=(11, max(4, 0.35 * len(nardl))))
            labels = nardl["product_label"] + " | " + nardl["link"] + " | " + nardl["data_variant"]
            plt.barh(labels, -np.log10(nardl["asymmetry_pvalue"].clip(lower=1e-6)), color="#8856a7")
            plt.xlabel("-log10 asymmetry p-value")
            plt.title("NARDL asymmetry strength across retained weekly links")
            save_plot(FIG_CH6_DIR / "05_nardl_asymmetry.png", [FIG_DIR / "23_nardl_asymmetry.png"])

        farmgate_links = chart[chart["link"].astype(str).str.contains("FarmGate")].copy()
        if not farmgate_links.empty:
            farmgate_links = farmgate_links.sort_values("coef_display")
            plt.figure(figsize=(11, max(4, 0.35 * len(farmgate_links))))
            labels = farmgate_links["product_label"] + " | " + farmgate_links["link"] + " | " + farmgate_links["model_family"]
            colors = farmgate_links["model_reliability"].map({"reliable": "#238b45", "conditionally_usable": "#fd8d3c"}).fillna("#969696")
            plt.barh(labels, farmgate_links["coef_display"], color=colors)
            plt.axvline(0, color="black", linewidth=0.8)
            plt.xlabel("Displayed coefficient")
            plt.title("Farm-gate links within the retained weekly model set")
            save_plot(FIG_CH6_DIR / "17_farmgate_chain_coefficients.png", [FIG_DIR / "32_farmgate_chain_coefficients.png"])

    if not vecm_results.empty:
        chart = vecm_results.groupby(["system_name", "status"], as_index=False).size().rename(columns={"size": "n"})
        if not chart.empty:
            pivot = chart.pivot_table(index="system_name", columns="status", values="n", aggfunc="sum", fill_value=0)
            plt.figure(figsize=(10, max(4, 0.5 * len(pivot))))
            left = np.zeros(len(pivot))
            color_map = {"ok": "#31a354", "infeasible": "#fdae6b", "failed": "#de2d26"}
            for col in pivot.columns:
                vals = pivot[col].to_numpy(dtype=float)
                plt.barh(pivot.index, vals, left=left, color=color_map.get(col, "#969696"), label=col)
                left += vals
            plt.xlabel("Number of product-specific system attempts")
            plt.title("VECM system feasibility on corrected weekly panels")
            plt.legend(frameon=False)
            save_plot(FIG_DIR / "16_vecm_feasibility.png", [FIG_SEQ_DIR / "09_vecm_system_feasibility.png", FIG_CH6_DIR / "10_vecm_feasibility.png"])

    if not scale_models.empty:
        chart = scale_models.sort_values("d_sum_current_coef").copy()
        if not chart.empty:
            plt.figure(figsize=(10, max(4, 0.45 * len(chart))))
            plt.barh(chart["product_label"], chart["d_sum_current_coef"], color="#9e9ac8")
            plt.axvline(0, color="black", linewidth=0.8)
            plt.xlabel("Coefficient on weekly change in current contract sum")
            plt.title("Procurement scale effects on ProZorro price changes")
            save_plot(FIG_DIR / "17_procurement_scale_effects.png", [FIG_SEQ_DIR / "12_procurement_scale_effects.png", FIG_CH6_DIR / "13_procurement_scale_effects.png"])

    if not discounts.empty:
        coef_cols = [c for c in ["lag_discount_coef", "lag_spread_coef", "retail_vs_consumer_gap_coef", "silpo_novus_gap_coef"] if c in discounts.columns]
        plot_df = discounts[["product_label"] + coef_cols].set_index("product_label")
        if not plot_df.empty:
            plt.figure(figsize=(11, max(4, 0.45 * len(plot_df))))
            plt.imshow(plot_df.fillna(0).values, aspect="auto", cmap="RdBu")
            plt.colorbar(label="Coefficient")
            plt.yticks(range(len(plot_df.index)), plot_df.index)
            plt.xticks(range(len(plot_df.columns)), plot_df.columns, rotation=45, ha="right")
            plt.title("Discount-model coefficient map")
            save_plot(FIG_CH6_DIR / "12_discount_coefficient_map.png", [FIG_DIR / "24_discount_coefficient_map.png"])

    if not margins_summary.empty:
        chart = margins_summary[["product_label", "stage", "mean_price_ratio", "sd_log_spread"]].copy()
        chart["label"] = chart["product_label"] + " | " + chart["stage"]
        chart = chart.sort_values("mean_price_ratio")
        plt.figure(figsize=(11, max(4, 0.30 * len(chart))))
        plt.barh(chart["label"], chart["mean_price_ratio"], color="#756bb1")
        plt.axvline(1, color="black", linewidth=0.8)
        plt.xlabel("Mean price ratio")
        plt.title("Average spread levels across chain segments")
        save_plot(FIG_CH6_DIR / "08_spread_levels.png", [FIG_DIR / "25_spread_levels.png"])

        heat = chart.pivot_table(index="product_label", columns="stage", values="sd_log_spread", aggfunc="mean", fill_value=0)
        if not heat.empty:
            plt.figure(figsize=(12, max(4, 0.45 * len(heat))))
            plt.imshow(heat.values, aspect="auto", cmap="PuRd")
            plt.colorbar(label="Spread volatility (sd of log spread)")
            plt.yticks(range(len(heat.index)), heat.index)
            plt.xticks(range(len(heat.columns)), heat.columns, rotation=45, ha="right")
            plt.title("Spread volatility by product and chain segment")
            save_plot(FIG_CH6_DIR / "09_spread_volatility.png", [FIG_DIR / "26_spread_volatility.png"])

    if not weekly.empty:
        idx = (
            weekly.groupby("week", as_index=False)
            .agg(
                farmgate_avg=("farmgate_combined_linear_weekly", "median"),
                producer_avg=("producer_linear_weekly", "mean"),
                producer_med=("producer_linear_weekly", "median"),
                procurement_avg=("prozorro_weekly", "mean"),
                procurement_med=("prozorro_weekly", "median"),
                retail_avg=("retail_category_baseline_weekly", "mean"),
                retail_med=("retail_category_baseline_weekly", "median"),
                consumer_avg=("consumer_linear_weekly", "mean"),
                consumer_med=("consumer_linear_weekly", "median"),
            )
            .sort_values("week")
        )
        fig, axes = plt.subplots(3, 2, figsize=(13, 11), sharex=False)
        comparisons = [
            ("producer_avg", "FarmGate vs Producer average"),
            ("procurement_avg", "FarmGate vs Procurement average"),
            ("retail_avg", "FarmGate vs Retail average"),
            ("consumer_avg", "FarmGate vs Consumer average"),
            ("producer_med", "FarmGate vs Producer median"),
            ("retail_med", "FarmGate vs Retail median"),
        ]
        for ax, (col, title) in zip(axes.flat, comparisons):
            sub = idx.loc[idx[col].notna(), ["week", "farmgate_avg", col]].copy()
            if sub.empty:
                ax.axis("off")
                continue
            ax.plot(sub["week"], sub["farmgate_avg"], label="FarmGate Ukraine average", color="#636363", linewidth=1.5)
            ax.plot(sub["week"], sub[col], label=title.replace("FarmGate vs ", ""), color="#3182bd", linewidth=1.4)
            ax.set_title(title)
            ax.legend(frameon=False, fontsize=7)
            ax.set_xlim(sub["week"].min(), sub["week"].max())
        fig.suptitle("Farm-gate benchmark block against chain-level dairy price indices", y=0.995)
        save_plot(FIG_SEQ_DIR / "13_farmgate_benchmark_block.png", [FIG_CH5_DIR / "13_farmgate_benchmark_block.png"])

        focus_products = [p for p in ["milk", "butter", "cheese", "sour_cream"] if p in set(weekly["product"])]
        if focus_products:
            fig, axes = plt.subplots(len(focus_products), 1, figsize=(12, 3.2 * len(focus_products)), sharex=False)
            if len(focus_products) == 1:
                axes = [axes]
            for ax, product in zip(axes, focus_products):
                g = weekly[weekly["product"].eq(product)].sort_values("week").copy()
                plotted = False
                for col, label, color in [
                    ("farmgate_combined_linear_weekly", "FarmGate Ukraine average", "#636363"),
                    ("producer_linear_weekly", "Producer", "#3182bd"),
                    ("prozorro_weekly", "ProZorro unit price", "#756bb1"),
                    ("retail_category_baseline_weekly", "Retail combined baseline", "#e6550d"),
                    ("consumer_linear_weekly", "ConsumerUA", "#31a354"),
                ]:
                    if col not in g.columns:
                        continue
                    sub = g.loc[g[col].notna(), ["week", col]].copy()
                    if sub.empty:
                        continue
                    base = float(sub[col].iloc[0])
                    if not np.isfinite(base) or base <= 0:
                        continue
                    sub["index100"] = 100.0 * sub[col] / base
                    ax.plot(sub["week"], sub["index100"], label=label, linewidth=1.6, color=color)
                    plotted = True
                if not plotted:
                    ax.axis("off")
                    continue
                ax.axhline(100.0, color="black", linewidth=0.7, alpha=0.5)
                ax.set_title(f"{PRODUCT_LABELS.get(product, product)}: normalized bridge from farm-gate to retail")
                ax.set_ylabel("First overlap week = 100")
                ax.legend(frameon=False, ncol=3, fontsize=8)
            fig.suptitle("Farm-gate-to-retail comparison on a common normalized scale", y=0.995)
            save_plot(FIG_CH5_DIR / "15_farmgate_to_chain_normalized.png", [FIG_DIR / "30_farmgate_to_chain_normalized.png"])

    europe = source_frames.get("europe_benchmark", pd.DataFrame())
    cme = source_frames.get("cme_benchmark", pd.DataFrame())
    producer = source_frames.get("producer", pd.DataFrame())
    if not producer.empty and (not europe.empty or not cme.empty):
        corr_rows = []
        for product, g in producer.groupby("product", dropna=False):
            pg = g[["date", "producer_linear"]].rename(columns={"producer_linear": "producer"}).copy()
            if not europe.empty:
                eg = europe[europe["product"].eq(product)][["date", "eu_price_uah"]].rename(columns={"eu_price_uah": "benchmark"})
                m = pg.merge(eg, on="date", how="inner").dropna()
                if len(m) >= 20:
                    corr_rows.append({"product": PRODUCT_LABELS.get(product, product), "benchmark": "Europe", "corr": float(m["producer"].corr(m["benchmark"]))})
            if not cme.empty:
                cg = cme.rename(columns={"cme_class3_uah": "benchmark"})
                m = pg.merge(cg, on="date", how="inner").dropna()
                if len(m) >= 20:
                    corr_rows.append({"product": PRODUCT_LABELS.get(product, product), "benchmark": "CME III", "corr": float(m["producer"].corr(m["benchmark"]))})
        corr_df = pd.DataFrame(corr_rows)
        if not corr_df.empty:
            pivot = corr_df.pivot_table(index="product", columns="benchmark", values="corr", aggfunc="mean", fill_value=0)
            plt.figure(figsize=(8, max(4, 0.45 * len(pivot))))
            plt.imshow(pivot.values, aspect="auto", cmap="RdBu", vmin=-1, vmax=1)
            plt.colorbar(label="Correlation")
            plt.yticks(range(len(pivot.index)), pivot.index)
            plt.xticks(range(len(pivot.columns)), pivot.columns, rotation=45, ha="right")
            plt.title("Producer correlations with external benchmarks")
            save_plot(FIG_SEQ_DIR / "14_benchmark_correlations.png", [FIG_CH5_DIR / "14_benchmark_correlations.png"])

    if not reliability_overview.empty:
        chart = reliability_overview.copy()
        plt.figure(figsize=(10, max(4, 0.45 * len(chart))))
        plt.barh(chart["block"] + " | " + chart["family"], chart["reliable_rows"], color="#2ca25f", label="reliable")
        plt.barh(chart["block"] + " | " + chart["family"], chart["conditionally_usable_rows"], left=chart["reliable_rows"], color="#feb24c", label="conditional")
        plt.xlabel("Model rows")
        plt.title("Reliability profile across model blocks")
        plt.legend(frameon=False)
        save_plot(FIG_CH6_DIR / "15_reliability_overview.png", [FIG_DIR / "27_reliability_overview.png"])

    index_weekly = source_frames.get("aggregate_index_weekly", pd.DataFrame())
    if not index_weekly.empty:
        plt.figure(figsize=(11, 5.5))
        for col, label, color in [
            ("farmgate_index_weekly", "Farm-gate index", "#636363"),
            ("producer_index_weekly", "Producer index", "#3182bd"),
            ("procurement_index_weekly", "Procurement index", "#756bb1"),
            ("retail_index_weekly", "Retail index", "#e6550d"),
            ("consumer_index_weekly", "Consumer index", "#31a354"),
            ("downstream_extension_weekly", "Downstream extension index", "#dd3497"),
        ]:
            if col in index_weekly.columns and index_weekly[col].notna().sum() > 0:
                plt.plot(index_weekly["week"], index_weekly[col], label=label, linewidth=1.7, color=color)
        plt.title("Aggregate dairy-chain indices at weekly frequency")
        plt.legend(frameon=False, ncol=3, fontsize=8)
        save_plot(FIG_CH6_DIR / "11_aggregate_index_overlay.png", [FIG_DIR / "28_aggregate_index_overlay.png"])

    if not index_models.empty:
        chart = index_models.copy()
        chart["display_coef"] = chart["lr_coef"].where(chart["lr_coef"].notna(), chart["sr_coef"])
        chart = chart.dropna(subset=["display_coef"]).sort_values("display_coef")
        if not chart.empty:
            plt.figure(figsize=(11, max(4, 0.45 * len(chart))))
            labels = chart["link"] + " | " + chart["model_family"] + " | " + chart["data_variant"]
            colors = chart["model_reliability"].map({"reliable": "#2ca25f", "conditionally_usable": "#feb24c"}).fillna("#969696")
            plt.barh(labels, chart["display_coef"], color=colors)
            plt.axvline(0, color="black", linewidth=0.8)
            plt.xlabel("Displayed coefficient")
            plt.title("Aggregate dairy-index model coefficients")
            save_plot(FIG_CH6_DIR / "14_aggregate_index_model_coefficients.png", [FIG_DIR / "29_aggregate_index_model_coefficients.png"])


def write_excel_outputs(
    panel: pd.DataFrame,
    weekly: pd.DataFrame,
    index_daily: pd.DataFrame,
    index_weekly: pd.DataFrame,
    index_weights: pd.DataFrame,
    inventory: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
    product_audit: pd.DataFrame,
    product_dictionary: pd.DataFrame,
    coverage: pd.DataFrame,
    intersections: pd.DataFrame,
    corr_scan: pd.DataFrame,
    series_tests: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    index_series_tests: pd.DataFrame,
    index_models: pd.DataFrame,
    index_vecm_summary: pd.DataFrame,
    link21_summary: pd.DataFrame,
    reliability_overview: pd.DataFrame,
    lp: pd.DataFrame,
    lp_summary: pd.DataFrame,
    margins: pd.DataFrame,
    margin_spreads: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
    findings: pd.DataFrame,
) -> None:
    out_xlsx = OUTPUT_DIR / "final_research_outputs.xlsx"
    readme = pd.DataFrame(
        [
            {
                "item": "analysis_scope",
                "value": "Full dairy price-transmission rerun on corrected governmental data with weekly medians, strict/adapted intersections, core ARDL/ECM/NARDL/VECM models, and retail discount mechanisms.",
            },
            {
                "item": "input_workbook",
                "value": str(FULL_UAH),
            },
            {
                "item": "added_dataset_workbook",
                "value": str(ADDED_DATASET_PATH),
            },
            {
                "item": "main_change",
                "value": "Uses corrected true monthly governmental layers from full_uah_final.xlsx and rebuilds the downstream retail stage from retail combined, Silpo, and Novus with explicit Silpo discount logic.",
            },
            {
                "item": "created",
                "value": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            },
        ]
    )
    with pd.ExcelWriter(out_xlsx, engine="openpyxl") as writer:
        readme.to_excel(writer, sheet_name="README", index=False)
        inventory.to_excel(writer, sheet_name="Data_Inventory", index=False)
        product_audit.to_excel(writer, sheet_name="Product_Audit_Long", index=False)
        product_dictionary.to_excel(writer, sheet_name="Product_Dictionary", index=False)
        weekly.to_excel(writer, sheet_name="Weekly_Panel", index=False)
        if not index_daily.empty:
            index_daily.to_excel(writer, sheet_name="Aggregate_Index_Daily", index=False)
        if not index_weekly.empty:
            index_weekly.to_excel(writer, sheet_name="Aggregate_Index_Weekly", index=False)
        if not index_weights.empty:
            index_weights.to_excel(writer, sheet_name="Aggregate_Index_Weights", index=False)
        coverage.to_excel(writer, sheet_name="Panel_Coverage", index=False)
        intersections.to_excel(writer, sheet_name="Intersections", index=False)
        if not link21_summary.empty:
            link21_summary.to_excel(writer, sheet_name="Link21_Summary", index=False)
        corr_scan.to_excel(writer, sheet_name="Lag_Correlations", index=False)
        series_tests.to_excel(writer, sheet_name="Series_Tests", index=False)
        chain_models.to_excel(writer, sheet_name="Core_Chain_Models", index=False)
        vecm_results.to_excel(writer, sheet_name="VECM_Systems", index=False)
        if not index_series_tests.empty:
            index_series_tests.to_excel(writer, sheet_name="Index_Series_Tests", index=False)
        if not index_models.empty:
            index_models.to_excel(writer, sheet_name="Index_Chain_Models", index=False)
        if not index_vecm_summary.empty:
            index_vecm_summary.to_excel(writer, sheet_name="Index_VECM", index=False)
        if not reliability_overview.empty:
            reliability_overview.to_excel(writer, sheet_name="Reliability_Overview", index=False)
        scale_models.to_excel(writer, sheet_name="Proc_Scale_Models", index=False)
        if "retail_catalog" in source_frames:
            source_frames["retail_catalog"].to_excel(writer, sheet_name="Retail_Item_Catalog", index=False)
        if "retail_match_audit" in source_frames:
            source_frames["retail_match_audit"].to_excel(writer, sheet_name="Retail_Match_Audit", index=False)
        if "retail_brand_support" in source_frames:
            source_frames["retail_brand_support"].to_excel(writer, sheet_name="Retail_Brand_Support", index=False)
        if "retail_literal_summary" in source_frames:
            source_frames["retail_literal_summary"].to_excel(writer, sheet_name="Retail_Literal_Summary", index=False)
        if "retail_name_reconciliation" in source_frames:
            source_frames["retail_name_reconciliation"].to_excel(writer, sheet_name="Retail_Name_Recon", index=False)
        if "retail_level_scores" in source_frames:
            source_frames["retail_level_scores"].to_excel(writer, sheet_name="Retail_Level_Scores", index=False)
        if "retail_level_selection" in source_frames:
            source_frames["retail_level_selection"].to_excel(writer, sheet_name="Retail_Level_Choice", index=False)
        if "consumer" in source_frames:
            source_frames["consumer"].to_excel(writer, sheet_name="ConsumerUA_Clean", index=False)
        if "europe_benchmark" in source_frames:
            source_frames["europe_benchmark"].to_excel(writer, sheet_name="Europe_Benchmark", index=False)
        if "cme_benchmark" in source_frames:
            source_frames["cme_benchmark"].to_excel(writer, sheet_name="CME_Benchmark", index=False)
        lp.to_excel(writer, sheet_name="Local_Projections", index=False)
        lp_summary.to_excel(writer, sheet_name="LP_Summary", index=False)
        margins.to_excel(writer, sheet_name="Margin_Models", index=False)
        margin_spreads.to_excel(writer, sheet_name="Spread_Summary", index=False)
        discounts.to_excel(writer, sheet_name="Discount_Models", index=False)
        findings.to_excel(writer, sheet_name="Robust_Findings", index=False)
    panel.to_csv(DATA_DIR / "final_daily_panel.csv", index=False)
    weekly.to_csv(DATA_DIR / "final_weekly_panel.csv", index=False)
    if not index_daily.empty:
        index_daily.to_csv(DATA_DIR / "aggregate_chain_index_daily.csv", index=False)
    if not index_weekly.empty:
        index_weekly.to_csv(DATA_DIR / "aggregate_chain_index_weekly.csv", index=False)
    if not index_weights.empty:
        index_weights.to_csv(DATA_DIR / "aggregate_chain_index_weights.csv", index=False)
    coverage.to_csv(DATA_DIR / "final_panel_coverage.csv", index=False)
    product_audit.to_csv(DATA_DIR / "product_audit_long.csv", index=False)
    product_dictionary.to_csv(DATA_DIR / "product_dictionary_standardized.csv", index=False)
    intersections.to_csv(DATA_DIR / "intersection_admissibility.csv", index=False)
    corr_scan.to_csv(OUTPUT_DIR / "lag_correlation_scan.csv", index=False)
    series_tests.to_csv(OUTPUT_DIR / "series_stationarity_tests.csv", index=False)
    chain_models.to_csv(OUTPUT_DIR / "core_chain_models.csv", index=False)
    vecm_results.to_csv(OUTPUT_DIR / "vecm_systems.csv", index=False)
    if not index_series_tests.empty:
        index_series_tests.to_csv(OUTPUT_DIR / "aggregate_index_series_tests.csv", index=False)
    if not index_models.empty:
        index_models.to_csv(OUTPUT_DIR / "aggregate_index_models.csv", index=False)
    if not index_vecm_summary.empty:
        index_vecm_summary.to_csv(OUTPUT_DIR / "aggregate_index_vecm.csv", index=False)
    if not link21_summary.empty:
        link21_summary.to_csv(OUTPUT_DIR / "link21_summary.csv", index=False)
    if not reliability_overview.empty:
        reliability_overview.to_csv(OUTPUT_DIR / "model_reliability_overview.csv", index=False)
    scale_models.to_csv(OUTPUT_DIR / "procurement_scale_models.csv", index=False)
    optimal_cols = [c for c in ["date", "product", "product_label", "retail_optimal_observed", "retail_optimal_baseline", "retail_optimal_observed_model", "retail_optimal_baseline_model", "retail_optimal_discount_share", "retail_optimal_n_item_keys"] if c in panel.columns]
    if optimal_cols:
        panel[optimal_cols].to_csv(DATA_DIR / "retail_optimal_daily.csv", index=False)
    if "retail_items_full" in source_frames:
        source_frames["retail_items_full"].to_csv(DATA_DIR / "retail_items_full_harmonized.csv", index=False)
    if "retail_catalog" in source_frames:
        source_frames["retail_catalog"].to_csv(DATA_DIR / "retail_item_catalog.csv", index=False)
    if "retail_match_audit" in source_frames:
        source_frames["retail_match_audit"].to_csv(DATA_DIR / "retail_match_audit.csv", index=False)
    if "retail_brand_daily" in source_frames:
        source_frames["retail_brand_daily"].to_csv(DATA_DIR / "retail_brand_daily.csv", index=False)
    if "retail_brand_support" in source_frames:
        source_frames["retail_brand_support"].to_csv(DATA_DIR / "retail_brand_support.csv", index=False)
    if "retail_literal_summary" in source_frames:
        source_frames["retail_literal_summary"].to_csv(DATA_DIR / "retail_literal_summary.csv", index=False)
    if "retail_name_reconciliation" in source_frames:
        source_frames["retail_name_reconciliation"].to_csv(DATA_DIR / "retail_name_reconciliation_examples.csv", index=False)
    if "retail_level_scores" in source_frames:
        source_frames["retail_level_scores"].to_csv(DATA_DIR / "retail_level_scores.csv", index=False)
    if "retail_level_selection" in source_frames:
        source_frames["retail_level_selection"].to_csv(DATA_DIR / "retail_level_selection.csv", index=False)
    if "consumer" in source_frames:
        source_frames["consumer"].to_csv(DATA_DIR / "consumerua_clean_daily.csv", index=False)
    if "europe_benchmark" in source_frames:
        source_frames["europe_benchmark"].to_csv(DATA_DIR / "europe_benchmark_daily.csv", index=False)
    if "cme_benchmark" in source_frames:
        source_frames["cme_benchmark"].to_csv(DATA_DIR / "cme_class3_daily.csv", index=False)
    lp.to_csv(OUTPUT_DIR / "local_projection_coefficients.csv", index=False)
    lp_summary.to_csv(OUTPUT_DIR / "local_projection_summary.csv", index=False)
    margins.to_csv(OUTPUT_DIR / "margin_market_power_models.csv", index=False)
    margin_spreads.to_csv(OUTPUT_DIR / "vertical_spread_summary.csv", index=False)
    discounts.to_csv(OUTPUT_DIR / "discount_strategy_models.csv", index=False)
    findings.to_csv(OUTPUT_DIR / "robust_findings.csv", index=False)

    metadata = pd.DataFrame(
        [
            {
                "sheet_name": "daily_panel_used",
                "what_added": "Corrected final daily modelling panel with retail combined, Silpo, Novus, and explicit procurement-scale fields.",
                "why_added": "Used in final identification and in retail discount / spread / local-projection models.",
                "dataset_role": "strict_core_plus_adapted_endpoint",
                "used_in_final_identification": "yes",
            },
            {
                "sheet_name": "weekly_panel_used",
                "what_added": "Weekly median modelling panel with raw and smoothed variants.",
                "why_added": "Used in long-run modelling, lag scans, and core ARDL/ECM/NARDL/VECM estimation.",
                "dataset_role": "strict_core_and_robustness",
                "used_in_final_identification": "yes",
            },
            {
                "sheet_name": "aggregate_index_weekly",
                "what_added": "Weighted aggregate dairy-chain index panel built from product-level series with structural proxy weights.",
                "why_added": "Used as a robustness system layer and for aggregate chain ARDL/ECM/NARDL/VECM checks.",
                "dataset_role": "robustness_extension",
                "used_in_final_identification": "supporting_evidence",
            },
            {
                "sheet_name": "retail_items_used",
                "what_added": "Retail item-level harmonised Silpo-Novus dataset with discount-aware effective prices.",
                "why_added": "Used to construct retail combined, retailer-specific, and matched downstream endpoints.",
                "dataset_role": "strict_core_plus_extension_inputs",
                "used_in_final_identification": "yes",
            },
            {
                "sheet_name": "intersections_used",
                "what_added": "Formal overlap and admissibility table across candidate links.",
                "why_added": "Used to separate strong core evidence from extension evidence.",
                "dataset_role": "admissibility_control",
                "used_in_final_identification": "yes",
            },
            {
                "sheet_name": "product_dictionary_used",
                "what_added": "Standardized product-definition audit and cross-dataset mapping table.",
                "why_added": "Used to verify comparability before modelling and to flag ambiguous mappings.",
                "dataset_role": "pre_model_audit",
                "used_in_final_identification": "yes",
            },
        ]
    )
    with pd.ExcelWriter(ADDED_DATASET_PATH, engine="openpyxl") as writer:
        metadata.to_excel(writer, sheet_name="metadata", index=False)
        panel.to_excel(writer, sheet_name="daily_panel_used", index=False)
        weekly.to_excel(writer, sheet_name="weekly_panel_used", index=False)
        if not index_weekly.empty:
            index_weekly.to_excel(writer, sheet_name="aggregate_index_weekly", index=False)
        if not index_weights.empty:
            index_weights.to_excel(writer, sheet_name="aggregate_index_weights", index=False)
        intersections.to_excel(writer, sheet_name="intersections_used", index=False)
        product_dictionary.to_excel(writer, sheet_name="product_dictionary_used", index=False)
        if "retail_items_full" in source_frames:
            source_frames["retail_items_full"].to_excel(writer, sheet_name="retail_items_used", index=False)


def slugify(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "_", str(text)).strip("_").lower()
    return text[:120] if text else "model"


def write_single_model_outputs(chain_models: pd.DataFrame, intersections: pd.DataFrame, scale_models: pd.DataFrame, discounts: pd.DataFrame, index_models: pd.DataFrame) -> None:
    if not intersections.empty:
        for row in intersections.itertuples():
            base_name = slugify(f"{row.product}_{row.link}")
            pd.DataFrame([row._asdict()]).to_csv(MODEL_OUTPUT_DIR / f"{base_name}_screening.csv", index=False)
            note_lines = [
                f"# {row.product_label} | {row.link}",
                "",
                f"- Admissibility status: {row.admissibility_status}",
                f"- Mapping type: {row.mapping_type}",
                f"- Overlap weeks: {row.overlap_weeks}",
                f"- Continuity share: {fmt_num(row.continuity_share, 3)}",
                f"- Raw-support shares: x={fmt_num(row.x_raw_support_share, 3)}, y={fmt_num(row.y_raw_support_share, 3)}",
            ]
            (MODEL_NOTE_DIR / f"{base_name}_screening.md").write_text("\n".join(note_lines), encoding="utf-8")
    if not chain_models.empty:
        for key, g in chain_models.groupby(["product", "link", "model_family"], dropna=False):
            product, link, family = key
            base_name = slugify(f"{product}_{link}_{family}")
            result_cols = ["product", "product_label", "link", "model_family", "data_variant", "n_obs", "sr_coef", "lr_coef", "ect_coef", "cointegration_p", "lag_order", "asymmetry_pvalue", "model_reliability"]
            diag_cols = ["product", "product_label", "link", "model_family", "data_variant", "ljungbox_p", "bp_p", "white_p", "jb_p", "stability_flag", "notes"]
            g[result_cols].to_csv(MODEL_OUTPUT_DIR / f"{base_name}_results.csv", index=False)
            g[diag_cols].to_csv(MODEL_DIAG_DIR / f"{base_name}_diagnostics.csv", index=False)
            best = g.sort_values(["model_reliability", "data_variant"], ascending=[True, True]).iloc[0]
            note_lines = [
                f"# {best['product_label']} | {best['link']} | {family}",
                "",
                f"- Variants estimated: {len(g)}",
                f"- Best retained reliability: {best['model_reliability']}",
                f"- Main displayed coefficient: {fmt_num(best['lr_coef'] if pd.notna(best['lr_coef']) else best['sr_coef'])}",
                f"- Error-correction coefficient: {fmt_num(best['ect_coef'])}",
                f"- Cointegration p-value: {fmt_num(best['cointegration_p'])}",
                f"- Diagnostic note: Ljung-Box p={fmt_num(best['ljungbox_p'])}, Breusch-Pagan p={fmt_num(best['bp_p'])}, White p={fmt_num(best['white_p'])}, Jarque-Bera p={fmt_num(best['jb_p'])}.",
                f"- Model note: {best['notes']}",
            ]
            (MODEL_NOTE_DIR / f"{base_name}_notes.md").write_text("\n".join(note_lines), encoding="utf-8")
    if not scale_models.empty:
        scale_models.to_csv(MODEL_OUTPUT_DIR / "procurement_scale_model_results.csv", index=False)
    if not discounts.empty:
        discounts.to_csv(MODEL_OUTPUT_DIR / "silpo_discount_model_results.csv", index=False)
    if not index_models.empty:
        for key, g in index_models.groupby(["link", "model_family"], dropna=False):
            link, family = key
            base_name = slugify(f"aggregate_{link}_{family}")
            g.to_csv(MODEL_OUTPUT_DIR / f"{base_name}_results.csv", index=False)


def write_vecm_detail_outputs(summary: pd.DataFrame, detail_tables: Dict[str, pd.DataFrame]) -> None:
    if not summary.empty:
        summary.to_csv(VECM_DETAIL_DIR / "vecm_detailed_summary.csv", index=False)
    for name, df in detail_tables.items():
        if df is None or df.empty:
            continue
        df.to_csv(VECM_DETAIL_DIR / f"{slugify(name)}.csv", index=False)


def fmt_pct(value: float, digits: int = 1) -> str:
    if pd.isna(value):
        return "n/a"
    return f"{100 * float(value):.{digits}f}%"


def fmt_num(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return "n/a"
    return f"{float(value):.{digits}f}"


def top_lp_sentence(lp_summary: pd.DataFrame) -> str:
    if lp_summary.empty:
        return "No local-projection models passed the minimum-overlap screen."
    h = lp_summary[lp_summary["horizon_days"].isin([7, 14])].copy()
    if h.empty:
        h = lp_summary.copy()
    # The farm-gate/producer diagnostic is long and smooth because both sides are reconstructed.
    # Keep it in the workbook, but do not make it the headline thesis signal.
    thesis_focus = h[~h["link"].isin(["FarmGateUA -> ProducerUA", "ProducerUA -> FarmGateUA"])].copy()
    if not thesis_focus.empty:
        h = thesis_focus
    best = h.sort_values(["core_share", "sig_share", "models"], ascending=False).head(5)
    parts = [
        f"{r.link} ({r.price_variant}, h={int(r.horizon_days)}): core share {fmt_pct(r.core_share)}, median coefficient {fmt_num(r.median_coef)}"
        for r in best.itertuples()
    ]
    return "; ".join(parts) + "."


def top_brand_sentence(brand_support: pd.DataFrame) -> str:
    if brand_support.empty:
        return "No brand-level retail support table was generated."
    brand_support = brand_support.copy()
    brand_support["brand_norm"] = brand_support["brand_norm"].fillna("").astype(str).str.strip()
    brand_support["brand_named_flag"] = brand_support["brand_norm"].ne("").astype(int)
    top = (
        brand_support.sort_values(["product", "retailer", "brand_named_flag", "n_rows"], ascending=[True, True, False, False])
        .groupby(["product", "retailer"], as_index=False)
        .head(1)
        .head(6)
    )
    parts = [
        f"{PRODUCT_LABELS.get(r.product, r.product)} in {r.retailer}: {r.brand_norm or 'unlabeled brand'} ({int(r.n_item_keys)} item keys, {int(r.n_dates)} days)"
        for r in top.itertuples()
    ]
    return "; ".join(parts) + "."


def top_literal_sentence(literal_summary: pd.DataFrame) -> str:
    if literal_summary.empty:
        return "No literal retail-type summary was generated."
    top = (
        literal_summary.groupby(["product", "product_literal"], as_index=False)["n_item_keys"]
        .sum()
        .sort_values("n_item_keys", ascending=False)
        .head(6)
    )
    parts = [
        f"{PRODUCT_LABELS.get(r.product, r.product)} -> {LITERAL_LABELS.get(r.product_literal, r.product_literal)} ({int(r.n_item_keys)} item keys)"
        for r in top.itertuples()
    ]
    return "; ".join(parts) + "."


def level_selection_sentence(level_selection: pd.DataFrame) -> str:
    if level_selection.empty:
        return "No retail-level candidate was selected."
    parts = [
        f"{r.product_label}: {r.candidate_label} (score {fmt_num(r.selection_score, 2)})"
        for r in level_selection.itertuples()
    ]
    return "; ".join(parts) + "."


def candidate_core_sentence(lp_summary: pd.DataFrame) -> str:
    if lp_summary.empty:
        return "Candidate downstream LP comparisons are not available."
    focus = lp_summary[
        lp_summary["link"].isin(
            [
                "ProZorro -> Retail",
                "ProZorro -> Retail matched",
                "ProZorro -> Silpo",
                "ProZorro -> Novus",
                "ProZorro -> Retail optimal",
            ]
        )
        & lp_summary["horizon_days"].isin([7, 14])
    ].copy()
    if focus.empty:
        return "Candidate downstream LP comparisons are not available."
    top = focus.groupby("link", as_index=False)["core_share"].mean().sort_values("core_share", ascending=False).head(5)
    parts = [f"{r.link}: mean 7/14-day core share {fmt_pct(r.core_share)}" for r in top.itertuples()]
    return "; ".join(parts) + "."


def write_summary_documents(
    inventory: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
    coverage: pd.DataFrame,
    lp: pd.DataFrame,
    lp_summary: pd.DataFrame,
    margins: pd.DataFrame,
    margin_spreads: pd.DataFrame,
    discounts: pd.DataFrame,
    findings: pd.DataFrame,
) -> None:
    lp_linear = lp[lp["model"].eq("LP_linear")] if not lp.empty else pd.DataFrame()
    total_models = len(lp_linear)
    core_models = int(lp_linear["core_signal"].sum()) if not lp_linear.empty else 0
    core_share = core_models / total_models if total_models else np.nan
    n_products = coverage["product"].nunique() if not coverage.empty else 0
    margin_flags = int(margins.get("persistent_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    asym_flags = int(margins.get("asymmetric_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    discount_signals = int(discounts.get("discount_strategy_signal", pd.Series(dtype=float)).fillna(0).sum()) if not discounts.empty else 0
    match_audit = source_frames.get("retail_match_audit", pd.DataFrame())
    matched_keys = int(match_audit["match_status"].eq("matched_both_shops").sum()) if not match_audit.empty else 0
    silpo_only = int(match_audit["match_status"].eq("silpo_only").sum()) if not match_audit.empty else 0
    novus_only = int(match_audit["match_status"].eq("novus_only").sum()) if not match_audit.empty else 0

    lines = [
        "# Second-Stage Estimation Summary",
        "",
        "## Analytical Purpose",
        "This second-stage run keeps the thesis chain as FarmGateUA -> ProducerUA -> ProZorro -> Retail, but estimates it with a different empirical design from the earlier RW4 ARDL/ECM/NARDL/VECM workflow. The aim is to check whether the same economic story survives when the model is rebuilt as a daily product-panel local-projection exercise and a vertical spread/discount market-power proxy exercise.",
        "",
        "The model uses the active `full_uah.xlsx`, `farm_gate_daily.xlsx`, and `farm_gate_all_missing_filled_daily.xlsx` workbooks from the nested Charniuk dairy research project. Farm-gate variants remain explicit (`initial` and `filled`, each with `linear` and `pchip` reconstruction), while retail is separated into observed and baseline price paths to keep discount behavior visible.",
        "",
        "## Data Coverage",
        f"The cleaned second-stage panel covers {n_products} standardized product groups. The strongest retail coverage remains in Silpo, while Novus is useful mainly as a thinner cross-retailer robustness layer. ProZorro is treated as an institutional procurement layer and is locally interpolated only across short gaps to avoid inventing long contract windows.",
        "",
        f"The retail-preparation block now works at item level before product-level aggregation. Silpo and Novus rows are harmonized with normalized brand names, canonical item names, standardized dairy product types, fat percentage, and pack-size fields. This produced {matched_keys} unique harmonized item keys observed in both shops, against {silpo_only} Silpo-only keys and {novus_only} Novus-only keys. Effective price keeps the discount inside the price, while discount amount, discount-type dummies, and markdown depth remain as separate covariates.",
        "",
        "![Panel coverage](../figures/01_panel_coverage.png)",
        "",
        "![Cross-shop match status](../figures/06_cross_shop_match_status.png)",
        "",
        "## Model Design",
        "The first model block uses local projections. For each admissible product and chain link, the cumulative change in the downstream log price over horizons 0, 1, 3, 7, 14, and 21 days is regressed on the current upstream log shock with HAC inference. A parallel asymmetric version separates positive and negative upstream shocks. This is intentionally different from the thesis draft's ECM/NARDL core and therefore functions as a robustness redesign rather than a simple rerun.",
        "",
        "The second model block estimates vertical spread equations. The dependent variable is the change in the log spread between downstream and upstream stages; regressors include the lagged spread, positive and negative upstream shocks, and retail discount measures where they vary. This is a market-power proxy, not a direct structural proof of conduct: persistent or asymmetric spreads and discount-responsive retail margins are interpreted as evidence consistent with timing control and category management.",
        "",
        "A separate consumer benchmark is kept alongside the harmonized merged-shop retail layer so the rerun can test whether conclusions remain coherent when retailer-facing prices are compared against the official consumer environment without collapsing both objects into one synthetic endpoint.",
        "",
        "## Main Local-Projection Results",
        f"The local-projection block estimated {total_models} linear pass-through equations, of which {core_models} met the p<0.10 and overlap screen ({fmt_pct(core_share)}). The strongest thesis-relevant 7/14-day signals, excluding the mechanically smooth reconstructed FarmGateUA-ProducerUA diagnostic, are: {top_lp_sentence(lp_summary)}",
        "",
        "The FarmGateUA -> ProducerUA and ProducerUA -> FarmGateUA local-projection rows are retained in the workbook, but they should not be promoted as headline causal evidence. Both sides are reconstructed over a long common daily horizon, so high fit can reflect shared smoothing and benchmark coherence rather than a clean product-level farm-to-processor response. This is consistent with the corrected-format draft's conservative farm-gate interpretation.",
        "",
        "![Local-projection pass-through](../figures/02_lp_pass_through_horizons.png)",
        "",
        "![Forward versus reverse evidence](../figures/03_forward_reverse_core_share.png)",
        "",
        "## Market-Power And Discount Results",
        f"The spread model produced {len(margins)} usable margin equations. Persistent-margin flags appear in {margin_flags} rows and asymmetric-margin flags appear in {asym_flags} rows. These flags should be read as market-power proxies: they indicate where spreads do not quickly mean-revert or where upstream increases and decreases affect margins differently.",
        "",
        f"The discount model produced {len(discounts)} usable product equations, with {discount_signals} discount-strategy signals. In this data structure, the discount result is strongest as a retail-pricing mechanism rather than as a farm-gate mechanism: discount variables are directly observed only for Silpo, while ProducerUA, ProZorro, and FarmGateUA do not contain comparable markdown variables.",
        "",
        "![Vertical spread proxy](../figures/04_vertical_spread_proxy.png)",
        "",
        "![Discount incidence](../figures/05_discount_incidence.png)",
        "",
        "## Thesis Interpretation",
        "The second-stage run supports a cautious but coherent interpretation. Vertical coordination remains visible, but it is not a single cost-plus coefficient. The local-projection redesign shows that short-horizon evidence is selective and product-dependent, while the spread regressions show where timing control, discount smoothing, and asymmetric margin behavior are plausible. This strengthens the thesis argument that downstream market power is better described as control over timing, visibility, and discount-mediated adjustment than as one static markup wedge.",
        "",
        "Farm-gate results should still be written conservatively. The farm-gate benchmark is economically necessary as the raw-milk origin of the chain, but it is much smoother and less product-specific than the retail and procurement data. The second-stage results therefore use FarmGateUA as an upstream benchmark and robustness dimension, not as a literal product-level farm-to-shelf elasticity for every dairy category.",
        "",
        "## Output Files",
        "- `outputs/second_stage_model_outputs.xlsx` contains all model tables.",
        "- `data/second_stage_daily_panel.csv` contains the cleaned daily panel used by the models.",
        "- `data/retail_items_full_harmonized.csv` contains the full harmonized Silpo-Novus item list.",
        "- `data/retail_match_audit.csv` contains the cross-shop item-match audit.",
        "- `data/consumerua_clean_daily.csv` contains the cleaned ConsumerUA daily layer used in the consumer-linked rerun.",
        "- `outputs/local_projection_coefficients.csv`, `outputs/margin_market_power_models.csv`, and `outputs/discount_strategy_models.csv` contain the main machine-readable estimates.",
    ]
    (DOC_DIR / "second_stage_estimation_summary.md").write_text("\n".join(lines), encoding="utf-8")
    write_html(DOC_DIR / "second_stage_estimation_summary.html", "Second-Stage Estimation Summary", lines)
    write_simple_docx(DOC_DIR / "second_stage_estimation_summary.docx", lines, force_bold=False)

    correction_lines = [
        "# **Bold Corrections And Additions For The Corrected-Format Draft**",
        "",
        "**Target file reviewed: `Charniuk_Maksym_MScThesis_Draft_correctedformat.docx`. These notes are intentionally separate from the main draft, as requested.**",
        "",
        "- **Correct the farm-gate share statement: the 4.8% core-finding share belongs to FarmGateUA -> ProZorro, not FarmGateUA -> ProducerUA. FarmGateUA -> ProducerUA has 8 core findings out of 709 rows, about 1.1%, so it should be described as the weakest direct farm-gate link.**",
        "- **Use `institutional procurement layer` for ProZorro instead of saying it can simply be interpreted as processors. ProducerUA is the processor-level domestic producer-price layer; ProZorro is a procurement-contract layer.**",
        "- **Keep EU, CME, and ConsumerUA outside the endogenous domestic chain. They should be described as external benchmarks, reconstruction anchors, diagnostics, or robustness checks, not as extra structural stages.**",
        "- **State explicitly that the farm-gate raw-milk benchmark is not a processed-product price. It is an upstream raw-milk anchor, so direct FarmGateUA-to-retail coefficients must be interpreted conservatively.**",
        "- **Write reverse-flow findings as price-leadership or coordination evidence, not as proof of literal reverse causality. Retail -> ProZorro and Retail -> FarmGateUA coefficients show that downstream price setting contains information reflected upstream, but they do not by themselves prove causal control.**",
        "- **For Silpo, keep the distinction between observed price and baseline price. Observed price is the effective shelf price after markdown; baseline price is the regular-price proxy. For Novus, do not infer a promotion baseline because the workbook has no comparable discount signal.**",
        "- **Correct the ConsumerUA caption in the methodology/data text where it says `Ukrainian producer prices by products`; it should say `Ukrainian consumer prices by products`.**",
        "- **Correct `U.S. CME III stock prices distribution` to `U.S. CME Class III milk futures price distribution`.**",
        "- **Correct visible typos in the figure list and prose: `balalnce` -> `balance`, `Ukraine(gross` -> `Ukraine (gross`, `frequenct` -> `frequency`, `bee hind` / `bee hind` style fragments -> `behind`, `presenet` -> `present`, and `standartisation` -> `standardization` or `standardisation` consistently.**",
        "- **The Casey thesis/reference was not found by filename or searchable DOCX text in this folder. I therefore aligned the separate documents to the available KSE thesis style and the latest corrected-format draft. If the Casey thesis exists under another filename or only as a scanned PDF, add/identify it before final formatting imitation.**",
        "- **Do not headline the second-stage FarmGateUA -> ProducerUA local-projection coefficient even when it is statistically strong. Both sides are reconstructed/smoothed over a long horizon, so this block should be treated as reconstruction coherence rather than direct causal farm-to-processor evidence.**",
        "- **Add the second-stage robustness sentence after the RW4 estimation discussion: `As a robustness redesign, the second-stage local-projection and spread-proxy models keep the same four-stage chain but avoid relying on the same ARDL/ECM/NARDL/VECM identification stack; the resulting evidence remains selective and product-specific, which supports a conservative interpretation of market power as timing and discount-mediated coordination rather than a single universal markup.`**",
    ]
    (DOC_DIR / "corrected_format_additions_bold.md").write_text("\n".join(correction_lines), encoding="utf-8")
    write_html(DOC_DIR / "corrected_format_additions_bold.html", "Bold Corrections And Additions", correction_lines)
    write_simple_docx(DOC_DIR / "corrected_format_additions_bold.docx", correction_lines, force_bold=True)


def write_thesis_style_chapters(
    inventory: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
    coverage: pd.DataFrame,
    lp: pd.DataFrame,
    lp_summary: pd.DataFrame,
    margins: pd.DataFrame,
    margin_spreads: pd.DataFrame,
    discounts: pd.DataFrame,
    findings: pd.DataFrame,
) -> None:
    retail_items = source_frames.get("retail_items_full", pd.DataFrame())
    match_audit = source_frames.get("retail_match_audit", pd.DataFrame())
    literal_summary = source_frames.get("retail_literal_summary", pd.DataFrame())
    brand_support = source_frames.get("retail_brand_support", pd.DataFrame())
    level_selection = source_frames.get("retail_level_selection", pd.DataFrame())
    level_scores = source_frames.get("retail_level_scores", pd.DataFrame())
    reconciliation = source_frames.get("retail_name_reconciliation", pd.DataFrame())

    matched_keys = int(match_audit["match_status"].eq("matched_both_shops").sum()) if not match_audit.empty else 0
    silpo_only = int(match_audit["match_status"].eq("silpo_only").sum()) if not match_audit.empty else 0
    novus_only = int(match_audit["match_status"].eq("novus_only").sum()) if not match_audit.empty else 0
    strict_matches = int(match_audit.get("strict_alignment_flag", pd.Series(dtype=float)).fillna(0).sum()) if not match_audit.empty else 0
    n_literal = int(retail_items["product_literal"].nunique()) if not retail_items.empty else 0
    n_brands = int(retail_items["brand_norm"].nunique()) if not retail_items.empty else 0
    n_products = int(coverage["product"].nunique()) if not coverage.empty else 0
    total_models = int(lp[lp["model"].eq("LP_linear")].shape[0]) if not lp.empty else 0
    core_models = int(lp[(lp["model"].eq("LP_linear")) & (lp["core_signal"].eq(1))].shape[0]) if not lp.empty else 0
    margin_flags = int(margins.get("persistent_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    asym_flags = int(margins.get("asymmetric_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    discount_signals = int(discounts.get("discount_strategy_signal", pd.Series(dtype=float)).fillna(0).sum()) if not discounts.empty else 0

    data_lines = [
        "# Chapter 5. Second-stage data and retail harmonisation",
        "",
        "This chapter documents the second-stage data architecture used to complement the corrected master-thesis draft. The logic is intentionally close to the main thesis: each source is still treated as an institutional pricing layer, but the downstream retail block is rebuilt more deeply at item level so that Novus and Silpo contribute directly to the stage-4 model rather than only through broad category pooling.",
        "",
        "The practical goal of the redesign is not to overturn the master-thesis structure. It is to improve it where the original bottleneck was strongest: retailer product naming, brand reconciliation, literal dairy-product typing, and explicit discount-aware construction of the effective retail price.",
        "",
        "## 5.1 Data sources and analytical layers",
        "The second-stage chain still follows FarmGateUA -> ProducerUA -> ProZorro -> Retail. FarmGateUA remains the raw-milk benchmark, ProducerUA remains the processor-level domestic price layer, ProZorro remains the institutional procurement layer, and the retail block is rebuilt from harmonised Novus and Silpo product records. ConsumerUA is added as a downstream robustness environment, not as a replacement for the retailer-facing stage.",
        "",
        f"The cleaned second-stage panel covers {n_products} thesis product groups. The retail input after dairy-only reconciliation contains {len(retail_items)} product-day observations, {n_brands} normalised brand identifiers, and {n_literal} literal dairy-product types. This is materially deeper downstream preparation than the summary-level retail block in the earlier second-stage run.",
        "",
        "Figure 5.1. Second-stage panel coverage by product and source.",
        "![Figure 5.1. Second-stage panel coverage by product and source.](../figures/01_panel_coverage.png)",
        "Source: author's calculations based on the second-stage panel coverage table.",
        "",
        "## 5.1.1 Retail reconstruction from Novus and Silpo",
        "Retail preparation now begins at the item level. Each row keeps the effective observed shelf price, meaning the discount is already embedded in the price actually faced by the buyer. At the same time, the discount amount, discount state, discount type, and markdown depth are retained as separate variables so that retail adjustment can be studied both through the price itself and through the promotional mechanism behind it.",
        "",
        "Literal dairy typing is also made explicit. Instead of trusting the raw shop category alone, product titles, product names, brands, and auxiliary entity labels are cleaned jointly and mapped into literal dairy types such as milk, kefir, yogurt, hard cheese, cottage cheese, sour cream, cream, butter, condensed milk, and milk powder. Plant-based analogues and other non-comparable items are screened out before aggregation, which keeps the downstream series closer to the real dairy chain addressed in the thesis.",
        "",
        f"The reconciled literal mix is selective rather than generic: {top_literal_sentence(literal_summary)}",
        "",
        "Figure 5.2. Retail literal-product mix after dairy-only reconciliation.",
        "![Figure 5.2. Retail literal-product mix after dairy-only reconciliation.](../figures/07_retail_literal_mix.png)",
        "Source: author's calculations based on the harmonised retail literal summary.",
        "",
        "## 5.1.2 Cross-shop name reconciliation and brand harmonisation",
        "The core reconciliation object is the harmonised item key. It combines the thesis product group, the cleaned brand, and a canonicalised item name stripped of redundant pack and percentage tokens. This step is necessary because Novus and Silpo often refer to the same item with slightly different wording, word order, or packaging notation.",
        "",
        f"The resulting audit identifies {matched_keys} matched cross-shop item keys, {silpo_only} Silpo-only keys, and {novus_only} Novus-only keys. Among the matched keys, {strict_matches} cases also align one-for-one on the stricter fat-and-pack diagnostic key. This is exactly the kind of retailer-name adaptation that was missing from the broad category view.",
        "",
        "Figure 5.3. Cross-shop retail item harmonisation status.",
        "![Figure 5.3. Cross-shop retail item harmonisation status.](../figures/06_cross_shop_match_status.png)",
        "Source: author's calculations based on the retail match audit.",
        "",
        f"Brand support remains economically meaningful after reconciliation: {top_brand_sentence(brand_support)}",
        "",
        "Figure 5.4. Dominant retailer-brand support by dairy product.",
        "![Figure 5.4. Dominant retailer-brand support by dairy product.](../figures/08_dominant_brand_support.png)",
        "Source: author's calculations based on the reconciled retail brand-support table.",
        "",
        "## 5.2 Downstream levels tested for the four-stage model",
        "The retail endpoint is not forced into one representation. Instead, the second-stage pipeline constructs and keeps several candidate downstream levels: merged full-list retail, matched cross-shop retail, Silpo-only retail, Novus-only retail, and a retail-plus-ConsumerUA endpoint. Each level is useful for a different reason. The merged series maximises assortment support, the matched series maximises cross-shop comparability, Silpo and Novus preserve retailer-specific pricing behaviour, and the ConsumerUA-linked variant extends the downstream environment when retailer support is thin.",
        "",
        "Candidate levels are compared product by product using a composite score that combines coverage, procurement alignment, consumer alignment, item-support depth, and discount variation. This complements the master-thesis draft because it keeps the same four-stage chain but tests whether stage 4 is better represented by a broader retail pool, a stricter matched subset, a retailer-specific panel, or a consumer-linked endpoint.",
        "",
        f"The selected level by product is as follows: {level_selection_sentence(level_selection)}",
        "",
        "Figure 5.5. Candidate downstream retail scores by product.",
        "![Figure 5.5. Candidate downstream retail scores by product.](../figures/09_retail_level_scores.png)",
        "Source: author's calculations based on the retail-level selection table.",
        "",
        "Figure 5.6. Chosen optimal retail level by product.",
        "![Figure 5.6. Chosen optimal retail level by product.](../figures/10_optimal_retail_level.png)",
        "Source: author's calculations based on the selected retail-level hierarchy.",
        "",
        "## 5.3 Discount-aware downstream environment",
        "Discount variables remain central because the user requested that the final effective price should include the markdown while the markdown mechanism itself remains modelled separately. This design matches the master-thesis interpretation more closely than a raw regular-price series would. It treats retail adjustment as a combination of baseline repricing and promotional smoothing, not as one undifferentiated observed price path.",
        "",
        "Figure 5.7. Retail discount environment by product and retailer.",
        "![Figure 5.7. Retail discount environment by product and retailer.](../figures/12_discount_environment.png)",
        "Source: author's calculations based on retailer-product discount states in the harmonised retail panel.",
        "",
        "## 5.4 What remains after second-stage preparation",
        "After reconciliation, the downstream data are no longer a loose category average. They are a hierarchy of comparable retail series supported by explicit item keys, literal product types, brand structure, and discount metadata. That makes the second-stage outputs much easier to defend as a robustness complement to the corrected thesis draft: the master thesis keeps the broader structural story, while this second-stage data chapter shows that the downstream block has been rebuilt in a more disciplined way.",
        "",
        "The most important dataset files produced by this chapter are `data/retail_items_full_harmonized.csv`, `data/retail_brand_daily.csv`, `data/retail_literal_summary.csv`, `data/retail_level_selection.csv`, and `data/retail_optimal_daily.csv`. Together they document exactly how the stage-4 retail block was built before the models were rerun.",
    ]

    estimation_lines = [
        "# Chapter 6. Second-stage estimation results",
        "",
        "The second-stage estimation chapter is designed to complement the corrected master-thesis draft rather than duplicate it. The draft already uses the richer RW4 ARDL/ECM/NARDL/VECM stack as the main structural evidence. This chapter therefore asks a narrower but important question: if the retail stage is rebuilt more carefully from Novus and Silpo item-level data, do the main economic conclusions survive under a different modelling design?",
        "",
        "The answer is broadly yes, but the evidence remains selective and product-dependent. That is precisely why the second-stage methods are useful: they stress-test the thesis argument without pretending to replace the draft's main identification logic.",
        "",
        "## 6.1 Model families and their role",
        "The first model family uses local projections. It estimates cumulative downstream responses over 0, 1, 3, 7, 14, and 21 days with HAC inference. This complements the main draft because it avoids forcing each link into one equilibrium-correction structure and instead asks where timing evidence is stable across horizons.",
        "",
        "The second model family uses vertical spread equations. These are not direct market-power proofs, but they are useful proxies for timing control, asymmetric margin adjustment, and discount-mediated smoothing. The third block keeps a focused discount model in which retail discount incidence is related to lagged discount states and upstream price variation.",
        "",
        f"In total, the second-stage local-projection block estimates {total_models} linear models, of which {core_models} pass the p<0.10 and overlap screen. The leading 7/14-day thesis-relevant signals are: {top_lp_sentence(lp_summary)}",
        "",
        "Figure 6.1. Local-projection pass-through by horizon.",
        "![Figure 6.1. Local-projection pass-through by horizon.](../figures/02_lp_pass_through_horizons.png)",
        "Source: author's calculations based on the second-stage local-projection summary.",
        "",
        "Figure 6.2. Forward versus reverse second-stage evidence.",
        "![Figure 6.2. Forward versus reverse second-stage evidence.](../figures/03_forward_reverse_core_share.png)",
        "Source: author's calculations based on the second-stage local-projection summary.",
        "",
        "## 6.2 Procurement to retail under alternative downstream levels",
        "A key improvement relative to the earlier second-stage run is that procurement-to-retail evidence is not evaluated only on one pooled retail line. The pipeline now tests merged retail, matched retail, retailer-specific series, the ConsumerUA-linked series, and the selected optimal series. This directly addresses the concern that stage-4 measurement can change the interpretation of market power.",
        "",
        f"The candidate comparison shows the following ranking of downstream LP evidence: {candidate_core_sentence(lp_summary)}",
        "",
        "Figure 6.3. Candidate downstream levels in procurement-to-retail local-projection tests.",
        "![Figure 6.3. Candidate downstream levels in procurement-to-retail local-projection tests.](../figures/11_candidate_downstream_core_share.png)",
        "Source: author's calculations based on 7/14-day LP core shares across retail candidates.",
        "",
        "This comparison is economically useful because it separates three ideas that are often conflated. First, more coverage is not always better if it comes from a weaker retailer-grounded series. Second, stricter matched-item support is cleaner but shorter. Third, the ConsumerUA-linked endpoint can help with continuity, but it should still be treated as a robustness extension rather than as the literal shelf-price stage. That logic mirrors the corrected master-thesis draft rather than competing with it.",
        "",
        "## 6.3 Vertical spreads, market power, and discount effects",
        f"The spread block produces {len(margins)} usable equations. Persistent-margin flags appear in {margin_flags} rows and asymmetric-margin flags appear in {asym_flags} rows. This does not prove structural market power on its own, but it is consistent with downstream timing control, selective margin adjustment, and retailer category management.",
        "",
        "Figure 6.4. Vertical spread and market-power proxy by chain segment.",
        "![Figure 6.4. Vertical spread and market-power proxy by chain segment.](../figures/04_vertical_spread_proxy.png)",
        "Source: author's calculations based on the vertical spread summary.",
        "",
        f"The discount block remains more cautious than the main RW4 draft: it yields {len(discounts)} usable equations and {discount_signals} formal discount-strategy signals. That weaker statistical result does not mean discounts are unimportant. Instead, it suggests that in this reduced-form second-stage design the discount mechanism is most visible as part of the retail data-generating process rather than as a stand-alone structural driver.",
        "",
        "Figure 6.5. Retail discount incidence by product.",
        "![Figure 6.5. Retail discount incidence by product.](../figures/05_discount_incidence.png)",
        "Source: author's calculations based on the second-stage discount model table.",
        "",
        "## 6.4 Interpretation relative to the corrected thesis draft",
        "The second-stage results support the same broad market interpretation as the master-thesis draft, but through a different route. Procurement still behaves like an institutional buffer rather than a frictionless conduit. Retail still behaves like a managed adjustment layer. The main economic story still points toward product-specific vertical coordination rather than toward one universal pass-through elasticity.",
        "",
        "What changes is the emphasis. The second-stage redesign gives more weight to the downstream data-generating process itself: how item names are reconciled, how brand structure survives aggregation, how discount states are retained inside and outside price, and how the choice of retail endpoint changes the strength of the stage-4 model. That is exactly why this chapter complements the corrected draft: it improves the credibility of the retail stage without forcing the whole thesis to depend on one extra modelling family.",
        "",
        "Farm-gate results should remain conservative here as well. The second-stage run still shows that direct FarmGateUA -> ProducerUA evidence can look mechanically strong because both sides are reconstructed and smooth. The more defensible interpretation remains the same as in the corrected draft: FarmGateUA is an upstream benchmark and robustness dimension, not a literal product-level farm-to-shelf elasticity generator.",
        "",
        "## 6.5 Conclusion and practical implication",
        "The strongest contribution of the second-stage chapter is methodological. It shows that once the retail block is rebuilt at item level, filtered to literal dairy products, reconciled across Novus and Silpo, and tested across multiple downstream levels, the core thesis still holds. The retail stage does not behave like a passive residual of upstream costs. It behaves like a strategic adjustment layer in which timing, discount use, brand structure, and category management matter for observed price transmission.",
        "",
        "This strengthens the corrected master-thesis draft rather than displacing it. The draft can continue to rely on the richer ECM/NARDL evidence as the main structural story, while the second-stage chapter demonstrates that the downstream interpretation survives deeper retail cleaning and a different empirical design.",
    ]

    combined_lines = data_lines + ["", ""] + estimation_lines

    outputs = [
        ("second_stage_data_chapter", "Second-stage data chapter", data_lines),
        ("second_stage_estimation_chapter", "Second-stage estimation chapter", estimation_lines),
        ("second_stage_data_estiamtion_updated_conclusion", "Second-stage data and estimation updated conclusion", combined_lines),
    ]
    for stem, title, lines in outputs:
        (DOC_DIR / f"{stem}.md").write_text("\n".join(lines), encoding="utf-8")
        write_html(DOC_DIR / f"{stem}.html", title, lines)
        write_simple_docx(DOC_DIR / f"{stem}.docx", lines, force_bold=False)


def write_html(path: Path, title: str, md_lines: Sequence[str]) -> None:
    body = []
    for line in md_lines:
        if line.startswith("# "):
            body.append(f"<h1>{markdown_inline(line[2:])}</h1>")
        elif line.startswith("## "):
            body.append(f"<h2>{markdown_inline(line[3:])}</h2>")
        elif line.startswith("### "):
            body.append(f"<h3>{markdown_inline(line[4:])}</h3>")
        elif line.startswith("- "):
            body.append(f"<p>&bull; {markdown_inline(line[2:])}</p>")
        elif re.match(r"^\d+\.\s", line):
            body.append(f"<p>{markdown_inline(line)}</p>")
        elif line.startswith("{{TABLE|") and line.endswith("}}"):
            payload = line[8:-2]
            parts = payload.split("|", 2)
            if len(parts) >= 2:
                caption = parts[0].strip()
                csv_ref = parts[1].strip()
                note = parts[2].strip() if len(parts) > 2 else ""
                csv_path = _resolve_doc_asset(path, csv_ref)
                if csv_path.exists():
                    df = pd.read_csv(csv_path)
                    head = "".join(f"<th>{html.escape(str(c))}</th>" for c in df.columns)
                    rows = []
                    for _, row in df.iterrows():
                        rows.append("<tr>" + "".join(f"<td>{html.escape(str(v))}</td>" for v in row.tolist()) + "</tr>")
                    body.append(f"<p><strong>{html.escape(caption)}</strong></p>")
                    body.append(f"<table border='1' cellspacing='0' cellpadding='4'><thead><tr>{head}</tr></thead><tbody>{''.join(rows)}</tbody></table>")
                    if note:
                        body.append(f"<p><em>{markdown_inline(note)}</em></p>")
        elif line.startswith("![") and "](" in line:
            alt = line[2 : line.find("]")]
            src = line[line.find("(") + 1 : line.rfind(")")]
            body.append(f'<figure><img src="{html.escape(src)}" alt="{html.escape(alt)}" style="max-width: 900px;"><figcaption>{html.escape(alt)}</figcaption></figure>')
        elif line.startswith("{{FORMULA|") and line.endswith("}}"):
            formula = line[10:-2].strip()
            body.append(
                "<div style=\"border:1px solid #666;padding:10px 14px;margin:10px auto;max-width:760px;text-align:center;font-weight:700;background:#fafafa;\">"
                + html.escape(formula)
                + "</div>"
            )
        elif not line.strip():
            body.append("")
        else:
            body.append(f"<p>{markdown_inline(line)}</p>")
    html_doc = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Georgia, 'Times New Roman', serif; max-width: 980px; margin: 32px auto; line-height: 1.45; }}
    h1, h2 {{ font-family: Arial, sans-serif; }}
    img {{ border: 1px solid #ddd; padding: 4px; }}
    code {{ background: #f6f6f6; padding: 1px 3px; }}
  </style>
</head>
<body>
{chr(10).join(body)}
</body>
</html>
"""
    path.write_text(html_doc, encoding="utf-8")


def markdown_inline(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", escaped)
    return escaped


def clean_md_text(line: str) -> Tuple[str, str]:
    stripped = line.strip()
    style = "normal"
    if stripped.startswith("# "):
        style = "title"
        stripped = stripped[2:].strip()
    elif stripped.startswith("## "):
        style = "heading"
        stripped = stripped[3:].strip()
    elif stripped.startswith("- "):
        style = "bullet"
        stripped = stripped[2:].strip()
    elif stripped.startswith("![") and "](" in stripped:
        alt = stripped[2 : stripped.find("]")]
        src = stripped[stripped.find("(") + 1 : stripped.rfind(")")]
        stripped = f"Figure: {alt} ({src})"
        style = "figure"
    stripped = stripped.replace("**", "").replace("`", "")
    return stripped, style


def docx_run(text: str, bold: bool = False, size: Optional[int] = None) -> str:
    props = []
    if bold:
        props.append("<w:b/>")
    if size:
        props.append(f'<w:sz w:val="{size}"/>')
    rpr = f"<w:rPr>{''.join(props)}</w:rPr>" if props else ""
    return f'<w:r>{rpr}<w:t xml:space="preserve">{html.escape(text)}</w:t></w:r>'


def docx_paragraph(text: str, style: str = "normal", force_bold: bool = False) -> str:
    if not text.strip():
        return "<w:p/>"
    bold = force_bold or style in {"title", "heading"}
    size = 32 if style == "title" else 26 if style == "heading" else None
    prefix = "• " if style == "bullet" else ""
    return f"<w:p>{docx_run(prefix + text, bold=bold, size=size)}</w:p>"


def write_simple_docx(path: Path, md_lines: Sequence[str], force_bold: bool = False) -> None:
    paragraphs = []
    for line in md_lines:
        text, style = clean_md_text(line)
        paragraphs.append(docx_paragraph(text, style=style, force_bold=force_bold))
    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {''.join(paragraphs)}
    <w:sectPr>
      <w:pgSz w:w="12240" w:h="15840"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>
'''
    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
'''
    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
'''
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def write_integrated_summary_workbook(
    intersections: pd.DataFrame,
    corr_scan: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    index_models: pd.DataFrame,
    index_vecm_summary: pd.DataFrame,
    link21_summary: pd.DataFrame,
    reliability_overview: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
    findings: pd.DataFrame,
) -> None:
    out_path = OUTPUT_DIR / "integrated_summary_workbook.xlsx"
    summary_rows: List[Dict[str, object]] = []
    if not intersections.empty:
        summary_rows.append({"block": "Intersections", "metric": "strong_links", "value": int(intersections["admissibility_status"].eq("strong").sum())})
        summary_rows.append({"block": "Intersections", "metric": "acceptable_links", "value": int(intersections["admissibility_status"].eq("acceptable").sum())})
    if not corr_scan.empty:
        top_corr = corr_scan.sort_values("best_abs_corr", ascending=False).head(1)
        if not top_corr.empty:
            r = top_corr.iloc[0]
            summary_rows.append({"block": "Lag scan", "metric": "top_signal", "value": f"{r['product_label']} | {r['link']} | corr={fmt_num(r['best_abs_corr'])} | lag={int(r['best_lag_weeks'])}"})
    if not chain_models.empty:
        for fam in ["ARDL", "ECM", "NARDL"]:
            sub = chain_models[chain_models["model_family"].eq(fam)]
            if not sub.empty:
                summary_rows.append({"block": fam, "metric": "rows", "value": int(len(sub))})
                summary_rows.append({"block": fam, "metric": "reliable_rows", "value": int(sub["model_reliability"].eq("reliable").sum())})
    if not vecm_results.empty:
        summary_rows.append({"block": "VECM", "metric": "feasible_systems", "value": int(vecm_results["status"].eq("ok").sum())})
    if not index_models.empty:
        summary_rows.append({"block": "Aggregate index models", "metric": "rows", "value": int(len(index_models))})
    if not index_vecm_summary.empty:
        summary_rows.append({"block": "Aggregate index VECM", "metric": "feasible_systems", "value": int(index_vecm_summary["status"].eq("ok").sum())})
    if not discounts.empty:
        summary_rows.append({"block": "Discounts", "metric": "strategy_signals", "value": int(discounts["discount_strategy_signal"].sum())})
    if not scale_models.empty:
        summary_rows.append({"block": "Procurement scale", "metric": "scale_signals", "value": int(scale_models["scale_signal_flag"].sum())})
    if not findings.empty:
        summary_rows.append({"block": "Integrated findings", "metric": "rows", "value": int(len(findings))})
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        pd.DataFrame(summary_rows).to_excel(writer, sheet_name="Summary", index=False)
        intersections.to_excel(writer, sheet_name="Intersections", index=False)
        corr_scan.to_excel(writer, sheet_name="LagScan", index=False)
        chain_models.to_excel(writer, sheet_name="ChainModels", index=False)
        vecm_results.to_excel(writer, sheet_name="VECM", index=False)
        if not index_models.empty:
            index_models.to_excel(writer, sheet_name="AggregateIndexModels", index=False)
        if not index_vecm_summary.empty:
            index_vecm_summary.to_excel(writer, sheet_name="AggregateIndexVECM", index=False)
        if not link21_summary.empty:
            link21_summary.to_excel(writer, sheet_name="Link21Summary", index=False)
        if not reliability_overview.empty:
            reliability_overview.to_excel(writer, sheet_name="ReliabilityOverview", index=False)
        discounts.to_excel(writer, sheet_name="Discounts", index=False)
        scale_models.to_excel(writer, sheet_name="ProcScale", index=False)
        findings.to_excel(writer, sheet_name="Findings", index=False)


def _resolve_doc_asset(doc_path: Path, asset_ref: str) -> Path:
    asset = Path(asset_ref)
    if asset.is_absolute():
        return asset
    return (doc_path.parent / asset).resolve()


def _apply_kse_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.079)
        section.bottom_margin = Inches(0.96)
        section.left_margin = Inches(1.441)
        section.right_margin = Inches(0.96)
    normal = doc.styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(14)
    if "Title" in doc.styles:
        title = doc.styles["Title"]
        title.font.name = "Garamond"
        title.font.size = Pt(16)
        title.font.bold = True
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Garamond"
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(4)
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Garamond"
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)


def _add_doc_table(doc: Document, df: pd.DataFrame, caption: str, note: str = "") -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    run.bold = True

    safe_df = df.copy()
    for col in safe_df.columns:
        if pd.api.types.is_float_dtype(safe_df[col]):
            safe_df[col] = safe_df[col].map(lambda x: "" if pd.isna(x) else f"{x:.3f}")
        else:
            safe_df[col] = safe_df[col].fillna("").astype(str)

    table = doc.add_table(rows=1, cols=len(safe_df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(safe_df.columns):
        hdr_cells[idx].text = str(col)
        for p in hdr_cells[idx].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for r in p.runs:
                r.bold = True
                r.font.name = "Garamond"
                r.font.size = Pt(9)
    for row in safe_df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.name = "Garamond"
                    r.font.size = Pt(9)
    if note:
        np = doc.add_paragraph()
        np.alignment = WD_ALIGN_PARAGRAPH.LEFT
        np.paragraph_format.space_before = Pt(3)
        np.paragraph_format.space_after = Pt(10)
        nr = np.add_run(note)
        nr.italic = True
        nr.font.name = "Garamond"
        nr.font.size = Pt(10)


def _add_formula_box(doc: Document, formula_text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(formula_text)
    run.font.name = "Garamond"
    run.font.size = Pt(11)
    run.bold = True


def write_kse_docx(path: Path, lines: Sequence[str]) -> None:
    doc = Document()
    _apply_kse_style(doc)
    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("{{TABLE|") and stripped.endswith("}}"):
            payload = stripped[8:-2]
            parts = payload.split("|", 2)
            if len(parts) >= 2:
                caption = parts[0].strip()
                csv_ref = parts[1].strip()
                note = parts[2].strip() if len(parts) > 2 else ""
                csv_path = _resolve_doc_asset(path, csv_ref)
                if csv_path.exists():
                    _add_doc_table(doc, pd.read_csv(csv_path), caption, note)
            continue
        if stripped.startswith("{{FORMULA|") and stripped.endswith("}}"):
            _add_formula_box(doc, stripped[10:-2].strip())
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(stripped[3:]).bold = True
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(stripped[4:]).bold = True
            continue
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.line_spacing = 1.5
            p.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.add_run(stripped[2:])
            continue
        m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if m:
            caption = m.group(1).strip()
            img_path = _resolve_doc_asset(path, m.group(2).strip())
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.2))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(10)
                run = cp.add_run(caption)
                run.italic = True
                run.font.name = "Garamond"
                run.font.size = Pt(11)
            continue
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        text = stripped.replace("`", "")
        p.add_run(text)
    doc.save(path)


def write_run_summary(
    product_dictionary: pd.DataFrame,
    intersections: pd.DataFrame,
    corr_scan: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
) -> None:
    strong = int(intersections["admissibility_status"].eq("strong").sum()) if not intersections.empty else 0
    acceptable = int(intersections["admissibility_status"].eq("acceptable").sum()) if not intersections.empty else 0
    weak_extension = int(intersections["admissibility_status"].eq("weak_extension").sum()) if not intersections.empty else 0
    reliable = int(chain_models["model_reliability"].eq("reliable").sum()) if not chain_models.empty else 0
    cond_usable = int(chain_models["model_reliability"].eq("conditionally_usable").sum()) if not chain_models.empty else 0
    vecm_ok = int(vecm_results["status"].eq("ok").sum()) if not vecm_results.empty else 0
    discount_signals = int(discounts["discount_strategy_signal"].sum()) if not discounts.empty else 0
    scale_signals = int(scale_models["scale_signal_flag"].sum()) if not scale_models.empty else 0
    model_family_counts = chain_models["model_family"].value_counts().to_dict() if not chain_models.empty else {}
    top_corr = corr_scan.sort_values("best_abs_corr", ascending=False).head(1) if not corr_scan.empty else pd.DataFrame()
    lines = [
        "# FINAL_RESEARCH Run Summary",
        "",
        "## Execution order",
        "- Load corrected full_uah_final.xlsx.",
        "- Run product-definition audit across government, procurement, and retail sources.",
        "- Build corrected daily panel with retail combined, Silpo, and Novus downstream variants.",
        "- Construct weekly median panels and controlled smoothing variants.",
        "- Score intersections and admissibility before core estimation.",
        "- Run the first strict weekly model pass.",
        "- Activate post-test adaptation where strict weekly links remain too thin, keeping adapted evidence labelled separately.",
        "- Estimate weekly ARDL, ECM, NARDL, and VECM system models.",
        "- Estimate daily local projections, margin models, discount models, and procurement-scale models.",
        "- Generate figures, workbooks, and thesis-style Chapter 5-6 outputs.",
        "- Save separate compact results, diagnostics, and notes for individual retained models.",
        "",
        "## Key counts",
        f"- Product dictionary rows: {len(product_dictionary)}",
        f"- Strong intersections: {strong}",
        f"- Acceptable intersections: {acceptable}",
        f"- Weak-extension intersections: {weak_extension}",
        f"- Reliable core models: {reliable}",
        f"- Conditionally usable core models: {cond_usable}",
        f"- Feasible VECM systems: {vecm_ok}",
        f"- Discount strategy signals: {discount_signals}",
        f"- Procurement-scale signals: {scale_signals}",
    ]
    if strong + acceptable == 0:
        lines.extend(
            [
                "",
                "## Admissibility note",
                "- The strict weekly stage did not retain any strong or acceptable links.",
                "- Weekly chain evidence therefore comes from the adapted direct-mapped extension layer, kept separate from strict identification.",
            ]
        )
    if model_family_counts:
        lines.extend(
            [
                "",
                "## Retained weekly model families",
                *(f"- {family}: {count}" for family, count in model_family_counts.items()),
            ]
        )
        if "ARDL" not in model_family_counts:
            lines.append("- ARDL was screened in the weekly pass but no ARDL row met the retained reporting standard on corrected data.")
    if not top_corr.empty:
        r = top_corr.iloc[0]
        lines.extend(
            [
                "",
                "## Strongest lag signal",
                f"- {r['product_label']} | {r['link']} | corr {fmt_num(r['best_abs_corr'])} | lag {int(r['best_lag_weeks'])} weeks",
            ]
        )
    md_path = LOG_DIR / "run_summary.md"
    txt_path = LOG_DIR / "run_summary.txt"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    txt_path.write_text("\n".join(line.replace("# ", "").replace("## ", "") for line in lines), encoding="utf-8")
    write_html(LOG_DIR / "run_summary.html", "FINAL_RESEARCH Run Summary", lines)


def write_final_readme() -> None:
    readme = [
        "# FINAL_RESEARCH",
        "",
        "This folder contains the corrected, reproducible rerun of the dairy price-transmission thesis pipeline built on `full_uah_final.xlsx`.",
        "",
        "The project is designed as one self-contained empirical system: source materials, processed data, econometric outputs, figures, run logs, and thesis-style writing are stored here so the analysis can be reopened without relying on scattered external files.",
        "",
        "## Research scope",
        "- Topic: vertical price transmission in Ukraine's dairy chain.",
        "- Core chain: FarmGate Ukraine average -> Producer Ukraine average -> ProZorro procurement price -> downstream retail endpoint.",
        "- Downstream mechanisms: Silpo discount behavior, Novus comparison, retail spread behavior, and consumer-benchmark comparison.",
        "- System extensions: aggregate chain indices and widened VECM robustness blocks.",
        "",
        "## Folder structure",
        "```text",
        "FINAL_RESEARCH/",
        "├── code/                     Reproducible Python scripts for the pipeline and companion documents",
        "├── data/                     Final processed datasets used by the modelling system",
        "├── documents/                Thesis-ready chapter text and workflow/result companions",
        "├── figures/                  Full graph system, split into chapter, sequence, and model-specific folders",
        "├── logs/                     Run summary and execution notes",
        "├── materials/                Inputs, references, and instruction documents used by the project",
        "├── outputs/                  Model outputs, summaries, chapter tables, and per-model exports",
        "└── README.md                 This guide",
        "```",
        "",
        "## Key entry points",
        "- `code/final_research_pipeline.py`: main end-to-end pipeline and output writer.",
        "- `code/generate_stepbystep_doc.py`: generates the workflow explanation document.",
        "- `code/generate_stepbystep_results.py`: generates the result-by-result companion and refreshes the chapter package.",
        "- `materials/inputs/full_uah_final.xlsx`: corrected governmental source workbook.",
        "- `materials/inputs/full_uah_final_whatadded_matched_smoothed.xlsx`: final matched/smoothed modelling workbook used by the pipeline.",
        "- `outputs/final_research_outputs.xlsx`: master workbook with the main processed tables and model summaries.",
        "- `outputs/integrated_summary_workbook.xlsx`: compact summary workbook for fast review.",
        "- `documents/Chapter5_6_analysis.docx`: thesis-style Chapter 5 and Chapter 6 document for the dissertation.",
        "- `documents/stepbystep.docx`: workflow explanation of transformations, datasets, and models.",
        "- `documents/stepbystep_results.docx`: result-by-result companion with figures and short interpretations.",
        "- `logs/run_summary.md`: compact run note with the main counts and retained model families.",
        "",
        "## Materials",
        "### `materials/inputs/`",
        "- `full_uah_final.xlsx`: corrected FarmGateUA, ProducerUA, and ConsumerUA source workbook used as the base truth.",
        "- `full_uah_final_whatadded_matched_smoothed.xlsx`: final modelling workbook with matched and smoothed panels used in the rerun.",
        "- `full_uah_final_whatadded_matched_smoothed.xlsx 2.xlsx`: duplicate save kept for traceability because it was part of the working project history.",
        "",
        "### `materials/references/`",
        "- `Charniuk_Maksym_MScThesis_Draft_correctedformat.docx`: thesis-format and writing reference for Chapters 5 and 6.",
        "- `data_estiamtion_updated_conclusion_fullversion.md`: earlier integrated chapter text used as a comparison and transition reference.",
        "",
        "### `materials/instructions/`",
        "- Correction and methodology notes used during the rebuild, including formatting guidance and data-method summaries.",
        "",
        "## Processed data (`data/`)",
        "- `product_dictionary_standardized.csv`: standardized product taxonomy across governmental, procurement, and retail layers.",
        "- `product_audit_long.csv`: long-form product-definition audit across sources.",
        "- `final_daily_panel.csv`: final modelling panel at daily frequency.",
        "- `final_weekly_panel.csv`: weekly median modelling panel used for long-run specifications.",
        "- `intersection_admissibility.csv`: strong / acceptable / weak / unusable overlap classification.",
        "- `final_panel_coverage.csv`: source coverage and overlap audit.",
        "- `aggregate_chain_index_daily.csv`, `aggregate_chain_index_weekly.csv`, `aggregate_chain_index_weights.csv`: aggregate dairy-chain index series and weights.",
        "- `retail_items_full_harmonized.csv`: cleaned SKU-level retail archive used to build category-level retail prices.",
        "- `retail_item_catalog.csv`, `retail_match_audit.csv`, `retail_name_reconciliation_examples.csv`: retail matching and harmonization audit outputs.",
        "- `retail_brand_daily.csv`, `retail_brand_support.csv`: brand-level retail panels and support measures.",
        "- `retail_level_scores.csv`, `retail_level_selection.csv`, `retail_optimal_daily.csv`: downstream endpoint selection and chosen retail layer.",
        "- `consumerua_clean_daily.csv`, `cme_class3_daily.csv`, `europe_benchmark_daily.csv`: benchmark and cleaned supporting series.",
        "",
        "## Outputs (`outputs/`)",
        "### Main workbooks and summaries",
        "- `final_research_outputs.xlsx`: main results workbook used for integrated review.",
        "- `integrated_summary_workbook.xlsx`: condensed cross-model summary workbook.",
        "- `core_chain_models.csv`, `discount_strategy_models.csv`, `procurement_scale_models.csv`, `aggregate_index_models.csv`, `aggregate_index_vecm.csv`: model-family exports.",
        "- `lag_correlation_scan.csv`, `link21_summary.csv`, `model_reliability_overview.csv`, `robust_findings.csv`: core screening and synthesis outputs.",
        "",
        "### Chapter tables",
        "- `outputs/chapter_tables/`: thesis-facing tables used by `Chapter5_6_analysis.docx`.",
        "",
        "### Per-model exports",
        "- `outputs/single_model_tables/`: one compact results file for each weekly specification and special model block.",
        "- `outputs/single_model_diagnostics/`: diagnostics saved per retained model.",
        "- `outputs/single_model_notes/`: plain-language notes and interpretation fragments per model.",
        "",
        "### VECM detail",
        "- `outputs/vecm_detail/`: table-by-table VECM exports, including stationarity, lag selection, Johansen tests, alpha/beta, speed of adjustment, IRF, and FEVD where feasible.",
        "",
        "## Figures (`figures/`)",
        "- Top-level numbered PNGs: integrated figure inventory covering the whole project.",
        "- `chapter5_data/`: figures used or considered for the data chapter.",
        "- `chapter6_results/`: figures used or considered for the results chapter.",
        "- `sequence/`: figures ordered by execution logic, from raw data to final model interpretation.",
        "- `model_specific/`: reserved for additional model-specific visuals.",
        "- `reliability/`: reserved for reliability-oriented visuals.",
        "",
        "## Documents (`documents/`)",
        "- `Chapter5_6_analysis.docx/.md/.html`: final thesis chapter package.",
        "- `stepbystep.docx/.md/.html`: methodological walkthrough of the final research sequence.",
        "- `stepbystep_results.docx/.md/.html`: comprehensive results companion with figures and short interpretations.",
        "",
        "## Logs (`logs/`)",
        "- `run_summary.md/.txt/.html`: compact execution note with the current model counts and main retained result signal.",
        "",
        "## Current run status",
        "- Product dictionary rows: 28",
        "- Strong intersections: 1",
        "- Acceptable intersections: 4",
        "- Weak-extension intersections: 10",
        "- Reliable core models: 30",
        "- Conditionally usable core models: 21",
        "- Feasible VECM systems: 3",
        "- Discount strategy signals: 2",
        "- Procurement-scale signals: 5",
        "- Strongest lag signal: Sour cream | FarmGate -> Producer | corr 0.932 | lag 0 weeks",
        "",
        "## Execution order",
        "1. Load the corrected governmental source workbook and validate the active sheets.",
        "2. Run the product-definition audit across governmental, procurement, and retail sources.",
        "3. Build cleaned retail product mappings, brand support tables, and the final daily panel.",
        "4. Construct weekly medians and controlled smoothing variants for long-run modelling.",
        "5. Score intersections and classify admissibility before estimation.",
        "6. Run the first strict weekly screening pass.",
        "7. Apply explicit post-test adaptation when strict overlap remains too thin.",
        "8. Estimate weekly ECM and NARDL retained models, while keeping ARDL screening outputs for traceability.",
        "9. Estimate daily local projections, vertical spread models, Silpo discount models, and procurement-scale models.",
        "10. Estimate aggregate chain index models and widened system/VECM robustness blocks where feasible.",
        "11. Export separate model tables, diagnostics, notes, chapter tables, and VECM detail tables.",
        "12. Generate figures, summaries, and the thesis-style chapter and companion documents.",
        "",
        "## How to use this folder",
        "1. Open `documents/Chapter5_6_analysis.docx` if you want the thesis-ready narrative first.",
        "2. Open `logs/run_summary.md` if you want the current high-level status in one page.",
        "3. Open `outputs/final_research_outputs.xlsx` for the integrated numerical output.",
        "4. Open `documents/stepbystep.docx` to understand the data-construction and modelling sequence.",
        "5. Open `documents/stepbystep_results.docx` if you want all interpretable outputs and the full figure appendix.",
        "6. Use `outputs/single_model_tables/`, `outputs/single_model_diagnostics/`, and `outputs/single_model_notes/` for model-by-model inspection.",
        "",
        "## Reproducibility note",
        "- The folder is organized so the core source materials used by the rerun are stored in `materials/` and the generated outputs are stored inside the project itself.",
        "- The modelling outputs currently present in `outputs/` are the active final rerun outputs; this README documents that state rather than promising a fresh rerun on every open.",
    ]
    (PROJECT_DIR / "README.md").write_text("\n".join(readme), encoding="utf-8")


def write_chapter56_legacy(
    inventory: pd.DataFrame,
    weekly: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
    product_dictionary: pd.DataFrame,
    intersections: pd.DataFrame,
    corr_scan: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    lp_summary: pd.DataFrame,
    margins: pd.DataFrame,
    margin_spreads: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
    link21_summary: pd.DataFrame,
    reliability_overview: pd.DataFrame,
    index_weekly: pd.DataFrame,
    index_weights: pd.DataFrame,
    index_models: pd.DataFrame,
    index_vecm_summary: pd.DataFrame,
) -> None:
    chapter_table_dir = OUTPUT_DIR / "chapter_tables"
    chapter_table_dir.mkdir(parents=True, exist_ok=True)

    n_products = int(product_dictionary["standardized_product"].nunique()) if not product_dictionary.empty else 0
    ambiguity = int(product_dictionary.get("ambiguity_rows", pd.Series(dtype=float)).fillna(0).sum()) if not product_dictionary.empty else 0
    strong = int(intersections["admissibility_status"].eq("strong").sum()) if not intersections.empty else 0
    acceptable = int(intersections["admissibility_status"].eq("acceptable").sum()) if not intersections.empty else 0
    weak_extension = int(intersections["admissibility_status"].eq("weak_extension").sum()) if not intersections.empty else 0
    unusable = int(intersections["admissibility_status"].eq("unusable").sum()) if not intersections.empty else 0
    reliable = chain_models[chain_models["model_reliability"].eq("reliable")].copy() if not chain_models.empty else pd.DataFrame()
    conditional = chain_models[chain_models["model_reliability"].eq("conditionally_usable")].copy() if not chain_models.empty else pd.DataFrame()
    retained = chain_models[chain_models["model_reliability"].isin(["reliable", "conditionally_usable"])].copy() if not chain_models.empty else pd.DataFrame()
    ardl_rows = chain_models[chain_models["model_family"].eq("ARDL")].copy() if not chain_models.empty else pd.DataFrame()
    ecm_rows = chain_models[chain_models["model_family"].eq("ECM")].copy() if not chain_models.empty else pd.DataFrame()
    nardl_rows = chain_models[chain_models["model_family"].eq("NARDL")].copy() if not chain_models.empty else pd.DataFrame()
    top_chain = reliable.assign(abs_coef=lambda d: d["lr_coef"].where(d["lr_coef"].notna(), d["sr_coef"]).abs()).sort_values("abs_coef", ascending=False).head(7) if not reliable.empty else pd.DataFrame()
    top_corr = corr_scan.sort_values("best_abs_corr", ascending=False).head(6) if not corr_scan.empty else pd.DataFrame()
    vecm_ok = vecm_results[vecm_results["status"].eq("ok")].copy() if not vecm_results.empty else pd.DataFrame()
    vecm_attempts = int(len(vecm_results)) if not vecm_results.empty else 0
    top_discount = discounts.sort_values("lag_discount_coef", ascending=False).head(5) if not discounts.empty else pd.DataFrame()
    top_scale = scale_models.sort_values("d_sum_current_coef", ascending=False).head(5) if not scale_models.empty else pd.DataFrame()
    discount_signals = int(discounts.get("discount_strategy_signal", pd.Series(dtype=float)).fillna(0).sum()) if not discounts.empty else 0
    scale_signals = int(scale_models.get("scale_signal_flag", pd.Series(dtype=float)).fillna(0).sum()) if not scale_models.empty else 0
    margin_flags = int(margins.get("persistent_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    asym_flags = int(margins.get("asymmetric_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    retail_items = source_frames.get("retail_items_full", pd.DataFrame())
    retail_match = source_frames.get("retail_match_audit", pd.DataFrame())
    brand_support = source_frames.get("retail_brand_support", pd.DataFrame())
    literal_summary = source_frames.get("retail_literal_summary", pd.DataFrame())
    level_choice = source_frames.get("retail_level_selection", pd.DataFrame())
    match_counts = retail_match["match_status"].value_counts() if not retail_match.empty and "match_status" in retail_match.columns else pd.Series(dtype=float)
    matched_both = int(match_counts.get("matched_both_shops", 0))
    silpo_only = int(match_counts.get("silpo_only", 0))
    novus_only = int(match_counts.get("novus_only", 0))
    strict_aligned = int(retail_match.get("strict_alignment_flag", pd.Series(dtype=float)).fillna(0).sum()) if not retail_match.empty else 0
    producer = source_frames.get("producer", pd.DataFrame())
    farm = source_frames.get("farm_gate", pd.DataFrame())
    europe = source_frames.get("europe_benchmark", pd.DataFrame())
    cme = source_frames.get("cme_benchmark", pd.DataFrame())
    producer_price_col = "producer_linear" if "producer_linear" in producer.columns else ("price_linear" if "price_linear" in producer.columns else None)
    farm_price_col = "farmgate_combined_linear" if "farmgate_combined_linear" in farm.columns else ("price_linear" if "price_linear" in farm.columns else None)
    eu_corr_rows: List[Tuple[str, int, float]] = []
    if not producer.empty and not europe.empty and producer_price_col is not None:
        for product in sorted(set(producer.get("product", [])).intersection(set(europe.get("product", [])))):
            left = producer[producer["product"].eq(product)][["date", producer_price_col]].rename(columns={producer_price_col: "producer_price"})
            right = europe[europe["product"].eq(product)][["date", "eu_price_uah"]]
            merged = left.merge(right, on="date", how="inner").dropna()
            if len(merged) >= 30:
                eu_corr_rows.append((PRODUCT_LABELS.get(product, product), int(len(merged)), float(merged["producer_price"].corr(merged["eu_price_uah"]))))
        eu_corr_rows = sorted(eu_corr_rows, key=lambda x: abs(x[2]), reverse=True)
    cme_corr = np.nan
    if not farm.empty and not cme.empty and farm_price_col is not None:
        merged = farm[["date", farm_price_col]].rename(columns={farm_price_col: "farm_price"}).merge(cme[["date", "cme_class3_uah"]], on="date", how="inner").dropna()
        if len(merged) >= 30:
            cme_corr = float(merged["farm_price"].corr(merged["cme_class3_uah"]))
    lp_focus = pd.DataFrame()
    if not lp_summary.empty:
        lp_focus = (
            lp_summary[lp_summary["horizon_days"].isin([7, 14])]
            .sort_values(["core_share", "sig_share", "median_coef"], ascending=[False, False, False])
            .head(8)
            .copy()
        )
    silpo_discount = (
        retail_items[retail_items["retailer"].eq("Silpo")].groupby("product", as_index=False).agg(
            mean_discount_share=("discount_present", "mean"),
            median_markdown=("markdown_rate", lambda s: float(s[s > 0].median()) if pd.Series(s).gt(0).any() else 0.0),
        )
        if not retail_items.empty and {"retailer", "discount_present", "markdown_rate", "product"}.issubset(retail_items.columns)
        else pd.DataFrame()
    )
    diag_text = "Diagnostic coverage is limited because only retained rows are discussed, but the reliability filter already removes models with the weakest overlap or unstable core diagnostics."
    if not retained.empty:
        lb_ok = int(retained["ljungbox_p"].ge(0.05).sum())
        bp_ok = int(retained["bp_p"].ge(0.05).sum())
        white_ok = int(retained["white_p"].ge(0.05).sum())
        jb_ok = int(retained["jb_p"].ge(0.05).sum())
        diag_text = (
            f"Within the retained weekly rows, {lb_ok} satisfy the Ljung-Box threshold, {bp_ok} satisfy the Breusch-Pagan threshold, "
            f"{white_ok} satisfy the White threshold, and {jb_ok} satisfy the Jarque-Bera threshold. These counts should not be read as a mechanical pass/fail score, "
            "but they do show that the final interpretation is already restricted to the cleaner part of the model set."
        )
    weight_text = "The aggregate dairy index uses a neutral geometric weighting rule because direct quantity weights are unavailable."
    if not index_weights.empty:
        top_weights = index_weights.sort_values("weight", ascending=False).head(5)
        weight_text = (
            "The aggregate dairy index is constructed as a geometric chain index with fixed structural proxy weights. "
            "Weights combine procurement participation, retail assortment support, and an equal-weight anchor "
            "according to w_i = 0.45*w_i^procurement + 0.45*w_i^retail + 0.10*w_i^equal. "
            "The largest final weights belong to "
            + "; ".join(f"{r.product_label} ({fmt_pct(r.weight)})" for r in top_weights.itertuples())
            + "."
        )

    link21_table = pd.DataFrame()
    if not link21_summary.empty:
        link21_table = link21_summary[
            [
                "link",
                "best_admissibility",
                "median_overlap_weeks",
                "best_corr",
                "best_lag_weeks",
                "best_model_family",
                "display_coef",
                "model_reliability",
                "reading",
            ]
        ].copy()
        link21_table = link21_table.rename(
            columns={
                "best_admissibility": "admissibility",
                "median_overlap_weeks": "overlap_weeks",
                "best_corr": "best_corr",
                "best_lag_weeks": "lag_weeks",
                "best_model_family": "best_model",
                "display_coef": "coef",
                "model_reliability": "reliability",
            }
        )
        link21_table.to_csv(chapter_table_dir / "table_6_1_link21_summary.csv", index=False)

    reliability_table = pd.DataFrame()
    if not reliability_overview.empty:
        reliability_table = reliability_overview.rename(
            columns={
                "block": "model_block",
                "family": "family",
                "rows": "rows_estimated",
                "reliable_rows": "reliable",
                "conditionally_usable_rows": "conditional",
            }
        )
        reliability_table.to_csv(chapter_table_dir / "table_6_2_reliability_overview.csv", index=False)

    index_table = pd.DataFrame()
    if not index_models.empty:
        index_table = index_models[
            ["link", "model_family", "data_variant", "n_obs", "lr_coef", "sr_coef", "ect_coef", "model_reliability"]
        ].copy()
        index_table["display_coef"] = index_table["lr_coef"].where(index_table["lr_coef"].notna(), index_table["sr_coef"])
        index_table = index_table.drop(columns=["lr_coef", "sr_coef"])
        index_table.to_csv(chapter_table_dir / "table_6_3_aggregate_index_models.csv", index=False)

    vecm_table = pd.DataFrame()
    if not index_vecm_summary.empty:
        vecm_table = index_vecm_summary.copy()
    elif not vecm_results.empty:
        vecm_table = vecm_results[["product_label", "system_name", "status", "reason", "n_obs"]].copy()
    if not vecm_table.empty:
        vecm_table.to_csv(chapter_table_dir / "table_6_4_vecm_summary.csv", index=False)

    mean_discount_text = ""
    if not silpo_discount.empty:
        top_disc = silpo_discount.sort_values("mean_discount_share", ascending=False).head(4)
        mean_discount_text = (
            "At Silpo, the deepest discount environment is concentrated in "
            + "; ".join(
                f"{PRODUCT_LABELS.get(r.product, r.product)} with mean discount share {fmt_pct(r.mean_discount_share)} and median markdown {fmt_pct(r.median_markdown)}"
                for r in top_disc.itertuples()
            )
            + "."
        )

    brand_text = top_brand_sentence(brand_support)
    literal_text = top_literal_sentence(literal_summary)
    level_text = level_selection_sentence(level_choice)

    lines: List[str] = [
        "# Chapter 5 and Chapter 6. Empirical analysis of vertical price transmission in Ukraine's dairy market",
        "",
        "## Chapter 5. Data, reconstruction logic, and admissibility of the empirical design",
        "",
        "The empirical strategy of this thesis follows the central idea developed in the introduction: price pressure reaches households through the consumer basket, but the shock does not become consumer-visible immediately or uniformly. In the Ukrainian dairy market, the same disturbance can be transformed at several points of the chain. It can be muted at the farm-gate stage, translated into processor prices only partially, filtered by institutional procurement rules, or re-timed at retail through category management and discounting. For that reason, the corrected final rerun treats data construction not as a preliminary technical step, but as part of the identification strategy itself.",
        "",
        "The current version of the analysis is built on corrected governmental monthly point-in-time prices and a stricter downstream design. It therefore differs from earlier internal runs in two important ways. First, the Ukraine-wide farm-gate, producer, and consumer layers now reflect true monthly prices rather than cumulative January-to-current-month aggregates. Second, the downstream endpoint is no longer treated as a generic retail average. It is reconstructed at category level from the Silpo and Novus item universe, and it is linked to ConsumerUA only after standardization of units, products, and price meaning.",
        "",
        "### 5.1 Final data sources, territorial logic, and product-definition audit",
        "",
        "The core chain is defined at the territory level of Ukraine. This is not an arbitrary simplification, but a necessary condition for internal consistency. The corrected farm-gate, producer, and consumer series are national averages by construction, so the main empirical chain must preserve the same territorial meaning. Regional heterogeneity remains relevant in procurement and retail, especially for ProZorro and retailer assortment, but regional evidence is treated as an additional layer rather than as the basis of the main equilibrium models.",
        "",
        "Accordingly, the main chain used in the thesis is Farm-gate Ukraine average, Producer Ukraine average, ProZorro all-Ukraine procurement unit price, and a downstream retail endpoint. Farm-gate remains the raw-milk benchmark. Producer represents domestic processor prices by dairy category. ProZorro captures institutional procurement prices from actual tenders and contracts. Retail is treated as the consumer-facing layer, but it is decomposed into observed prices and strategy variables where retailer data allow this.",
        "",
        f"The final inventory contains {len(inventory):,} stored source blocks or derivative datasets. The harmonised retail item universe alone contains {len(retail_items):,} item-date observations, while the product dictionary consolidates {len(product_dictionary)} mapping rows into {n_products} standardized dairy groups. The audit flags {ambiguity} ambiguous or approximate mappings. This is important because the model does not assume that products are comparable merely because names look similar. Each product can be direct, approximate, inferred, constructed, widened, or invalid depending on how the underlying source defines the economic object.",
        "",
        "The product-definition audit is therefore the first formal stage of the pipeline. Governmental producer and consumer sources are relatively transparent because products are carried in explicit official labels. Farm-gate is simpler but narrower because it refers to raw milk rather than processed products. The most difficult cases are ProZorro, where one product may be embedded in profile titles or institutional wording, and retail, where the product meaning must be inferred from the joint content of title, brand, fat content, package size, and category. The final model never assumes that these naming systems are equivalent by default.",
        "",
        "Figure 5.1 shows the corrected governmental layers before transformation. It is useful as an opening figure because it makes visible the basic fact on which the whole rerun rests: the national producer and consumer price layers move over a long horizon with economically plausible level changes, while the farm-gate benchmark remains the smoother raw-milk series against which later product prices are interpreted rather than mechanically matched.",
        "",
        "![Figure 5.1. Raw corrected governmental series used in the core chain.](../figures/chapter5_data/01_raw_government_layers.png)",
        "",
        "Source: author's calculations based on the corrected `full_uah_final.xlsx` governmental layers.",
        "",
        "Figure 5.2 then places the raw retail unit-price series alongside these upstream layers at product level. The purpose is descriptive rather than inferential. Before any smoothing, interpolation, or modelling, it is necessary to see how much of the retail variation is genuine category movement and how much is driven by assortment rotation, retailer-specific pricing, and discount activity.",
        "",
        "![Figure 5.2. Raw retail observed-price series before transformation.](../figures/chapter5_data/02_raw_retail_observed_series.png)",
        "",
        "Source: author's calculations based on the harmonised Silpo and Novus item data.",
        "",
        "### 5.2 Reconstruction logic, external benchmarks, and transformation choices",
        "",
        "The corrected governmental layers remain the empirical base truth. Their interpolation to daily frequency is preserved from the corrected input workbook, but in the estimation stage they are not used blindly at the highest possible frequency. Instead, daily versions are kept for short-run mechanism models, while weekly medians are constructed for long-run and equilibrium-style modelling. This distinction matters because a daily interpolation can support alignment and descriptive continuity, but it should not be mistaken for direct day-by-day market observation.",
        "",
        "European dairy prices and CME Class III remain in the research, but only as external benchmarks. They are not structural stages of the Ukrainian domestic chain. Their role is threefold. First, they locate the Ukrainian dairy market within the broader global and European price environment. Second, they make transparent the fact that benchmark information was used to shape the weekly path of reconstructed governmental series. Third, they help us understand whether the corrected domestic series continue to move in a way that is economically coherent once the previous aggregation error is removed.",
        "",
        "The benchmark comparison remains relevant precisely because the corrected rerun is stricter. If the corrected Ukrainian series were to drift away from any plausible benchmark structure, that would weaken confidence in the reconstruction. Instead, the benchmark comparison shows economically recognizable co-movement without collapsing domestic prices into foreign ones. The domestic interpretation therefore remains structurally valid, while the benchmark block clarifies the wider dairy environment in which Ukraine operates.",
        "",
        *(
            [
                f"The strongest producer-to-Europe correlations are observed for {eu_corr_rows[0][0]} ({fmt_num(eu_corr_rows[0][2])}, {eu_corr_rows[0][1]} aligned observations)"
                + (f", {eu_corr_rows[1][0]} ({fmt_num(eu_corr_rows[1][2])}, {eu_corr_rows[1][1]} observations)" if len(eu_corr_rows) > 1 else "")
                + (f", and {eu_corr_rows[2][0]} ({fmt_num(eu_corr_rows[2][2])}, {eu_corr_rows[2][1]} observations)." if len(eu_corr_rows) > 2 else ".")
            ]
            if eu_corr_rows
            else ["The benchmark role is retained even where product-specific correlations differ, because the benchmark block is mainly descriptive and reconstructive rather than identificational."]
        ),
        "",
        (f"At the raw-milk end of the chain, the corrected farm-gate benchmark has a correlation of {fmt_num(cme_corr)} with CME Class III over the overlapping sample. This supports the view that Ukraine's raw-milk market is exposed to the broader dairy cycle without behaving as a simple external-price replica." if pd.notna(cme_corr) else "The CME comparison is preserved as a contextual raw-milk benchmark rather than as a structural explanatory variable in the domestic chain."),
        "",
        "Figure 5.3 shows the external benchmark block itself, while Figures 5.4 and 5.5 make the reconstruction issue more concrete by plotting product-level domestic series against aggregate chain indices. This sequencing is important for interpretation. The reader first sees the external context, and only after that the transformed domestic structure that will be used in the models.",
        "",
        "![Figure 5.3. External benchmark block for European dairy prices and CME Class III.](../figures/chapter5_data/03_raw_external_benchmarks.png)",
        "",
        "Source: author's calculations based on European benchmark and CME Class III series converted into UAH-equivalent units.",
        "",
        "![Figure 5.4. Product-level price paths and aggregate dairy-chain indices.](../figures/chapter5_data/04_dataset_product_lines_and_indices.png)",
        "",
        "Source: author's calculations based on product-level producer, procurement, retail, and consumer panels.",
        "",
        "![Figure 5.5. Aggregate dairy price indices by chain level.](../figures/chapter5_data/05_aggregate_chain_indices.png)",
        "",
        "Source: author's calculations based on geometric chain indices with structural proxy weights.",
        "",
        "The aggregate index layer is used only as a robustness system. Because direct quantity weights are unavailable, the thesis does not use estimated transmission coefficients as weights. That would be circular. Instead, product weights are constructed from exogenous or weakly endogenous structural proxies: procurement participation, retail assortment support, and a small equal-weight anchor. In log-linear form the aggregate index is defined as ln(P_t^agg) = Σ_i w_i ln(P_it), where the weights are fixed at the product level rather than allowed to move endogenously with the estimated pass-through. " + weight_text,
        "",
        "### 5.3 Retail harmonisation, discount logic, and downstream endpoint construction",
        "",
        "The downstream block is the part of the data architecture that changed most relative to earlier runs. The key improvement is that retail prices are now built from harmonised product-type series instead of being treated as a broad pool of interchangeable SKUs. Every Silpo and Novus item is mapped into a standardized dairy taxonomy, with package size, fat content, and product meaning checked before aggregation. This step was necessary because a large part of retail volatility can otherwise come from composition changes rather than from price adjustment itself.",
        "",
        f"After harmonisation, the audit identifies {matched_both} item keys observed in both shops, {silpo_only} Silpo-only keys, and {novus_only} Novus-only keys. Only {strict_aligned} of the matched keys are strict one-to-one alignments on the tighter diagnostic specification. This is why the final downstream endpoint uses category-level medians and baseline series instead of pretending that the raw item universe is stable enough to be pooled without adjustment.",
        "",
        brand_text,
        "",
        literal_text,
        "",
        "Silpo is modelled in two linked ways. First, `price_current` is the actual post-discount price observed by the consumer and therefore the correct observed downstream price for transmission analysis. Second, discount existence and discount depth are kept as a separate behavioural layer. This distinction matters economically. It allows the thesis to ask not only whether costs reach the shelf, but also whether the retailer chooses to reveal that adjustment through visible price changes or to redistribute it through markdowns.",
        "",
        mean_discount_text or "Silpo discounts remain product-specific, which confirms that promotional behaviour should be analysed as a structured component of retail adjustment rather than as random noise.",
        "",
        "ConsumerUA is linked to retail only after both retailers are transformed to the same product-type level. It remains a separate benchmark rather than a synthetic retail endpoint. This preserves a clean distinction between the observed downstream shelf price and the broader consumer-price environment against which retail adjustment is interpreted.",
        "",
        "Figure 5.6 shows the coverage of the transformed product panel, while Figures 5.7 through 5.10 describe the internal structure of the retail block. This is the point at which the chapter moves from raw data to properly constructed price objects. The order is intentional: coverage first, harmonisation second, literal product structure third, and discount behaviour only after price objects themselves have been stabilized.",
        "",
        "![Figure 5.6. Coverage of the transformed product panel by chain step.](../figures/chapter5_data/06_panel_coverage.png)",
        "",
        "Source: author's calculations based on the final daily modelling panel.",
        "",
        "![Figure 5.7. Cross-shop retail item harmonisation status.](../figures/chapter5_data/07_cross_shop_match_status.png)",
        "",
        "Source: author's calculations based on the retail match audit.",
        "",
        "![Figure 5.8. Retail literal-product mix after dairy-only reconciliation.](../figures/chapter5_data/08_retail_literal_mix.png)",
        "",
        "Source: author's calculations based on the harmonised retail literal-type mapping.",
        "",
        "![Figure 5.9. Dominant retailer-brand support by dairy product.](../figures/chapter5_data/09_dominant_brand_support.png)",
        "",
        "Source: author's calculations based on the reconciled retail brand-support table.",
        "",
        "![Figure 5.10. Regional procurement-price profile for leading ProZorro regions.](../figures/chapter5_data/10_prozorro_region_profile.png)",
        "",
        "Source: author's calculations based on ProZorro region-product aggregates.",
        "",
        "Figure 5.11 isolates the Silpo discount environment. Novus is excluded from this figure deliberately because it does not provide a comparable discount layer. The discount figure therefore serves a clean purpose: it documents how often discounts appear in Silpo and how deep they are by product. In the previous version, markdown depth was understated because the logic relied too narrowly on one value field. In the corrected rerun, median markdown depth differs by product because discount depth is reconstructed from the proper relationship between observed price, baseline price, discount amount, and discount percent.",
        "",
        "![Figure 5.11. Silpo discount environment by dairy product.](../figures/chapter5_data/11_silpo_discount_environment.png)",
        "",
        "Source: author's calculations based on Silpo discount states and corrected markdown-depth reconstruction.",
        "",
        "### 5.4 Weekly alignment, smoothing strategy, and the long-run baseline",
        "",
        "The long-run modelling layer is built on weekly medians. This step is not a presentational simplification. It is a substantive filter intended to remove volatility that does not carry a stable equilibrium meaning. In procurement, weekly aggregation is especially important because individual tenders and contract revisions can generate large local jumps that are institutional rather than market-clearing in character. In retail, smoothing is deliberately lighter because promotions are part of the mechanism being studied. Governmental series are smoothed only where needed for long-run stability checks, not because the thesis assumes that smoother data are always better data.",
        "",
        "The weekly layer is therefore a baseline rather than a replacement. Raw and smoothed versions are both kept. Long-run ARDL, ECM, and NARDL models are estimated on weekly-aligned baseline series, while the daily panel remains available for short-run local projections, discount regressions, and spread equations. This two-layer strategy is one of the main methodological improvements of the final rerun because it separates equilibrium identification from tactical high-frequency adjustment without breaking the chain logic.",
        "",
        "Figure 5.12 shows the weekly chain overlays for the core dairy products. The graph is especially informative because it now includes the farm-gate benchmark on every relevant panel. This is the correct way to keep farm-gate in the story: not by forcing raw milk to behave like a processed product price, but by showing how the raw-milk benchmark relates to the product-level producer, procurement, retail, and consumer series over time.",
        "",
        "![Figure 5.12. Weekly median chain overlays on corrected data.](../figures/chapter5_data/12_weekly_chain_overlay.png)",
        "",
        "Source: author's calculations based on weekly medians of the corrected daily panel.",
        "",
        "The farm-gate benchmark block is then shown again in Figure 5.13, now against average and median chain-level dairy indices. This step is useful because it translates the product panels into an integrated market-level picture. The resulting aggregate index should not be read as one literal observed product. It is a latent chain-level price index that captures common dairy dynamics across heterogeneous products. Its role in the thesis is robustness and system interpretation rather than replacement of the product-level evidence.",
        "",
        "![Figure 5.13. Farm-gate benchmark against chain-level dairy indices.](../figures/chapter5_data/13_farmgate_benchmark_block.png)",
        "",
        "Source: author's calculations based on farm-gate, producer, procurement, retail, and consumer aggregate indices.",
        "",
        "### 5.5 Admissibility rules, strict versus adapted intersections, and readiness for estimation",
        "",
        "The final stage of Chapter 5 is the admissibility screen. Every candidate chain link is evaluated by overlap window, continuity, unit comparability, product mapping quality, and the extent to which the series is dominated by interpolation rather than by observed variation. This matters especially for the downstream part of the chain, where procurement, retail, and consumer series overlap only late in the sample and where synthetic endpoints can otherwise look stronger than they really are.",
        "",
        f"In the strict weekly pass, the corrected data retain {strong} strong intersections and {acceptable} acceptable intersections. Because this strict layer is empty, the post-test adaptation stage is activated explicitly. The final weekly evidence therefore relies on {weak_extension} weak-but-usable extension links, while {unusable} candidate links remain unusable. This is not a failure of the research design. On the contrary, it is an important empirical finding. It shows that once the corrected monthly logic and tighter downstream construction are applied, the full chain does not support indiscriminate weekly estimation.",
        "",
        f"{level_text}",
        "",
        "Figures 5.14 and 5.15 summarize the downstream selection stage by product. They show that the preferred downstream representation is product-specific rather than universal. In some products, the broader merged retail endpoint is more stable. In others, a matched or consumer-linked endpoint performs better. This is exactly why the thesis now treats downstream endpoint choice as part of the empirical design instead of as a trivial data-management decision.",
        "",
        "![Figure 5.14. Candidate downstream retail scores by product.](../figures/chapter5_data/13_retail_level_scores.png)",
        "",
        "Source: author's calculations based on the retail-level scoring framework.",
        "",
        "![Figure 5.15. Chosen downstream retail level by product.](../figures/chapter5_data/14_optimal_retail_level.png)",
        "",
        "Source: author's calculations based on the final retail-level selection table.",
        "",
        "At this point the dataset is ready for estimation. The empirical chain is internally consistent, the distinction between price objects and behavioural objects is explicit, the long-run and short-run layers are separated, and weak intersections are labelled openly. Chapter 6 therefore begins from a properly staged design rather than from a mechanically generated model list.",
        "",
        "## Chapter 6. Estimation results, model reliability, and economic interpretation",
        "",
        "The estimation chapter follows the logic of the research question rather than the internal order of code execution. It begins with the movement of shocks from upstream to downstream, because this is the economically primary direction in vertical price transmission. It then turns to backward retailer-driven effects, which are interpreted only as bargaining or robustness evidence. After that it examines the mechanisms through which retail modifies visible pass-through, especially discounts and spread management. Only then does it return to the system level, where aggregate chain indices and VECM feasibility clarify what can and cannot be defended as a whole-chain dynamic system.",
        "",
        "### 6.1 Model strategy, diagnostics, and reliability of the retained evidence",
        "",
        "Four model families structure the final analysis. First, weekly ARDL, ECM, and NARDL models test pairwise chain links where overlap is sufficient for long-run interpretation. Second, VECM is used as a system-feasibility test rather than as an assumption that a clean full-chain system must exist. Third, daily local projections describe horizon-by-horizon short-run adjustment. Fourth, discount, spread, and procurement-scale equations explain how observed pass-through is filtered by the institutional and strategic environment.",
        "",
        f"On corrected data, the retained weekly evidence is selective. ARDL remains in the methodological screen, but it yields {len(ardl_rows)} retained weekly rows. ECM contributes {len(ecm_rows)} rows and NARDL contributes {len(nardl_rows)} rows. Product-specific VECM attempts number {vecm_attempts}, but only {len(vecm_ok)} are feasible under the final admissibility and overlap rules. This pattern already suggests that the Ukrainian dairy chain is better described as a staged and partially filtered transmission process than as a single stable equilibrium relation observed equally well at every level.",
        "",
        diag_text,
        "",
        "Figure 6.1 shows the strongest lag-correlation signals before formal estimation. Figure 6.2 then summarizes the full set of 21 directional chain links requested for the empirical design. This matrix is especially useful because it keeps screening, model reliability, and interpretation in one place, which prevents the chapter from overstating links that are economically interesting but statistically or data-wise too weak.",
        "",
        "![Figure 6.1. Top weekly lag-correlation signals across admissible links.](../figures/chapter6_results/01_weekly_corr_scan.png)",
        "",
        "Source: author's calculations based on the weekly lag-profile scan.",
        "",
        "![Figure 6.2. Summary of the 21 directional chain links.](../figures/chapter6_results/02_link21_status_matrix.png)",
        "",
        "Source: author's calculations based on the screening and retained weekly model set.",
        "",
        "{{TABLE|Table 6.1. Integrated summary of the 21 directional chain links|../outputs/chapter_tables/table_6_1_link21_summary.csv|Source: author's calculations based on the intersection screen, lag scan, and retained weekly model outputs.}}",
        "",
        "{{TABLE|Table 6.2. Reliability overview by model block|../outputs/chapter_tables/table_6_2_reliability_overview.csv|Source: author's calculations based on final model-block reliability counts.}}",
        "",
        "### 6.2 Upstream-to-downstream transmission and the pairwise weekly models",
        "",
        "The cleanest weekly evidence continues to sit in the upstream-to-procurement block. This is economically intuitive. Producer prices are sufficiently close in meaning to procurement prices for several dairy categories, and the overlap window between these two stages is longer and cleaner than the overlap between procurement and the late downstream retail layer. The corrected rerun therefore reinforces a cautious but clear conclusion: the most defensible long-run transmission relation is not the full farm-to-shelf chain but the narrower producer-to-procurement step.",
        "",
        "The lag scan shows that delayed rather than contemporaneous co-movement is the more relevant signal. This already suggests that the Ukrainian dairy chain does not operate as a frictionless weekly arbitrage system. Instead, price changes appear to travel with short institutional lags. Among the clearest retained lag signals are:",
        "",
    ]
    for r in top_corr.itertuples():
        lines.append(f"- {r.product_label}: {r.link}, correlation {fmt_num(r.best_abs_corr)}, lag {int(r.best_lag_weeks)} weeks.")
    lines.extend(
        [
            "",
            "Figure 6.3 plots the retained weekly coefficients. It shows that the corrected data do not produce one universal elasticity that can be repeated across all products and chain stages. The strongest reliable rows are selective, concentrated in the producer-to-procurement stage, and product-specific in magnitude.",
            "",
            "![Figure 6.3. Retained weekly model coefficients on corrected data.](../figures/chapter6_results/03_core_model_coefficients.png)",
            "",
            "Source: author's calculations based on retained ARDL, ECM, and NARDL results.",
            "",
            "Figures 6.4 and 6.5 deepen this weekly interpretation by separating error-correction speed from asymmetry. This matters for the market-power question. A negative and economically plausible error-correction term means that prices return toward a long-run relation after a disturbance. A nonlinear asymmetry result means that positive and negative upstream changes do not pass through in the same way.",
            "",
            "![Figure 6.4. ECM speed of adjustment across retained weekly links.](../figures/chapter6_results/04_ecm_speed_of_adjustment.png)",
            "",
            "Source: author's calculations based on retained ECM rows.",
            "",
            "![Figure 6.5. NARDL asymmetry strength across retained weekly links.](../figures/chapter6_results/05_nardl_asymmetry.png)",
            "",
            "Source: author's calculations based on retained NARDL rows.",
            "",
        ]
    )
    if not top_chain.empty:
        lines.append("The strongest reliable pairwise weekly rows are:")
        for r in top_chain.itertuples():
            coef = r.lr_coef if pd.notna(r.lr_coef) else r.sr_coef
            lines.append(f"- {r.product_label}: {r.model_family} for {r.link} with displayed coefficient {fmt_num(coef)} and reliability status {r.model_reliability}.")
        lines.append("")
    lines.extend(
        [
            "Substantively, these results imply that vertical price transmission exists, but its cleanest long-run form appears before the consumer-facing stage. Procurement behaves as an institutional filter rather than as a passive relay. In that sense, procurement is central to the thesis argument: it is the point at which producer costs are translated into institutional transactions, but not yet into observed consumer-facing shelf prices. This is exactly where buffering and delay can emerge.",
            "",
            "The downstream weekly links remain much weaker because the retail layer is both late and strategically shaped. That does not mean downstream transmission is absent. It means that downstream transmission is better interpreted through the short-run mechanism models than through a long-run weekly equilibrium relation that the data cannot support cleanly for every link.",
            "",
            "### 6.3 Downstream extensions, retail behaviour, discounts, and procurement scale",
            "",
            "Once the analysis moves beyond the producer-to-procurement core, interpretation must become more disciplined. Producer-to-retail, farm-gate-to-retail, and backward retail-driven links are informative, but mostly as extension evidence. They can reveal bargaining conditions, timing control, or common strategic adjustment, yet they should not be presented as if they carried the same equilibrium credibility as the upstream weekly core. This distinction is one of the main corrections introduced by the final rerun.",
            "",
            "Daily local projections provide the most useful bridge between the weekly equilibrium layer and the retailer mechanism layer. Because they keep daily timing, they can capture responses that are too horizon-specific or too tactical to survive in weekly cointegration-style models. Figure 6.6 reports the horizon profiles, while Figure 6.7 contrasts forward and reverse evidence shares.",
            "",
            "![Figure 6.6. Local-projection pass-through by horizon.](../figures/chapter6_results/06_lp_pass_through_horizons.png)",
            "",
            "Source: author's calculations based on daily local projections.",
            "",
            "![Figure 6.7. Forward versus reverse second-stage evidence.](../figures/chapter6_results/07_forward_reverse_core_share.png)",
            "",
            "Source: author's calculations based on 7-day and 14-day local-projection summaries.",
            "",
            "The downstream-endpoint comparison is particularly important for this thesis because it answers a substantive measurement question: what exactly is the correct downstream object when the shelf price is shaped by retailer-specific assortment and discounting, while ConsumerUA remains a smoother official benchmark? The final design does not answer this by choosing one endpoint a priori. It tests several candidates and keeps the choice explicit.",
            "",
            "![Figure 6.8. Candidate downstream levels in procurement-to-retail local-projection tests.](../figures/chapter6_results/08_candidate_downstream_core_share.png)",
            "",
            "Source: author's calculations based on candidate downstream endpoint comparisons.",
            "",
            "The resulting pattern is clear. The broader downstream object can improve overlap and descriptive continuity, but it also becomes more constructed. The stricter endpoint is cleaner but thinner. This is why the chapter continues to distinguish strict evidence from adapted evidence even after the final rerun.",
            "",
            "The spread and margin block then asks where observed price differences are stable, where they mean-revert, and where they respond asymmetrically to upstream changes. These regressions do not prove market power on their own. What they do provide is structured evidence on whether downstream adjustment behaves like a passive markup or like a managed margin. The distinction is economically meaningful in the context of Ukrainian food retail, where promotional timing and category-level management are central competitive tools.",
            "",
            f"In the corrected rerun, the spread block estimates {len(margins)} usable equations. Persistent-margin flags appear in {margin_flags} rows and asymmetric-margin flags in {asym_flags} rows. Those counts are not large enough to justify sweeping claims, but they are large enough to reject the idea that downstream prices are merely transparent reflections of upstream costs.",
            "",
            "![Figure 6.9. Average spread levels across chain segments.](../figures/chapter6_results/08_spread_levels.png)",
            "",
            "Source: author's calculations based on product-level spread summaries.",
            "",
            "![Figure 6.10. Spread volatility by product and chain segment.](../figures/chapter6_results/09_spread_volatility.png)",
            "",
            "Source: author's calculations based on the standard deviation of log spreads.",
            "",
            "![Figure 6.11. Vertical spread and market-power proxy by chain segment.](../figures/chapter6_results/10_vertical_spread_proxy.png)",
            "",
            "Source: author's calculations based on spread-regression coefficients.",
            "",
            "The discount block gives the retailer mechanism its clearest behavioural form. Here the corrected definition of `price_current`, `discount_value`, and `discount_%` matters directly. Discounts are no longer treated as noise or as an imprecise afterthought. They are a structural part of downstream price adjustment. The observed retail price is the effective post-discount price. Discount variables then explain how that observed price is managed over time.",
            "",
            f"Across the estimated discount equations, {discount_signals} of {len(discounts)} are flagged as strategic signals under the conservative rule. The strongest lagged discount coefficients appear in:",
            "",
        ]
    )
    for r in top_discount.itertuples():
        lines.append(f"- {r.product_label}: lagged discount coefficient {fmt_num(r.lag_discount_coef)}, strategy signal flag {int(r.discount_strategy_signal)}.")
    lines.extend(
        [
            "",
            "These coefficients support a specific interpretation of downstream market power. Retail influence does not appear primarily as one constant markup wedge. It appears through timing control, discount-mediated smoothing, and selective exposure of upstream pressure to the final shelf price. Some pass-through may be present in the background while the visible price is temporarily managed through markdown policy.",
            "",
            "![Figure 6.12. Discount incidence by product.](../figures/chapter6_results/11_discount_incidence.png)",
            "",
            "Source: author's calculations based on discount-incidence models.",
            "",
            "![Figure 6.13. Discount-model coefficient map.](../figures/chapter6_results/12_discount_coefficient_map.png)",
            "",
            "Source: author's calculations based on product-level discount regressions.",
            "",
            f"Procurement trade-scale variables further reinforce the interpretation that the chain is institutionally filtered. The final pipeline estimates {len(scale_models)} procurement-scale equations, of which {scale_signals} are flagged as signal-bearing. This means that procurement is shaped not only by the price per unit, but also by trade scale, contract size, and revision dynamics. The current contract sum is especially important because it can capture both bargaining power and changes in lot composition.",
            "",
        ]
    )
    for r in top_scale.itertuples():
        lines.append(f"- {r.product_label}: coefficient on weekly change in current contract sum {fmt_num(r.d_sum_current_coef)}.")
    lines.extend(
        [
            "",
            "![Figure 6.14. Procurement-scale effects on ProZorro price changes.](../figures/chapter6_results/13_procurement_scale_effects.png)",
            "",
            "Source: author's calculations based on procurement-scale model estimates.",
            "",
            "### 6.4 Aggregate dairy indices, VECM system evidence, and the limits of full-chain modelling",
            "",
            "The final rerun adds an aggregate dairy-index block as a robustness check. This block answers an important question: if product-specific idiosyncrasies are partially smoothed into a market-level latent price index, does the transmission story change materially? The answer is informative. Aggregate indices can clarify the system-level trend, but they do not erase the basic structural bottleneck of the corrected data. The full chain still has limited common support once the strict and adapted endpoint logic is respected.",
            "",
            "{{TABLE|Table 6.3. Aggregate dairy-index models|../outputs/chapter_tables/table_6_3_aggregate_index_models.csv|Source: author's calculations based on weighted aggregate dairy-chain indices.}}",
            "",
            "Figure 6.15 plots the aggregate chain indices directly, while Figure 6.16 summarizes the aggregate-index model coefficients. These results should be read as robustness evidence. They show whether the general shape of pass-through survives aggregation, not whether aggregation should replace the product-level evidence that remains central to the thesis.",
            "",
            "![Figure 6.15. Aggregate dairy-chain indices at weekly frequency.](../figures/chapter6_results/11_aggregate_index_overlay.png)",
            "",
            "Source: author's calculations based on weighted weekly chain indices.",
            "",
            "![Figure 6.16. Aggregate dairy-index model coefficients.](../figures/chapter6_results/14_aggregate_index_model_coefficients.png)",
            "",
            "Source: author's calculations based on aggregate-index ARDL, ECM, and NARDL estimates.",
            "",
            "The VECM block remains important, but mainly as a feasibility test. Product-level VECM attempts are numerous, yet the final corrected chain does not sustain a credible general system for most products because support becomes too thin once all stages are aligned honestly. The aggregate-index VECM is therefore useful not because it proves a strong whole-chain equilibrium, but because it shows how far a system interpretation can be pushed before it becomes more synthetic than empirical.",
            "",
            "{{TABLE|Table 6.4. VECM feasibility and system summary|../outputs/chapter_tables/table_6_4_vecm_summary.csv|Source: author's calculations based on product-level and aggregate-index VECM feasibility checks.}}",
            "",
            "![Figure 6.17. VECM system feasibility on corrected weekly panels.](../figures/chapter6_results/10_vecm_feasibility.png)",
            "",
            "Source: author's calculations based on product-level VECM feasibility tests.",
            "",
            "Figure 6.18 summarizes the reliability profile across all model blocks. This is an important closing figure for the empirical chapter because it makes the interpretive hierarchy visible. Product-level weekly models provide the main equilibrium evidence where they are reliable. Daily mechanism models explain the downstream behaviour that weekly equilibrium models cannot identify cleanly. Aggregate-index models and VECM provide robustness and system perspective, but not a license to overstate whole-chain certainty.",
            "",
            "![Figure 6.18. Reliability profile across model blocks.](../figures/chapter6_results/15_reliability_overview.png)",
            "",
            "Source: author's calculations based on model-block reliability counts.",
            "",
            "### 6.5 Integrated interpretation, limitations, and conclusion",
            "",
            "Taken together, the corrected rerun strengthens rather than weakens the thesis argument. The central result is not that every chain link is now strong. It is that once the data are corrected and the downstream layer is rebuilt more honestly, the economically strongest evidence remains where one would expect it in a filtered vertical chain: from producer to procurement, and then from procurement to retail through selective, timed, and strategy-dependent mechanisms rather than through one frictionless coefficient.",
            "",
            "Three conclusions follow. First, vertical price transmission in the Ukrainian dairy market is real, but it is delayed and filtered. Second, procurement acts as an institutional buffer and moderator, not merely as a neutral passage point. Third, downstream market power is best interpreted through timing control, discount-mediated smoothing, and category management rather than through a single constant markup wedge. This interpretation aligns with the institutional structure of the Ukrainian food consumer market and the retailer behaviour visible in the corrected Silpo data.",
            "",
            "The limitations remain important. The corrected monthly-to-daily logic improves internal consistency, but some weekly intersections remain thin and force the empirical design to rely on adapted downstream evidence. Farm-gate is still a raw-milk benchmark rather than a literal processed-product price. Retail-consumer is useful, but it remains a constructed endpoint. VECM feasibility is limited. None of these limitations overturn the results, but they do define the boundary of what can be claimed with confidence.",
            "",
            "For the purposes of the thesis, this is the correct concluding position. The corrected empirical system does not claim more than the data allow. It shows where transmission is strongest, where it is delayed, where procurement buffers it, and where retail strategy changes what the consumer sees. In that sense, the final rerun produces a more credible and more academically defensible Chapter 5 and Chapter 6 than the earlier internal versions of the project.",
        ]
    )

    md_path = DOC_DIR / "Chapter5_6_analysis.md"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    write_html(DOC_DIR / "Chapter5_6_analysis.html", "Chapter 5 and 6 Analysis", lines)
    write_kse_docx(DOC_DIR / "Chapter5_6_analysis.docx", lines)


def write_chapter56(
    inventory: pd.DataFrame,
    weekly: pd.DataFrame,
    source_frames: Dict[str, pd.DataFrame],
    product_dictionary: pd.DataFrame,
    intersections: pd.DataFrame,
    corr_scan: pd.DataFrame,
    chain_models: pd.DataFrame,
    vecm_results: pd.DataFrame,
    lp_summary: pd.DataFrame,
    margins: pd.DataFrame,
    margin_spreads: pd.DataFrame,
    discounts: pd.DataFrame,
    scale_models: pd.DataFrame,
    link21_summary: pd.DataFrame,
    reliability_overview: pd.DataFrame,
    index_weekly: pd.DataFrame,
    index_weights: pd.DataFrame,
    index_models: pd.DataFrame,
    index_vecm_summary: pd.DataFrame,
) -> None:
    chapter_table_dir = OUTPUT_DIR / "chapter_tables"
    chapter_table_dir.mkdir(parents=True, exist_ok=True)

    def _ordered_products(values: Sequence[str]) -> List[str]:
        seen: List[str] = []
        for product in values:
            code = harmonize_product_code(product)
            if code and code not in seen and code != "eggs":
                seen.append(code)
        return sorted(seen, key=lambda p: PRODUCT_ORDER.index(p) if p in PRODUCT_ORDER else 999)

    retail_items = source_frames.get("retail_items_full", pd.DataFrame())
    retail_match = source_frames.get("retail_match_audit", pd.DataFrame())
    brand_support = source_frames.get("retail_brand_support", pd.DataFrame())
    literal_summary = source_frames.get("retail_literal_summary", pd.DataFrame())
    producer = source_frames.get("producer", pd.DataFrame())
    farm = source_frames.get("farm_gate", pd.DataFrame())
    europe = source_frames.get("europe_benchmark", pd.DataFrame())
    cme = source_frames.get("cme_benchmark", pd.DataFrame())
    region_profile = source_frames.get("prozorro_region_profile", pd.DataFrame())

    n_products = int(product_dictionary["standardized_product"].map(harmonize_product_code).nunique()) if not product_dictionary.empty else 0
    ambiguity = int(product_dictionary.get("ambiguity_rows", pd.Series(dtype=float)).fillna(0).sum()) if not product_dictionary.empty else 0
    strong = int(intersections["admissibility_status"].eq("strong").sum()) if not intersections.empty else 0
    acceptable = int(intersections["admissibility_status"].eq("acceptable").sum()) if not intersections.empty else 0
    weak_extension = int(intersections["admissibility_status"].eq("weak_extension").sum()) if not intersections.empty else 0
    unusable = int(intersections["admissibility_status"].eq("unusable").sum()) if not intersections.empty else 0
    reliable = chain_models[chain_models["model_reliability"].eq("reliable")].copy() if not chain_models.empty else pd.DataFrame()
    retained = chain_models[chain_models["model_reliability"].isin(["reliable", "conditionally_usable"])].copy() if not chain_models.empty else pd.DataFrame()
    ardl_rows = chain_models[chain_models["model_family"].eq("ARDL")].copy() if not chain_models.empty else pd.DataFrame()
    ecm_rows = chain_models[chain_models["model_family"].eq("ECM")].copy() if not chain_models.empty else pd.DataFrame()
    nardl_rows = chain_models[chain_models["model_family"].eq("NARDL")].copy() if not chain_models.empty else pd.DataFrame()
    vecm_ok = vecm_results[vecm_results["status"].eq("ok")].copy() if not vecm_results.empty else pd.DataFrame()
    vecm_attempts = int(len(vecm_results)) if not vecm_results.empty else 0
    discount_signals = int(discounts.get("discount_strategy_signal", pd.Series(dtype=float)).fillna(0).sum()) if not discounts.empty else 0
    scale_signals = int(scale_models.get("scale_signal_flag", pd.Series(dtype=float)).fillna(0).sum()) if not scale_models.empty else 0
    margin_flags = int(margins.get("persistent_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0
    asym_flags = int(margins.get("asymmetric_margin_flag", pd.Series(dtype=float)).fillna(0).sum()) if not margins.empty else 0

    match_counts = retail_match["match_status"].value_counts() if not retail_match.empty and "match_status" in retail_match.columns else pd.Series(dtype=float)
    matched_both = int(match_counts.get("matched_both_shops", 0))
    silpo_only = int(match_counts.get("silpo_only", 0))
    novus_only = int(match_counts.get("novus_only", 0))
    strict_aligned = int(retail_match.get("strict_alignment_flag", pd.Series(dtype=float)).fillna(0).sum()) if not retail_match.empty else 0

    product_list = _ordered_products(product_dictionary["standardized_product"].tolist()) if not product_dictionary.empty else []
    product_list_text = ", ".join(PRODUCT_LABELS.get(p, p) for p in product_list)
    retail_focus_products = focus_products_from_counts(retail_items, limit=6)
    retail_focus_text = ", ".join(PRODUCT_LABELS.get(p, p) for p in retail_focus_products) if retail_focus_products else "Drinking milk / fermented milk, Butter, CHEESE"
    brand_text = top_brand_sentence(brand_support)
    literal_text = top_literal_sentence(literal_summary)

    producer_price_col = "producer_linear" if "producer_linear" in producer.columns else ("price_linear" if "price_linear" in producer.columns else None)
    farm_price_col = "farmgate_combined_linear" if "farmgate_combined_linear" in farm.columns else ("price_linear" if "price_linear" in farm.columns else None)
    eu_corr_rows: List[Tuple[str, int, float]] = []
    if not producer.empty and not europe.empty and producer_price_col is not None:
        for product in sorted(set(producer.get("product", [])).intersection(set(europe.get("product", [])))):
            left = producer[producer["product"].eq(product)][["date", producer_price_col]].rename(columns={producer_price_col: "producer_price"})
            right = europe[europe["product"].eq(product)][["date", "eu_price_uah"]]
            merged = left.merge(right, on="date", how="inner").dropna()
            if len(merged) >= 30:
                eu_corr_rows.append((PRODUCT_LABELS.get(product, product), int(len(merged)), float(merged["producer_price"].corr(merged["eu_price_uah"]))))
        eu_corr_rows = sorted(eu_corr_rows, key=lambda x: abs(x[2]), reverse=True)

    cme_corr = np.nan
    if not farm.empty and not cme.empty and farm_price_col is not None:
        merged = farm[["date", farm_price_col]].rename(columns={farm_price_col: "farm_price"}).merge(cme[["date", "cme_class3_uah"]], on="date", how="inner").dropna()
        if len(merged) >= 30:
            cme_corr = float(merged["farm_price"].corr(merged["cme_class3_uah"]))

    weight_text = "The aggregate dairy index uses fixed structural proxy weights instead of endogenous pass-through weights."
    if not index_weights.empty:
        top_weights = index_weights.sort_values("weight", ascending=False).head(5)
        weight_text = (
            "The aggregate dairy index uses a fixed-weight geometric rule with weights based on procurement participation, retail assortment support, "
            "and an equal-weight anchor. The largest final weights belong to "
            + "; ".join(f"{row.product_label} ({fmt_pct(row.weight)})" for row in top_weights.itertuples())
            + "."
        )

    if not retained.empty:
        lb_ok = int(retained["ljungbox_p"].ge(0.05).sum())
        bp_ok = int(retained["bp_p"].ge(0.05).sum())
        white_ok = int(retained["white_p"].ge(0.05).sum())
        jb_ok = int(retained["jb_p"].ge(0.05).sum())
        diag_text = (
            f"Across the retained weekly equations, {lb_ok} satisfy the Ljung-Box threshold, {bp_ok} satisfy the Breusch-Pagan threshold, "
            f"{white_ok} satisfy the White threshold, and {jb_ok} satisfy the Jarque-Bera threshold. These counts help order the evidence rather than replace model-by-model judgement."
        )
    else:
        diag_text = "The corrected overlap remains selective, which is why the retained weekly interpretation is intentionally narrow."

    strongest_weekly = "No retained weekly equation met the reporting threshold."
    if not reliable.empty:
        best = reliable.assign(abs_coef=lambda d: d["lr_coef"].where(d["lr_coef"].notna(), d["sr_coef"]).abs()).sort_values("abs_coef", ascending=False).iloc[0]
        strongest_weekly = (
            f"The strongest retained weekly equation is {best.product_label}: {best.link} in the {best.data_variant} specification, "
            f"estimated with {best.model_family}, reported coefficient {fmt_num(best.lr_coef if pd.notna(best.lr_coef) else best.sr_coef)}, "
            f"and error-correction term {fmt_num(best.ect_coef)}."
        )

    weekly_product_sentences: List[str] = []
    if not retained.empty:
        rel_rank = {"reliable": 2, "conditionally_usable": 1}
        for product_code in ["milk", "cheese", "sour_cream", "butter"]:
            subset = retained[retained["product"].eq(product_code)].copy()
            if subset.empty:
                continue
            subset["rel_rank"] = subset["model_reliability"].map(rel_rank).fillna(0)
            subset["coef_for_rank"] = subset["lr_coef"].where(subset["lr_coef"].notna(), subset["sr_coef"]).abs()
            best_row = subset.sort_values(["rel_rank", "coef_for_rank", "n_obs"], ascending=[False, False, False]).iloc[0]
            coef = best_row["lr_coef"] if pd.notna(best_row["lr_coef"]) else best_row["sr_coef"]
            detail = (
                f"For {best_row.product_label.lower()}, the best retained weekly relation is {best_row.model_family} in {best_row.data_variant}, "
                f"with coefficient {fmt_num(coef)} and error-correction term {fmt_num(best_row.ect_coef)}."
            )
            if pd.notna(best_row.get("cointegration_p", np.nan)):
                detail += f" The cointegration p-value equals {fmt_num(best_row.cointegration_p)}."
            if pd.notna(best_row.get("asymmetry_pvalue", np.nan)):
                detail += f" The asymmetry p-value equals {fmt_num(best_row.asymmetry_pvalue)}."
            weekly_product_sentences.append(detail)
    weekly_product_text = " ".join(weekly_product_sentences)

    farmgate_detail_text = "The farm-gate block is treated as the raw-milk benchmark, so its strongest reading is about directional pressure rather than literal processed-product equivalence."
    farmgate_corr = corr_scan[corr_scan["link"].astype(str).str.contains("FarmGate")] if not corr_scan.empty else pd.DataFrame()
    farmgate_retained = retained[retained["link"].astype(str).str.contains("FarmGate")] if not retained.empty else pd.DataFrame()
    if not farmgate_corr.empty:
        corr_parts = []
        top_corr = farmgate_corr.sort_values("best_abs_corr", ascending=False).head(3)
        for row in top_corr.itertuples():
            corr_parts.append(
                f"{row.product_label} in {row.link} shows best weekly lag correlation {fmt_num(row.best_abs_corr)} at lag {int(row.best_lag_weeks)}"
            )
        farmgate_detail_text = "The lag profile confirms that farm-gate pressure is visible first where product mapping remains closest to raw-milk content. " + "; ".join(corr_parts) + "."
    if not farmgate_retained.empty:
        fg_parts = []
        unique_fg = (
            farmgate_retained.sort_values(["model_reliability", "n_obs"], ascending=[False, False])
            .drop_duplicates(subset=["product_label", "link"])
            .head(3)
        )
        for row in unique_fg.itertuples():
            coef = row.lr_coef if pd.notna(row.lr_coef) else row.sr_coef
            fg_parts.append(
                f"{row.product_label} in {row.link} retains {row.model_family} with coefficient {fmt_num(coef)} and error-correction term {fmt_num(row.ect_coef)}"
            )
        farmgate_detail_text += " In the retained weekly equations, " + "; ".join(fg_parts) + "."

    top_region_text = "Regional procurement heterogeneity remains visible but secondary in the core chain."
    if not region_profile.empty:
        top_regions = region_profile.groupby("region", as_index=False)["n_obs"].sum().sort_values("n_obs", ascending=False).head(5)
        top_region_text = "The most active ProZorro regions in the dairy sample are " + "; ".join(
            f"{row.region} ({int(row.n_obs)} observations)" for row in top_regions.itertuples()
        ) + "."

    benchmark_sentence = (
        f"The correlation between the raw-milk benchmark and CME Class III equals {fmt_num(cme_corr)} over the overlapping sample."
        if pd.notna(cme_corr)
        else "CME Class III is retained as a contextual raw-milk benchmark."
    )
    europe_sentence = (
        "The strongest producer-to-Europe co-movement appears in "
        + "; ".join(f"{label} ({fmt_num(corr)}, {n} aligned observations)" for label, n, corr in eu_corr_rows[:3])
        + "."
        if eu_corr_rows
        else "European benchmarks are retained as context for the corrected domestic series."
    )

    discount_detail_text = "Discount evidence is retained mainly for Silpo, where markdown information is directly observed."
    if not discounts.empty:
        discount_focus_texts = []
        for row in (
            discounts.sort_values(["discount_strategy_signal", "lag_discount_p", "r2"], ascending=[False, True, False])
            .head(2)
            .itertuples()
        ):
            sentence = (
                f"{row.product_label} shows a discount-persistence coefficient of {fmt_num(row.lag_discount_coef)} "
                f"(p = {fmt_num(row.lag_discount_p)})"
            )
            if pd.notna(getattr(row, "retail_vs_consumer_gap_coef", np.nan)):
                sentence += (
                    f", while the retail-versus-consumer gap coefficient equals "
                    f"{fmt_num(getattr(row, 'retail_vs_consumer_gap_coef'))} "
                    f"(p = {fmt_num(getattr(row, 'retail_vs_consumer_gap_p', np.nan))})"
                )
            if pd.notna(getattr(row, "silpo_novus_gap_coef", np.nan)):
                sentence += (
                    f", and the Silpo-versus-Novus gap coefficient equals "
                    f"{fmt_num(getattr(row, 'silpo_novus_gap_coef'))} "
                    f"(p = {fmt_num(getattr(row, 'silpo_novus_gap_p', np.nan))})"
                )
            discount_focus_texts.append(sentence + ".")
        if discount_focus_texts:
            discount_detail_text = " ".join(discount_focus_texts)

    scale_detail_text = "Procurement-scale effects remain product-specific rather than uniform across the chain."
    if not scale_models.empty:
        scale_focus_texts = []
        scale_focus = scale_models.sort_values(["scale_signal_flag", "d_sum_current_p", "d_expected_p"], ascending=[False, True, True]).head(3)
        for row in scale_focus.itertuples():
            pieces = []
            if pd.notna(row.d_expected_coef):
                pieces.append(f"expected value {fmt_num(row.d_expected_coef)} (p = {fmt_num(row.d_expected_p)})")
            if pd.notna(row.d_sum_initial_coef):
                pieces.append(f"initial contract sum {fmt_num(row.d_sum_initial_coef)} (p = {fmt_num(row.d_sum_initial_p)})")
            if pd.notna(row.d_sum_current_coef):
                pieces.append(f"current contract sum {fmt_num(row.d_sum_current_coef)} (p = {fmt_num(row.d_sum_current_p)})")
            if pieces:
                scale_focus_texts.append(f"For {row.product_label.lower()}, the strongest procurement-scale terms are " + "; ".join(pieces) + ".")
        if scale_focus_texts:
            scale_detail_text = " ".join(scale_focus_texts)

    index_detail_text = "The aggregate-index block is informative mainly as a robustness check."
    if not index_models.empty:
        idx_focus = index_models.assign(display_coef=lambda d: d["lr_coef"].where(d["lr_coef"].notna(), d["sr_coef"])).sort_values(
            ["model_reliability", "data_variant", "model_family"], ascending=[False, True, True]
        ).head(2)
        idx_sentences = []
        for row in idx_focus.itertuples():
            idx_sentences.append(
                f"{row.model_family} on {row.data_variant} data reports coefficient {fmt_num(row.display_coef)} with error-correction term {fmt_num(row.ect_coef)}"
                + (
                    f" and cointegration p-value {fmt_num(row.cointegration_p)}."
                    if pd.notna(row.cointegration_p)
                    else "."
                )
            )
        if idx_sentences:
            index_detail_text = " ".join(idx_sentences)
    if not index_vecm_summary.empty:
        vecm_ok_rows = index_vecm_summary[index_vecm_summary["status"].eq("ok")].copy()
        if not vecm_ok_rows.empty:
            ok_sentence = "; ".join(
                f"{row.system_name} is feasible with {int(row.n_obs)} weekly observations and rank {int(row.vecm_rank)}"
                for row in vecm_ok_rows.itertuples()
                if pd.notna(getattr(row, "vecm_rank", np.nan))
            )
            if ok_sentence:
                index_detail_text += " The widened system block improves feasibility: " + ok_sentence + "."

    reliability_table = pd.DataFrame()
    if not reliability_overview.empty:
        reliability_table = reliability_overview.rename(
            columns={
                "block": "model_block",
                "family": "family",
                "rows": "rows_estimated",
                "reliable_rows": "reliable_rows",
                "conditionally_usable_rows": "conditionally_usable_rows",
            }
        )
        reliability_table.to_csv(chapter_table_dir / "table_6_1_model_reliability.csv", index=False)

    pairwise_table = pd.DataFrame()
    if not link21_summary.empty:
        detail_source = chain_models.copy() if not chain_models.empty else pd.DataFrame()
        best_details = pd.DataFrame()
        if not detail_source.empty:
            detail_source["rel_rank"] = detail_source["model_reliability"].map({"reliable": 2, "conditionally_usable": 1}).fillna(0)
            detail_source["coef_rank"] = detail_source["lr_coef"].where(detail_source["lr_coef"].notna(), detail_source["sr_coef"]).abs()
            best_details = (
                detail_source.sort_values(["link", "rel_rank", "coef_rank", "n_obs"], ascending=[True, False, False, False])
                .groupby("link", as_index=False)
                .head(1)[["link", "product_label", "n_obs", "ect_pvalue", "cointegration_p", "asymmetry_pvalue"]]
                .rename(columns={"product_label": "retained_product"})
            )
        pairwise_table = link21_summary[
            [
                "link",
                "best_admissibility",
                "median_overlap_weeks",
                "best_corr",
                "best_lag_weeks",
                "best_model_product",
                "best_model_family",
                "display_coef",
                "ect_coef",
                "model_reliability",
                "reading",
            ]
        ].copy()
        pairwise_table = pairwise_table.rename(
            columns={
                "best_admissibility": "admissibility",
                "median_overlap_weeks": "overlap_weeks",
                "best_corr": "best_lag_corr",
                "best_model_product": "retained_product",
                "best_lag_weeks": "best_lag_weeks",
                "best_model_family": "retained_model",
                "display_coef": "reported_coef",
                "ect_coef": "ect",
                "model_reliability": "reliability",
            }
        )
    if not best_details.empty:
        pairwise_table = pairwise_table.drop(columns=["retained_product"], errors="ignore").merge(best_details, on="link", how="left")
    pairwise_table.to_csv(chapter_table_dir / "table_6_2_pairwise_weekly_models.csv", index=False)

    key_results_rows: List[Dict[str, object]] = []
    if not reliable.empty:
        weekly_focus = reliable.assign(display_coef=lambda d: d["lr_coef"].where(d["lr_coef"].notna(), d["sr_coef"]))
        weekly_focus = weekly_focus.sort_values(["link", "n_obs"], ascending=[True, False])
        for link in ["FarmGate -> Producer", "FarmGate -> Procurement", "Producer -> Procurement"]:
            sub = weekly_focus[weekly_focus["link"].eq(link)].head(1)
            if sub.empty:
                continue
            row = sub.iloc[0]
            key_results_rows.append(
                {
                    "evidence_block": "Weekly core link",
                    "scope": link,
                    "estimate": f"{row['model_family']} {fmt_num(row['display_coef'])}; ECT {fmt_num(row['ect_coef'])}",
                    "reading": "Retained weekly evidence for the main chain.",
                }
            )
    if not discounts.empty:
        top_discount = discounts.sort_values(["discount_strategy_signal", "lag_discount_p"], ascending=[False, True]).head(1)
        if not top_discount.empty:
            row = top_discount.iloc[0]
            key_results_rows.append(
                {
                    "evidence_block": "Silpo discount strategy",
                    "scope": row["product_label"],
                    "estimate": f"lag discount {fmt_num(row['lag_discount_coef'])}; p {fmt_num(row['lag_discount_p'])}",
                    "reading": "Observed shelf adjustment is partly managed through markdown policy.",
                }
            )
    if not scale_models.empty:
        top_scale = scale_models.sort_values(["scale_signal_flag", "d_sum_current_p"], ascending=[False, True]).head(1)
        if not top_scale.empty:
            row = top_scale.iloc[0]
            key_results_rows.append(
                {
                    "evidence_block": "Procurement scale",
                    "scope": row["product_label"],
                    "estimate": f"current contract sum {fmt_num(row['d_sum_current_coef'])}; p {fmt_num(row['d_sum_current_p'])}",
                    "reading": "Procurement intensity modifies observed price adjustment.",
                }
            )
    if not index_models.empty:
        idx_focus = index_models[index_models["link"].astype(str).isin(["Index FarmGate -> Producer", "Index FarmGate -> Downstream extension", "Index Producer -> Procurement"])].copy()
        idx_focus["display_coef"] = idx_focus["lr_coef"].where(idx_focus["lr_coef"].notna(), idx_focus["sr_coef"])
        for _, row in idx_focus.sort_values(["model_reliability", "n_obs"], ascending=[False, False]).drop_duplicates(subset=["link"]).iterrows():
            key_results_rows.append(
                {
                    "evidence_block": "Aggregate index robustness",
                    "scope": row["link"],
                    "estimate": f"{row['model_family']} {fmt_num(row['display_coef'])}; ECT {fmt_num(row['ect_coef'])}",
                    "reading": "The staged transmission pattern survives aggregation.",
                }
            )
    if not index_vecm_summary.empty:
        for _, row in index_vecm_summary.iterrows():
            if str(row.get("status", "")) != "ok":
                continue
            key_results_rows.append(
                {
                    "evidence_block": "Aggregate system VECM",
                    "scope": row.get("system_name", "aggregate_chain"),
                    "estimate": f"rank {int(row['vecm_rank'])}; n {int(row['n_obs'])}",
                    "reading": "The widened downstream extension makes a compact system model feasible as robustness.",
                }
            )
    key_results_table = pd.DataFrame(key_results_rows)
    if not key_results_table.empty:
        key_results_table.to_csv(chapter_table_dir / "table_6_1_key_results_synthesis.csv", index=False)

    daily_rows: List[Dict[str, object]] = []
    if not lp_summary.empty:
        lp_focus = (
            lp_summary[lp_summary["horizon_days"].isin([7, 14])]
            .sort_values(["core_share", "sig_share", "mean_abs_coef"], ascending=[False, False, False])
            .drop_duplicates(subset=["link", "horizon_days"])
            .head(5)
        )
        for row in lp_focus.itertuples():
            daily_rows.append(
                {
                    "block": "Local projections",
                    "product_or_scope": "Multi-product",
                    "equation": f"{row.link}, h={int(row.horizon_days)}",
                    "estimate": fmt_num(row.median_coef),
                    "significance_or_support": f"core {fmt_pct(row.core_share)}, sig {fmt_pct(row.sig_share)}",
                    "reading": "Daily timing evidence for short-run pass-through.",
                }
            )
    if not margins.empty:
        margin_focus = margins.assign(abs_lag=lambda d: d["lag_spread_coef"].abs()).sort_values(["persistent_margin_flag", "asymmetric_margin_flag", "abs_lag"], ascending=[False, False, False]).head(4)
        for row in margin_focus.itertuples():
            daily_rows.append(
                {
                    "block": "Spread models",
                    "product_or_scope": row.product_label,
                    "equation": row.stage,
                    "estimate": fmt_num(row.lag_spread_coef),
                    "significance_or_support": f"p={fmt_num(row.lag_spread_p)}; asym p={fmt_num(row.asymmetry_pvalue)}",
                    "reading": "Tests whether vertical spreads behave like managed margins.",
                }
            )
    if not discounts.empty:
        discount_focus = discounts.assign(abs_disc=lambda d: d["lag_discount_coef"].abs()).sort_values(["discount_strategy_signal", "abs_disc"], ascending=[False, False]).head(4)
        gap_coef_col = "retail_vs_consumer_gap_coef" if "retail_vs_consumer_gap_coef" in discount_focus.columns else ("retail_consumer_gap_coef" if "retail_consumer_gap_coef" in discount_focus.columns else "")
        for row in discount_focus.itertuples():
            gap_text = ""
            if gap_coef_col:
                gap_text = f"; gap {fmt_num(getattr(row, gap_coef_col, np.nan))}"
            daily_rows.append(
                {
                    "block": "Silpo discount models",
                    "product_or_scope": row.product_label,
                    "equation": "Discount incidence and depth",
                    "estimate": fmt_num(row.lag_discount_coef),
                    "significance_or_support": f"p={fmt_num(row.lag_discount_p)}{gap_text}",
                    "reading": "Shows whether markdown policy smooths or exposes pass-through.",
                }
            )
    if not scale_models.empty:
        scale_focus = scale_models.assign(abs_current=lambda d: d["d_sum_current_coef"].abs()).sort_values(["scale_signal_flag", "abs_current"], ascending=[False, False]).head(4)
        for row in scale_focus.itertuples():
            daily_rows.append(
                {
                    "block": "Procurement scale models",
                    "product_or_scope": row.product_label,
                    "equation": "Change in current contract sum",
                    "estimate": fmt_num(row.d_sum_current_coef),
                    "significance_or_support": f"p={fmt_num(row.d_sum_current_p)}",
                    "reading": "Captures whether contract scale modifies price adjustment.",
                }
            )
    daily_table = pd.DataFrame(daily_rows)
    if not daily_table.empty:
        daily_table = daily_table.drop_duplicates(subset=["block", "product_or_scope", "equation", "estimate"]).reset_index(drop=True)
        daily_table.to_csv(chapter_table_dir / "table_6_3_daily_mechanism_models.csv", index=False)

    system_rows: List[Dict[str, object]] = []
    if not index_models.empty:
        index_focus = index_models.assign(display_coef=lambda d: d["lr_coef"].where(d["lr_coef"].notna(), d["sr_coef"])).sort_values(["model_reliability", "link", "data_variant"], ascending=[False, True, True]).head(6)
        for row in index_focus.itertuples():
            system_rows.append(
                {
                    "block": "Aggregate index models",
                    "system_or_link": row.link,
                    "family": row.model_family,
                    "variant": row.data_variant,
                    "estimate": fmt_num(row.display_coef),
                    "status": row.model_reliability,
                    "diagnostic_note": f"ECT={fmt_num(row.ect_coef)}; cointegration p={fmt_num(row.cointegration_p)}",
                }
            )
    system_rows.append(
        {
            "block": "Product-level VECM",
            "system_or_link": "Product systems",
            "family": "VECM",
            "variant": "weekly",
            "estimate": f"{len(vecm_ok)} feasible / {vecm_attempts} attempted",
            "status": "screening outcome",
            "diagnostic_note": "No product-level system is treated as the main identification layer unless overlap, rank, and diagnostics are jointly admissible.",
        }
    )
    if not index_vecm_summary.empty:
        for row in index_vecm_summary.itertuples():
            system_rows.append(
                {
                    "block": "Aggregate index VECM",
                    "system_or_link": getattr(row, "system_name", "aggregate_chain"),
                    "family": "VECM",
                    "variant": "weekly",
                    "estimate": getattr(row, "status", ""),
                    "status": getattr(row, "reason", ""),
                    "diagnostic_note": f"n={getattr(row, 'n_obs', 'n/a')}",
                }
            )
    system_table = pd.DataFrame(system_rows)
    if not system_table.empty:
        system_table.to_csv(chapter_table_dir / "table_6_4_system_models.csv", index=False)

    synthesis_table = pd.DataFrame(
        [
            {
                "evidence_block": "Product audit and matching",
                "key_result": f"{n_products} unified domestic categories, with CHEESE collapsed and Other retained.",
                "interpretation": "The domestic product language is now consistent enough for cross-dataset modelling.",
                "limitation_or_next_step": "EU benchmark groupings remain unchanged and are used comparatively only.",
            },
            {
                "evidence_block": "Weekly pairwise core",
                "key_result": "Producer -> Procurement remains the strongest defensible long-run segment.",
                "interpretation": "Pass-through is clearest before the chain reaches retailer strategy.",
                "limitation_or_next_step": "Several downstream weekly links remain weak because overlap begins late.",
            },
            {
                "evidence_block": "Daily short-run timing",
                "key_result": "Local projections retain daily timing signals that weekly models cannot capture cleanly.",
                "interpretation": "Short-run transmission is present but horizon-specific.",
                "limitation_or_next_step": "Daily coefficients should be interpreted as mechanism evidence rather than equilibrium evidence.",
            },
            {
                "evidence_block": "Retail behaviour",
                "key_result": f"{discount_signals} discount-strategy signals and {margin_flags} persistent-margin signals are retained.",
                "interpretation": "Retail prices are managed through markdowns and spread policy, not only through passive markup.",
                "limitation_or_next_step": "This evidence is strongest for Silpo because markdown variables are explicit there.",
            },
            {
                "evidence_block": "Procurement scale",
                "key_result": f"{scale_signals} procurement-scale signals are retained.",
                "interpretation": "Procurement works as an institutional filter, not only as a price relay.",
                "limitation_or_next_step": "Scale variables should remain separate from price variables in interpretation.",
            },
            {
                "evidence_block": "System robustness",
                "key_result": f"Aggregate-index models are feasible, while product-level VECM feasibility is {len(vecm_ok)}/{vecm_attempts}.",
                "interpretation": "System evidence is informative as robustness but does not replace the staged interpretation.",
                "limitation_or_next_step": "The thesis should keep a staged-chain reading as the main conclusion.",
            },
        ]
    )
    synthesis_table.to_csv(chapter_table_dir / "table_6_5_integrated_synthesis.csv", index=False)

    lines: List[str] = [
        "# Chapter 5 and Chapter 6. Empirical analysis of vertical price transmission in Ukraine's dairy market",
        "",
        "## Chapter 5. Data and empirical design",
        "",
        "This chapter explains how the empirical system is constructed after correcting the governmental price logic and after tightening the downstream retail definition. The purpose is not only to describe the datasets, but also to show why some data combinations are valid for the main chain and why others can be used only as supporting evidence. This order follows the logic of the thesis as a whole: the question is not simply whether prices co-move, but where in the Ukrainian dairy chain shocks are translated, filtered, delayed, or strategically managed.",
        "",
        "### 5.1 Final data sources description",
        "",
        "The empirical chain is defined at the territorial level of Ukraine for the core governmental stages. Farm-gate, producer, and consumer prices are national averages in the corrected official series, so the main chain must preserve the same territorial meaning. By contrast, ProZorro is retained as an all-Ukraine transactional layer because procurement prices are observed through realized tenders and contracts rather than through national averaging. Retail is modeled through Silpo, Novus, and a combined retail category series built from harmonized product-level observations.",
        "",
        f"The final inventory contains {len(inventory):,} source or derivative blocks. The harmonized retail item universe contains {len(retail_items):,} retailer-date observations. The standardized product dictionary compresses raw product definitions into {n_products} common domestic categories: {product_list_text}. CHEESE is intentionally treated as one domestic category across producer, consumer, ProZorro, Silpo, and Novus so that the main tables compare like with like rather than several partially overlapping cheese subgroups under different names. Other is kept as an explicit residual dairy category instead of being dropped silently.",
        "",
        f"The product-definition audit flags {ambiguity} ambiguous or approximate mappings. This matters because the model does not assume that products are comparable merely because names look similar. Governmental datasets define products directly through official row labels, farm-gate through the agricultural product label, ProZorro through procurement profiles and tender titles, and retail through product names, item titles, normalized brands, fat-content cues, and package descriptions. The audit therefore belongs to the identification strategy itself.",
        "",
        "Farm-gate remains the raw-milk benchmark. This is economically necessary because farm-gate captures the origin of milk-supply pressure, but it is not a literal processed-product price. For that reason, the thesis treats farm-gate as the upstream raw-milk anchor and compares it both to product-level series and to chain-level dairy indices rather than forcing it into a false one-to-one equivalence with processed goods.",
        "",
        "Figure 5.1 opens the empirical chapter with the corrected governmental layers. Figure 5.2 then shows the observed retail series at product level. Here the retail price object is the observed package price based on `price_current`, not a unit-price transformation. This is an important correction because the consumer-facing shelf price is the relevant downstream object for transmission and discount analysis.",
        "",
        "![Figure 5.1. Raw corrected governmental series used in the core chain.](../figures/chapter5_data/01_raw_government_layers.png)",
        "",
        "Source: author's calculations based on the corrected FarmGateUA, ProducerUA, and ConsumerUA layers from `full_uah_final.xlsx`.",
        "",
        "![Figure 5.2. Raw retail observed-price series before transformation.](../figures/chapter5_data/02_raw_retail_observed_series.png)",
        "",
        "Source: author's calculations based on harmonized Silpo and Novus item data. The figure focuses on the main categories and the most frequent additional retail categories in the sample: " + retail_focus_text + ".",
        "",
        "### 5.2 Reconstruction logic, external benchmarks, and transformation choices",
        "",
        "The corrected monthly official series remain the base truth. Their widening to daily frequency is inherited from the corrected source workbook, but the thesis does not treat daily interpolation as if it were direct market observation. Instead, the design is explicitly two-layered. Weekly medians provide the opening long-run layer for equilibrium-style modelling. Daily data provide the main layer for short-run timing, procurement discreteness, and discount-driven retail behaviour. This separation is a modelling choice with economic meaning, not a presentation trick.",
        "",
        "{{FORMULA|p_week(w) = median{ p_d : d belongs to week w }}}",
        "",
        "Weekly medians are used because they suppress local procurement jumps and transient retail noise without erasing the broader price path. This is particularly important for ProZorro, where tender timing and contract revisions can create sharp day-specific moves that are institutional rather than market-clearing in character. Retail smoothing is lighter, because promotions and markdowns are part of the mechanism under study and should not be over-smoothed away.",
        "",
        "The same distinction also disciplines interpretation. A weekly coefficient is read as evidence about the baseline transmission path after high-frequency noise is compressed. A daily coefficient is read as evidence about timing, tactical adjustment, or temporary buffering. This difference matters for a consumer-market thesis: households observe shelf prices daily, but strategic pass-through is easier to identify once the high-frequency price environment has first been stabilized.",
        "",
        "European dairy prices and CME Class III are retained as external benchmarks rather than as domestic structural stages. Their function is to anchor the corrected Ukrainian series within a broader dairy cycle, to make the reconstruction logic explicit, and to show that the corrected domestic paths remain economically coherent after the former aggregation error is removed. " + europe_sentence,
        "",
        benchmark_sentence,
        "",
        "{{FORMULA|ln(P_t^agg) = SUM_i w_i * ln(P_it),     SUM_i w_i = 1}}",
        "",
        weight_text,
        "",
        "Figures 5.3 to 5.5 follow the descriptive logic used in strong empirical theses: first external context, then transformed product-level domestic series, and only after that the aggregate market-level robustness layer. Each figure therefore answers a different question in sequence: where the broader dairy cycle comes from, what the corrected domestic product paths look like, and how those product paths translate into a latent chain-level dairy price index.",
        "",
        "![Figure 5.3. External benchmark block for European dairy prices and CME Class III.](../figures/chapter5_data/03_raw_external_benchmarks.png)",
        "",
        "Source: author's calculations based on European benchmark series and CME Class III in UAH-equivalent terms.",
        "",
        "![Figure 5.4. Product-level price paths and aggregate dairy-chain indices.](../figures/chapter5_data/04_dataset_product_lines_and_indices.png)",
        "",
        "Source: author's calculations based on producer, procurement, retail combined, and ConsumerUA product-level series.",
        "",
        "![Figure 5.5. Aggregate dairy price indices by chain level.](../figures/chapter5_data/05_aggregate_chain_indices.png)",
        "",
        "Source: author's calculations based on fixed-weight geometric dairy indices.",
        "",
        "### 5.3 Retail harmonisation, discount logic, and downstream endpoint construction",
        "",
        "The downstream block is reconstructed from the product level upward. Each Silpo and Novus SKU is standardized by product type, retailer, brand, and item identity. Package size, fat content, and title-level wording are used to make the mapping more precise, but the final price object remains the observed package price. The models do not rely on `unit_price` for Silpo or Novus as the main downstream price. Unit-based fields are retained only as support metadata, while the observed retail series used for modelling are based on `price_current` and on the reconstructed non-promotional baseline built from the markdown information.",
        "",
        "{{FORMULA|p_observed = price_current     ;     p_baseline = price_current + discount_value, if discount_value > 0}}",
        "",
        "{{FORMULA|discount_dummy = 1(discount active), 0(otherwise)     ;     markdown_rate = (p_baseline - p_observed) / p_baseline}}",
        "",
        f"After harmonisation, the audit records {matched_both} item keys matched across both shops, {silpo_only} Silpo-only keys, and {novus_only} Novus-only keys. Only {strict_aligned} matched keys survive the strictest one-to-one alignment rule. This is why the final downstream category series are built from stable product-type medians rather than from a naive pooling of all available SKUs. The aim is to measure price movement, not assortment replacement.",
        "",
        "In practical terms, the retail layer therefore contains two linked price objects. The first is the observed consumer-facing package price, which is the right object for pass-through and consumer-visibility analysis. The second is the reconstructed baseline price, which approximates the non-promotional shelf level and is more suitable for long-run comparison. This distinction makes it possible to separate price transmission from promotional masking instead of conflating both in one downstream series.",
        "",
        brand_text,
        "",
        literal_text,
        "",
        top_region_text,
        "",
        "Figures 5.6 to 5.10 summarize the downstream structure that the models actually use. The product-distribution and brand-distribution figures show where retailer data are dense enough to support category medians. The procurement-region figure clarifies where institutional price information is concentrated. The Silpo discount figure then isolates the behavioural layer that later matters in Chapter 6. Novus is not included in that discount-environment figure because it does not provide a comparable explicit markdown structure.",
        "",
        "![Figure 5.6. Retail product distribution by retailer.](../figures/chapter5_data/07_retail_product_distribution.png)",
        "",
        "Source: author's calculations based on harmonized Silpo and Novus item keys.",
        "",
        "![Figure 5.7. Retail brand distribution by retailer.](../figures/chapter5_data/08_retail_brand_distribution.png)",
        "",
        "Source: author's calculations based on normalized retailer-brand support.",
        "",
        "![Figure 5.8. Regional procurement-price profile for leading ProZorro regions.](../figures/chapter5_data/10_prozorro_region_profile.png)",
        "",
        "Source: author's calculations based on ProZorro unit-price observations across regions.",
        "",
        "![Figure 5.9. Silpo discount environment by dairy product.](../figures/chapter5_data/11_silpo_discount_environment.png)",
        "",
        "Source: author's calculations based on Silpo discount incidence and positive markdown depth only.",
        "",
        "![Figure 5.10. Weekly chain overlays on the transformed data.](../figures/chapter5_data/12_weekly_chain_overlay.png)",
        "",
        "Source: author's calculations based on weekly medians for farm-gate, producer, procurement, retail combined, and ConsumerUA.",
        "",
        "The farm-gate benchmark is then separated into a dedicated comparison block. This is important because raw milk is not a literal processed-product price, but it is the upstream pressure point that later feeds into processors, procurement, and finally the consumer market. The benchmark figures therefore help distinguish a true upstream signal from a misleading same-product comparison.",
        "",
        farmgate_detail_text,
        "",
        "![Figure 5.11. Farm-gate benchmark block against chain-level dairy price indices.](../figures/chapter5_data/13_farmgate_benchmark_block.png)",
        "",
        "Source: author's calculations based on weekly chain-level medians and averages.",
        "",
        "![Figure 5.12. Farm-gate-to-retail comparison on a common normalized scale.](../figures/chapter5_data/15_farmgate_to_chain_normalized.png)",
        "",
        "Source: author's calculations based on normalized weekly product-level paths, with the first common observation scaled to 100.",
        "",
        "### 5.4 Admissibility rules, dataset intersections, and readiness for estimation",
        "",
        "Before estimation, every candidate link is screened by overlap length, continuity, mapping quality, and support from non-interpolated observations. This is crucial because the corrected data are cleaner, but the valid intersection is not uniformly wide across all chain stages. The model does not treat every mechanically aligned pair as equally admissible.",
        "",
        f"The formal screening yields {strong} strong intersections, {acceptable} acceptable intersections, {weak_extension} weak-but-usable extension intersections, and {unusable} unusable links. The scarcity of strong intersections is itself informative. It shows that honest product matching and honest overlap rules impose discipline on the empirical system. In the corrected data, the main chain can be estimated, but not every potential link can be interpreted with the same degree of confidence.",
        "",
        "This screen also clarifies the role of weekly versus daily data. Weekly medians are the first layer because they are the right starting point for long-run and error-correction reasoning. Daily data then become the main layer when the research turns to short-run price reaction, procurement discreteness, discount management, and retailer-specific timing. In that sense, the chapter ends where the estimation chapter begins: with a structured statement of what is estimable, what is only supportive, and why.",
        "",
        "## Chapter 6. Estimation results and interpretation",
        "",
        "The estimation chapter is organized in the same economic order as the chain itself. It begins upstream, where the cleanest long-run evidence should appear if vertical transmission exists. It then moves downstream, where transmission becomes more strategically filtered. Only after that does it turn to backward robustness, retail discount mechanisms, procurement-scale effects, and system-level robustness. This sequencing matters because the thesis is not trying to maximize the number of estimated coefficients. It is trying to show where shocks travel credibly and where they are absorbed or strategically managed.",
        "",
        "### 6.1 Model strategy, diagnostics, and reliability of the retained evidence",
        "",
        "The final design uses weekly ARDL, ECM, and NARDL specifications as the long-run opening layer, followed by daily local projections, spread models, discount models, and procurement-scale equations as the main short-run mechanism layer. VECM is treated as a system-feasibility and robustness block. This strategy reflects the corrected data architecture itself. Weekly aggregation is appropriate for equilibrium-style relations. Daily frequency is necessary once the analysis reaches procurement discreteness and retail behaviour.",
        "",
        f"On the corrected final data, ARDL remains part of the screening design, but only {len(ardl_rows)} retained ARDL rows survive the reporting threshold. ECM contributes {len(ecm_rows)} retained rows and NARDL contributes {len(nardl_rows)} retained rows. Product-level VECM attempts number {vecm_attempts}, of which {len(vecm_ok)} are feasible. The hierarchy of evidence is therefore explicit rather than assumed: the weekly core is selective, the daily mechanism block is central for downstream interpretation, and the system block is informative mainly as robustness.",
        "",
        "This ordering is close to the logic used in strong empirical KSE theses: first establish which equations remain admissible after diagnostics, then interpret only those parts of the empirical system that survive the filter, and finally separate the system-level robustness exercise from the main identification story. The present chapter follows that logic deliberately so that the narrative strength of the results does not exceed the empirical support behind them.",
        "",
        diag_text,
        "",
        "{{TABLE|Table 6.1. Key synthesis of the retained empirical results|../outputs/chapter_tables/table_6_1_key_results_synthesis.csv|Source: author's calculations based on the strongest retained weekly, daily, discount, procurement-scale, aggregate-index, and system-robustness outputs.}}",
        "",
        "![Figure 6.1. Top weekly lag-correlation signals across admissible links.](../figures/chapter6_results/01_weekly_corr_scan.png)",
        "",
        "Source: author's calculations based on the weekly lag-profile scan.",
        "",
        "### 6.2 Upstream-to-downstream transmission and the pairwise weekly models",
        "",
        "The weekly pairwise block starts from the upstream end of the chain. This is where the most defensible long-run evidence should appear if the corrected data are working properly. That expectation is borne out by the results. The cleanest retained relations lie in the producer-to-procurement segment, not in the late downstream segment. In economic terms, this means that the processor-to-procurement transition carries the most stable equilibrium signal, while the retail-facing segment is more institutionally and strategically filtered.",
        "",
        strongest_weekly,
        "",
        "Procurement is therefore not a trivial middle layer. It is the point where processor price pressure becomes institutional transaction price pressure. That is exactly where buffering, contract timing, and lot-specific selection can change the form of transmission before it reaches the shelf. The weekly models support the thesis argument that the Ukrainian dairy chain is vertically connected, but not frictionless.",
        "",
        weekly_product_text,
        "",
        farmgate_detail_text,
        "",
        "The same logic also explains why several downstream weekly links remain weak. Once retail enters the picture, observed pass-through becomes more sensitive to assortment, promotional timing, and retailer-specific price management. The weekly results do not show that downstream transmission is absent. They show that downstream transmission should not be reduced to one common long-run coefficient when empirical support is late and behaviourally filtered.",
        "",
        "![Figure 6.2. Summary of the 21 directional chain links.](../figures/chapter6_results/02_link21_status_matrix.png)",
        "",
        "Source: author's calculations based on the admissibility screen and retained weekly model set.",
        "",
        "![Figure 6.3. Retained weekly model coefficients on corrected data.](../figures/chapter6_results/03_core_model_coefficients.png)",
        "",
        "Source: author's calculations based on retained weekly ARDL, ECM, and NARDL models.",
        "",
        "![Figure 6.4. ECM speed of adjustment across retained weekly links.](../figures/chapter6_results/04_ecm_speed_of_adjustment.png)",
        "",
        "Source: author's calculations based on retained ECM equations.",
        "",
        "![Figure 6.5. NARDL asymmetry strength across retained weekly links.](../figures/chapter6_results/05_nardl_asymmetry.png)",
        "",
        "Source: author's calculations based on retained NARDL equations.",
        "",
        "### 6.3 Downstream extensions, retail behavior, discounts, and procurement scale",
        "",
        "Once the analysis reaches the downstream end of the chain, weekly equilibrium evidence becomes less informative than daily mechanism evidence. This is not a weakness of the design. It is one of the main results of the corrected empirical system. Retail prices are managed through category policy, markdowns, and timing choices. Procurement prices are shaped not only by unit-price movement, but also by contract scale and revisions. These are precisely the kinds of mechanisms that are easier to detect at daily frequency and harder to summarize as one stable long-run slope.",
        "",
        f"The corrected results retain {margin_flags} persistent-margin flags, {asym_flags} asymmetric-margin flags, {discount_signals} discount-strategy signals, and {scale_signals} procurement-scale signals. Jointly, these imply that downstream pricing is not a passive continuation of upstream cost movement. It is a managed layer in which institutional procurement and retail strategy both matter.",
        "",
        discount_detail_text,
        "",
        scale_detail_text,
        "",
        "![Figure 6.6. Local-projection pass-through by horizon.](../figures/chapter6_results/06_lp_pass_through_horizons.png)",
        "",
        "Source: author's calculations based on daily local projections.",
        "",
        "![Figure 6.7. Forward versus reverse evidence shares in the daily mechanism block.](../figures/chapter6_results/07_forward_reverse_core_share.png)",
        "",
        "Source: author's calculations based on 7-day and 14-day local-projection summaries.",
        "",
        "![Figure 6.8. Average spread levels across chain segments.](../figures/chapter6_results/08_spread_levels.png)",
        "",
        "Source: author's calculations based on product-level spread summaries.",
        "",
        "![Figure 6.9. Vertical spread and market-power proxy by chain segment.](../figures/chapter6_results/10_vertical_spread_proxy.png)",
        "",
        "Source: author's calculations based on spread-regression coefficients.",
        "",
        "![Figure 6.10. Discount incidence by product.](../figures/chapter6_results/11_discount_incidence.png)",
        "",
        "Source: author's calculations based on discount-incidence models.",
        "",
        "![Figure 6.11. Discount-model coefficient map.](../figures/chapter6_results/12_discount_coefficient_map.png)",
        "",
        "Source: author's calculations based on product-level Silpo discount equations.",
        "",
        "![Figure 6.12. Procurement-scale effects on ProZorro price changes.](../figures/chapter6_results/13_procurement_scale_effects.png)",
        "",
        "Source: author's calculations based on procurement-scale models with expected value and contract-sum variables.",
        "",
        "{{TABLE|Table 6.2. Daily downstream mechanism models|../outputs/chapter_tables/table_6_3_daily_mechanism_models.csv|Source: author's calculations based on local projections, spread models, Silpo discount models, and procurement-scale equations.}}",
        "",
        "In substantive terms, the downstream evidence supports a market-power interpretation that is more nuanced than a single markup story. Retail influence appears through timing control, through the use of discounts to keep observed prices temporarily below the non-promotional baseline, and through selective exposure of upstream pressure to the final shelf price. ProZorro, in turn, acts as a buffer: it passes part of upstream movement downstream, but it also absorbs shocks through institutional timing and scale effects before those shocks reach retail.",
        "",
        "### 6.4 VECM system evidence, aggregate dairy indices, and the limits of full-chain modelling",
        "",
        "The aggregate-index block is introduced as a robustness system, not as a substitute for the product-level results. Because direct quantity weights are unavailable, the aggregate dairy indices are built from fixed structural proxy weights. That makes them suitable for checking whether the broad transmission story survives aggregation, but not for replacing the product-level evidence that remains central to the thesis.",
        "",
        "The aggregate-index results support the same general conclusion as the product-level models. Producer-to-procurement remains the cleanest structural step, while the downstream end remains more filtered. At the same time, the system block also makes clear why a fully unified chain model should be treated cautiously. Product-level VECM feasibility remains limited once the corrected overlap and the honest downstream construction are imposed. The system perspective is therefore useful, but it does not overturn the staged interpretation of the chain.",
        "",
        index_detail_text,
        "",
        "![Figure 6.13. Aggregate dairy-chain indices at weekly frequency.](../figures/chapter6_results/11_aggregate_index_overlay.png)",
        "",
        "Source: author's calculations based on fixed-weight weekly dairy indices.",
        "",
        "![Figure 6.14. Aggregate dairy-index model coefficients.](../figures/chapter6_results/14_aggregate_index_model_coefficients.png)",
        "",
        "Source: author's calculations based on aggregate-index ECM and NARDL models.",
        "",
        "![Figure 6.15. VECM system feasibility on corrected weekly panels.](../figures/chapter6_results/10_vecm_feasibility.png)",
        "",
        "Source: author's calculations based on product-level and aggregate-index VECM feasibility checks.",
        "",
        "![Figure 6.16. Farm-gate transmission lag map across the chain.](../figures/chapter6_results/16_farmgate_lag_map.png)",
        "",
        "Source: author's calculations based on farm-gate-specific weekly lag-correlation profiles.",
        "",
        "![Figure 6.17. Farm-gate links within the retained weekly model set.](../figures/chapter6_results/17_farmgate_chain_coefficients.png)",
        "",
        "Source: author's calculations based on retained farm-gate-to-downstream weekly equations.",
        "",
        "{{TABLE|Table 6.3. Aggregate-index and VECM system results|../outputs/chapter_tables/table_6_4_system_models.csv|Source: author's calculations based on aggregate-index chain models and VECM feasibility outputs.}}",
        "",
        "### 6.5 Interpretation of the key results, limitations and ways for improvement",
        "",
        "Taken together, the corrected empirical system leads to a coherent interpretation of the Ukrainian dairy market. First, vertical price transmission exists, but the strongest evidence is not equally distributed across the chain. It is strongest before the consumer-facing stage, particularly in the producer-to-procurement segment. Second, procurement is not a neutral corridor. It is a filtering institution whose price and scale dynamics shape what later reaches the shelf. Third, retail market power is visible less as one universal markup coefficient and more as a set of strategic adjustment tools: timing, markdown policy, category management, and selective asymmetry.",
        "",
        "This interpretation is consistent with the broader logic of Chapters 1 to 4. The thesis asks how the Ukrainian food consumer market, and especially the dairy market, converts upstream pressure into household-visible prices. The empirical answer is now more precise. Farm-gate remains the raw-milk anchor. Producer prices transmit upstream pressure into the processed-product stage. Procurement converts this pressure into institutional transaction prices, sometimes buffering and sometimes amplifying it through scale and contract conditions. Retail then determines how much of that pressure becomes visible to the consumer and how much is temporarily managed away through pricing strategy.",
        "",
        "The limitations also remain important. Strict weekly overlap is still scarce. Product-level VECM feasibility is low. Farm-gate is necessarily a raw-milk benchmark rather than a literal processed-product series. Retail data remain retailer-specific, and the strongest discount evidence is concentrated in Silpo because its markdown logic is observed directly. None of these limitations invalidate the findings, but they set the boundary of what can be claimed with confidence.",
        "",
        "The most productive extensions are therefore clear. Better quantity weights would strengthen the aggregate-index block. Longer downstream overlap would improve full-chain system estimation. Richer regional retail coverage could make the market-power interpretation more spatially explicit. Yet even without those extensions, the corrected empirical system already improves the thesis materially because it makes the data logic, the admissibility logic, and the hierarchy of evidence explicit.",
        "",
        "For the main thesis argument, the decisive point is that the empirical pattern now fits the institutional structure described in the earlier chapters. Raw-milk pressure is visible, but it does not pass to the shelf mechanically. Processor prices remain the first strong carrier of the shock. Procurement both relays and filters. Retail then decides how much of the pressure becomes immediately visible to consumers and how much is delayed, softened, or redistributed across products through markdowns and category policy.",
        "",
        "{{TABLE|Table 6.4. Integrated interpretation, limitations, and research implications|../outputs/chapter_tables/table_6_5_integrated_synthesis.csv|Source: author's synthesis based on the final empirical results.}}",
        "",
        "### 6.0 Conclusion",
        "",
        "The final result of the empirical rebuild is not a claim that every segment of the dairy chain behaves identically. It is a more disciplined and more informative statement about where transmission is strongest, where it becomes delayed, and where it is strategically reshaped before the price reaches the consumer. On the corrected data, the most defensible long-run evidence still lies before the retail stage, especially in the producer-to-procurement link. Once the chain reaches procurement and retail, transmission does not disappear, but it becomes institutionally and strategically filtered rather than purely mechanical.",
        "",
        "The added farm-gate block sharpens that conclusion. Farm-gate does not behave like a processed-product series, and the thesis now treats it honestly as a raw-milk benchmark rather than forcing a false one-to-one comparison. Even so, the lag profiles, retained farm-gate equations, and widened aggregate-index system all point in the same direction: upstream milk-supply pressure is visible, but it reaches the shelf through intermediate filters. Producer prices become the first processed-product carrier of the shock. Procurement relays and conditions it through contracting and scale. Retail then decides how much of that pressure becomes immediately consumer-visible and how much is delayed, softened, or selectively exposed through markdowns and category strategy. That is the central economic result of the thesis-grade rebuild, and it fits the institutional structure of Ukraine's dairy market more closely than the earlier empirical versions did.",
    ]

    md_path = DOC_DIR / "Chapter5_6_analysis.md"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    write_html(DOC_DIR / "Chapter5_6_analysis.html", "Chapter 5 and 6 Analysis", lines)
    write_kse_docx(DOC_DIR / "Chapter5_6_analysis.docx", lines)


def main() -> None:
    ensure_dirs()
    clear_generated_outputs()
    if not FULL_UAH.exists():
        raise FileNotFoundError(FULL_UAH)

    product_audit, product_dictionary = build_product_definition_audit()
    panel, inventory, source_frames = build_panel()
    weekly = build_weekly_panels(panel)
    source_frames["weekly_panel"] = weekly
    index_daily, index_weekly, index_weights = build_aggregate_chain_indices(panel, source_frames)
    source_frames["aggregate_index_daily"] = index_daily
    source_frames["aggregate_index_weekly"] = index_weekly
    source_frames["aggregate_index_weights"] = index_weights
    inventory = pd.concat(
        [
            inventory,
            pd.DataFrame(
                [
                    {
                        "source": "weekly_panel",
                        "rows": len(weekly),
                        "products": weekly["product"].nunique() if "product" in weekly.columns else np.nan,
                        "date_min": weekly["week"].min() if "week" in weekly.columns else pd.NaT,
                        "date_max": weekly["week"].max() if "week" in weekly.columns else pd.NaT,
                        "columns": ", ".join(map(str, weekly.columns[:20])),
                    }
                ]
                + (
                    [
                        {
                            "source": "aggregate_index_daily",
                            "rows": len(index_daily),
                            "products": np.nan,
                            "date_min": index_daily["date"].min() if not index_daily.empty else pd.NaT,
                            "date_max": index_daily["date"].max() if not index_daily.empty else pd.NaT,
                            "columns": ", ".join(map(str, index_daily.columns[:20])) if not index_daily.empty else "",
                        },
                        {
                            "source": "aggregate_index_weekly",
                            "rows": len(index_weekly),
                            "products": np.nan,
                            "date_min": index_weekly["week"].min() if not index_weekly.empty else pd.NaT,
                            "date_max": index_weekly["week"].max() if not index_weekly.empty else pd.NaT,
                            "columns": ", ".join(map(str, index_weekly.columns[:20])) if not index_weekly.empty else "",
                        },
                    ]
                    if not index_weekly.empty
                    else []
                )
            ),
        ],
        ignore_index=True,
    )
    coverage = coverage_table(panel)
    intersections = build_intersection_table(weekly)
    corr_scan = run_weekly_correlation_scan(weekly, intersections)
    series_tests, chain_models, vecm_results = run_core_chain_models(weekly, intersections)
    vecm_detail_summary, vecm_detail_tables = run_vecm_detailed_tables(weekly)
    index_series_tests, index_models, index_vecm_details = run_aggregate_index_models(index_weekly)
    lp = run_local_projections(panel)
    lp_summary = summarize_lp(lp)
    margins, margin_spreads = run_margin_models(panel)
    discounts = run_discount_models(panel)
    scale_models = run_procurement_scale_models(weekly)
    findings = robust_findings(lp, margins, discounts)
    link21_summary = build_link21_summary(intersections, corr_scan, chain_models)
    index_vecm_summary = index_vecm_details.get("summary", pd.DataFrame())
    reliability_overview = build_model_reliability_overview(chain_models, discounts, scale_models, vecm_results, index_models, index_vecm_summary)

    write_excel_outputs(
        panel,
        weekly,
        index_daily,
        index_weekly,
        index_weights,
        inventory,
        source_frames,
        product_audit,
        product_dictionary,
        coverage,
        intersections,
        corr_scan,
        series_tests,
        chain_models,
        vecm_results,
        index_series_tests,
        index_models,
        index_vecm_summary,
        link21_summary,
        reliability_overview,
        lp,
        lp_summary,
        margins,
        margin_spreads,
        discounts,
        scale_models,
        findings,
    )
    plot_outputs(panel, coverage, lp_summary, margin_spreads, discounts, source_frames)
    plot_final_outputs(weekly, corr_scan, chain_models, vecm_results, scale_models, index_models, link21_summary, reliability_overview, discounts, margin_spreads, source_frames)
    write_single_model_outputs(chain_models, intersections, scale_models, discounts, index_models)
    write_vecm_detail_outputs(vecm_detail_summary, {**vecm_detail_tables, **{k: v for k, v in index_vecm_details.items() if k != "summary"}})
    write_integrated_summary_workbook(intersections, corr_scan, chain_models, vecm_results, index_models, index_vecm_summary, link21_summary, reliability_overview, discounts, scale_models, findings)
    write_run_summary(product_dictionary, intersections, corr_scan, chain_models, vecm_results, discounts, scale_models)
    write_chapter56(
        inventory,
        weekly,
        source_frames,
        product_dictionary,
        intersections,
        corr_scan,
        chain_models,
        vecm_results,
        lp_summary,
        margins,
        margin_spreads,
        discounts,
        scale_models,
        link21_summary,
        reliability_overview,
        index_weekly,
        index_weights,
        index_models,
        index_vecm_summary,
    )
    write_final_readme()

    print("FINAL_RESEARCH pipeline completed.")
    print(f"Detailed workbook: {OUTPUT_DIR / 'final_research_outputs.xlsx'}")
    print(f"Integrated summary workbook: {OUTPUT_DIR / 'integrated_summary_workbook.xlsx'}")
    print(f"Chapter document: {DOC_DIR / 'Chapter5_6_analysis.docx'}")
    print(f"Added dataset workbook: {ADDED_DATASET_PATH}")


if __name__ == "__main__":
    main()
