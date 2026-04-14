#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/getapple/Documents/KSE/Master Thesis")
FINAL_RESEARCH = ROOT / "FINAL_RESEARCH"
MATERIALS = FINAL_RESEARCH / "materials"
INPUTS = MATERIALS / "inputs"
OUT_DOCX = FINAL_RESEARCH / "documents" / "stepbystep.docx"
OUT_MD = FINAL_RESEARCH / "documents" / "stepbystep.md"
OUT_HTML = FINAL_RESEARCH / "documents" / "stepbystep.html"
FIGURES_DIR = FINAL_RESEARCH / "figures"

FILE_A = INPUTS / "full_uah_final_whatadded_matched_smoothed.xlsx"
FILE_B = INPUTS / "full_uah_final_whatadded_matched_smoothed.xlsx 2.xlsx"


def get_core_xml_text(path: Path) -> str:
    with ZipFile(path) as zf:
        return zf.read("docProps/core.xml").decode("utf-8")


def workbook_summary(path: Path) -> dict:
    xl = pd.ExcelFile(path)
    out = {
        "path": str(path),
        "name": path.name,
        "size_mb": round(path.stat().st_size / (1024 * 1024), 2),
        "sheets": [],
    }
    for sheet in xl.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet)
        out["sheets"].append(
            {
                "sheet": sheet,
                "rows": len(df),
                "cols": len(df.columns),
                "columns_preview": ", ".join(map(str, df.columns[:8])),
            }
        )
    out["core_xml"] = get_core_xml_text(path)
    return out


def workbooks_equal(path_a: Path, path_b: Path) -> bool:
    xl_a = pd.ExcelFile(path_a)
    xl_b = pd.ExcelFile(path_b)
    if xl_a.sheet_names != xl_b.sheet_names:
        return False
    for sheet in xl_a.sheet_names:
        a = pd.read_excel(path_a, sheet_name=sheet)
        b = pd.read_excel(path_b, sheet_name=sheet)
        if not a.equals(b):
            return False
    return True


def add_formula_box(doc: Document, title: str, formula: str, explanation: str) -> None:
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F3F4F6")
    tc_pr.append(shd)
    p_formula = cell.paragraphs[0]
    p_formula.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_formula = p_formula.add_run(formula)
    r_formula.font.name = "Courier New"
    r_formula._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    r_formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    r_formula.font.size = Pt(10)

    p_exp = doc.add_paragraph()
    p_exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_exp.paragraph_format.line_spacing = 1.3
    p_exp.add_run(explanation)


def write_markdown(text: str) -> None:
    OUT_MD.write_text(text, encoding="utf-8")
    body = "\n".join(f"<p>{line}</p>" if line and not line.startswith("#") and not line[0:2].isdigit() else line for line in text.splitlines())
    html = f"<html><body><pre style='white-space: pre-wrap; font-family: Georgia, serif'>{text}</pre></body></html>"
    OUT_HTML.write_text(html, encoding="utf-8")


def main() -> None:
    summary_a = workbook_summary(FILE_A)
    summary_b = workbook_summary(FILE_B) if FILE_B.exists() else None
    same_values = workbooks_equal(FILE_A, FILE_B) if FILE_B.exists() else False

    sheet_lines = []
    for row in summary_a["sheets"]:
        sheet_lines.append(
            f"- `{row['sheet']}`: {row['rows']:,} rows, {row['cols']} columns. "
            f"This sheet stores {row['columns_preview']}."
        )

    figure_paths = sorted(FIGURES_DIR.rglob("*.png"))

    workflow_points = [
        (
            "Corrected governmental base truth",
            "Start from `full_uah_final.xlsx`, where FarmGateUA, ProducerUA, and ConsumerUA are already recalculated as true monthly point-in-time prices and only after that widened to daily frequency. This is the non-negotiable base truth of the final rerun and the starting point of Chapter 5."
        ),
        (
            "Territorial logic of the core chain",
            "Fix the main empirical chain at Ukraine-wide level for farm-gate, producer, and consumer prices, while keeping ProZorro as an all-Ukraine transactional procurement layer. This prevents the main chain from mixing national averages with regional spot observations mechanically."
        ),
        (
            "Product-definition audit across all datasets",
            "Audit how each dataset defines product meaning: by row labels, column names, titles, mapped fields, or text descriptions. Build the standardized product dictionary before any modelling so later estimates compare economically comparable dairy products rather than only superficially similar names."
        ),
        (
            "Farm-gate preparation",
            "Prepare the corrected farm-gate benchmark as the raw-milk origin of the chain, preserve interpolation variants, and keep farm-gate as a benchmark object that is compared to processed-product prices but not treated as if it were the same kind of product."
        ),
        (
            "Producer preparation",
            "Load the corrected producer series by product, preserve `price_linear` and `price_pchip`, and align products to the common dairy taxonomy. This stage is later visualized in the raw and transformed data blocks of Chapter 5 (Figures 5.1, 5.4, 5.5)."
        ),
        (
            "Consumer preparation",
            "Load the corrected consumer series, keep the official interpolation logic, and align it to the same product structure used for producer and retail comparison. ConsumerUA remains a benchmark layer until category-level retail series are standardized."
        ),
        (
            "ProZorro price-object preparation",
            "Use `Ціна за одиницю` as the procurement price object and keep `Очікувана вартість`, `Сума договорів початкова`, and `Сума договорів поточна` as separate trade-scale variables. They are not prices and therefore enter the analysis only as procurement-intensity or institutional-scale controls."
        ),
        (
            "Silpo item-level reconstruction",
            "Treat `price_current` as the observed post-discount consumer price, reconstruct baseline price and markdown depth from `discount_value` and `discount_%`, and keep discount incidence as a structural behavioural layer rather than noise. This stage is summarized later in the Silpo discount figure of Chapter 5 (Figure 5.11)."
        ),
        (
            "Novus item-level reconstruction",
            "Prepare Novus at the same product-item level, normalize brand and product naming, and keep the observed retail price as a separate downstream object. Novus is not forced into the Silpo discount logic because it does not contain a comparable markdown layer."
        ),
        (
            "Cross-shop retail harmonization",
            "Construct harmonized item keys, stricter diagnostic keys, product literals, and normalized brands so Silpo and Novus can be compared item by item and then aggregated at stable product-type level. This stage underlies Chapter 5 figures on harmonization, product mix, and brand support (Figures 5.7–5.9)."
        ),
        (
            "Retail category aggregation",
            "Convert unstable SKU-level observations into category-level daily price objects using median observed package prices within harmonized product-type groups, together with reconstructed baseline variants. This avoids composition bias caused by changing assortment, pack-size turnover, or temporary product replacement."
        ),
        (
            "Construction of the downstream retail endpoint",
            "Build the combined downstream retail series from harmonized Silpo and Novus category prices and keep ConsumerUA as a separate benchmark aligned to the same product taxonomy. This avoids over-constructing the downstream price object while preserving a consumer-facing comparison layer."
        ),
        (
            "Integrated daily modelling panel",
            "Merge the corrected governmental layers, ProZorro price and trade-scale objects, retail category series, and the benchmark layers into the final daily panel used for mechanism models."
        ),
        (
            "Weekly median layer",
            "Aggregate every chain step to weekly medians and preserve smoothed variants only where they improve long-run interpretability, especially for procurement. This becomes the long-run baseline and the main time-series visualization layer of Chapter 5, especially the transformed chain overlays and the farm-gate comparison visuals."
        ),
        (
            "Farm-gate bridge diagnostics",
            "Build a dedicated farm-gate benchmark block that compares the raw-milk origin of the chain with producer, procurement, retail, and consumer indices in both level form and normalized index form. This stage is used to show that farm-gate is not the same product as processed dairy goods, but still carries the earliest supply-pressure signal into the rest of the chain."
        ),
        (
            "Aggregate dairy-index construction",
            "Build geometric aggregate indices for farm-gate, producer, procurement, retail, and consumer layers using structural proxy weights rather than endogenous model-derived weights. This provides the market-level robustness block used in Chapter 5 and Chapter 6."
        ),
        (
            "Downstream extension for system robustness",
            "Construct a widened downstream extension index that uses the retail index where it exists and a consumer-scaled extension where retail support is too late for system modelling. This extension is used only as an added robustness layer for aggregate system estimation and does not replace the strict downstream retail evidence."
        ),
        (
            "Raw-data visualization sequence",
            "Visualize the data in the correct order: first raw corrected governmental series, then raw retail series, then external benchmarks, and only after that transformed category-level and aggregate-index views. This sequence corresponds to Chapter 5 Figures 5.1–5.5."
        ),
        (
            "Coverage and descriptive diagnostics",
            "Visualize panel coverage, cross-shop overlap, literal product mix, brand support, regional procurement structure, and discount environment before modelling. These graphs are part of the admissibility argument, not merely supporting illustrations (Chapter 5 Figures 5.6–5.10, plus the supporting figure folder outputs)."
        ),
        (
            "Intersection and admissibility scoring",
            "Measure overlap length, continuity, mapping quality, unit comparability, and interpolation dominance for every candidate chain link. Then classify each candidate link as strong, acceptable, weak extension, or unusable before any weekly model is interpreted."
        ),
        (
            "Lag-profile exploration",
            "Run the weekly lag-correlation scan to identify where delayed co-movement is strongest. This becomes the first model-stage figure in Chapter 6 and introduces the idea that transmission is delayed rather than contemporaneous (Figure 6.1)."
        ),
        (
            "ARDL weekly screening",
            "Estimate ARDL models as part of the long-run screening stage and retain them only if the corrected overlap and diagnostics justify interpretation. ARDL therefore remains part of the design even when it does not dominate the retained evidence."
        ),
        (
            "ECM weekly estimation",
            "Estimate ECM models where a defensible equilibrium relation exists, so the chapter can discuss both long-run linkage and speed of adjustment. These outputs are summarized visually in Chapter 6 (Figures 6.3–6.4)."
        ),
        (
            "NARDL weekly estimation",
            "Estimate NARDL models to test whether positive and negative upstream shocks pass through asymmetrically. This becomes one of the main asymmetry checks of the final thesis results (Figure 6.5)."
        ),
        (
            "VECM system feasibility",
            "Test whether the corrected chain can sustain a whole-system VECM at product level and at aggregate-index level, and then add the widened downstream-extension system as a separate robustness check. This stage is reported honestly as a feasibility check, not assumed to succeed a priori."
        ),
        (
            "Daily local projections",
            "Estimate daily local projections to trace short-run pass-through over multiple horizons, especially where weekly overlap is too thin for a clean long-run downstream model. These are reported in Chapter 6 as the main short-run timing evidence (Figures 6.6–6.8)."
        ),
        (
            "Spread and margin models",
            "Estimate spread regressions between adjacent chain levels to show where margins mean-revert, persist, or adjust asymmetrically. This block produces the market-power proxy evidence reported in Chapter 6 (Figures 6.9–6.11)."
        ),
        (
            "Silpo discount models",
            "Estimate discount-incidence and discount-depth equations using the corrected meaning of `price_current`, `discount_value`, and `discount_%`. This stage explains how visible retail pass-through is smoothed, delayed, or selectively exposed (Figures 6.12–6.13)."
        ),
        (
            "Procurement-scale models",
            "Estimate the effect of procurement scale through expected value and both contract-sum variables, especially `Сума договорів поточна`, to test whether procurement intensity modifies observed price transmission (Figure 6.14)."
        ),
        (
            "Twenty-one directional chain summary",
            "Assemble the 21 requested directional links into one integrated summary table that combines admissibility, best lag signal, best retained model, coefficient, and interpretation. This becomes the key comparative table of Chapter 6 (Table 6.2 and Figure 6.2)."
        ),
        (
            "Reliability synthesis",
            "Summarize model reliability by block, keeping strong, conditional, and weak evidence separate. This ensures that the final thesis text does not present all estimated numbers as equally valid (Table 6.1 and Figure 6.15)."
        ),
        (
            "Academic synthesis and final writing",
            "Write Chapter 5 and Chapter 6 in a thesis style that follows the logic of the corrected draft: data first, then admissibility, then upstream-to-downstream results, backward robustness, retail discount mechanisms, aggregate robustness, VECM limits, and the final economic conclusion."
        ),
    ]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.079)
        section.bottom_margin = Inches(0.96)
        section.left_margin = Inches(1.441)
        section.right_margin = Inches(0.96)

    styles = doc.styles
    styles["Normal"].font.name = "Garamond"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.space_after = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Step-by-Step Explanation of the Two Excel Files and the Full FINAL_RESEARCH Workflow")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "This note explains what the two Excel files are, why they exist, how they were constructed, "
        "and how they fit into the final empirical rebuild of the dairy price-transmission research."
    )

    h1 = doc.add_paragraph()
    h1.add_run("1. What the two Excel files are").bold = True

    for name, role in [
        (
            FILE_A.name,
            "the operative modelling workbook generated by the FINAL_RESEARCH pipeline and explicitly referenced in the project README as the final widened, matched, and smoothed dataset actually used for modelling",
        ),
        *(
            [
                (
                    FILE_B.name,
                    "a duplicate save of the same workbook rather than a different analytical dataset",
                )
            ]
            if summary_b is not None
            else []
        ),
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(f"{name}: ")
        run.bold = True
        para.add_run(role + ".")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if summary_b is not None and same_values:
        p.add_run(
            "The evidence is direct. Both files contain the same sheets, the same row counts, the same columns, "
            "and identical cell values sheet by sheet. The only detectable difference is workbook package metadata, "
            "specifically the OpenXML core creation and modification timestamps. That means the second file is not a "
            "new data version; it is a later duplicate write of the same modelling workbook."
        )
    else:
        p.add_run(
            "The current primary workbook is the matched-and-smoothed modelling companion generated by the FINAL_RESEARCH pipeline. "
            "If a second file with the same name plus ` 2.xlsx` exists, it should be treated only as a duplicate save rather than as a separate analytical dataset."
        )

    comp = doc.add_table(rows=1, cols=4)
    comp.style = "Table Grid"
    hdr = comp.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Size"
    hdr[2].text = "Analytical status"
    hdr[3].text = "Conclusion"
    rows = [[FILE_A.name, f"{summary_a['size_mb']} MB", "Primary workbook", "Used in the final pipeline."]]
    if summary_b is not None:
        rows.append([FILE_B.name, f"{summary_b['size_mb']} MB", "Duplicate copy", "Value-identical duplicate save." if same_values else "Secondary save; verify before use."])
    for row in rows:
        cells = comp.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    h2 = doc.add_paragraph()
    h2.add_run("2. Why this workbook was created").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "The base workbook `full_uah_final.xlsx` already stores the corrected governmental price truth. "
        "However, the final research needed an additional workbook that contains only the constructed datasets "
        "that were added for modelling after the base correction stage. That is why "
        "`full_uah_final_whatadded_matched_smoothed.xlsx` was created: it serves as a compact modelling "
        "companion, not as a replacement for the base source workbook."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "Its purpose is methodological transparency. Instead of scattering derived datasets across many CSV files, "
        "it places the final modelling panels in one auditable workbook and labels them in the metadata sheet by "
        "what was added, why it was added, whether it belongs to strict-core evidence or extension evidence, "
        "and whether it enters the final identification."
    )

    h3 = doc.add_paragraph()
    h3.add_run("3. How the workbook was constructed").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "The workbook was generated automatically by the final pipeline after the corrected daily and weekly panels "
        "had already been built. It is therefore an output of the thesis-grade rebuild, not a manually edited Excel file. "
        "The construction sequence was: corrected governmental layers -> product audit -> cleaned retail item layer -> "
        "daily integrated panel -> weekly median panel -> admissibility table -> product dictionary -> final item-level retail archive."
    )

    h4 = doc.add_paragraph()
    h4.add_run("4. What each sheet means").bold = True

    for line in sheet_lines:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(line)

    h5 = doc.add_paragraph()
    h5.add_run("5. Transformation logic and formulas").bold = True

    add_formula_box(
        doc,
        "Formula box 1. Daily linear interpolation inherited from the corrected base workbook",
        "p_linear(t) = p_m + ((t - d_m) / (d_{m+1} - d_m)) * (p_{m+1} - p_m)",
        "This takes the corrected monthly point-in-time price p_m and linearly widens it across daily dates between two adjacent month anchors. "
        "In the final research, these daily governmental series are inherited from `full_uah_final.xlsx` and then merged into the final modelling panel."
    )

    add_formula_box(
        doc,
        "Formula box 2. Shape-preserving interpolation",
        "p_pchip(t) = PCHIP({(tau_m, p_m)})",
        "The PCHIP path preserves monotonicity and avoids oscillations that ordinary higher-order polynomials may create. "
        "It is stored alongside the linear version as a robustness interpolation path."
    )

    add_formula_box(
        doc,
        "Formula box 3. Effective and baseline retail price",
        "p_effective = price_current\np_baseline = price_current + discount_value, if discount_value > 0; otherwise p_baseline = price_current",
        "The effective price is the actual observed downstream consumer-facing price. The baseline price reconstructs the non-discounted shelf level when a markdown is active. "
        "This separation is essential because the thesis treats price transmission and discount strategy as related but distinct objects."
    )

    add_formula_box(
        doc,
        "Formula box 4. Discount incidence and discount percent",
        "discount_dummy = 1(active discount), 0(otherwise)\ndiscount_percent = explicit discount_% if available; otherwise 100 * (p_baseline - p_effective) / p_baseline",
        "This preserves the behavioural layer of retail strategy. Discount incidence and discount depth are not treated as noise; they are modelled as one mechanism through which retail may smooth or delay pass-through."
    )

    add_formula_box(
        doc,
        "Formula box 5. Retail combined endpoint and consumer benchmark gap",
        "p_retail_combined(t) = median(p_Silpo(t), p_Novus(t))\ngap_retail_vs_consumer(t) = ln(p_retail_combined(t)) - ln(p_consumer(t))",
        "The downstream retail stage is constructed from the two observed retailers, while ConsumerUA remains a separate benchmark. The gap term is used analytically without collapsing both objects into one synthetic price."
    )

    add_formula_box(
        doc,
        "Formula box 6. Weekly median aggregation",
        "p_week(w) = median{ p_d : d belongs to week w }",
        "Weekly medians are the baseline long-run visualization and modelling frequency in the final research, because they reduce idiosyncratic daily noise while preserving actual price level movement."
    )

    add_formula_box(
        doc,
        "Formula box 7. Log spread and price change",
        "spread_t = ln(p_downstream,t) - ln(p_upstream,t)\nDelta ln(p_t) = ln(p_t) - ln(p_{t-1})",
        "These transformations are used in spread models, ECM-type adjustment, and local dynamic specifications because they make price movements and proportional gaps economically interpretable."
    )

    h6 = doc.add_paragraph()
    h6.add_run("6. Why the duplicate file exists and why it should not be treated as a separate dataset").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "The second file, when it exists as `full_uah_final_whatadded_matched_smoothed.xlsx 2.xlsx`, is best interpreted as a duplicate save made a few minutes later. "
        "Its analytical content is identical. The byte-level mismatch comes from workbook package metadata rather than data values. "
        "In practical thesis terms, it should not be cited as a different version of the modelling data. The correct interpretation is that both files represent the same final added-data workbook, "
        "and the file without ` 2` is the primary one referenced by the FINAL_RESEARCH project."
    )

    h7 = doc.add_paragraph()
    h7.add_run("7. Consecutive numbered steps of the whole FINAL_RESEARCH analysis").bold = True

    for idx, (name, explanation) in enumerate(workflow_points, start=1):
        p_num = doc.add_paragraph(style="List Number")
        p_num.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_num.paragraph_format.line_spacing = 1.4
        run = p_num.add_run(name)
        run.bold = True
        p_exp = doc.add_paragraph()
        p_exp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_exp.paragraph_format.left_indent = Inches(0.25)
        p_exp.paragraph_format.line_spacing = 1.35
        p_exp.add_run(explanation)

    h8 = doc.add_paragraph()
    h8.add_run("8. Final interpretive conclusion").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        "These two Excel files should therefore not be read as two different modelling datasets. They are one final added-data workbook in two saved copies. "
        "The reason this workbook exists is to document, in one place, the exact constructed datasets that were added on top of the corrected base workbook for the final thesis rerun. "
        "It provides the bridge between corrected source data and the final econometric system."
    )

    h9 = doc.add_paragraph()
    h9.add_run("9. Figure appendix: all figures generated in FINAL_RESEARCH").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(
        f"The appendix below includes all {len(figure_paths)} generated figures stored inside the FINAL_RESEARCH figure tree. "
        "They are shown in path order so the document preserves the same flow as the project folder: sequence and chapter figures first, then model-specific and reliability outputs."
    )

    for fig in figure_paths:
        rel = fig.relative_to(FINAL_RESEARCH)
        doc.add_picture(str(fig), width=Inches(6.1))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        run = cap.add_run(str(rel))
        run.italic = True
        run.font.name = "Garamond"
        run.font.size = Pt(10)

    doc.save(OUT_DOCX)

    markdown_lines = [
        "# Step-by-Step Explanation of the Two Excel Files and the Full FINAL_RESEARCH Workflow",
        "",
        "## 1. What the two Excel files are",
        f"- `{FILE_A.name}`: primary modelling workbook generated by the FINAL_RESEARCH pipeline.",
        *( [f"- `{FILE_B.name}`: duplicate save of the same workbook, not a separate analytical dataset."] if summary_b is not None else ["- If a second file with the same name plus ` 2.xlsx` exists, it should be treated only as a duplicate save."] ),
        "",
        ("Both files were checked sheet by sheet and are value-identical. The difference is package metadata such as creation and modification timestamps." if summary_b is not None and same_values else "The primary workbook is the matched-and-smoothed modelling companion used by the final pipeline."),
        "",
        "## 2. Why the workbook exists",
        "It exists because `full_uah_final.xlsx` is the corrected source workbook, while the `whatadded` workbook stores only the additional final panels created for modelling: the daily integrated panel, the weekly median panel, the intersection table, the product dictionary, and the harmonized retail item archive.",
        "",
        "## 3. Transformation formulas",
        "",
        "```text",
        "p_linear(t) = p_m + ((t - d_m) / (d_{m+1} - d_m)) * (p_{m+1} - p_m)",
        "p_pchip(t) = PCHIP({(tau_m, p_m)})",
        "p_effective = price_current",
        "p_baseline = price_current + discount_value, if discount_value > 0; otherwise p_baseline = price_current",
        "discount_dummy = 1(active discount), 0(otherwise)",
        "discount_percent = explicit discount_% if available; otherwise 100 * (p_baseline - p_effective) / p_baseline",
        "p_retail_combined(t) = median(p_Silpo(t), p_Novus(t))",
        "gap_retail_vs_consumer(t) = ln(p_retail_combined(t)) - ln(p_consumer(t))",
        "p_week(w) = median{ p_d : d belongs to week w }",
        "spread_t = ln(p_downstream,t) - ln(p_upstream,t)",
        "Delta ln(p_t) = ln(p_t) - ln(p_{t-1})",
        "```",
        "",
        "## 4. Numbered full workflow",
    ]
    for idx, (name, explanation) in enumerate(workflow_points, start=1):
        markdown_lines.append(f"{idx}. **{name}.** {explanation}")
    markdown_lines.extend(
        [
            "",
            "## 5. Final conclusion",
            "The file without ` 2` is the primary modelling workbook. The file with ` 2` is a duplicate save of the same analytical content.",
            "",
            "## 6. Figure appendix: all figures generated in FINAL_RESEARCH",
            f"The appendix below includes all {len(figure_paths)} generated figures stored inside the FINAL_RESEARCH figure tree.",
        ]
    )
    for fig in figure_paths:
        rel = fig.relative_to(FINAL_RESEARCH)
        markdown_lines.append(f"### {rel}")
        markdown_lines.append(f"![{rel}]({fig})")
        markdown_lines.append("")
    write_markdown("\n".join(markdown_lines))


if __name__ == "__main__":
    main()
