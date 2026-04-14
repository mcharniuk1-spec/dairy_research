from __future__ import annotations

import csv
import math
import re
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import final_research_pipeline as frp


ROOT = Path("/Users/getapple/Documents/KSE/Master Thesis")
FINAL = ROOT / "FINAL_RESEARCH"
OUTPUT = FINAL / "outputs"
DOCS = FINAL / "documents"
FIG_ROOT = FINAL / "figures"
FIG5 = FIG_ROOT / "chapter5_data"
FIG6 = FIG_ROOT / "chapter6_results"

OUT_DOCX = DOCS / "stepbystep_results.docx"
OUT_MD = DOCS / "stepbystep_results.md"
OUT_HTML = DOCS / "stepbystep_results.html"

CHAPTER_MD = DOCS / "Chapter5_6_analysis.md"
CHAPTER_DOCX = DOCS / "Chapter5_6_analysis.docx"
CHAPTER_HTML = DOCS / "Chapter5_6_analysis.html"


GREEN_FILL = "2E7D32"
RED_FILL = "C62828"
WHITE_TEXT = RGBColor(255, 255, 255)


def fmt_num(value: object, digits: int = 3) -> str:
    if value is None or value == "":
        return ""
    try:
        val = float(value)
    except Exception:
        return str(value)
    if math.isnan(val):
        return ""
    return f"{val:.{digits}f}"


def fmt_p(value: object) -> str:
    if value is None or value == "":
        return ""
    try:
        val = float(value)
    except Exception:
        return str(value)
    if math.isnan(val):
        return ""
    if val < 0.001:
        return "<0.001"
    return f"{val:.3f}"


def sig_label(p: object) -> str:
    if p is None or p == "":
        return "n/a"
    try:
        val = float(p)
    except Exception:
        return "n/a"
    if math.isnan(val):
        return "n/a"
    if val < 0.01:
        return "strongly significant"
    if val < 0.05:
        return "significant"
    if val < 0.10:
        return "marginally significant"
    return "not significant"


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_cell_text(cell, bold: bool = False, color: RGBColor | None = None) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Garamond"
            run.font.size = Pt(9)
            run.bold = bold or run.bold
            if color is not None:
                run.font.color.rgb = color


def set_run_style(run, *, bold: bool = False, italic: bool = False, size: int = 12, color: RGBColor | None = None) -> None:
    run.font.name = "Garamond"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text).bold = True


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_style(run, size=12)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_style(run, italic=True, size=11)


def add_table(doc: Document, caption: str, columns: Sequence[str], rows: Sequence[Dict[str, str]]) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(10)
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    set_run_style(run, bold=True, size=12)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, col in enumerate(columns):
        cell = table.rows[0].cells[idx]
        cell.text = col
        style_cell_text(cell, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row.get(col, "")
            cell = cells[idx]
            cell.text = str(value)
            style_cell_text(cell)
        highlight = row.get("__highlight__", "")
        if highlight == "green":
            for key in ["Result", "Interpretation"]:
                if key in columns:
                    cell = cells[columns.index(key)]
                    add_cell_shading(cell, GREEN_FILL)
                    style_cell_text(cell, color=WHITE_TEXT)
        elif highlight == "red":
            for key in ["Result", "Interpretation"]:
                if key in columns:
                    cell = cells[columns.index(key)]
                    add_cell_shading(cell, RED_FILL)
                    style_cell_text(cell, color=WHITE_TEXT)


def csv_rows(path: Path) -> List[Dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def classify_highlight(is_important: bool, is_positive: bool) -> str:
    if not is_important:
        return ""
    return "green" if is_positive else "red"


def interpret_pairwise(row: Dict[str, str]) -> Tuple[str, str]:
    link = row["link"]
    model = row.get("retained_model", "") or "none retained"
    coef = row.get("reported_coef", "")
    ect = row.get("ect", "")
    rel = row.get("reliability", "") or "not retained"
    if model == "none retained":
        result = f"{link}: no retained weekly coefficient; overlap weeks {row.get('overlap_weeks','')}."
        interp = "**The weekly overlap is too thin for a defensible retained equilibrium reading, so this link remains descriptive only.**"
        return result, interp
    result = f"{link}: {model}, coef {fmt_num(coef)}, ECT {fmt_num(ect)}, reliability {rel}."
    interp = "**This link provides retained weekly evidence, but its economic weight depends on the overlap quality and the retained reliability label rather than on the coefficient alone.**"
    return result, interp


def pairwise_rows() -> List[Dict[str, str]]:
    rows = []
    for row in csv_rows(OUTPUT / "chapter_tables" / "table_6_2_pairwise_weekly_models.csv"):
        result, interp = interpret_pairwise(row)
        important = any(k in row["link"] for k in ["Producer -> Procurement", "Procurement -> Retail", "FarmGate -> Producer"])
        positive = row.get("reliability", "") in {"reliable", "conditionally_usable"}
        rows.append(
            {
                "Scope": row["link"],
                "Model": row.get("retained_model", "") or "None retained",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": classify_highlight(important, positive),
            }
        )
    return rows


def weekly_model_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    coef = row.get("lr_coef") or row.get("sr_coef")
    ect = row.get("ect_coef", "")
    pval = row.get("ect_pvalue") or row.get("cointegration_p") or row.get("asymmetry_pvalue")
    sig = sig_label(pval)
    result = (
        f"{row['product_label']} | {row['link']} | {row['model_family']} | {row['data_variant']}: "
        f"LR/SR coef {fmt_num(coef)}, ECT {fmt_num(ect)}, p {fmt_p(pval)}, n {row.get('n_obs','')}."
    )
    interp = f"**This weekly equation is {row.get('model_reliability','')} and the retained coefficient is {sig}; it should be read as long-run evidence only if the diagnostics remain supportive.**"
    important = row["link"] == "Producer -> Procurement" or row["product"] in {"milk", "butter", "cheese"}
    positive = row.get("model_reliability", "") == "reliable" and sig != "not significant"
    return result, interp, classify_highlight(important, positive)


def weekly_model_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "core_chain_models.csv"):
        result, interp, highlight = weekly_model_interpretation(row)
        out.append(
            {
                "Scope": f"{row['product_label']} | {row['link']}",
                "Model": f"{row['model_family']} | {row['data_variant']}",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def lp_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    sig_share = float(row.get("sig_share") or 0)
    core_share = float(row.get("core_share") or 0)
    result = (
        f"{row['link']} | horizon {row['horizon_days']}d | mean coef {fmt_num(row['mean_coef'])}, "
        f"median coef {fmt_num(row['median_coef'])}, sig share {fmt_num(sig_share)}, core share {fmt_num(core_share)}, "
        f"median n {fmt_num(row['median_n_obs'], 0)}."
    )
    if core_share > 0:
        interp = "**This horizon retains short-run timing evidence, so it helps identify when pass-through becomes visible even if weekly equilibrium support is thin.**"
    elif sig_share > 0:
        interp = "**This horizon shows some significant short-run movement, but the evidence is not strong enough to treat it as a stable core signal across products.**"
    else:
        interp = "**This horizon does not retain stable short-run evidence, which itself supports a delayed or weak transmission reading for this link.**"
    important = any(k in row["link"] for k in ["ProducerUA -> ProZorro", "ProZorro -> Retail", "ProZorro -> Novus", "ProZorro -> Silpo"])
    positive = core_share > 0 or sig_share >= 0.5
    return result, interp, classify_highlight(important, positive)


def lp_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "local_projection_summary.csv"):
        result, interp, highlight = lp_interpretation(row)
        out.append(
            {
                "Scope": f"{row['link']} | {row['price_variant']}",
                "Model": "Local projections",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def margin_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    lag_p = row.get("lag_spread_p")
    asym_p = row.get("asymmetry_pvalue")
    result = (
        f"{row['product_label']} | {row['spread']} | {row['reconstruction_variant']}: "
        f"lag-spread {fmt_num(row['lag_spread_coef'])} (p {fmt_p(lag_p)}), "
        f"upstream+ {fmt_num(row['upstream_pos_coef'])} (p {fmt_p(row['upstream_pos_p'])}), "
        f"upstream- {fmt_num(row['upstream_neg_abs_coef'])} (p {fmt_p(row['upstream_neg_abs_p'])}), "
        f"asym p {fmt_p(asym_p)}, R2 {fmt_num(row['r2'])}."
    )
    if row.get("asymmetric_margin_flag") == "1":
        interp = "**This spread behaves asymmetrically, which is consistent with managed downstream adjustment rather than a purely passive markup.**"
    elif sig_label(lag_p) != "not significant":
        interp = "**This spread mean-reverts significantly, so the margin adjusts systematically even without strong asymmetry.**"
    else:
        interp = "**This spread does not deliver strong margin-management evidence in statistical terms, so it should be treated as weak or descriptive.**"
    important = row["product"] in {"milk", "butter", "cheese"} and "procurement" in row["stage"]
    positive = row.get("asymmetric_margin_flag") == "1" or sig_label(lag_p) != "not significant"
    return result, interp, classify_highlight(important, positive)


def margin_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "margin_market_power_models.csv"):
        result, interp, highlight = margin_interpretation(row)
        out.append(
            {
                "Scope": f"{row['product_label']} | {row['stage']}",
                "Model": "Spread / margin",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def discount_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    lag_p = row.get("lag_discount_p")
    gap_p = row.get("retail_vs_consumer_gap_p")
    result = (
        f"{row['product_label']}: discount share {fmt_num(row['mean_discount_share'])}, "
        f"lag discount {fmt_num(row['lag_discount_coef'])} (p {fmt_p(lag_p)}), "
        f"retail-consumer gap {fmt_num(row['retail_vs_consumer_gap_coef'])} (p {fmt_p(gap_p)}), "
        f"Silpo-Novus gap {fmt_num(row['silpo_novus_gap_coef'])} (p {fmt_p(row['silpo_novus_gap_p'])}), "
        f"R2 {fmt_num(row['r2'])}."
    )
    if row.get("discount_strategy_signal") == "1":
        interp = "**This product shows a retained discount-strategy signal, meaning markdown policy helps explain how visible retail pass-through is managed.**"
    else:
        interp = "**This product does not retain a full discount-strategy signal, so discounting is present but not strongly systematic in the final specification.**"
    important = row["product"] in {"milk", "butter", "cheese"}
    positive = row.get("discount_strategy_signal") == "1"
    return result, interp, classify_highlight(important, positive)


def discount_rows() -> List[Dict[str, str]]:
    return [
        {
            "Scope": row["product_label"],
            "Model": "Silpo discount strategy",
            "Result": discount_interpretation(row)[0],
            "Interpretation": discount_interpretation(row)[1],
            "__highlight__": discount_interpretation(row)[2],
        }
        for row in csv_rows(OUTPUT / "discount_strategy_models.csv")
    ]


def scale_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    result = (
        f"{row['product_label']}: lag price {fmt_num(row['lag_price_coef'])} (p {fmt_p(row['lag_price_p'])}), "
        f"d producer {fmt_num(row['d_producer_coef'])} (p {fmt_p(row['d_producer_p'])}), "
        f"d expected {fmt_num(row['d_expected_coef'])} (p {fmt_p(row['d_expected_p'])}), "
        f"d sum initial {fmt_num(row['d_sum_initial_coef'])} (p {fmt_p(row['d_sum_initial_p'])}), "
        f"d sum current {fmt_num(row['d_sum_current_coef'])} (p {fmt_p(row['d_sum_current_p'])}), "
        f"R2 {fmt_num(row['r2'])}."
    )
    if row.get("scale_signal_flag") == "1":
        interp = "**This product retains a procurement-scale signal, meaning contract scale modifies procurement price adjustment rather than merely accompanying it.**"
    else:
        interp = "**This product does not retain a strong procurement-scale signal, so the trade-scale variables remain informative but not decisive here.**"
    important = row["product"] in {"milk", "butter", "cheese", "sour_cream"}
    positive = row.get("scale_signal_flag") == "1"
    return result, interp, classify_highlight(important, positive)


def scale_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "procurement_scale_models.csv"):
        result, interp, highlight = scale_interpretation(row)
        out.append(
            {
                "Scope": row["product_label"],
                "Model": "Procurement scale",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def index_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    pval = row.get("ect_pvalue") or row.get("cointegration_p") or row.get("asymmetry_pvalue")
    result = (
        f"{row['link']} | {row['model_family']} | {row['data_variant']}: "
        f"coef {fmt_num(row.get('lr_coef') or row.get('sr_coef'))}, ECT {fmt_num(row['ect_coef'])}, "
        f"p {fmt_p(pval)}, reliability {row['model_reliability']}."
    )
    interp = "**This aggregate-index equation is a robustness check on system-wide dairy dynamics rather than the main identification result, but it helps verify whether the broad transmission story survives aggregation.**"
    important = "Producer -> Procurement" in row["link"]
    positive = row.get("model_reliability") == "reliable"
    return result, interp, classify_highlight(important, positive)


def index_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "aggregate_index_models.csv"):
        result, interp, highlight = index_interpretation(row)
        out.append(
            {
                "Scope": row["link"],
                "Model": f"Aggregate index {row['model_family']}",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def vecm_interpretation(row: Dict[str, str]) -> Tuple[str, str, str]:
    result = (
        f"{row['product_label']} | {row['system_name']}: status {row['status']}, "
        f"reason {row['reason']}, n {row['n_obs']}."
    )
    if row["status"] == "ok":
        interp = "**This system is feasible and can be used as interdependent chain evidence, although it still remains secondary to the retained pairwise design.**"
    else:
        interp = "**This system is not feasible on the corrected overlap, which is itself a substantive result limiting any full-chain equilibrium claim at product level.**"
    important = row["product"] in {"milk", "butter", "cheese"} and row["system_name"] == "full_chain"
    positive = row["status"] == "ok"
    return result, interp, classify_highlight(important, positive)


def vecm_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "vecm_systems.csv"):
        result, interp, highlight = vecm_interpretation(row)
        out.append(
            {
                "Scope": f"{row['product_label']} | {row['system_name']}",
                "Model": "VECM feasibility",
                "Result": result,
                "Interpretation": interp,
                "__highlight__": highlight,
            }
        )
    return out


def reliability_rows() -> List[Dict[str, str]]:
    out = []
    for row in csv_rows(OUTPUT / "chapter_tables" / "table_6_1_model_reliability.csv"):
        result = (
            f"{row['model_block']} | {row['family']}: rows {row['rows_estimated']}, "
            f"reliable {row['reliable_rows']}, conditional {row['conditionally_usable_rows']}."
        )
        reliable = float(row["reliable_rows"] or 0)
        conditional = float(row["conditionally_usable_rows"] or 0)
        if reliable > 0:
            interp = "**This block contributes retained evidence and can enter the final interpretation hierarchy.**"
        elif conditional > 0:
            interp = "**This block is usable only conditionally, so it supports interpretation but should not anchor the main claim on its own.**"
        else:
            interp = "**This block does not contribute retained evidence in the final hierarchy, which is important for honesty about model strength.**"
        important = row["family"] in {"ECM", "NARDL", "VECM", "ARDL"}
        positive = reliable > 0
        out.append(
            {
                "Scope": row["model_block"],
                "Model": row["family"],
                "Result": result,
                "Interpretation": interp,
                "__highlight__": classify_highlight(important, positive),
            }
        )
    return out


FIGURE_NOTES = {
    "01_raw_government_layers.png": "This figure establishes the corrected national governmental price paths before any modelling transformation.",
    "02_raw_retail_observed_series.png": "This figure shows observed retail package prices and makes the downstream price object explicit.",
    "03_raw_external_benchmarks.png": "This figure places the Ukrainian dairy chain against European and CME benchmark dynamics.",
    "04_dataset_product_lines_and_indices.png": "This figure connects product-level series with aggregate dairy indices.",
    "05_aggregate_chain_indices.png": "This figure shows the latent chain-level dairy indices used as system robustness checks.",
    "06_panel_coverage.png": "This figure summarizes how much aligned support exists for each chain stage before estimation.",
    "07_cross_shop_match_status.png": "This figure shows how much of the retail universe can be matched across Silpo and Novus.",
    "07_retail_product_distribution.png": "This figure shows the product distribution across the two retailers.",
    "08_retail_brand_distribution.png": "This figure shows how brand support is distributed across Silpo and Novus.",
    "10_prozorro_region_profile.png": "This figure shows regional procurement concentration inside ProZorro.",
    "11_silpo_discount_environment.png": "This figure isolates the explicit Silpo markdown environment.",
    "12_weekly_chain_overlay.png": "This figure shows the weekly median chain paths for the main dairy products.",
    "13_farmgate_benchmark_block.png": "This figure compares farm-gate raw milk with chain-level dairy indices.",
    "01_weekly_corr_scan.png": "This figure summarizes the strongest lag-correlation signals before weekly model retention.",
    "02_link21_status_matrix.png": "This figure synthesizes the full 21-link directional model design.",
    "03_core_model_coefficients.png": "This figure compares retained weekly coefficients across model families and products.",
    "04_ecm_speed_of_adjustment.png": "This figure shows how quickly retained ECM equations move back toward equilibrium.",
    "05_nardl_asymmetry.png": "This figure shows asymmetric long-run evidence in retained weekly NARDL models.",
    "06_lp_pass_through_horizons.png": "This figure summarizes daily pass-through by response horizon.",
    "07_forward_reverse_core_share.png": "This figure compares forward and reverse daily timing evidence.",
    "08_candidate_downstream_core_share.png": "This figure compares downstream endpoint candidates on daily core evidence.",
    "08_spread_levels.png": "This figure shows average vertical spread levels across chain segments.",
    "09_spread_volatility.png": "This figure shows volatility differences across chain spreads.",
    "10_vecm_feasibility.png": "This figure reports the feasibility boundary for VECM system estimation.",
    "10_vertical_spread_proxy.png": "This figure summarizes the market-power proxy evidence from spread equations.",
    "11_aggregate_index_overlay.png": "This figure overlays the aggregate dairy indices used in the system-robustness block.",
    "11_discount_incidence.png": "This figure shows discount incidence by product where markdown states are observed.",
    "12_discount_coefficient_map.png": "This figure summarizes the retained discount-model coefficients.",
    "13_procurement_scale_effects.png": "This figure visualizes the procurement-scale coefficient block.",
    "14_aggregate_index_model_coefficients.png": "This figure summarizes aggregate-index model coefficients.",
    "15_reliability_overview.png": "This figure shows the reliability distribution across model blocks.",
}


def figure_groups() -> List[Tuple[str, List[Path]]]:
    groups: List[Tuple[str, List[Path]]] = []

    root_files = sorted([p for p in FIG_ROOT.glob("*.png") if p.is_file()])
    if root_files:
        groups.append(("Project-level figure inventory", root_files))

    for subdir in sorted([p for p in FIG_ROOT.iterdir() if p.is_dir()]):
        pngs = sorted(subdir.glob("*.png"))
        if not pngs:
            continue
        if subdir.name == "chapter5_data":
            label = "Chapter 5 figures"
        elif subdir.name == "chapter6_results":
            label = "Chapter 6 figures"
        elif subdir.name == "sequence":
            label = "Execution-sequence figures"
        elif subdir.name == "model_specific":
            label = "Model-specific figures"
        elif subdir.name == "reliability":
            label = "Reliability figures"
        else:
            label = f"{subdir.name} figures"
        groups.append((label, pngs))
    return groups


def generate_docx() -> None:
    doc = Document()
    frp._apply_kse_style(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Step-by-Step Results of FINAL_RESEARCH")
    set_run_style(run, bold=True, size=16)

    add_body_paragraph(
        doc,
        "This document consolidates the interpretable final outputs of the dairy price transmission project. "
        "It includes retained and non-retained results, but it distinguishes them explicitly through short interpretation sentences and color logic. "
        "Green cells mark the most important retained or supportive results. Red cells mark the most important non-significant or infeasible results in areas that are central to the thesis question."
    )

    sections = [
        ("1. Model reliability and diagnostic hierarchy", reliability_rows()),
        ("2. Weekly pairwise transmission across the 21 directional links", pairwise_rows()),
        ("3. Weekly retained core-chain model rows", weekly_model_rows()),
        ("4. Daily local-projection results", lp_rows()),
        ("5. Spread and market-power proxy results", margin_rows()),
        ("6. Silpo discount-strategy results", discount_rows()),
        ("7. Procurement-scale results", scale_rows()),
        ("8. Aggregate dairy-index robustness results", index_rows()),
        ("9. VECM and system-feasibility results", vecm_rows()),
    ]

    columns = ["Scope", "Model", "Result", "Interpretation"]
    for heading, rows in sections:
        add_section_heading(doc, heading, level=1)
        add_table(doc, heading, columns, rows)

    add_section_heading(doc, "10. Figures used in FINAL_RESEARCH", level=1)
    add_body_paragraph(
        doc,
        "The figures below are included in the same order as the final research logic: first data and transformed datasets, then model evidence, then integrated reliability visuals. "
        "Each figure is followed by a one-sentence reliability-oriented note."
    )

    for chapter_name, images in figure_groups():
        add_section_heading(doc, chapter_name, level=2)
        for img in images:
            doc.add_picture(str(img), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, str(img.relative_to(FIG_ROOT)))
            note = FIGURE_NOTES.get(img.name, "This figure forms part of the final empirical evidence set and should be read together with the chapter interpretation.")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(note)
            set_run_style(r, bold=True, size=11)

    doc.save(OUT_DOCX)


def generate_markdown() -> None:
    lines: List[str] = [
        "# Step-by-Step Results of FINAL_RESEARCH",
        "",
        "This document consolidates the interpretable final outputs of the dairy price transmission project. Green highlights in the DOCX mark the most important retained or supportive results. Red highlights mark the most important non-significant or infeasible results in central parts of the thesis.",
        "",
    ]
    sections = [
        ("Model reliability and diagnostic hierarchy", reliability_rows()),
        ("Weekly pairwise transmission across the 21 directional links", pairwise_rows()),
        ("Weekly retained core-chain model rows", weekly_model_rows()),
        ("Daily local-projection results", lp_rows()),
        ("Spread and market-power proxy results", margin_rows()),
        ("Silpo discount-strategy results", discount_rows()),
        ("Procurement-scale results", scale_rows()),
        ("Aggregate dairy-index robustness results", index_rows()),
        ("VECM and system-feasibility results", vecm_rows()),
    ]
    for title, rows in sections:
        lines += [f"## {title}", ""]
        for row in rows:
            lines.append(f"1. **{row['Scope']} | {row['Model']}**")
            lines.append(f"Result: {row['Result']}")
            lines.append(f"{row['Interpretation']}")
            lines.append("")

    lines += ["## Figures used in FINAL_RESEARCH", ""]
    for chapter_name, images in figure_groups():
        lines += [f"### {chapter_name}", ""]
        for img in images:
            lines.append(f"![{img.name}]({img})")
            lines.append(f"**{FIGURE_NOTES.get(img.name, 'This figure forms part of the final empirical evidence set and should be read together with the chapter interpretation.')}**")
            lines.append("")

    OUT_MD.write_text("\n".join(lines), encoding="utf-8")
    frp.write_html(OUT_HTML, "Step-by-Step Results of FINAL_RESEARCH", lines)


def improve_chapter_conclusion() -> None:
    text = CHAPTER_MD.read_text(encoding="utf-8")
    replacement = """### 6.0 Conclusion

The final result of the empirical rebuild is not a claim that every segment of the dairy chain behaves in the same way. It is a more disciplined statement about where transmission is structurally strongest, where it weakens, and where it becomes strategically managed. On the corrected data, the most defensible long-run transmission still appears before the retail stage, especially between producers and procurement. That is the part of the chain where price meaning is closest, overlap is cleaner, and weekly equilibrium-style estimation remains most credible.

At the same time, the absence of equally strong downstream weekly evidence should not be read as evidence of no downstream transmission. The daily mechanism models show why. Once price pressure reaches procurement and retail, adjustment becomes more tactical. Procurement does not simply forward processor prices. It filters them through contract timing, expected value, and the current and initial contract sums. Retail does not simply add a constant markup. It decides how quickly pressure becomes visible to the consumer, whether it is softened through markdowns, and whether different categories are adjusted in the same way or selectively.

This is precisely why the corrected Silpo discount interpretation matters. Observed shelf price is the post-discount price that the consumer actually sees, whereas discount incidence and discount depth form a separate behavioural layer. The final evidence therefore supports a market-power reading in which downstream influence is visible not only through price levels, but also through the timing and form of adjustment. In the strongest cases, retailers appear to smooth, delay, or selectively expose pass-through rather than simply mirror upstream movement.

The aggregate dairy-index block reinforces rather than overturns this conclusion. It shows that the broad transmission narrative survives aggregation, but it also confirms the limits of a single full-chain system estimate. Product-level VECM remains infeasible on the corrected overlap, and that infeasibility is itself informative. It indicates that the Ukrainian dairy chain is integrated enough to show staged transmission, but not clean enough to justify a single unified equilibrium claim across all products and all downstream stages.

For the broader thesis argument, this matters economically. The Ukrainian food consumer market, and especially its dairy segment, does not pass shocks from the farm to the shelf mechanically. Raw-milk pressure is visible at the farm-gate benchmark. Producer prices are the first clear carrier of the shock into processed-product markets. Procurement then acts as both relay and institutional filter. Retail finally governs how much of this pressure becomes household-visible and how much is temporarily absorbed through discounts, category management, and relative-price strategy. The conclusion is therefore stronger than a simple pass-through finding: the chain is vertically connected, but the transmission mechanism is staged, selective, and shaped by downstream market structure.

The remaining limitations are also substantive rather than cosmetic. Strict weekly overlap remains scarce, especially late in the chain. Regional procurement evidence is richer than regional evidence elsewhere, which means spatial heterogeneity can be discussed more confidently for ProZorro than for the corrected governmental averages. Retail evidence remains richer for Silpo than for Novus because explicit markdown states are observed directly only in Silpo. These limits do not invalidate the final results, but they define the boundary of what can be claimed without overstatement.

The most valuable next step is therefore not to abandon the current design, but to deepen it. Longer downstream overlap, richer quantity or expenditure weights, broader retailer coverage, and more stable regional downstream series would strengthen the full-chain system block and make the market-power interpretation even sharper. Even without those additions, however, the corrected final empirical system now supports a coherent and academically defensible conclusion: price shocks in Ukraine’s dairy market do travel vertically, but their speed, symmetry, and consumer visibility depend on where in the chain the adjustment occurs and on which institution controls that stage of price formation.
"""
    text = re.sub(r"### 6\.0 Conclusion[\s\S]*$", replacement, text)
    CHAPTER_MD.write_text(text, encoding="utf-8")
    lines = text.splitlines()
    frp.write_html(CHAPTER_HTML, "Chapter 5 and 6 Analysis", lines)
    frp.write_kse_docx(CHAPTER_DOCX, lines)


def main() -> None:
    generate_docx()
    generate_markdown()
    improve_chapter_conclusion()
    print(f"Created: {OUT_DOCX}")
    print(f"Created: {OUT_MD}")
    print(f"Created: {OUT_HTML}")
    print(f"Updated chapter: {CHAPTER_DOCX}")


if __name__ == "__main__":
    main()
