#!/usr/bin/env python3
"""Build a corrected full thesis draft from Draft 2 formatting.

This builder deliberately avoids LibreOffice/soffice because the local soffice
binary crashes during conversion. Validation is done through the DOCX ZIP
package, python-docx reopening, embedded-image checks, and a text QA pass.
"""

from __future__ import annotations

import json
import math
import os
import re
import textwrap
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image
import matplotlib.pyplot as plt


SCRIPT_DIR = Path(__file__).resolve().parent
PACKAGE_ROOT = SCRIPT_DIR.parent
LEGACY_ROOT = Path("/Users/getapple/Documents/KSE/Master Thesis")


def resolve_root() -> Path:
    """Prefer the portable GitHub package; allow THESIS_ROOT for local reruns."""
    env_root = os.environ.get("THESIS_ROOT")
    if env_root:
        return Path(env_root).expanduser().resolve()
    if (PACKAGE_ROOT / "data" / "Newmodel_data" / "newmodel.xlsx").exists():
        return PACKAGE_ROOT
    return LEGACY_ROOT


ROOT = resolve_root()
DATA_ROOT = ROOT / "data" / "Newmodel_data" if (ROOT / "data" / "Newmodel_data").exists() else ROOT / "Newmodel_data"
SOURCE_DOCS = ROOT / "doc" / "source" if (ROOT / "doc" / "source").exists() else ROOT
DRAFT2 = SOURCE_DOCS / "Maksym_Charniuk_MSc_thesis_draft_2.docx"
COMMENTED = SOURCE_DOCS / "Commented_draft2.docx"
TRANSCRIPT = SOURCE_DOCS / "Nivievskyi_5_05_transcript.docx"
NEWMODEL = DATA_ROOT / "newmodel.xlsx"
V2 = ROOT / "outputs" / "newmodel_deep_rebuild_v2"
V1 = ROOT / "outputs" / "newmodel_rebuild"
V3 = ROOT / "outputs" / "market_power_final_v3"
OUT = ROOT / "outputs" / "market_power_final_v10"
FIG_OUT = OUT / "figures"
TABLE_OUT = OUT / "tables"
REPORT_OUT = OUT / "reports"
DOC_OUT = ROOT / "doc"

DOCX_OUT = DOC_OUT / "Maksym_Charniuk_MSc_thesis_market_power_rewritten_v10.docx"


def ensure_dirs() -> None:
    for p in [OUT, FIG_OUT, TABLE_OUT, REPORT_OUT, DOC_OUT]:
        p.mkdir(parents=True, exist_ok=True)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def style_exists(doc: Document, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except Exception:
        return False


def add_p(doc: Document, text: str = "", style: str = "Normal", bold_first: bool = False):
    if style and not style_exists(doc, style):
        style = "Normal"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_first and ":" in text:
        first, rest = text.split(":", 1)
        run = p.add_run(first + ":")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_title(doc: Document, text: str) -> None:
    p = add_p(doc, text.upper(), style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_title(doc: Document, text: str) -> None:
    doc.add_page_break()
    add_title(doc, text)


def add_many(doc: Document, paras: list[str], style: str = "Normal") -> None:
    for para in paras:
        add_p(doc, para, style=style)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_p(doc, item, style="List Paragraph")


def clean_display_text(value: object) -> str:
    """Convert internal pipeline labels into thesis-facing wording."""
    if isinstance(value, float):
        return "" if math.isnan(value) else f"{value:.4g}"
    if pd.isna(value):
        return ""
    text = str(value)
    replacements = {
        "newmodel_v2_retail": "cleaned ProZorro-retail mechanism evidence",
        "newmodel_v1_ecm_ardl_nardl": "earlier high-frequency reconstruction package",
        "newmodel_v2": "observed monthly official data",
        "newmodel.xlsx": "the main integrated workbook",
        "outputs/newmodel_deep_rebuild_v2/clean_data": "the cleaned data output folder",
        "cleaned from newmodel.xlsx": "cleaned from the main integrated workbook",
        "observed newmodel data": "observed cleaned official data",
        "cleaned newmodel data": "cleaned observed data",
        "newmodel": "main integrated workbook",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def add_table(doc: Document, df: pd.DataFrame, title: str, note: str = "", max_rows: int = 12) -> None:
    add_p(doc, title, style="Heading 3")
    if df.empty:
        add_p(doc, "No observations available after the final reliability screen.")
        return
    view = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(view.columns))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for j, col in enumerate(view.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(view.columns):
            text = clean_display_text(row[col])
            cells[j].text = text[:140]
            for paragraph in cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    if note:
        add_p(doc, clean_display_text(note), style="source")


def add_picture(doc: Document, path: Path, caption: str, source: str, width: float = 6.3) -> None:
    if not path.exists():
        add_p(doc, f"[Missing figure: {path.name}]", style="figures")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = add_p(doc, caption, style="figures")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src = add_p(doc, source, style="source")
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_formula(doc: Document, number: int, title: str, formula: str, variables: str) -> None:
    """Insert a numbered formula box with a separate variable-description row."""
    title = f"Formula ({number}). {title}"
    add_p(doc, title, style="Heading 3")
    table = doc.add_table(rows=2, cols=1)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    formula_cell = table.rows[0].cells[0]
    formula_cell.text = formula
    variables_cell = table.rows[1].cells[0]
    variables_cell.text = f"Variables: {variables}"
    for p in formula_cell.paragraphs:
        for r in p.runs:
            r.font.name = "Cambria Math"
            r.font.size = Pt(10)
    for p in variables_cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)


def polish_literature_text(text: str) -> str:
    """Keep Draft 2 literature content, but repair citation punctuation and Chicago author-date form."""
    replacements = {
        "The literature relevant to this thesis follows the logic": (
            "The literature relevant to this thesis follows the logic of Nivievskyi and von Cramon-Taubadel’s supply-chain work. "
            "It first explains chain organization and bottlenecks, and only then turns to transmission metrics and retail pricing. "
            "In the Ukrainian dairy market, coefficients become meaningful only after the market architecture and institutional frictions are explicit (Nivievskyi 2012; Nivievskyi and von Cramon-Taubadel 2015)."
        ),
        "The literature review first examines vertical coordination": (
            "The literature review first examines vertical coordination and buyer power, then price transmission and asymmetry, then retail pricing and promotions, and finally transition-economy instability together with the econometric tools used for non-stationary chain data. "
            "The chain is defined first, shock transmission second, and model choice only after the relevant relationships are clear."
        ),
        "Prices in dairy are not formed in one market": (
            "Prices in dairy are not formed in one market. Farms sell raw milk, processors transform it into final dairy products, and retailers determine shelf access, final prices, and promotional regimes. "
            "For this reason, price transmission is tied to vertical coordination, because coordination shapes risk allocation and bargaining power along the chain. "
            "Martinez’s evidence from United States livestock and poultry markets shows that contracts and integration often replace spot exchange to stabilize supply and solve quality and transaction-cost problems. At the same time, they can push bargaining power toward downstream actors that control market access and commercial conditions in front of consumers (Martinez 2002; USDA Economic Research Service 1999)."
        ),
        "Taken together wider range of literature": (
            "A wider range of literature on vertical coordination reaches the same conclusion: downstream control over assets, information, or market access can generate efficiency gains for consumers, but it can also reshape surplus distribution (Ba et al. 2019). "
            "This is especially relevant for dairy because it is a perishable supply chain. The buyer with shelf access can shape adjustment not only through price, but also through collection terms, processing capacity, quality requirements, and contractual timing (Sexton 2013)."
        ),
        "Perishability is critical factor": (
            "Perishability is critical because raw milk spoils quickly and quality is distributed unevenly across suppliers. Processors can work efficiently only with stable, regular, and traceable inflows, so price formation depends on logistics and chain organization as much as on supply and demand. "
            "The Food and Agriculture Organization of the United Nations (FAO) describes these cold-chain and quality constraints as central to dairy supply-chain performance (Food and Agriculture Organization of the United Nations 2013). "
            "In Ukraine, household and industrial farms enter the formal chain on unequal terms. The procurement question is therefore not only price, but also who can supply milk with the scale, cooling, traceability, and compliance capacity required by formal processors (Litvinov 2025a)."
        ),
        "For the Ukrainian case, Nivievskyi’s earlier work": (
            "For the Ukrainian case, Nivievskyi’s earlier work is especially valuable because it connects general vertical-coordination theory to the institutional realities of the local dairy sector. "
            "Across this work, competitiveness is shaped jointly by farm structure, quality constraints, contracts, trade exposure, and policy settings, not by price incentives alone (Nivievskyi and von Cramon-Taubadel 2008; Nivievskyi 2009; Nivievskyi 2012; Nivievskyi and von Cramon-Taubadel 2015). "
            "This thesis follows that view: it treats the dairy chain first as a coordination problem and only then as a transmission problem."
        ),
        "For the further understanding of the topic": (
            "For the purposes of this thesis, price transmission is used to assess how price or cost changes at one chain stage are reflected at other stages. "
            "Vavra and Goodwin’s OECD survey shows that transmission has both long-run and short-run dimensions and that observed patterns depend on technology, margins, market structure, and institutions, not only on concentration (Vavra and Goodwin 2005). "
            "Evidence of market power is therefore not simply a gap between upstream and downstream prices. The empirical question is how the system returns to equilibrium after a shock and which stage buffers the adjustment."
        ),
        "Meyer and von Cramon-Taubadel make an important correction": (
            "Meyer and von Cramon-Taubadel make an important correction: asymmetric adjustment is often observed, but it is easy to overstate when the econometric form does not match the actual adjustment process (Meyer and von Cramon-Taubadel 2004). "
            "The broader agricultural-pricing literature reaches a similar conclusion from another angle. Prices along the food chain do not reflect farm value alone; by the time a product reaches the consumer, it also includes processing, storage, transport, packaging, and retail services, each of which can alter the apparent transmission pattern (Koester and von Cramon-Taubadel 2023)."
        ),
        "One more important topic is the classic": (
            "The classic “rockets and feathers” finding of Kinnucan and Forker remains influential because it shows that retail dairy prices can rise faster than they fall after farm prices decline (Kinnucan and Forker 1987). "
            "At the same time, Kinnucan and Zhang caution against reading farm-retail coefficients too literally when stages are built from differently aggregated data and carry different marketing-cost structures; in that case, the coefficient may capture more than pure transmission alone (Kinnucan and Zhang 2015)."
        ),
        "That warning is extremely relevant here": (
            "That warning is directly relevant here because farm-gate, processor, procurement, and retail prices are not prices for the same physical object. "
            "Processing transforms raw milk into differentiated products, so asymmetry must be interpreted as a question about adjustment within a chain of related but non-identical products, not as a direct raw-milk-to-shelf comparison."
        ),
        "Retail pricing introduces another layer of complexity": (
            "Retail pricing introduces another layer of complexity because the shelf is a managed commercial space, not a passive reflection of upstream costs. "
            "In concentrated, multiproduct, promotion-heavy grocery retail, consumer prices may move through baseline repricing, temporary discounts, assortment shifts, or packaging changes (Loy 2023; Koester and von Cramon-Taubadel 2023). "
            "Apparent retail rigidity may therefore reflect bargaining power, but it may also reflect menu costs, consumer search, stock-out risk, category management, or product-specific promotion schedules."
        ),
        "The promotions literature is especially critical here": (
            "The promotions literature is especially important because markdowns are part of the pricing mechanism itself, not only statistical noise. "
            "Retailers may hold the baseline price while changing the effective consumer price through temporary discounts, and menu-cost models show why baseline prices may remain rigid even when market conditions change (Chevalier, Kashyap, and Rossi 2003; Nakamura and Steinsson 2008). "
            "This thesis therefore treats promotions as a mechanism of adjustment, while avoiding the claim that every promotion directly reveals market power."
        ),
        "Instability is reason for an additional complication": (
            "Instability adds another complication because price transmission can become regime-dependent under policy uncertainty, trade disruption, and war. "
            "In Ukraine, the dairy sector has faced repeated trade shocks, regulatory adjustment, and wartime disruption. Evidence from the Ukrainian wheat-flour chain shows that political and economic instability can coincide with changes in transmission itself, either strengthening, weakening, or distorting adjustment patterns (Brümmer, von Cramon-Taubadel, and Zorya 2009; Litvinov 2025a). "
            "Dairy is a different food chain, but the methodological warning is the same: transmission should not be assumed stable or frictionless across all periods."
        ),
        "The theoretical methodological foundations": (
            "The methodological foundations of this thesis follow the standard treatment of non-stationary price systems. "
            "Engle and Granger established the cointegration and error-correction framework, Johansen extended the logic to systems of linked variables, Pesaran, Shin, and Smith developed the autoregressive distributed lag approach for mixed integration orders, and Enders and Siklos together with Shin, Yu, and Greenwood-Nimmo formalized nonlinear and asymmetric adjustment (Engle and Granger 1987; Johansen 1988; Pesaran, Shin, and Smith 2001; Enders and Siklos 2001; Shin, Yu, and Greenwood-Nimmo 2014). "
            "The dynamic question is whether prices share a long-run relationship, how quickly disequilibria are corrected, and whether positive and negative shocks are corrected differently."
        ),
        "Taken together, all these literatures have a blank spot": (
            "Taken together, these literatures leave a clear gap. There is limited Ukrainian dairy evidence that combines official monthly chain prices, ProZorro procurement records, SKU-level retail observations, discount mechanisms, and a reliability-ranked model hierarchy under one market-power question. "
            "This thesis addresses that gap by keeping the object narrow: market power in the Ukrainian dairy value chain, studied through price-transmission evidence."
        ),
    }
    for prefix, replacement in replacements.items():
        if text.startswith(prefix):
            return replacement
    return text


def extract_literature_review() -> list[tuple[str, str]]:
    """Preserve the Draft 2 literature review content and styles."""
    src = Document(DRAFT2)
    starts = {}
    for i, par in enumerate(src.paragraphs):
        text = par.text.strip().upper()
        if text in {
            "INTRODUCTION",
            "MARKET ANALYSIS AND INSTITUTIONAL BACKGROUND",
            "LITERATURE REVIEW",
            "METHODOLOGY",
            "DATA",
            "ESTIMATION RESULTS",
            "CONCLUSIONS",
            "WORKS CITED",
        }:
            starts[text] = i
    out = []
    for par in src.paragraphs[starts["LITERATURE REVIEW"] + 1 : starts["METHODOLOGY"]]:
        text = par.text.strip()
        if text:
            if text.startswith("The frequency difference across the dataset creates another issues"):
                text = (
                    "The frequency differences across the datasets create a methodological challenge. "
                    "Because retail prices are available only for a short period, synthetic daily and weekly datasets are used only as supporting material and are reconciled with monthly information through Chow-Lin and Denton-type procedures (Chow and Lin 1971; Denton 1971). "
                    "External benchmarks also matter: European Union dairy-price monitoring and Chicago Mercantile Exchange (CME) Class III milk signals help separate chain-specific transmission from common external shocks. "
                    "In a period of exchange-rate volatility and changing trade routes, that distinction is necessary when downstream rigidity must be distinguished from domestic bargaining power (European Commission Directorate-General for Agriculture and Rural Development 2026)."
                )
            text = polish_literature_text(text)
            out.append((text, par.style.name if par.style else "Normal"))
    return out


def extract_works_cited() -> list[str]:
    """Recover Draft 2 references so preserved literature citations remain backed."""
    src = Document(DRAFT2)
    starts = {}
    for i, par in enumerate(src.paragraphs):
        text = par.text.strip().upper()
        if text in {
            "WORKS CITED",
            "APPENDICES",
            "APPENDIX",
        }:
            starts[text] = i
    if "WORKS CITED" not in starts:
        return []
    end_candidates = [v for k, v in starts.items() if k != "WORKS CITED" and v > starts["WORKS CITED"]]
    end = min(end_candidates) if end_candidates else len(src.paragraphs)
    refs = []
    for par in src.paragraphs[starts["WORKS CITED"] + 1 : end]:
        text = par.text.strip()
        if text:
            refs.append(text)
    return refs


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, engine="python")


def load_evidence() -> dict[str, pd.DataFrame]:
    tables = V2 / "tables"
    clean = V2 / "clean_data"
    return {
        "selection": read_csv(tables / "model_selection_H1_H2_v2.csv"),
        "register": read_csv(tables / "integrated_evidence_register_v2.csv"),
        "validation": read_csv(tables / "old_new_dataset_validation_v2.csv"),
        "workbook_audit": read_csv(tables / "newmodel_workbook_audit_v2.csv"),
        "retail_audit": read_csv(tables / "retail_classification_repair_audit_v2.csv"),
        "loy_first": read_csv(tables / "loy_style_first_stage_v2.csv"),
        "loy_second": read_csv(tables / "loy_style_second_stage_v2.csv"),
        "additional": read_csv(tables / "additional_models_dols_lp_threshold_v2.csv"),
        "farm": read_csv(clean / "clean_farmgate_monthly_ua_region_v2.csv"),
        "proc": read_csv(clean / "clean_processor_monthly_ua_v2.csv"),
        "consumer": read_csv(clean / "clean_consumer_monthly_ua_region_v2.csv"),
        "prozorro": read_csv(clean / "clean_prozorro_lot_level_v2.csv"),
        "retail": read_csv(clean / "clean_retail_sku_day_title_repaired_v2.csv"),
    }


def dataset_table(ev: dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = [
        {
            "thesis name": "farm-gate raw milk price",
            "raw source name": "Середня ціна продукції сільського господарства, реалізованої підприємствами; product row: Молоко",
            "source": "State Statistics Service of Ukraine (SSSU), agricultural sales dataset",
            "level": "farm-gate",
            "frequency": "monthly",
            "period": "2015-M01 to 2026-M03",
            "unit after audit": "Ukrainian hryvnia per kilogram (UAH/kg)",
            "role": "main H1 upstream variable",
        },
        {
            "thesis name": "processor-level dairy price",
            "raw source name": "Середні ціни виробників промислової продукції; commodity rows: butter, hard cheese, processed milk, kefir, skimmed milk powder, sour cream",
            "source": "SSSU industrial producer-price dataset",
            "level": "processor",
            "frequency": "monthly",
            "period": "2013-M01 to 2026-M03",
            "unit after audit": "UAH/kg",
            "role": "main H1 downstream and H2 upstream variable",
        },
        {
            "thesis name": "official consumer dairy price",
            "raw source name": "Середні споживчі ціни на товари (послуги); rows: pasteurized milk, sour cream, soft cheese",
            "source": "SSSU consumer-price dataset",
            "level": "official consumer",
            "frequency": "monthly",
            "period": "2017-M01 to 2026-M03",
            "unit after audit": "UAH/kg",
            "role": "main H2 official downstream benchmark",
        },
        {
            "thesis name": "ProZorro procurement lot price",
            "raw source name": "ProzorroM(full): Товар, Кількість, Одиниця виміру, Ціна за одиницю, Регіон організатора",
            "source": "ProZorro public procurement records",
            "level": "institutional procurement",
            "frequency": "lot/event; aggregated to day, week, month",
            "period": "2023-01-02 to 2026-04-29",
            "unit after audit": "UAH/kg where unit and quantity allow reliable conversion",
            "role": "H1 and H2 procurement mechanism; not a national farm or retail price",
        },
        {
            "thesis name": "Silpo and Novus retail price",
            "raw source name": "product_title, product_name, price_current, unit_price, discount fields",
            "source": "retail web-scraped observations",
            "level": "retail",
            "frequency": "Stock Keeping Unit (SKU)-day",
            "period": "2025-10-21 to 2026-01-08",
            "unit after audit": "UAH/kg if package conversion is reliable; otherwise package-price mechanism only",
            "role": "H2 short-window retail and discount mechanism",
        },
    ]
    return pd.DataFrame(rows)


def audit_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "audit item": "Processor unit",
                "risk": "Processor_price is stored as Гривні за тону.",
                "correction": "Divide by 1000 and report as UAH/kg.",
                "status": "mandatory before modeling",
            },
            {
                "audit item": "Farm-gate unit",
                "risk": "Farm_price uses the short unit label Гривня, but magnitudes are tonne-level prices.",
                "correction": "Treat as UAH/tonne, divide by 1000, and label as UAH/kg.",
                "status": "mandatory before modeling",
            },
            {
                "audit item": "ProZorro numeric fields",
                "risk": "ProzorroM(full) includes text numbers and non-breaking spaces.",
                "correction": "Strip non-breaking spaces, normalize decimal signs, and recalculate reliable UAH/kg lot prices.",
                "status": "mandatory before aggregation",
            },
            {
                "audit item": "Retail classification",
                "risk": "Legacy product labels misclassify some dairy and non-dairy items.",
                "correction": "Rebuild product from product_title and product_name; keep a controlled product column in every table.",
                "status": "mandatory before H2 retail models",
            },
            {
                "audit item": "Integrated consumer sheet",
                "risk": "Combined Consumer_price sheet contains implausible values above 1000 UAH/kg.",
                "correction": "Use the raw Consumer_2017_2026 component workbook and cleaned national/region series.",
                "status": "mandatory before official consumer models",
            },
        ]
    )


def evidence_reading(row: pd.Series) -> str:
    """Human-readable reliability reading for result tables."""
    beta_p = pd.to_numeric(row.get("p(β)"), errors="coerce")
    ect = pd.to_numeric(row.get("λ error-correction"), errors="coerce")
    ect_p = pd.to_numeric(row.get("p(λ)"), errors="coerce")
    coint_p = pd.to_numeric(row.get("cointegration p"), errors="coerce")
    if pd.notna(ect) and ect >= 0:
        return "price link only; adjustment sign is weak"
    if pd.notna(ect) and ect < 0 and pd.notna(ect_p) and ect_p < 0.10 and pd.notna(coint_p) and coint_p < 0.10:
        return "price link and adjustment evidence"
    if pd.notna(coint_p) and coint_p < 0.10 and pd.notna(beta_p) and beta_p < 0.05:
        return "strong price link; adjustment weak or slow"
    if pd.notna(beta_p) and beta_p < 0.05:
        return "price link only; long-run diagnostics weak"
    return "diagnostics too weak for headline use"


def selected_tables(ev: dict[str, pd.DataFrame]) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    sel = ev["selection"].copy()
    # Strict thesis rule: newmodel national observed monthly data first.
    h1 = sel[
        (sel["hypothesis"] == "H1")
        & (sel["source_block"] == "newmodel_v2")
        & (sel["link"] == "FarmGate -> Processor")
    ].copy()
    h1 = h1.sort_values(["integrated_reliability", "product"])
    h1_table = h1[
        [
            "link",
            "product",
            "model_family",
            "n_obs",
            "coef",
            "pvalue",
            "ect_coef",
            "ect_pvalue",
            "cointegration_p",
            "diagnostic_ljungbox_p",
            "integrated_reliability",
        ]
    ].rename(
        columns={
            "coef": "β long-run pass-through",
            "pvalue": "p(β)",
            "ect_coef": "λ error-correction",
            "ect_pvalue": "p(λ)",
            "cointegration_p": "cointegration p",
            "diagnostic_ljungbox_p": "serial-correlation p",
            "integrated_reliability": "use",
        }
    )
    if not h1_table.empty:
        h1_table.loc[h1_table["product"].isin(["butter", "skim_milk_powder"]), "use"] = "probable / supporting"
        h1_table.loc[h1_table["product"].isin(["kefir"]), "use"] = "probable / supporting"
        h1_table["evidence reading"] = h1_table.apply(evidence_reading, axis=1)
    h2 = sel[
        (sel["hypothesis"] == "H2")
        & (sel["source_block"] == "newmodel_v2")
        & (sel["link"].isin(["Processor -> Consumer", "ProZorro -> Consumer"]))
    ].copy()
    h2 = h2.sort_values(["link", "product"])
    h2_table = h2[
        [
            "link",
            "product",
            "model_family",
            "n_obs",
            "coef",
            "pvalue",
            "ect_coef",
            "ect_pvalue",
            "cointegration_p",
            "diagnostic_ljungbox_p",
            "integrated_reliability",
        ]
    ].rename(
        columns={
            "coef": "β long-run pass-through",
            "pvalue": "p(β)",
            "ect_coef": "λ error-correction",
            "ect_pvalue": "p(λ)",
            "cointegration_p": "cointegration p",
            "diagnostic_ljungbox_p": "serial-correlation p",
            "integrated_reliability": "use",
        }
    )
    if not h2_table.empty:
        h2_table.loc[
            (h2_table["link"] == "Processor -> Consumer") & (h2_table["product"] == "drinking_milk"),
            "use",
        ] = "probable / supporting"
        h2_table.loc[
            h2_table["link"].str.contains("ProZorro", na=False),
            "use",
        ] = "probable / supporting"
        h2_table["evidence reading"] = h2_table.apply(evidence_reading, axis=1)
    retail = sel[
        (sel["hypothesis"] == "H2")
        & (sel["link"].str.contains("ProZorro -> Retail", na=False))
        & (sel["source_block"] == "newmodel_v2_retail")
        & (sel["integrated_reliability"] != "appendix or discard")
    ].copy()
    retail = retail.sort_values(["source_block", "product", "n_obs"], ascending=[True, True, False])
    retail_table = retail[
        [
            "source_block",
            "link",
            "product",
            "model_family",
            "n_obs",
            "coef",
            "pvalue",
            "integrated_reliability",
            "thesis_use",
        ]
    ].head(12).rename(
        columns={
            "source_block": "source block",
            "coef": "coefficient",
            "pvalue": "p-value",
            "integrated_reliability": "use",
        }
    )
    validation = ev["validation"].copy()
    return h1_table, h2_table, retail_table, validation


def loy_mechanism_table(ev: dict[str, pd.DataFrame]) -> pd.DataFrame:
    loy = ev["loy_second"].copy()
    if loy.empty:
        return pd.DataFrame()
    out = loy[loy["term"].isin(["perishable", "retail_link"])].copy()
    out["interpretation"] = out.apply(
        lambda r: (
            "Retail-linked pairs show faster correction in the Loy-style first-stage screen."
            if r["dependent"] == "speed_measure" and r["term"] == "retail_link" and r["pvalue"] < 0.1 and r["coef"] > 0
            else "No statistically strong explanatory role in the current short retail/procurement sample."
        ),
        axis=1,
    )
    return out.rename(
        columns={
            "dependent": "first-stage measure",
            "term": "mechanism variable",
            "coef": "coefficient",
            "pvalue": "p-value",
            "n_obs": "observations",
        }
    )[["first-stage measure", "mechanism variable", "coefficient", "p-value", "observations", "interpretation"]]


def evidence_counts(ev: dict[str, pd.DataFrame]) -> pd.DataFrame:
    sel = ev["selection"]
    if sel.empty:
        return pd.DataFrame()
    return (
        sel.groupby(["hypothesis", "source_block", "integrated_reliability"])
        .size()
        .reset_index(name="models")
        .sort_values(["hypothesis", "source_block", "integrated_reliability"])
    )


def data_summary(ev: dict[str, pd.DataFrame]) -> pd.DataFrame:
    def fmt_num(x: float | int | None, digits: int = 2) -> str:
        if x is None or pd.isna(x):
            return ""
        return f"{float(x):,.{digits}f}"

    def period(df: pd.DataFrame) -> str:
        d = pd.to_datetime(df["date"], errors="coerce") if "date" in df.columns else pd.Series(dtype="datetime64[ns]")
        if not d.notna().any():
            return ""
        return f"{d.min().date()} to {d.max().date()}"

    def stat_row(
        *,
        layer: str,
        df: pd.DataFrame,
        price_col: str,
        price_label: str,
        unit: str = "UAH/kg",
        product_col: str = "product",
        coverage_note: str = "",
        actor_note: str = "",
    ) -> dict[str, object]:
        prices = pd.to_numeric(df.get(price_col, pd.Series(dtype=float)), errors="coerce").dropna()
        products = sorted(map(str, df.get(product_col, pd.Series(dtype=object)).dropna().unique()))
        return {
            "chain layer": layer,
            "analytical price": price_label,
            "unit": unit,
            "period": period(df),
            "observations": int(len(df)),
            "products": ", ".join(products[:8]),
            "mean": fmt_num(prices.mean()),
            "sd": fmt_num(prices.std()),
            "min-max": f"{fmt_num(prices.min())}-{fmt_num(prices.max())}" if len(prices) else "",
            "brands / companies / buyers": actor_note,
            "model role": coverage_note,
        }

    rows: list[dict[str, object]] = []

    farm = ev["farm"].copy()
    if not farm.empty:
        f = farm[(farm["quality_flag"] == "ok") & (farm["territory"] == "Україна") & (farm["farm_type"] == "enterprises")].copy()
        rows.append(
            stat_row(
                layer="farm-gate raw milk",
                df=f,
                price_col="price_uah_kg",
                price_label="farm-gate raw milk price",
                actor_note="enterprise farm-gate series; regional rows kept for dispersion only",
                coverage_note="main H1 upstream variable",
            )
        )

    proc = ev["proc"].copy()
    if not proc.empty:
        p = proc[(proc["quality_flag"] == "ok") & (proc["territory"] == "Україна")].copy()
        rows.append(
            stat_row(
                layer="processor-level sale",
                df=p,
                price_col="price_uah_kg",
                price_label="industrial producer price converted to processor-level price",
                actor_note="six SSSU processor product rows; firm names not observed",
                coverage_note="main H1 downstream and H2 upstream variable",
            )
        )

    consumer = ev["consumer"].copy()
    if not consumer.empty:
        c = consumer[(consumer["quality_flag"] == "ok") & (consumer["territory"] == "Україна")].copy()
        rows.append(
            stat_row(
                layer="official consumer sale",
                df=c,
                price_col="price_uah_kg",
                price_label="official average consumer price",
                actor_note="national consumer average; no brand or retailer identifiers",
                coverage_note="main H2 downstream benchmark",
            )
        )

    proz = ev["prozorro"].copy()
    if not proz.empty:
        z = proz[(proz["quality_flag"] == "ok") & (proz["product"].notna()) & (proz["product"] != "other_dairy")].copy()
        winners = z["Переможець"].nunique(dropna=True) if "Переможець" in z.columns else 0
        organizers = z["Організатор"].nunique(dropna=True) if "Організатор" in z.columns else 0
        rows.append(
            stat_row(
                layer="institutional procurement",
                df=z,
                price_col="price_uah_kg",
                price_label="repaired ProZorro lot unit price",
                actor_note=f"{winners:,} winning suppliers; {organizers:,} procuring organizers",
                coverage_note="H1/H2 procurement mechanism and wartime extension",
            )
        )

    retail = ev["retail"].copy()
    if not retail.empty:
        r = retail[(retail["quality_flag"] == "ok") & (retail["product"] != "exclude_non_dairy")].copy()
        kg_rows = r[pd.to_numeric(r.get("price_uah_kg_v2", pd.Series(dtype=float)), errors="coerce").notna()]
        brands = r["brand"].nunique(dropna=True) if "brand" in r.columns else 0
        skus = r["sku"].nunique(dropna=True) if "sku" in r.columns else 0
        retailers = ", ".join(sorted(map(str, r["retailer"].dropna().unique()))) if "retailer" in r.columns else ""
        retail_row = stat_row(
            layer="retail SKU-day sale",
            df=r,
                price_col="price_package_uah",
                price_label="observed package price; UAH/kg used only where reliable",
                unit="UAH/package; UAH/kg subset",
                actor_note=f"{brands:,} brands; {skus:,} SKUs; retailers: {retailers}",
                coverage_note=f"H2 retail/discount mechanism; {len(kg_rows):,} reliable UAH/kg rows",
            )
        rows.append(retail_row)

    return pd.DataFrame(rows)


def regional_facts(ev: dict[str, pd.DataFrame]) -> dict[str, float | int | str]:
    farm = ev["farm"].copy()
    cons = ev["consumer"].copy()
    facts = {}
    if not farm.empty:
        f = farm[(farm["quality_flag"] == "ok") & (farm["territory"] != "Україна")].copy()
        f["date"] = pd.to_datetime(f["date"], errors="coerce")
        last = f[f["date"] == f["date"].max()]
        facts["farm_regions"] = int(f["territory"].nunique())
        facts["farm_last_month"] = str(f["date"].max().date())
        facts["farm_last_min"] = float(last["price_uah_kg"].min())
        facts["farm_last_max"] = float(last["price_uah_kg"].max())
    if not cons.empty:
        c = cons[(cons["quality_flag"] == "ok") & (cons["territory"] != "Україна")].copy()
        c["date"] = pd.to_datetime(c["date"], errors="coerce")
        last = c[c["date"] == c["date"].max()]
        facts["consumer_regions"] = int(c["territory"].nunique())
        facts["consumer_last_month"] = str(c["date"].max().date())
        facts["consumer_last_min"] = float(last["price_uah_kg"].min())
        facts["consumer_last_max"] = float(last["price_uah_kg"].max())
    return facts


def trade_facts() -> dict[str, float]:
    facts = {}
    for kind, file in [
        ("exports", DATA_ROOT / "additioanl materials" / "share_dairy_exp.xlsx"),
        ("imports", DATA_ROOT / "additioanl materials" / "share_dairy_imp.xlsx"),
    ]:
        raw = pd.read_excel(file, header=None)
        if kind == "exports":
            row = raw[raw.astype(str).apply(lambda r: r.str.contains("Dairy EX", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            share = raw[raw.astype(str).apply(lambda r: r.str.contains("Share of dairy in animal-origin products", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            total = raw[raw.astype(str).apply(lambda r: r.str.contains("'TOTAL", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            facts["dairy_exports_2024_eur_m"] = float(row.iloc[11])
            facts["dairy_exports_animal_share_2024"] = float(share.iloc[11])
            facts["dairy_exports_total_share_2024"] = float(row.iloc[11]) / float(total.iloc[11])
        else:
            row = raw[raw.astype(str).apply(lambda r: r.str.contains("Dairy imports", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            share = raw[raw.astype(str).apply(lambda r: r.str.contains("Share of dairy in animal-origin products", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            total = raw[raw.astype(str).apply(lambda r: r.str.contains("'TOTAL", regex=False, case=False, na=False).any(), axis=1)].iloc[0]
            facts["dairy_imports_2024_eur_m"] = float(row.iloc[11])
            facts["dairy_imports_animal_share_2024"] = float(share.iloc[11])
            facts["dairy_imports_total_share_2024"] = float(row.iloc[11]) / float(total.iloc[11])
    return facts


def read_trade_inputs() -> tuple[pd.DataFrame, pd.DataFrame]:
    """Read trade-value/share and partner-flow workbooks used for Chapter 2."""
    base = DATA_ROOT / "additioanl materials"
    exp = pd.read_excel(base / "share_dairy_exp.xlsx", header=None)
    imp = pd.read_excel(base / "share_dairy_imp.xlsx", header=None)
    years = [int(x) for x in exp.iloc[13, 2:12]]

    def row(df: pd.DataFrame, label: str) -> list[float]:
        mask = df.astype(str).apply(lambda r: r.str.contains(label, regex=False, case=False, na=False).any(), axis=1)
        return [float(x) for x in df.loc[mask].iloc[0, 2:12]]

    trade = pd.DataFrame(
        {
            "year": years,
            "dairy_exports_eur_m": row(exp, "Dairy EX"),
            "dairy_imports_eur_m": row(imp, "Dairy imports"),
            "export_share_animal_origin_pct": [x * 100 for x in row(exp, "Share of dairy in animal-origin products")],
            "import_share_animal_origin_pct": [x * 100 for x in row(imp, "Share of dairy in animal-origin products")],
        }
    )
    trade["dairy_balance_eur_m"] = trade["dairy_exports_eur_m"] - trade["dairy_imports_eur_m"]

    raw = pd.read_excel(base / "biggest_trade.xlsx", sheet_name="Sheet1", header=None)
    partner_rows = []
    for _, r in raw.iloc[7:].iterrows():
        label = r.iloc[0]
        if not isinstance(label, str) or not label.strip():
            continue
        val = r.iloc[10]
        if pd.isna(val):
            continue
        flow = "Import from partner" if " IM" in label else "Export to partner" if "EX" in label else "Other"
        if flow == "Other":
            continue
        partner = label.replace(" IM", "").replace(" EX", "").strip()
        partner_rows.append({"partner": partner, "flow": flow, "value_2024_eur_m": float(val)})
    partners = pd.DataFrame(partner_rows)
    if not partners.empty:
        partners = partners.reindex(partners["value_2024_eur_m"].abs().sort_values(ascending=False).index).head(8)
    return trade, partners


def create_trade_figure() -> Path:
    """Create a Chapter 2 trade figure from share_dairy_exp, share_dairy_imp, and biggest_trade."""
    trade, partners = read_trade_inputs()
    trade.to_csv(TABLE_OUT / "trade_summary_v10.csv", index=False)
    partners.to_csv(TABLE_OUT / "trade_partner_flows_2024_v10.csv", index=False)

    out = FIG_OUT / "fig_09_dairy_trade_hs0401_0406.png"
    fig, axes = plt.subplots(2, 2, figsize=(13.8, 8.2))
    ax = axes[0, 0]
    ax.plot(trade["year"], trade["dairy_exports_eur_m"], marker="o", linewidth=2.4, label="Dairy exports")
    ax.plot(trade["year"], trade["dairy_imports_eur_m"], marker="o", linewidth=2.4, label="Dairy imports")
    ax.set_title("Dairy-only trade values")
    ax.set_ylabel("EUR million")
    ax.grid(True, alpha=0.25)
    ax.legend(frameon=False)

    ax = axes[0, 1]
    ax.plot(trade["year"], trade["export_share_animal_origin_pct"], marker="o", linewidth=2.4, label="Export share")
    ax.plot(trade["year"], trade["import_share_animal_origin_pct"], marker="o", linewidth=2.4, label="Import share")
    ax.set_title("Dairy share in animal-origin trade")
    ax.set_ylabel("Percent")
    ax.grid(True, alpha=0.25)
    ax.legend(frameon=False)

    ax = axes[1, 0]
    colors = ["#2a9d8f" if x >= 0 else "#d95f02" for x in trade["dairy_balance_eur_m"]]
    ax.bar(trade["year"], trade["dairy_balance_eur_m"], color=colors)
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Dairy trade balance")
    ax.set_ylabel("EUR million")
    ax.grid(True, axis="y", alpha=0.25)

    ax = axes[1, 1]
    if not partners.empty:
        p = partners.sort_values("value_2024_eur_m")
        colors = ["#d95f02" if v < 0 else "#2a9d8f" for v in p["value_2024_eur_m"]]
        ax.barh(p["partner"], p["value_2024_eur_m"], color=colors)
        ax.axvline(0, color="black", linewidth=0.8)
    ax.set_title("Largest Harmonized System 04 partner flows, 2024")
    ax.set_xlabel("EUR million; imports negative, exports positive")
    ax.grid(True, axis="x", alpha=0.25)

    fig.suptitle("Ukraine Dairy Trade Exposure and Partner Structure", fontsize=18, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(out, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out


def create_coefficients_figure(h1_table: pd.DataFrame, h2_table: pd.DataFrame) -> Path:
    """Create a clean coefficient figure from the final selected H1/H2 tables."""
    rows = []
    for _, r in h1_table.iterrows():
        rows.append(
            {
                "label": f"H1 Farm-gate -> processor: {str(r['product']).replace('_', ' ')}",
                "coef": float(r["β long-run pass-through"]),
                "use": r["use"],
                "block": "H1",
            }
        )
    for _, r in h2_table.iterrows():
        rows.append(
            {
                "label": f"H2 {str(r['link']).replace(' -> ', ' -> ')}: {str(r['product']).replace('_', ' ')}",
                "coef": float(r["β long-run pass-through"]),
                "use": r["use"],
                "block": "H2",
            }
        )
    df = pd.DataFrame(rows)
    df["rank"] = range(len(df))
    df.to_csv(TABLE_OUT / "selected_coefficients_v10.csv", index=False)

    colors = df.apply(
        lambda r: "#2a9d8f" if r["use"] == "main reliable" and r["block"] == "H1" else "#4575b4" if r["use"] == "main reliable" else "#8f8f8f",
        axis=1,
    )
    out = FIG_OUT / "fig_07_selected_coefficients_large.png"
    fig, ax = plt.subplots(figsize=(12.8, 7.0))
    y = range(len(df))
    ax.barh(y, df["coef"], color=list(colors))
    ax.axvline(1.0, color="black", linestyle="--", linewidth=1.0, alpha=0.7)
    ax.set_yticks(list(y))
    ax.set_yticklabels(df["label"], fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("Long-run pass-through coefficient")
    ax.set_title("Selected H1 and H2 Pass-Through Coefficients after Reliability Screening", fontsize=15, fontweight="bold")
    ax.grid(True, axis="x", alpha=0.25)
    for i, val in enumerate(df["coef"]):
        ax.text(val + 0.02, i, f"{val:.2f}", va="center", fontsize=8)
    ax.text(1.02, -0.7, "one-to-one benchmark", fontsize=8, color="black")
    fig.tight_layout()
    fig.savefig(out, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return out


def copy_figures() -> dict[str, Path]:
    figures = {
        "chain": [V3 / "figures" / "fig_01_value_chain_market_power.png", FIG_OUT / "fig_01_value_chain_market_power.png"],
        "h1_prices": [V3 / "figures" / "fig_02_h1_monthly_prices_large.png", FIG_OUT / "fig_02_h1_monthly_prices_large.png"],
        "h2_bridge": [V3 / "figures" / "fig_03_h2_processor_consumer_large.png", FIG_OUT / "fig_03_h2_processor_consumer_large.png"],
        "prozorro": [V3 / "figures" / "fig_04_prozorro_weekly_large.png", FIG_OUT / "fig_04_prozorro_weekly_large.png"],
        "discounts": [V3 / "figures" / "fig_05_retail_discounts_large.png", FIG_OUT / "fig_05_retail_discounts_large.png"],
        "reliability": [V3 / "figures" / "fig_06_reliability_screen_large.png", FIG_OUT / "fig_06_reliability_screen_large.png"],
        "regional": [V1 / "figures" / "fig_08_regional_farmgate_dispersion.png", FIG_OUT / "fig_08_regional_farmgate_dispersion.png"],
        "cost": [V1 / "figures" / "fig_10_livestock_cost_index.png", FIG_OUT / "fig_10_livestock_cost_index.png"],
    }
    out = {}
    for key, candidates in figures.items():
        for src in candidates:
            if src.exists():
                dst = FIG_OUT / src.name
                if src.resolve() != dst.resolve():
                    dst.write_bytes(src.read_bytes())
                out[key] = dst
                break
    return out


def build_doc() -> dict:
    ensure_dirs()
    ev = load_evidence()
    figs = copy_figures()
    h1_table, h2_table, retail_table, validation_table = selected_tables(ev)
    loy_table = loy_mechanism_table(ev)
    ds_table = dataset_table(ev)
    audit = audit_table()
    counts = evidence_counts(ev)
    summary = data_summary(ev)
    regions = regional_facts(ev)
    trade = trade_facts()
    figs["trade"] = create_trade_figure()
    figs["coefficients"] = create_coefficients_figure(h1_table, h2_table)

    h1_table.to_csv(TABLE_OUT / "main_h1_models_v10.csv", index=False)
    h2_table.to_csv(TABLE_OUT / "main_h2_models_v10.csv", index=False)
    retail_table.to_csv(TABLE_OUT / "h2_retail_support_v10.csv", index=False)
    loy_table.to_csv(TABLE_OUT / "loy_mechanism_check_v10.csv", index=False)
    validation_table.to_csv(TABLE_OUT / "old_new_validation_v10.csv", index=False)
    ds_table.to_csv(TABLE_OUT / "dataset_dictionary_v10.csv", index=False)
    audit.to_csv(TABLE_OUT / "mandatory_data_audit_v10.csv", index=False)
    summary.to_csv(TABLE_OUT / "descriptive_statistics_v10.csv", index=False)
    counts.to_csv(TABLE_OUT / "evidence_reliability_counts_v10.csv", index=False)

    doc = Document(DRAFT2)
    clear_body(doc)
    set_update_fields(doc)

    # Title page, using Draft 2 styles.
    p = add_p(doc, "MARKET POWER IN THE DAIRY VALUE CHAIN IN UKRAINE", style="Author")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, "", style="Author")
    p = add_p(doc, "by", style="Author")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_p(doc, "Maksym Charniuk", style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, "", style="Normal")
    p = add_p(doc, "A thesis submitted in partial fulfillment of the requirements for the degree of MA in Economic Analysis.", style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_p(doc, "Kyiv School of Economics", style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_p(doc, "2026", style="Date")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = add_p(doc, "Thesis Supervisor: Professor Oleg Nivievskyi", style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_title(doc, "Abstract")
    add_many(
        doc,
        [
            "This thesis studies market power in the dairy value chain in Ukraine. The research question is whether price adjustment along the chain reveals bargaining asymmetry between farm-gate raw milk producers, processors, institutional procurement, retailers, and final consumers. Price transmission is therefore not treated as the final research objective. It is used as the empirical method for detecting delayed, incomplete, asymmetric, or strategically managed adjustment.",
            "The empirical design is organized around two hypotheses. Hypothesis 1 (H1) states that market power exists between farm-gate raw milk producers and processors. Hypothesis 2 (H2) states that market power exists between processors or procurement channels and downstream retail actors. The thesis follows the logic of Loy, Weiss, and Glauben (2016): first estimate pass-through and adjustment, then interpret heterogeneity through market structure, perishability, menu costs, search costs, Stock Keeping Unit (SKU) support, and promotions. This logic is adapted to Ukrainian data, where confidential margins are not observed and direct proof of legal abuse is outside the scope of the thesis.",
            "The main dataset is the cleaned integrated workbook built from official monthly data, ProZorro records, and retailer observations. The national Ukraine monthly series from the State Statistics Service of Ukraine (SSSU) provide the core evidence: farm-gate raw milk prices, processor-level industrial producer prices, and official consumer prices. ProZorro public procurement system (ProZorro) records and Silpo/Novus retail observations are used as institutional and retail mechanism evidence. Before estimation, the data are audited and corrected: farm-gate and processor prices are converted from tonne-level magnitudes to Ukrainian hryvnia per kilogram (UAH/kg), ProZorro numeric fields are repaired from text values with non-breaking spaces, and retail products are reclassified from product_title and product_name.",
            "The results support a cautious market-power interpretation. Farm-gate and processor-level prices are linked across several dairy products, but adjustment is not cleanly complete or immediate. Processor-to-consumer models show stronger official downstream connection for drinking milk and sour cream, while ProZorro-to-retail evidence is shorter and more fragile. The policy implication is not that the thesis proves unlawful conduct. It is that the dairy chain contains stages where perishability, procurement design, payment delays, and retail promotions can shift adjustment risk away from the actors with stronger bargaining positions and toward farmers, suppliers, or consumers.",
        ],
    )

    doc.add_page_break()
    add_title(doc, "Table of Contents")
    add_field(doc.add_paragraph(), r'TOC \o "1-2" \h \z \u')
    add_title(doc, "List of Figures")
    for i, fig in enumerate(
        [
            "Dairy value chain and empirical hypotheses",
            "Ukraine dairy trade exposure and Harmonized System chapter 04 (HS 04) partner structure",
            "Regional farm-gate raw milk price dispersion",
            "Model-use classification after reliability screening",
            "Farm-gate raw milk and processor-level prices in Ukraine",
            "Official processor-consumer bridge for common dairy products",
            "ProZorro weekly median procurement prices by product",
            "Retail discount incidence in Silpo and Novus observations",
            "Selected main pass-through coefficients",
        ],
        1,
    ):
        add_p(doc, f"Figure {i}. {fig}", style="toc 1")
    add_title(doc, "List of Tables")
    for i, tbl in enumerate(
        [
            "Empirical model families and interpretation rules",
            "Dataset map and source terminology",
            "Mandatory data-audit corrections before modelling",
            "Cleaned dataset descriptive statistics and coverage",
            "Main H1 model selection",
            "Main H2 model selection",
            "H2 retail-support evidence",
            "Loy-style mechanism check",
            "Old-new dataset validation",
            "Reliability screen by source block",
            "Abbreviation discipline and variable definitions",
        ],
        1,
    ):
        add_p(doc, f"Table {i}. {tbl}", style="toc 1")
    add_title(doc, "List of Abbreviations")
    for item in [
        "State Statistics Service of Ukraine (SSSU)",
        "Kyiv School of Economics (KSE)",
        "ProZorro public procurement system (ProZorro)",
        "Stock Keeping Unit (SKU)",
        "Ukrainian hryvnia (UAH)",
        "European Union (EU)",
        "Harmonized System chapter 04 (HS 04)",
        "Error Correction Model (ECM)",
        "Vector Error Correction Model (VECM)",
        "Autoregressive Distributed Lag model (ARDL)",
        "Nonlinear Autoregressive Distributed Lag model (NARDL)",
        "Dynamic Ordinary Least Squares (DOLS)",
        "Local Projection (LP)",
        "Ordinary Least Squares (OLS)",
        "Instrumental Variables (IV)",
        "Lerner Index (LI)",
        "Cost Pass-Through (CPT)",
        "Asymmetric Cost Pass-Through (aCPT)",
        "Unfair Trading Practices (UTP)",
        "Antimonopoly Committee of Ukraine (AMCU)",
        "Herfindahl-Hirschman Index (HHI)",
    ]:
        add_p(doc, item)
    add_title(doc, "Acknowledgments")
    add_p(doc, "The author is grateful to Professor Oleg Nivievskyi for detailed guidance on the empirical design and for pushing the thesis away from a generic price-transmission framing toward a clearer market-power research question. The author also thanks the Kyiv School of Economics faculty and classmates for comments during the research-workshop process.")

    # Introduction.
    add_page_title(doc, "Introduction")
    add_many(
        doc,
        [
            "The starting point of the thesis is a concrete Ukrainian policy problem. Since the full-scale invasion, food suppliers and large grocery chains have debated settlement periods, delayed payments, liquidity pressure, and the boundary between self-regulation and state intervention. The dairy chain is a useful case for this problem because milk is perishable, processors depend on regular raw-milk supply, and retailers can shape the consumer-visible price through assortment, payment terms, and promotions.",
            "Market power in this setting is not observed only through high shelf prices. It can appear when one actor controls the timing, form, or visibility of adjustment. A farm-gate price shock can be absorbed by processors, translated into public procurement prices, passed to official consumer prices, or reshaped through retail discounts before consumers observe it. Price transmission is therefore used as the empirical method, while market power is the research object.",
            "The Ukrainian dairy chain is also policy-relevant. The Food Security Strategy of Ukraine to 2027 identifies milk and dairy products as part of the food-security basket, and the research note in the project materials reports that milk and dairy accounted for 13.7 percent of household monetary food expenditure in 2021. The additional Trade Map workbooks show that dairy-only exports reached about EUR {:.1f} million in 2024, while dairy imports reached about EUR {:.1f} million. These are not the largest trade flows in Ukrainian agriculture, but they show that the sector is connected both to domestic food affordability and to external market pressure (Cabinet of Ministers of Ukraine 2024; International Trade Centre 2024).".format(
                trade["dairy_exports_2024_eur_m"], trade["dairy_imports_2024_eur_m"]
            ),
            "The institutional context makes this problem more than a theoretical possibility. The Antimonopoly Committee of Ukraine (AMCU) has discussed payment delays between large retailers and food suppliers in the context of the 2023 retailer memorandum, including the policy relevance of a 30-day payment limit for basic food products. The AMCU has also reported dairy-related upstream cases involving raw-milk procurement, including the Yagotynsky butter plant case and the Rozhyshche cheese plant case. These materials do not prove that the behavior in the model sample is unlawful. They justify why market-power mechanisms should be studied empirically (Antimonopoly Committee of Ukraine 2023; AMCU Northern Interregional Territorial Department n.d.; AMCU Western Interregional Territorial Department 2023).",
            "The central research question is: does price adjustment in the Ukrainian dairy value chain reveal market power between farm-gate raw milk producers, processors, institutional procurement channels, and downstream retail actors?",
            "The thesis is built around two hypotheses. H1 states that market power exists between farm-gate raw milk producers and processors. H2 states that market power exists between processors or procurement channels and downstream retail actors. This two-hypothesis structure makes every model answer one of the two economic questions instead of treating each data channel as a separate claim.",
            "The novelty of the thesis is fourfold. First, it reframes Ukrainian dairy price transmission as a market-power problem rather than as a generic co-movement exercise. Second, it separates the three value-chain layers: farm-gate raw milk, processor-level sale, and retail or consumer sale. Third, it combines official monthly SSSU data with cleaned ProZorro and retailer observations while explicitly ranking the reliability of each source. Fourth, it adapts the Loy, Weiss, and Glauben (2016) logic by treating pass-through speed and asymmetry as first-stage empirical objects and then interpreting them through market structure, perishability, and promotion mechanisms.",
            "The contribution is deliberately cautious. The thesis does not estimate confidential margins, does not observe bilateral contracts, and cannot prove abuse of dominance. It can, however, identify where price adjustment is incomplete, delayed, asymmetric, or promotion-mediated. Such evidence is economically meaningful because market power in perishable food chains is often exercised through control over timing, shelf access, payment terms, and risk allocation.",
        ],
    )
    add_picture(
        doc,
        figs["chain"],
        "Figure 1. Dairy value chain and empirical hypotheses.",
        "Source: author’s design based on the final thesis data architecture.",
    )
    add_many(
        doc,
        [
            "Figure 1 summarizes the revised architecture. The farm-gate level is the raw milk market. The processor level is measured by industrial producer prices for processed dairy products and, in a separate institutional channel, by ProZorro procurement prices. The retail and consumer level is measured by official consumer prices and by Silpo/Novus SKU-day observations. Farm-gate-to-retail comparisons are therefore not treated as the main identification route, because raw milk is not the same product as a retail dairy SKU. The chain must pass through processing.",
            "The rest of the thesis follows this logic. Chapter 2 explains the market and institutional background. Chapter 3 reviews the literature on market power and vertical price transmission. Chapter 4 presents the methodology and the evidence hierarchy. Chapter 5 documents the data audit and the cleaned datasets. Chapter 6 reports the empirical results by H1 and H2. Chapter 7 concludes with policy implications and limitations.",
        ],
    )
    add_p(doc, "Argument, Novelty, and Evaluation Logic", style="Heading 2")
    add_many(
        doc,
        [
            "The argument starts from the policy problem. In a perishable chain, market power can be exercised even when final prices do not visibly explode. It can appear when one layer controls when adjustment happens, which shocks are passed forward, and which shocks are absorbed by weaker actors. This is why price transmission is useful: it allows the thesis to observe adjustment behavior when contracts and margins are not available.",
            "The research question is reasonable because it is narrow enough to be estimated and broad enough to matter. It does not ask whether the entire Ukrainian dairy sector is competitive in a legal sense. It asks whether the observed price dynamics are consistent with market-power mechanisms at two places in the chain. The first place is the relationship between raw milk sellers and processors. The second place is the relationship between processors or procurement channels and retail or consumer prices. This two-link structure is the empirical backbone of the thesis.",
            "The novelty is not simply the use of several datasets. The novelty is the disciplined combination of datasets with different reliability levels. The thesis uses observed SSSU monthly data as the core, ProZorro and retail data as mechanism evidence, and reconstructed datasets as robustness after validation. This matters because a large empirical file can make a thesis weaker if every output is treated as equally credible. The evidence hierarchy therefore tells the reader which evidence is structural, which evidence is supporting, and which evidence belongs only in the appendix.",
            "The two hypotheses are also a writing discipline. H1 is about processor power over farm-gate raw milk producers. H2 is about downstream power between processors/procurement channels and retail or consumer prices. ProZorro, discounts, nonlinear adjustment, and regional dispersion are not additional hypotheses. They are mechanisms used to interpret H1 and H2. This prevents the results chapter from becoming a catalogue of coefficients and keeps the thesis aligned with the market-power research question.",
            "The thesis also avoids a one-sided industry story. A decline in household milk production, a rise in processor prices, or a widening price spread can be consistent with market power, but each can also reflect quality improvement, wartime costs, energy shocks, logistics, changing product mix, import competition, or seasonality. The empirical task is therefore not to assume market power in advance. It is to test whether the adjustment patterns are consistent with it after alternative explanations and data limitations are made visible.",
            "The intended contribution for policy is a screening framework. If official prices, procurement prices, and retail observations show delayed or asymmetric adjustment, this can guide where regulators, industry associations, or public buyers should look more carefully. A thesis cannot replace confidential investigations, but it can identify where the price data suggest that bargaining conditions deserve attention.",
            "The argument also evolves earlier Kyiv School of Economics (KSE) Agrocenter work on Ukraine’s dairy sector. Litvinov’s first volume describes the sector’s EU-integration and structural transformation, while the second volume identifies supply-chain bottlenecks such as retailer delayed payments and institutional constraints. The present thesis uses those sector insights as background, but adds product-level econometric evidence on where price adjustment is delayed, incomplete, or promotion-mediated (Litvinov 2025a; Litvinov 2025b).",
        ],
    )
    add_p(doc, "Operational Definition of the Two Hypotheses", style="Heading 2")
    add_many(
        doc,
        [
            "H1 is operationalized as a set of models in which the upstream variable is the farm-gate raw milk price and the downstream variable is either a processor-level product price or a ProZorro processed-product procurement price. The strongest H1 evidence is the monthly farm-gate-to-processor block because it uses observed Ukrainian data and has the longest overlap. The ProZorro H1 block is useful because it shows whether farm-gate pressure appears in institutional procurement prices after processing, but it is less reliable because the ProZorro period is shorter and the data are lot-based.",
            "H1 is supported when farm-gate and processor-level prices have an economically admissible long-run relation and when the adjustment pattern is consistent with processor control over timing or completeness. H1 is not supported merely because processor prices are higher than farm-gate prices. A higher level is expected after processing. The relevant evidence is whether price changes and deviations are transmitted in a way that suggests bargaining asymmetry.",
            "H2 is operationalized as a set of models in which the upstream side is the processor-level or procurement price and the downstream side is the official consumer price or observed retail price. The strongest H2 evidence is the processor-to-official-consumer block because it uses national monthly SSSU data. The ProZorro-to-retail and discount blocks explain mechanisms, but their short coverage means that they cannot carry the whole H2 conclusion alone.",
            "H2 is supported when downstream prices do not simply mirror upstream processor or procurement prices, especially where adjustment is incomplete, delayed, asymmetric, or mediated by discounts. The retail layer can exercise market power through shelf access, payment terms, promotion rules, and the ability to manage consumer-visible prices. However, a retail discount is not automatically evidence of market power; it becomes relevant when it systematically changes the pass-through relation.",
            "This operational definition also clarifies what the thesis excludes. The thesis does not estimate farm profitability, processor cost functions, retailer margins, or confidential payment terms. Those are important but unavailable. Instead, it uses price dynamics as observable evidence and interprets them with transparent limits.",
        ],
    )

    # Market analysis.
    add_page_title(doc, "Market Analysis and Institutional Background")
    add_p(doc, "Why Dairy Matters in the Ukrainian Food Economy", style="Heading 2")
    add_many(
        doc,
        [
            "The dairy sector matters because it connects household food expenditure, farm incomes, processor utilization, public procurement, and retail pricing. This chapter does not argue that dairy is the largest sector in Ukrainian agriculture. The more precise point is that dairy is a compact value chain with high perishability and visible policy tension. That makes it a suitable case for studying how market power can shape the transmission of costs and prices.",
            "The sector is also structurally uneven. In 2024, Ukraine produced about 7.2 million tonnes of milk; households produced around 4.2 million tonnes, while the industrial segment produced around 3.0 million tonnes. Yet the formal processing channel is much more enterprise-based: of the milk purchased for processing in 2024, agricultural enterprises supplied 89.8 percent and households supplied 10.2 percent. This gap is central for H1 because biological production and bargaining access to processors are not the same thing (AgroTimes 2025; MilkUA.info 2025).",
            "Recent sector reports describe the same structural pressure from a different angle. The United States Department of Agriculture (USDA) Foreign Agricultural Service expects Ukraine’s milk production and several dairy products to remain constrained by cow-herd decline, population outflows, depressed disposable income, labor shortages, and attacks on energy infrastructure. Those shocks do not by themselves prove market power, but they make the timing of price adjustment commercially important because processors, farms, and retailers face different abilities to absorb disruption (Tarassevych 2024).",
            "The additional trade workbooks show that dairy-only exports were about EUR {:.1f} million in 2024 and represented roughly {:.1f} percent of animal-origin product exports. Dairy imports were about EUR {:.1f} million and represented roughly {:.1f} percent of animal-origin product imports. The trade balance is therefore not the main empirical object, but it matters for the market story: domestic dairy prices are formed in a sector exposed to import competition in cheese and other processed products while still relying on domestic raw milk supply.".format(
                trade["dairy_exports_2024_eur_m"],
                trade["dairy_exports_animal_share_2024"] * 100,
                trade["dairy_imports_2024_eur_m"],
                trade["dairy_imports_animal_share_2024"] * 100,
            ),
            "The sector also faces cost pressure. The SSSU agricultural cost-index workbook included in the project materials confirms that livestock production costs moved sharply during the pre-war and wartime period, but the series mixes monthly, quarterly, and year-on-year bases. For that reason, the cost-index evidence is used only as a minor contextual fact, not as a model variable. The empirical models focus on observed prices in comparable UAH/kg units (State Statistics Service of Ukraine 2025).",
        ],
    )
    add_picture(
        doc,
        figs["trade"],
        "Figure 2. Ukraine dairy trade exposure and Harmonized System chapter 04 (HS 04) partner structure.",
        "Source: International Trade Centre (ITC) Trade Map calculations based on SSSU and United Nations Comtrade data; author’s processing of share_dairy_exp.xlsx, share_dairy_imp.xlsx, and biggest_trade.xlsx.",
    )
    add_p(
        doc,
        "Figure 2 uses three project trade workbooks. The first two panels show dairy-only export and import values and their shares in animal-origin trade. The balance panel shows that the dairy position has been negative in recent years. The partner panel is broader because the source workbook is Harmonized System chapter 04 (HS 04), which includes dairy produce, eggs, honey, and related animal-origin products; it is used only to show external market exposure, not as a model variable.",
    )
    add_p(doc, "Chain Architecture and Stakeholders", style="Heading 2")
    add_many(
        doc,
        [
            "The key conceptual correction is that the dairy chain should not be treated as a direct farm-gate-to-retail chain. Raw milk is sold by farms, processed by dairy enterprises, and only then sold as drinking milk, sour cream, kefir, cheese, butter, or other dairy products. The processing layer is the hinge of the chain. It transforms the physical product and also changes the bargaining environment.",
            "Layer 1 is farm-gate raw milk production and sale. The SSSU source does not use the English term farm gate. In the raw dataset, the relevant indicator is “Середня ціна продукції сільського господарства, реалізованої підприємствами” for the product row “Молоко.” In the thesis, this is called the farm-gate raw milk price because it represents the price received by agricultural enterprises for milk sold at the upstream stage.",
            "Layer 2 is processing and processor-level sale. The raw SSSU source calls the series “Середні ціни виробників промислової продукції.” In the thesis, these observations are called processor-level prices, not generic producer prices, because the term producer can otherwise mean the agricultural raw-milk producer. This distinction is essential for avoiding a conceptual error in the empirical chapter.",
            "Layer 3 is retail or consumer sale. The official SSSU source calls the series “Середні споживчі ціни на товари (послуги).” These are official consumer prices, not the same object as Silpo or Novus online prices. Silpo and Novus observations are SKU-day retail data collected for specific products and promotions. They are useful for H2, but their short time coverage means that they are mechanism evidence rather than the main structural proof.",
            "The market structure behind these layers is asymmetric. KSE sector work reports that the top eight processors accounted for about one-third of the dairy-products market in 2024 and the top twenty for slightly more than half of sector revenue, which suggests moderate processor concentration rather than a simple monopoly story. Retail is more concentrated at the gatekeeper level: the 2025 Opendatabot retail index reports UAH 529.24 billion of revenue for the top ten retailers, with ATB-Market first, Silpo-Food second, and Novus also among the leading chains. This does not give a dairy-category retail share, but it supports the bargaining-power logic behind H2 (Litvinov 2025a; Opendatabot 2025).",
            "ProZorro is a separate institutional procurement channel. It is not farm-gate raw milk and it is not a direct retail shelf price. Economically, it can proxy processed-product realization prices in public procurement contracts, especially where the product title and unit conversion are reliable. The thesis therefore uses ProZorro cautiously: as a bridge between processor/procurement behavior and downstream institutional demand.",
        ],
    )
    add_p(doc, "Market Power Mechanisms", style="Heading 2")
    add_many(
        doc,
        [
            "Market power in this thesis means the ability of one chain layer to influence the timing, completeness, or form of price adjustment. This is narrower than a legal finding of abuse and broader than a simple markup. A processor may exercise bargaining power over farms if raw milk price changes are transmitted selectively to processor prices or if deviations from the long-run relation are corrected slowly. A retailer may exercise downstream power if procurement or processor-price changes are not fully reflected in shelf prices, or if promotions are used to manage price visibility while shifting pressure to suppliers.",
            "Perishability is the central economic mechanism. Raw milk cannot wait for a better price in the same way that grain or butter can. Farmers with limited storage and fewer buyer options are more exposed to monopsony-like bargaining conditions. Processed products differ: butter and hard cheese are more storable, while drinking milk, kefir, and sour cream are closer to fresh-product constraints. This difference matters when interpreting asymmetry and adjustment speed.",
            "Payment delays are the second mechanism. The AMCU retailer memorandum material indicates that deferred payment terms in grocery supply chains are a recognized policy issue. If retailers can delay payment to suppliers while maintaining shelf access power, they may reduce their own liquidity risk at the expense of processors. That is not directly estimated in the price model, but it motivates why H2 includes retail and procurement channels.",
            "Promotions are the third mechanism. Discounts can weaken the visible relationship between procurement prices and observed shelf prices. A retail promotion does not necessarily mean lower market power. It can be a tactical instrument that changes when consumers observe a price response and how suppliers share the cost of that response. That is why discounts support H2 but are not a separate hypothesis.",
        ],
    )
    add_p(doc, "Legal and Institutional Motivation", style="Heading 2")
    add_many(
        doc,
        [
            "The legal and quasi-legal evidence is used only as motivation. This distinction is important. The AMCU retailer memorandum material concerns payment terms in food supply chains and explains why delayed payment can be a market-power mechanism. It does not by itself identify the pass-through coefficients estimated in the model. The upstream dairy cases involving raw-milk procurement show that Ukrainian competition authorities have treated buyer power in milk procurement as a real issue. They do not prove that all processors exercise market power in the sample period (Antimonopoly Committee of Ukraine 2023; AMCU Northern Interregional Territorial Department n.d.; AMCU Western Interregional Territorial Department 2023).",
            "This careful use of legal evidence improves the research design. Without it, the thesis would risk sounding like a purely technical time-series exercise. With it, the empirical question becomes policy-relevant: if public authorities have already faced disputes over raw-milk purchase conditions and retailer-supplier payment terms, then it is reasonable to ask whether public price data also show adjustment patterns consistent with bargaining asymmetry.",
            "The AMCU evidence also helps explain why the thesis treats market power as control over adjustment rather than only as a high price level. In retailer-supplier relationships, a delayed payment term can matter even if the shelf price is temporarily discounted. In farm-processor relationships, a processor can have bargaining power even if the processor price later rises. The timing and allocation of risk are therefore central to the empirical interpretation.",
            "The thesis does not use court or AMCU materials as dependent variables because they are not systematic time series. They are cited in the introduction and market-analysis chapter to motivate the institutional setting, then the empirical chapters return to prices. This separation keeps the research reliable: institutional facts motivate the question, while price data answer it.",
        ],
    )
    add_p(doc, "Possible Outcomes and Non-Market-Power Explanations", style="Heading 2")
    add_many(
        doc,
        [
            "The market-analysis chapter must not conclude too quickly that every unfavorable farmer or consumer outcome is caused by market power. There are at least four alternative explanations. First, dairy products are not homogeneous. A raw-milk price shock may not transmit one-to-one into butter, hard cheese, drinking milk, or sour cream because the milk content, fat content, shelf life, and processing technology differ. Second, processors face energy, labor, packaging, logistics, and credit costs that are not fully captured by raw milk prices. Third, wartime disruption can change routes, demand, and inventories. Fourth, official consumer prices and online retail observations are not the same price object.",
            "These alternative explanations do not weaken the thesis; they make the interpretation more credible. If the thesis acknowledges them before presenting the results, then the market-power conclusion becomes conditional and economic rather than rhetorical. A result supports market power when the product link is admissible, the data are clean, the adjustment pattern is persistent, and the interpretation cannot be reduced to a simple unit or matching error.",
            "The market structure can also produce mixed outcomes. Processor bargaining power over farms may coexist with retailer bargaining power over processors. Consumers may sometimes benefit from short-term promotions even while suppliers carry more risk. Public procurement may stabilize prices in some products and amplify pressure in others. The revised thesis therefore does not search for a single winner or loser. It maps where adjustment power appears along the chain.",
        ],
    )
    add_p(doc, "Actor-Level Bargaining Logic", style="Heading 2")
    add_many(
        doc,
        [
            "Farmers are the upstream actors with the strongest perishability constraint. Raw milk must be cooled, collected, and processed quickly. Larger agricultural enterprises may have better quality control and more stable processor relationships, but they still sell a perishable input. Household producers are important for rural livelihoods but are less central to industrial processing if processors require regular volumes and standardized quality. This upstream duality explains why the thesis uses enterprise farm-gate prices as the main upstream model series.",
            "Processors are the chain hinge. They buy raw milk, transform it into differentiated products, and sell into retail, public procurement, and other channels. Their bargaining position depends on local raw milk supply, plant capacity, product specialization, brand strength, and relationships with buyers. A processor can be powerful relative to farms and simultaneously constrained by retailers or public buyers. This is why the thesis separates H1 and H2 instead of treating the chain as one direct price line.",
            "Public procurement is an institutional buyer, not a retailer in the standard grocery sense. ProZorro observations reflect contract prices, organizer decisions, procurement categories, quantities, and tender timing. They are valuable because they reveal processed-product transaction prices outside ordinary supermarket shelves. They are risky because lot composition and unit definitions vary. The thesis therefore uses ProZorro as a procurement bridge with strong data-audit requirements.",
            "Retailers are downstream actors with consumer access, shelf allocation, and promotion capacity. Their market power is not only a question of final price level. It can involve payment delays, listing requirements, private-label strategy, discount campaigns, and the ability to decide which price is visible to consumers at a given moment. This is why the Silpo and Novus data matter even though the time window is short.",
            "Current market commentary gives a concrete example of this bargaining logic. MilkUA reported in January 2026 that Ukrainian raw-milk purchase prices declined across quality grades despite higher farm production costs, with weak demand, surplus dairy products, export pressure, blackouts, and frost named as market pressures. That type of situation is exactly where bargaining asymmetry matters: if processors and retailers can clear inventories or protect margins while farm prices fall quickly, the burden of adjustment moves upstream (MilkUA.info 2026).",
            "Consumers are the final observers of the chain but not necessarily the actors who see the cost shock first. Official consumer prices show the broader market outcome. Retail SKU prices show specific shelf-level decisions. The thesis uses both because a single consumer-price index would miss promotion behavior, while retailer web data alone would be too short and retailer-specific for structural inference.",
            "The policy relevance comes from the interaction of these actors. If processors delay adjustment to farmers, raw milk producers bear more risk. If retailers delay payment or manage promotions strategically, processors and suppliers bear liquidity and inventory risk. If consumers see only temporary discounts, the final price signal may hide upstream pressure. A market-power thesis should therefore analyze the chain as a sequence of bargaining problems.",
        ],
    )
    add_picture(
        doc,
        figs["regional"],
        "Figure 3. Regional farm-gate raw milk price dispersion.",
        "Source: SSSU agricultural sales data; author’s conversion to UAH/kg and regional aggregation.",
    )
    add_many(
        doc,
        [
            "The territorial dimension is included as a short extension. The main H1 and H2 models use the national Ukraine series because the processor-level data are national and because the research question is the national dairy value chain. Still, regional dispersion matters for market-power interpretation. In the cleaned farm-gate data, the regional extension covers {} regions, and in the last observed month ({}) regional farm-gate prices range from about {:.2f} to {:.2f} UAH/kg. This dispersion does not replace the national model; it shows that a national average can hide local bargaining conditions.".format(
                regions.get("farm_regions", 0),
                regions.get("farm_last_month", ""),
                regions.get("farm_last_min", 0),
                regions.get("farm_last_max", 0),
            ),
            "The market-analysis chapter therefore motivates the empirical strategy. The thesis should not ask whether every dairy price moves together with every other dairy price. It should ask which links are economically admissible, which links are supported by observed data, and which links indicate delayed, incomplete, asymmetric, or strategically managed adjustment.",
        ],
    )

    # Literature review.
    add_page_title(doc, "Literature Review")
    add_p(
        doc,
        "The literature review follows the established structure of the thesis and reads the price-transmission literature through the market-power frame: price transmission is the empirical method, while market power is the thesis object.",
    )
    for text, style in extract_literature_review():
        if text.strip().upper() == "LITERATURE REVIEW":
            continue
        add_p(doc, text, style=style)

    # Methodology.
    add_page_title(doc, "Methodology")
    add_p(doc, "Evidence Hierarchy and Research Design", style="Heading 2")
    add_many(
        doc,
        [
            "The methodology follows an evidence hierarchy. The main evidence comes from observed national monthly SSSU data. ProZorro and retail observations are used to explain institutional and promotional mechanisms. Reconstructed daily and weekly datasets are reviewed, validated against the observed official data, and kept only as supporting or appendix evidence. This hierarchy prevents European Union (EU)-based reconstruction from carrying the main market-power argument.",
            "Loy, Weiss, and Glauben (2016) provide the methodological discipline for interpreting asymmetry. Their German milk-retailing study estimates cost pass-through at store-product level, then relates pass-through heterogeneity to market power, search costs, menu costs, and stock-out risk. The important lesson for Ukraine is not to copy their data structure mechanically. It is to avoid interpreting every asymmetric adjustment as market power without considering product-specific costs, perishability, and retail-management mechanisms.",
            "The empirical design follows two hypotheses. H1 is tested through farm-gate raw milk to processor-level price models and, with lower reliability, farm-gate to ProZorro procurement models. H2 is tested through processor-level to official consumer price models, ProZorro to official consumer models, and short-window ProZorro-to-retail and discount models. The retail block is interpreted cautiously because Silpo and Novus coverage is from late 2025 to early 2026.",
            "The national Ukraine series is the main model object. Territorial models are not used as full replacements because processor-level SSSU prices are national. Regional information is reported as a descriptive extension for farm-gate and consumer prices and as a procurement-region profile for ProZorro where relevant.",
        ],
    )
    add_table(
        doc,
        pd.DataFrame(
            [
                {
                    "family": "Error Correction Model (ECM)",
                    "use": "main long-run adjustment model where cointegration is plausible",
                    "interpretation": "speed and direction of correction after a disequilibrium shock",
                },
                {
                    "family": "Autoregressive Distributed Lag model (ARDL)",
                    "use": "lag-selection and bounds-testing benchmark",
                    "interpretation": "short-run dynamics and long-run pass-through when sample supports lag structure",
                },
                {
                    "family": "Vector Error Correction Model (VECM)",
                    "use": "three-layer system only if cointegration rank and sample size support it",
                    "interpretation": "which layer adjusts after a whole-chain disequilibrium",
                },
                {
                    "family": "Nonlinear Autoregressive Distributed Lag model (NARDL)",
                    "use": "asymmetry test when positive and negative shocks may differ",
                    "interpretation": "whether price increases and decreases are transmitted differently",
                },
                {
                    "family": "Dynamic Ordinary Least Squares (DOLS)",
                    "use": "additional long-run robustness model",
                    "interpretation": "long-run relation corrected for leads/lags of upstream changes",
                },
                {
                    "family": "Local Projection (LP)",
                    "use": "additional dynamic response check",
                    "interpretation": "response of downstream price changes over chosen horizons",
                },
                {
                    "family": "Threshold ECM in the style of Loy, Weiss, and Glauben (2016)",
                    "use": "first-stage pass-through speed and asymmetry screen",
                    "interpretation": "separate correction after positive and negative deviations",
                },
                {
                    "family": "Ordinary Least Squares (OLS) second-stage screen",
                    "use": "Loy-style comparison of first-stage speed/asymmetry measures with channel characteristics",
                    "interpretation": "associational mechanism evidence, not the main time-series identification",
                },
                {
                    "family": "Instrumental variables (IV) and Lerner Index boundary",
                    "use": "not estimated as a headline model because product-level margins, wholesale costs, and credible instruments are not observed",
                    "interpretation": "market power is inferred from pass-through and adjustment patterns rather than direct markups",
                },
            ]
        ),
        "Table 1. Empirical model families and interpretation rules",
        "Source: author’s methodology design based on the final evidence package and Loy, Weiss, and Glauben (2016).",
        max_rows=10,
    )
    add_many(
        doc,
        [
            "The comparison with Loy, Weiss, and Glauben (2016) also defines what is not estimated as a headline model. Their paper can compare Ordinary Least Squares (OLS) and instrumental-variable estimates because it observes retail and wholesale prices that allow a Lerner-style margin proxy and because the second-stage problem is a cross-sectional explanation of pass-through heterogeneity. The Ukrainian dataset does not observe confidential processor-retailer margins, marginal costs, or a credible external instrument for product-level bargaining power. For that reason, OLS is used only as a limited second-stage mechanism screen, while Instrumental Variables (IV) and the Lerner Index are discussed as identification boundaries rather than reported as main results.",
            "This restriction improves reliability. A direct Lerner Index without observed marginal cost or a consistent wholesale-retail margin would look precise but would not measure market power cleanly. The thesis therefore infers market-power mechanisms from long-run pass-through, error-correction speed, asymmetry, procurement behavior, discount incidence, and source reliability. This is narrower than the German scanner-data design, but it fits the Ukrainian data better.",
        ],
    )
    add_p(doc, "Core Model Equations", style="Heading 2")
    add_formula(
        doc,
        1,
        "Long-run price relation",
        "ln(P_downstream,t) = alpha + beta ln(P_upstream,t) + u_t",
        "P_downstream,t is the processor-level, ProZorro procurement, official consumer, or retail price depending on the hypothesis; P_upstream,t is the farm-gate raw milk price for H1 and the processor/procurement price for H2; beta is the long-run pass-through coefficient; u_t is the disequilibrium residual.",
    )
    add_formula(
        doc,
        2,
        "Autoregressive Distributed Lag model",
        "ln(P_d,t)=a_0+sum_{i=1}^{p} a_i ln(P_d,t-i)+sum_{j=0}^{q} b_j ln(P_u,t-j)+e_t",
        "P_d,t is the downstream price; P_u,t is the upstream price; p and q are lag orders selected by information criteria and diagnostics; b_j measures current and lagged pass-through; the implied long-run multiplier is sum b_j divided by one minus sum a_i.",
    )
    add_formula(
        doc,
        3,
        "Error-correction model",
        "Delta ln(P_downstream,t) = gamma + lambda u_{t-1} + sum phi_i Delta ln(P_downstream,t-i) + sum theta_j Delta ln(P_upstream,t-j) + eps_t",
        "Delta is the first-difference operator; u_{t-1} is the lagged residual from Formula (1); lambda is the error-correction coefficient; phi_i and theta_j are short-run dynamics; eps_t is the error term. A negative lambda means that deviations are corrected over time.",
    )
    add_formula(
        doc,
        4,
        "Vector Error Correction Model",
        "Delta y_t = Pi y_{t-1} + sum_{i=1}^{k-1} Gamma_i Delta y_{t-i} + c + eps_t, where y_t=[ln(P_farmgate,t), ln(P_processor,t), ln(P_consumer,t)]'",
        "y_t is the vector of farm-gate, processor-level, and consumer prices; Pi contains long-run cointegration information; Gamma_i captures short-run dynamics; k is the lag length; c is a deterministic component. VECM is used only when the cointegration rank and sample support a three-layer system.",
    )
    add_formula(
        doc,
        5,
        "Nonlinear or asymmetric adjustment",
        "u_{t-1}^{+}=max(u_{t-1},0), u_{t-1}^{-}=min(u_{t-1},0); Delta ln(P_downstream,t)=...+lambda^{+}u_{t-1}^{+}+lambda^{-}u_{t-1}^{-}+eps_t",
        "u_{t-1}^{+} and u_{t-1}^{-} split positive and negative disequilibrium; lambda^{+} and lambda^{-} measure correction after positive and negative deviations. Different values indicate asymmetric pass-through under H1 or H2.",
    )
    add_formula(
        doc,
        6,
        "Dynamic Ordinary Least Squares robustness model",
        "ln(P_d,t)=alpha+beta ln(P_u,t)+sum_{r=-R}^{R} psi_r Delta ln(P_u,t+r)+eta_t",
        "P_d,t is the downstream price; P_u,t is the upstream price; beta is the long-run coefficient; leads and lags of upstream changes correct for endogeneity and serial correlation in the long-run relation; R is the lead-lag window.",
    )
    add_formula(
        doc,
        7,
        "Local Projection response",
        "Delta_h ln(P_d,t+h)=alpha_h+beta_h Delta ln(P_u,t)+controls_t+eps_{t+h}",
        "Delta_h ln(P_d,t+h) is the downstream price change over horizon h; Delta ln(P_u,t) is the current upstream shock; beta_h is the pass-through response at horizon h. Local projections are used as robustness checks, not as the main long-run model.",
    )
    add_formula(
        doc,
        8,
        "Threshold error-correction pass-through in the style of Loy, Weiss, and Glauben",
        "Delta ln(P_d,t)=alpha+lambda^{+}u_{t-1}^{+}+lambda^{-}u_{t-1}^{-}+short-run terms+eps_t; speed=(|lambda^{+}|+|lambda^{-}|)/2, asymmetry=lambda^{+}-lambda^{-}",
        "P_d,t is the downstream price; u_{t-1}^{+} and u_{t-1}^{-} are positive and negative disequilibrium terms; speed summarizes correction strength; asymmetry summarizes whether positive and negative deviations are corrected differently. This follows the first-stage logic of Loy, Weiss, and Glauben (2016).",
    )
    add_formula(
        doc,
        9,
        "ProZorro aggregation",
        "P_{p,T}^{ProZorro}=median(unit_price_{lot,p,t}) for all reliable lots in product p and period T",
        "P_{p,T}^{ProZorro} is the product-period procurement price; p is the controlled dairy product category; T is day, week, or month; unit_price_{lot,p,t} is the repaired UAH/kg lot price. Medians are preferred because lots differ by organizer, size, and product specification.",
    )
    add_formula(
        doc,
        10,
        "Retail discount mechanism",
        "RetailPrice_{sku,t}=alpha+beta ProcurementPrice_{p,t}+delta DiscountIncidence_{sku,t}+rho DiscountDepth_{sku,t}+controls+eps_{sku,t}",
        "RetailPrice_{sku,t} is the observed Silpo or Novus SKU-day price; ProcurementPrice_{p,t} is the matched product-level ProZorro price; DiscountIncidence equals one when the SKU is promoted; DiscountDepth measures the markdown size where available; controls include product, retailer, or time controls when supported by sample size.",
    )
    add_p(doc, "Diagnostics and Reliability Rules", style="Heading 2")
    add_many(
        doc,
        [
            "Every model is screened before it enters the thesis narrative. The minimum checks are product admissibility, unit consistency, time overlap, stationarity, cointegration or long-run plausibility, residual serial correlation, heteroskedasticity, coefficient sign and magnitude, peak-coefficient flags, and economic interpretation. A statistically significant but economically incoherent coefficient is not treated as a headline result.",
            "The reconstructed datasets are not discarded, but their role is restricted. FarmGate(oldmodel), Producer(oldmodel), Consumer(oldmodel), and the old ProZorro and retail sheets are compared with observed official data where the periods overlap. If the correlation and sign agreement are strong, the old result can be used as supporting evidence. If the validation is weak or the intersection is too short, the result is moved to the appendix or omitted from the main story.",
            "The EU reconstruction is especially restricted. It can support data-construction discussion only after monthly aggregation back to observed Ukrainian series, correlation checks, and error metrics. It cannot be the core proof of market power because EU price patterns cannot be extrapolated to Ukraine without validation.",
        ],
    )
    add_p(doc, "Step-by-Step Estimation Pipeline", style="Heading 2")
    add_many(
        doc,
        [
            "The estimation pipeline begins with source-level cleaning. The script reads the raw integrated workbook and component workbooks, identifies the chain level, standardizes product labels, parses dates, converts all admissible prices to UAH/kg, and records quality flags. No model is estimated from raw workbook cells directly. This is essential because the same workbook contains observed SSSU series, reconstructed series, ProZorro records, EU benchmarks, and retail web-scraped observations.",
            "The second step is product admissibility. For H1, farm-gate raw milk can be linked to processor-level dairy products because raw milk is the upstream input, but the interpretation is raw-material pass-through rather than identical-product pass-through. For H2, only product links with a meaningful downstream counterpart are considered as main evidence. Drinking milk and sour cream are the cleanest official processor-to-consumer links. Cheese, butter, kefir, and skimmed milk powder are used where the downstream source supports a comparable product or as supporting evidence where it does not.",
            "The third step is frequency alignment. The main official models are monthly. ProZorro is event-level and is aggregated to daily, weekly, and monthly product medians. Retail is SKU-day and is aggregated to retailer-product-day or pooled product-day only after product repair. Weekly and daily models are not automatically preferred simply because they have more rows. A high-frequency series based on reconstruction or short retail overlap may be less reliable than a monthly observed series with fewer but cleaner observations.",
            "The fourth step is stationarity and long-run testing. The pipeline checks unit-root behavior using Augmented Dickey-Fuller and related tests where available, then tests long-run relations through Engle-Granger, ARDL-style bounds logic, or Johansen/VECM checks when the system model is feasible. The thesis does not force a VECM if the rank test does not support it or if the product overlap is weak.",
            "The fifth step is model estimation. ECM and ARDL-family models are used for the main pass-through logic. NARDL and threshold ECM models are used when asymmetric adjustment is economically meaningful. DOLS and LP are used as additional model families beyond ECM, VECM, ARDL, and NARDL. Their role is robustness: DOLS checks whether the long-run slope is stable after adding leads and lags of upstream changes, while LP checks whether downstream responses appear over short horizons.",
            "The sixth step is reliability screening. The model output is not accepted simply because a p-value is small. The screen checks the observation count, overlap, product mapping, coefficient sign, coefficient magnitude, peak-coefficient flags, residual diagnostics, and whether the model is based on observed data or reconstructed data. This is why some statistically significant retail and old-model results are not used as headline evidence.",
            "The final step is economic interpretation. Each retained model is translated into a sentence about bargaining power, adjustment speed, or strategic pricing. The thesis avoids mechanical coefficient narration. A negative and significant adjustment term is not the conclusion; the conclusion is what that adjustment speed means for farmers, processors, retailers, public buyers, or consumers.",
        ],
    )
    add_p(doc, "Why the Final Model Prioritizes Monthly SSSU Evidence", style="Heading 2")
    add_many(
        doc,
        [
            "The final model prioritizes monthly SSSU evidence because it is observed, official, national, and long enough for time-series interpretation. This choice is not made because monthly data are perfect. Monthly data smooth short-run behavior and cannot identify day-by-day promotion strategies. They are chosen because the thesis needs a defensible structural base before adding higher-frequency mechanisms.",
            "The processor-level series are especially important. The term producer price can create confusion because agricultural producers and industrial producers are different actors. The industrial producer-price data are therefore called processor-level prices in the economic interpretation. This naming is not stylistic. It is the condition that allows the reader to understand where H1 ends and where H2 begins.",
            "The monthly SSSU consumer series are also important because they provide a bridge from processor prices to official downstream consumer prices. Silpo and Novus observations are valuable, but they are not official national consumer prices. A reliable H2 story should therefore begin with the processor-to-consumer bridge and then use retail SKU evidence to explain the promotional mechanism.",
        ],
    )

    # Data chapter.
    add_page_title(doc, "Data")
    add_p(doc, "Data Sources and Source Terminology", style="Heading 2")
    add_many(
        doc,
        [
            "The data chapter is written before the results because the reliability of the thesis depends on correctly naming the chain levels. The raw source names are not always the same as the thesis concepts. The SSSU farm-gate source does not say farm gate; it reports average sales prices of agricultural products sold by enterprises. The SSSU processor source does not say processor gate; it reports average industrial producer prices. The SSSU consumer source reports average consumer prices. The thesis uses economic labels only after explaining these source names.",
            "All price series used for model comparison are harmonized to UAH/kg. This is not a cosmetic step. Without it, processor prices in hryvnia per tonne and farm-gate tonne-level magnitudes would be mechanically incomparable with consumer and retail prices in kilogram or package units.",
        ],
    )
    add_table(doc, ds_table, "Table 2. Dataset map and source terminology", "Source: author’s compilation from the main integrated workbook and component raw workbooks.", max_rows=6)
    add_p(doc, "Mandatory Data Audit", style="Heading 2")
    add_many(
        doc,
        [
            "The main integrated workbook is the primary model file, but it is not model-ready in raw form. Four risks are mandatory audit items before any econometric estimation. First, Processor_price is reported in hryvnia per tonne. Second, Farm_price appears as a tonne-level farm-gate price despite the short unit label Гривня. Third, ProzorroM(full) contains numeric values with non-breaking spaces and text formatting. Fourth, retail classification contains visible errors and must be rebuilt from product_title and product_name.",
            "After correction, every analytical table contains a controlled product column. Farm-gate observations are identified as raw_milk. Processor, ProZorro, official consumer, and retail observations are assigned to comparable product groups only where the product content supports comparison. Non-dairy and ambiguous retail items are excluded or moved to other_dairy.",
        ],
    )
    add_table(doc, audit, "Table 3. Mandatory data-audit corrections before modelling", "Source: author’s audit of the main integrated workbook.", max_rows=10)
    add_p(
        doc,
        "The descriptive-statistics table turns the cleaned files into the data layer used for Chapter 6. It reports only the analytical price that is actually used for each chain stage. The official farm-gate, processor, and consumer rows are directly comparable in UAH/kg. ProZorro is comparable only after lot-level unit repair. Retail is reported mainly as SKU-day package-price evidence because reliable UAH/kg conversion is available only for a smaller subset of observations.",
    )
    add_table(doc, summary, "Table 4. Cleaned dataset descriptive statistics and coverage", "Source: author’s cleaned data output folder.", max_rows=10)
    add_p(doc, "Product Harmonization", style="Heading 2")
    add_many(
        doc,
        [
            "Product harmonization follows the value-chain logic. Farm-gate raw milk is not the same product as butter, cheese, sour cream, or yogurt. Therefore, H1 is interpreted as raw-material-to-processed-product transmission, not as a physical one-to-one product comparison. The model asks whether farm-gate raw milk prices are transmitted into processor-level output prices in a way consistent with market power.",
            "For H2, product comparability is stricter. Processor-level drinking milk can be compared with official consumer drinking milk. Processor-level sour cream can be compared with official consumer sour cream. ProZorro and retail comparisons are kept at product or SKU level where product_title and product_name support matching. Where the match is broad or the overlap is short, the result is marked as mechanism evidence.",
            "The old model sheets are explicitly named as old-model sheets. Producer(oldmodel) is conceptually the same chain level as processor-level prices, because it corresponds to industrial producer prices of processed products. It is not the same as farm-gate raw milk producers. The thesis therefore avoids the ambiguous term producer prices unless the source name is being described.",
        ],
    )
    add_p(doc, "Retail and ProZorro Preparation", style="Heading 2")
    add_many(
        doc,
        [
            "The H2 retail model is prepared at the SKU-day level. Each Silpo and Novus observation is classified by product using product_title and product_name. The product column is then used to aggregate to retailer-product-day and pooled product-day series. This preserves the product identity needed for comparison with ProZorro lots.",
            "ProZorro is prepared at lot level before aggregation. The raw fields Товар, Кількість, Одиниця виміру, and Ціна за одиницю are parsed and converted to reliable unit prices. Lots that cannot be converted to UAH/kg are not used for price-level econometrics. The main aggregation is the product-level median by week or month, with lot counts and regional counts retained as reliability indicators.",
            "The main models use the national Ukraine level. Regional ProZorro and regional official consumer information are used only as a short territorial extension. This protects the thesis from mixing national processor prices with regional downstream prices in a way that would look precise but would not be economically clean.",
        ],
    )
    add_p(doc, "Source-by-Source Reliability Assessment", style="Heading 2")
    add_many(
        doc,
        [
            "The farm-gate raw milk series is reliable after unit correction. The main risk is the misleading short unit label, not the economic content. The row “Середня ціна продукції сільського господарства, реалізованої підприємствами” for “Молоко” corresponds to the upstream price received by agricultural enterprises. The thesis uses enterprise farm-gate prices as the main upstream series because they are most relevant for processors that require regular, quality-controlled raw milk supply.",
            "The processor-level series is reliable after conversion from hryvnia per tonne to UAH/kg. The raw source name “Середні ціни виробників промислової продукції” is kept in the data chapter, but the thesis label is processor-level price. This avoids confusing industrial producers of processed products with farm producers of raw milk. The processor series covers six product groups and is the central bridge between H1 and H2.",
            "The official consumer series is reliable for product groups available in the SSSU consumer-price dataset. It should not be confused with retailer web prices. The official consumer series measures average consumer prices for selected goods, including pasteurized milk, sour cream, and soft cheese. It is the strongest downstream benchmark for H2 because it is national and monthly.",
            "ProZorro data are reliable as procurement records but weaker as a market-price series. Each lot is an institutional purchase with its own organizer, contract conditions, region, quantity, unit, and product title. The thesis therefore treats ProZorro as an institutional procurement channel and a proxy for processed-product realization prices, not as a universal processor or retail price. Robust median aggregation is used to reduce outlier influence.",
            "Silpo and Novus data are reliable for observed SKU-day retail behavior after product repair, but their time span is short. Their value is not that they can prove long-run retail market power. Their value is that they show how retail prices, package sizes, product identities, and discounts are managed at the shelf. This makes them relevant for H2 mechanism analysis.",
            "The old-model and EU-based reconstruction sheets are treated as secondary evidence. Some of them validate well against observed official data and can support the interpretation. Others validate weakly or are too dependent on interpolation assumptions. This is why the data chapter includes the old-new validation table before the results chapter uses any reconstructed evidence.",
        ],
    )
    add_p(doc, "Descriptive Patterns Before Modelling", style="Heading 2")
    add_many(
        doc,
        [
            "The descriptive patterns already show why the model should not be a simple farm-to-retail comparison. Farm-gate raw milk prices are much lower in level than processor-level and consumer prices because raw milk is an input, not the final product. Processor-level products differ widely: butter and cheese have different raw milk content, storability, and market exposure than drinking milk or kefir. Consumer prices include retail, distribution, and product-specific costs. Level comparisons therefore need economic interpretation before any model is estimated.",
            "The time series also contain wartime structural breaks. After 2022, price levels, volatility, procurement behavior, and retail conditions changed. The thesis does not remove the wartime period because it is central to policy relevance. Instead, it treats the period as part of the interpretation and avoids over-claiming precision from short wartime retail windows.",
            "Regional descriptive evidence matters mainly for the farm-gate stage. If farm-gate prices differ substantially across regions, then processors and local buyers may face different procurement environments. However, because processor prices are national in the SSSU processor dataset, regional farm-gate variation is not turned into a full regional H1 model. It is used as context and as a warning against overinterpreting national averages.",
        ],
    )
    add_p(doc, "Reproducibility and Data Processing Notes", style="Heading 2")
    add_many(
        doc,
        [
            "The empirical package is reproducible from the project folders. The main integrated workbook is the primary input. The cleaning scripts write separate clean data files for farm-gate prices, processor prices, official consumer prices, ProZorro lots, ProZorro aggregates, retail SKU-day observations, retail aggregates, and validation tables.",
            "The data processing uses Python for workbook parsing, product classification, unit conversion, aggregation, visualization, and document generation. The modelling scripts use standard time-series and regression routines from the Python scientific stack. The data chapter states this explicitly because software choices affect the transparency of scraping, cleaning, and estimation.",
            "Retail scraping is not treated as an official source. It is a constructed dataset based on observed Silpo and Novus product pages. The reliability of this source depends on whether the product title, product name, package size, unit price, current price, and discount fields are correctly parsed. Because retail classification errors were visible in the raw sheets, the final pipeline reclassifies products from text rather than trusting the legacy product column.",
            "The ProZorro parser is conservative. It removes non-breaking spaces from numbers, normalizes comma and dot decimal signs, standardizes units, and excludes observations that cannot produce reliable UAH/kg prices. The cleaned lot table keeps the original product title and organizer-region fields so that questionable product mappings can be audited later.",
            "The official SSSU data are not mechanically standardized without interpretation. For farm-gate prices, the raw unit label is short and misleading, so the magnitude is checked before conversion. For processor prices, the unit clearly states hryvnia per tonne. For consumer prices, the unit is already kilogram or 1000 grams for the main products. The data chapter explains each case because unit mistakes would invalidate the model.",
            "The final tables are designed to be readable in the thesis. Full model registers remain in the output folder. The thesis main text contains only the selected models, reliability counts, and validation tables needed for the argument. This prevents the empirical chapter from becoming a technical appendix.",
        ],
    )

    # Results.
    add_page_title(doc, "Estimation Results")
    add_p(doc, "How to Read the Results", style="Heading 2")
    add_many(
        doc,
        [
            "The results chapter is organized by hypotheses rather than by model family. The goal is not to show every estimated equation. The goal is to identify which pieces of evidence are reliable enough to support an economic interpretation of market power.",
            "A model is called main evidence only if it uses observed national data, has adequate overlap, passes the product-admissibility screen, and gives an economically interpretable coefficient. A model is called supporting evidence if it is useful but short, reconstructed, or sensitive to product matching. Appendix-only evidence is still disclosed in the reliability screen, but it does not drive the conclusion.",
        ],
    )
    add_picture(
        doc,
        figs["reliability"],
        "Figure 4. Model-use classification after reliability screening.",
        "Source: author’s integrated evidence register.",
    )
    add_p(doc, "Hypothesis 1: Farm-Gate Raw Milk Producers and Processors", style="Heading 2")
    add_many(
        doc,
        [
            "H1 states that market power exists between farm-gate raw milk producers and processors. The main empirical link is farm-gate raw milk to processor-level dairy prices. This is the strongest part of the thesis because it uses observed national monthly data over a long period, with farm-gate prices from 2015 and processor-level prices from 2013.",
            "The main H1 model is a restricted threshold error-correction design in the style of Loy, Weiss, and Glauben (2016), supported by standard ECM, ARDL, DOLS, LP, and VECM checks where feasible. The threshold model is useful because it separates the long-run slope from the speed and direction of adjustment. The long-run slope says whether processor-level prices move with raw milk prices. The adjustment term says whether deviations are corrected quickly or slowly.",
            "The main H1 results show economically meaningful long-run slopes for drinking milk, hard cheese, kefir, and sour cream. The strongest H1 interpretation is attached to products where the price link is supported by cointegration or plausible adjustment evidence. Kefir, skimmed milk powder, and butter are kept as supporting products because their adjustment signs, diagnostics, or validation screens are weaker. The coefficients are not interpreted mechanically as margins. They show that processor-level prices are linked to raw milk prices, while the adjustment terms are weak or slow for several products. This pattern is consistent with a processing layer that transmits raw milk costs selectively rather than immediately.",
            "The H1 model table follows the same reading logic as the methodology equations. The β column is the long-run pass-through coefficient from Formula (1). The λ column is the error-correction coefficient from Formula (3) or its threshold version in Formula (8). β shows the long-run connection; λ shows how quickly deviations are corrected. The evidence-reading column separates a price-link result from a stronger adjustment result, which is essential for avoiding overstatement.",
        ],
    )
    add_picture(
        doc,
        figs["h1_prices"],
        "Figure 5. Farm-gate raw milk and processor-level prices in Ukraine.",
        "Source: SSSU agricultural sales and industrial producer-price datasets; author’s conversion to UAH/kg.",
    )
    add_table(doc, h1_table, "Table 5. Main H1 model selection", "Source: author’s model estimates from cleaned observed data.", max_rows=8)
    add_many(
        doc,
        [
            "The economic interpretation is cautious but clear. Farm-gate raw milk prices are not irrelevant for processor-level prices; the long-run slopes are large and statistically precise. However, only part of the product set gives strong adjustment evidence. This matters because the market-power claim is stronger when prices are linked and deviations are corrected in a way that indicates downstream timing control. In a perishable raw-milk market, even weak or slow correction is informative because farmers cannot store raw milk while waiting for stronger bargaining conditions. The processing layer can therefore become the point where upstream price pressure is absorbed or delayed.",
            "The ProZorro extension supports H1 only weakly. Farm-gate-to-ProZorro monthly models have short samples of about 32 to 34 monthly observations. Some coefficients are statistically significant, but the overlap is too short for headline inference. Old daily ProZorro models have more observations but depend on reconstructed daily farm-gate series. Because the old farm-gate reconstruction validates well against observed data, these models can support the mechanism, but they should not replace the observed monthly H1 block.",
            "The VECM system evidence is also limited. A full three-chain VECM is attractive conceptually because it can connect farm-gate, processor, and official consumer prices. Statistically, however, the sample supports only a small number of feasible systems. The thesis therefore treats VECM as a robustness check and uses pairwise H1 and H2 blocks as the main specification. This is more defensible than forcing a system model onto weak product overlap.",
        ],
    )
    add_p(doc, "H1 Product-Level Interpretation", style="Heading 3")
    add_many(
        doc,
        [
            "Drinking milk is one of the cleanest H1 products because the processor-level series corresponds to processed liquid milk and the downstream official consumer series also contains pasteurized milk. The H1 result indicates strong long-run connection between farm-gate raw milk and processor-level drinking milk. Economically, this means the processor price is not detached from raw milk costs. The market-power question is therefore about adjustment speed and completeness, not about whether the link exists at all.",
            "Sour cream is also informative because it is a fresh processed product with strong perishability and a clear official consumer counterpart. The H1 coefficient is economically meaningful, but the adjustment speed is not strong enough to imply immediate correction. This fits the idea that processors can smooth raw milk shocks before they become visible in processed-product prices.",
            "Kefir has a strong product identity at the processor level but does not have the same official consumer counterpart in the SSSU consumer file. It is therefore more useful for H1 than for H2. Its H1 result supports the broader conclusion that processor-level fresh dairy prices move with farm-gate raw milk prices, but it is not used as strong adjustment proof because the correction evidence is weaker than the long-run price link.",
            "Hard cheese and butter require a more careful interpretation because they are more storable and have different raw milk and fat content. A long-run relation with farm-gate raw milk is economically plausible, but market power cannot be inferred from the slope alone. Storability can change the timing of adjustment and may reduce the immediate pressure to transmit raw milk price changes.",
            "Skimmed milk powder is the weakest processor product for headline interpretation because old-model validation is weaker and because the product is more exposed to industrial and external-market dynamics. It remains useful as a processor-level product in the observed monthly data, but the thesis should not build its main conclusion around it.",
        ],
    )
    add_p(doc, "H1 Robustness and Old-Model Evidence", style="Heading 3")
    add_many(
        doc,
        [
            "The model archive contains many H1 models, including weekly and daily ECM, NARDL, LP, margin, aggregate-index, and VECM outputs. This evidence is reorganized rather than discarded. A high-frequency result enters the main discussion only if the product link is admissible, the intersection is representative, and the reconstructed series validates against the observed official series.",
            "The validation screen supports using the old FarmGate reconstruction as robustness because its correlation with observed farm-gate data is high. It is less supportive for every processor product. For example, kefir, butter, and sour cream validate better than skimmed milk powder. This means that old high-frequency H1 models can reinforce the main monthly story for selected products, but they should not be presented as equally reliable across all products.",
            "The H1 ProZorro extension is treated as a procurement mechanism. ProZorro prices are processed-product procurement prices, not farm-gate raw milk prices. A farm-gate-to-ProZorro model therefore tests whether raw milk costs are visible in institutional procurement prices after processing and contracting. Because the ProZorro period is short, this block supports the existence of procurement-channel adjustment but does not replace the core farm-gate-to-processor model.",
            "The three-chain monthly model is conceptually important but empirically restricted. A system model is appropriate only where farm-gate, processor, and consumer series overlap for a comparable product and where cointegration rank supports the system. System evidence is therefore reported briefly, while the main interpretation remains in the pairwise H1 and H2 structure.",
        ],
    )
    add_p(doc, "Hypothesis 2: Processors, Procurement, and Downstream Retail Actors", style="Heading 2")
    add_many(
        doc,
        [
            "H2 states that market power exists between processors or procurement channels and downstream retail actors. The main official H2 evidence uses processor-level prices and official SSSU consumer prices. This bridge is more reliable than the retail web-scraped block because it covers 2017 to 2026 and uses national monthly data.",
            "The processor-to-consumer models for drinking milk and sour cream provide the main H2 official price-link evidence, but their reliability is not identical. Sour cream gives the cleaner official downstream result, while drinking milk has a strong slope but weaker long-run diagnostic support. The correct interpretation is therefore not that H2 is mechanically proven by two coefficients. It is that downstream prices are linked to processor prices, while the timing and strength of correction remain product-specific.",
            "The ProZorro-to-consumer model for drinking milk is supporting evidence. It links an institutional procurement price to the official consumer benchmark, but the ProZorro period begins only in 2023. It is therefore useful for the wartime and public-procurement mechanism, not for a full long-run structural conclusion.",
            "The H2 model table uses the same β and λ notation as the H1 table so the reader can compare the two hypotheses directly. A strong β with weak or slow λ means that prices are connected over time but the downstream layer may still control the timing of correction. The evidence-reading column prevents the text from treating a slope result as if it were automatically a full cointegration-and-adjustment result.",
        ],
    )
    add_picture(
        doc,
        figs["h2_bridge"],
        "Figure 6. Official processor-consumer bridge for common dairy products.",
        "Source: SSSU industrial producer-price and consumer-price datasets; author’s conversion to UAH/kg.",
    )
    add_table(doc, h2_table, "Table 6. Main H2 model selection", "Source: author’s model estimates from cleaned observed data.", max_rows=8)
    add_many(
        doc,
        [
            "The H2 retail block is deliberately narrower. Silpo and Novus data are rich at SKU level, but the calendar window is short. That means the retail evidence can explain promotions, timing, and product matching, but it cannot carry the whole downstream market-power claim by itself. The retail block is therefore written as a mechanism section rather than as a separate hypothesis.",
            "The cleaned retail data show that discounts are common and product-specific. Promotions can make the observed shelf price move differently from procurement prices, especially when retailers manage temporary markdowns. This is why the thesis separates baseline price, observed price, discount incidence, and discount depth where the variables exist. A discount coefficient is not interpreted as consumer welfare by itself; it is interpreted as evidence that retail pricing has an active strategic layer.",
        ],
    )
    add_p(doc, "H2 Official Consumer Bridge", style="Heading 3")
    add_many(
        doc,
        [
            "The processor-to-consumer official bridge is the cleanest H2 model because both series are national and monthly. For drinking milk, the processor-level product and the official consumer product are close enough to support a direct price-transmission interpretation. For sour cream, the match is also strong. Soft cheese is available in the official consumer series, but the processor-level hard cheese series is not the same product. That mismatch is why cheese evidence is handled carefully.",
            "The H2 official results should be interpreted through adjustment, not only through slope. A high long-run slope may reflect cost pass-through, product margins, or downstream price setting. The error-correction term asks a more relevant market-power question: when the downstream price is away from its long-run relation with the processor price, how quickly does it return? Slow correction can be consistent with downstream price-setting capacity.",
            "The official consumer bridge also helps evaluate the retail web-scraped data. If official consumer prices and Silpo/Novus observed prices move differently, the difference may reflect retailer-specific promotions, package mix, online availability, or product selection. That is why the thesis does not merge official consumer prices and retail SKU prices without explanation.",
        ],
    )
    add_p(doc, "H2 Retail and Discount Mechanism", style="Heading 3")
    add_many(
        doc,
        [
            "The H2 retail mechanism is built at product and SKU level. A SKU is a specific product identifier, not only a product category. This distinction matters because a discount on one 850-gram bottle of kefir is not the same observation as a regular price on a different brand or package size. The cleaned retail data therefore preserve product_title, product_name, brand, package information, discount incidence, and discount depth where possible.",
            "The retail classification repair changes the empirical interpretation. In the raw retail sheets, some items are visibly misclassified. After reclassification, non-dairy items are excluded and ambiguous dairy items are separated. This reduces sample size but increases reliability. A smaller corrected sample is preferable to a larger sample that mixes dairy products with unrelated goods.",
            "The ProZorro-to-retail models show why retail results are supporting rather than headline evidence. In the cleaned SKU-day retail block, the direct ProZorro-to-retail coefficients are weak and short-window. Some older archive models show significant coefficients, but their signs and magnitudes are sensitive to matching rules and peak-coefficient flags. The reliable interpretation is that procurement and retail prices interact through a short-run pricing environment shaped by discounts and product matching, not that retail pass-through has been structurally identified with precision.",
            "Discounts are especially relevant for policy because they can obscure pass-through. A consumer may observe a lower promotional price at the same time that suppliers face higher procurement or input costs. Conversely, a retailer may maintain a high regular price while using promotions to manage demand. Both cases mean that observed retail prices are strategic objects, not passive reflections of upstream costs.",
            "The retail block also explains why the thesis should not conclude that all downstream power belongs to retailers. Processors may also have brand power, product specialization, or supply constraints. Public procurement may create its own price discipline. H2 is therefore written as a downstream market-power hypothesis, not as a claim that retailers alone determine every downstream price.",
        ],
    )
    add_picture(
        doc,
        figs["prozorro"],
        "Figure 7. ProZorro weekly median procurement prices by product.",
        "Source: ProZorro lot-level records from ProzorroM(full); author’s parsing and aggregation.",
    )
    add_picture(
        doc,
        figs["discounts"],
        "Figure 8. Retail discount incidence in Silpo and Novus observations.",
        "Source: author’s cleaned Silpo and Novus SKU-day retail data.",
    )
    add_table(doc, retail_table, "Table 7. H2 retail-support evidence", "Source: integrated evidence register. These models are supporting evidence because the retail window is short.", max_rows=12)
    add_many(
        doc,
        [
            "The retail-support models are mixed, which is itself informative. The cleaned retail rows are retained in the main table even when they are weak, because they are the most reliable product-repaired retail evidence. Several old and second-stage ProZorro-to-retail models are statistically significant, but many coefficients are large, negative, or sensitive to the matching definition. These are peak-coefficient and intersection risks. The thesis therefore does not write that retail prices mechanically follow procurement prices. It writes that the downstream retail layer has a separate pricing regime, where procurement prices, discounts, and product identity jointly shape observed retail prices.",
            "This interpretation follows the logic of Loy, Weiss, and Glauben (2016). They do not stop after finding asymmetric pass-through. They ask what explains heterogeneity in pass-through and warn that market power is not the only explanation. Retail asymmetry may reflect bargaining power, but it may also reflect menu costs, stock-out risk, product perishability, and promotion schedules. The policy interpretation therefore focuses on risk allocation and transparency rather than on direct accusations.",
        ],
    )
    add_table(doc, loy_table, "Table 8. Loy-style mechanism check", "Source: author’s second-stage summary based on first-stage threshold error-correction measures.", max_rows=8)
    add_many(
        doc,
        [
        "The Loy-style table is the explicit bridge between the price-transmission estimates and the market-power interpretation. The first-stage models generate speed and asymmetry measures. The second-stage screen asks whether these measures are associated with product or channel characteristics. The retail-link coefficient is positive and statistically meaningful for the speed measure, while the asymmetry equation is weak. This supports a cautious reading: retail-linked pairs differ in adjustment speed, but the available Ukrainian data do not support a strong claim that product perishability or retail linkage alone explains all asymmetry.",
            "This is why the thesis keeps Loy-style evidence as a mechanism check. It strengthens the argument that downstream pricing is managed differently across channels, but it does not replace the official monthly H1 and H2 models.",
        ],
    )
    add_p(doc, "Supporting Evidence from Old Model and Extra Folders", style="Heading 2")
    add_many(
        doc,
        [
            "The earlier model archive, final research folder, extra folder, and second-stage outputs are used as validation and mechanism evidence. Their strongest contribution is not to replace the observed monthly results, but to test whether similar mechanisms appear when higher-frequency reconstructed panels are used. The validation table is the gatekeeper for this decision.",
            "The old FarmGate reconstruction validates strongly against the observed Farm_milk_2015 series, with correlation above 0.96 over the overlap. Several old processor reconstructions also validate well, especially kefir, butter, and sour cream, while skimmed milk powder validates weakly. The thesis therefore uses old reconstructed results selectively: strong validation can support a mechanism; weak validation cannot support a main result.",
            "The extra equilibrium and retail short-run models are treated as robustness and appendix evidence. They are valuable for understanding possible dynamics, but they often rely on short intersections, reconstructed data, or broad product matches. Keeping them outside the headline results makes the thesis more credible and easier to grade.",
        ],
    )
    add_table(doc, validation_table, "Table 9. Old-new dataset validation", "Source: author’s comparison of reconstructed sheets with observed official data.", max_rows=10)
    add_picture(
        doc,
        figs["coefficients"],
        "Figure 9. Selected main pass-through coefficients.",
        "Source: author’s model selection table after reliability screening.",
    )
    add_p(
        doc,
        "Figure 9 deliberately keeps only selected coefficients after the reliability screen. Teal bars show main H1 farm-gate-to-processor evidence, blue bars show main H2 processor-to-consumer evidence, and grey bars show supporting evidence that is informative but weaker. The one-to-one line is not a target value; it is a reading benchmark. Coefficients below it imply incomplete long-run pass-through, while coefficients above it require cautious interpretation because they may reflect product composition, margins, or short-sample sensitivity rather than pure market power.",
    )
    add_p(doc, "Regional Extension", style="Heading 2")
    add_many(
        doc,
        [
            "The regional extension is short by design. It is included because market power may be local, especially in raw-milk procurement. However, the main processor-level SSSU data are national, so a full regional H1 model would mix regional upstream prices with national processor prices. The correct use is descriptive: show regional dispersion, then return to national models for inference.",
            "For H1, regional farm-gate dispersion indicates that farmers face different local price environments. This can reflect transport costs, processor density, local quality composition, and bargaining conditions. For H2, regional official consumer prices and ProZorro organizer regions can be used to describe downstream heterogeneity, but not as a substitute for a national processor-to-consumer model.",
            "The extension therefore supports the policy message rather than the econometric proof: national dairy policy should consider that market power can be spatially uneven. The same national average may hide stronger monopsony-like conditions in some raw-milk procurement areas and different retail pass-through patterns in consumer markets.",
        ],
    )
    add_p(doc, "Integrated Interpretation", style="Heading 2")
    add_many(
        doc,
        [
            "Taken together, the results support a market-power interpretation, but not an overclaim. H1 is better supported than H2 because the official monthly farm-gate and processor-level data are longer and cleaner. H2 is supported through official processor-consumer price links and through shorter ProZorro and retail mechanism evidence. The retail block is important for the story, but it is not strong enough to be the only proof of downstream market power.",
            "The strongest economic finding is that the processing layer is not just a technical transformation stage. It is the hinge where raw milk price pressure becomes product-level price pressure. The second finding is that downstream prices are managed through institutional procurement and retail promotions, not transmitted mechanically. The third finding is that data quality determines how strongly each result can be interpreted.",
        ],
    )
    add_p(doc, "Reliability Assessment by Evidence Block", style="Heading 2")
    add_many(
        doc,
        [
            "The monthly Ukrainian official data models are reliable as the main empirical base. Their main strength is that they are observed, national, and long enough for long-run adjustment analysis. Their limitation is frequency: monthly data cannot observe within-month promotions or tender timing. The thesis therefore treats them as structural evidence, not as the only description of price behavior.",
            "The ProZorro aggregation models are probable and need validation. Their strength is that they observe institutional procurement transactions with product titles, quantities, organizers, and regions. Their weakness is heterogeneity: tenders differ in size, product specification, contract conditions, and unit reliability. Median aggregation reduces outlier risk, but it cannot eliminate all contract heterogeneity.",
            "The Silpo and Novus retail models are probable mechanism evidence. Their strength is SKU-level detail and direct observation of discounts. Their weakness is the short time period and classification risk. After repair, the retail data are good for explaining promotion behavior, but they remain too short for a confident long-run market-power estimate.",
            "The discount-effect models are probable mechanism evidence. They are important because promotions are a real retail pricing instrument, but they are not a separate hypothesis. The thesis uses them to explain why retail prices may not fully or immediately follow procurement prices.",
            "The official consumer price comparison is reliable for national monthly H2 models where the product match is clean. It is less reliable where processor and consumer product categories differ. Drinking milk and sour cream are therefore stronger than cheese comparisons.",
            "The EU-based interpolation and reconstruction models are appendix-only unless they pass correlation and error checks against observed Ukrainian monthly data. A validated reconstruction can support robustness; an unvalidated reconstruction cannot be a main result.",
            "The ARDL, ECM, VECM, NARDL, daily, and weekly archive outputs are useful as a model archive. Their role in the thesis depends on validation, overlap, and interpretability. A model with many daily observations is not automatically better if the underlying daily series is reconstructed. A model with a significant coefficient is not useful if the coefficient is economically implausible or product mapping is weak.",
            "This reliability assessment is the reason the final results chapter is narrower than the full output folder. A high-quality master’s thesis should prefer fewer defensible results over many fragile ones. The appendix preserves transparency, while the main text protects the research argument.",
        ],
    )
    add_p(doc, "What the Results Do Not Show", style="Heading 2")
    add_many(
        doc,
        [
            "The results do not show that every actor in the dairy chain exercises market power in the same way. They also do not show that market power is the only cause of delayed or asymmetric adjustment. The models cannot separate every cost shock, contract term, quality change, logistics disruption, and retail strategy. The thesis therefore uses the phrase evidence consistent with market power rather than proof of abuse.",
            "The results do not justify treating farm-gate-to-retail models as the main evidence. Raw milk and retail dairy products are different economic objects. Direct farm-to-retail models can be useful as broad robustness checks, but the final thesis correctly routes the main argument through the processor layer.",
            "The results do not justify using EU reconstruction as a main empirical base. EU prices can help with interpolation experiments, but without strong validation, EU movements cannot be assumed to represent Ukrainian prices. The thesis therefore uses observed Ukrainian data first and reconstruction only as appendix or supporting material.",
            "The results do not justify overloading the reader with every table. Many archived outputs are statistically interesting but narratively harmful because they are weakly mapped, short, or hard to interpret. The main text stays focused and documents the broader screen in the appendix.",
        ],
    )

    # Conclusions.
    add_page_title(doc, "Conclusions")
    add_many(
        doc,
        [
            "This thesis reframed the research from generic vertical price transmission to market power in the dairy value chain in Ukraine. That reframing changes the logic of the whole paper. The question is no longer whether dairy prices move together. The question is whether price adjustment reveals bargaining asymmetry between farms, processors, procurement channels, retailers, and consumers.",
            "The first hypothesis is supported with caution. Farm-gate raw milk and processor-level dairy prices are economically linked in the observed national monthly data. The long-run slopes are meaningful for several processed dairy products, but adjustment is not uniformly fast or complete. The strongest H1 evidence therefore concerns the existence of a processor-stage price link and slow or selective correction, not a direct legal proof of monopsony. In a perishable raw-milk market, this pattern is consistent with processor bargaining power because farmers have limited ability to delay sale or search for alternative buyers.",
            "The second hypothesis is supported more moderately. Processor-to-consumer official models provide the strongest downstream price-link evidence, especially where the product match is clean. ProZorro-to-consumer evidence is useful but short. ProZorro-to-retail and retail-discount models show that downstream price setting is active and product-specific, but the retail window is too short for a full structural conclusion. The correct conclusion is not that every retail coefficient proves market power; it is that downstream actors can shape how procurement shocks become visible at the shelf.",
            "The policy implication is that market-power analysis in dairy should pay attention to timing and risk allocation. For farmers, the main concern is whether raw milk price adjustments reflect competitive demand or whether processors can delay and smooth payments in their favor. For processors, the concern is whether retailers and institutional buyers use procurement terms, payment delays, and shelf power to shift adjustment costs upstream. For consumers, the concern is whether promotions and delayed pass-through obscure the relationship between costs and final prices.",
            "The legal and institutional materials are important but must be interpreted carefully. AMCU cases and retailer-payment discussions motivate the research because they show that bargaining power, payment terms, and procurement conditions are real policy issues in Ukrainian food chains. They do not prove that the firms in the empirical sample violated competition law. The thesis provides economic evidence consistent with market-power mechanisms, not a legal judgment.",
            "The main limitation is data. Official monthly data are reliable but not high-frequency. Retail data are high-frequency but short. ProZorro data are detailed but irregular and product-heterogeneous. Old reconstructed daily and weekly datasets are useful only after validation against observed data. Future research should combine official price data with confidential contract terms, processor margins, payment-delay information, and longer retailer panels.",
            "The final contribution is methodological discipline. A narrower thesis with two hypotheses, a clear value-chain architecture, an explicit data audit, and a reliability screen is more defensible than a larger thesis that reports every possible model. The cleaned evidence shows where market-power mechanisms are most credible and where the evidence should remain supporting or appendix-only. That is the basis for a clearer, more policy-relevant master’s thesis.",
        ],
    )
    add_p(doc, "Policy Recommendations", style="Heading 2")
    add_bullets(
        doc,
        [
            "Recommendation 1. Improve public monitoring of raw-milk procurement prices by region and buyer type, because national averages can hide local monopsony-like conditions.",
            "Recommendation 2. Develop clearer reporting on payment terms between retailers and food suppliers, especially for perishable products, because delayed payment can transfer liquidity risk upstream.",
            "Recommendation 3. Treat ProZorro dairy procurement as a useful institutional price signal, but improve unit and product standardization so that public procurement data can be used more reliably for market monitoring.",
            "Recommendation 4. Encourage competition-policy analysis that separates legal abuse from economic evidence of bargaining power. Price-transmission models are useful screening tools, not substitutes for case-specific investigations.",
            "Recommendation 5. Build longer retail price panels with SKU identifiers, package sizes, and discount variables. Without longer retail panels, the retail part of H2 will remain mechanism evidence rather than structural proof.",
        ],
    )

    # Works cited.
    add_page_title(doc, "Works Cited")
    works = [
        "AMCU Northern Interregional Territorial Department. n.d. “PJSC Yagotynsky Butter Plant Fined for Abuse of Monopoly Position During Milk Procurement from Households.” Accessed May 8, 2026. https://northmtv.amcu.gov.ua/news/pat-yagotinskiy-maslozavod-oshtrafovano-za-zlovzhivannya-monopolnim-stanovishchem-pid-chas-zakupivli-moloka-u-naselennya-2.",
        "AMCU Western Interregional Territorial Department. 2023. “Milk Procurement at Understated Prices: Western Territorial Department Fined a Cheese Plant in Volyn.” https://westernmtv.amcu.gov.ua/news/zakupilya-moloka-za-zanizhenimi-cinami-zahidne-mtv-oshtrafuvalo-sirzavod-na-volini.",
        "AgroTimes. 2025. “Industrial Segment Produced 3 Million Tonnes of Milk in 2024.” May 27. https://agrotimes.ua/tvarinnitstvo/promyslovyj-segment-u-2024-roczi-vyrobyv-3-mln-tonn-moloka/.",
        "Antimonopoly Committee of Ukraine. 2023. “AMCU Provided Its Conclusion Regarding the Actions of the Largest Retail Chains under the Memorandum.” https://amcu.gov.ua/news/amku-nadav-svij-visnovok-shchodo-dij-najbilshih-torgovelnih-merezh-v-ramkah-vikonannya-polozhen-ukladenogo-mizh-nimi-memorandumu.",
        "Ba, Helene A., Yann de Mey, Sylvie Thoron, and Matty Demont. 2019. “Inclusiveness of Contract Farming along the Vertical Coordination Continuum: Evidence from the Vietnamese Rice Sector.” Land Use Policy 87: 104050. https://doi.org/10.1016/j.landusepol.2019.104050.",
        "Brümmer, Bernhard, Stephan von Cramon-Taubadel, and Sergiy Zorya. 2009. “The Impact of Market and Policy Instability on Price Transmission between Wheat and Flour in Ukraine.” European Review of Agricultural Economics 36 (2): 203-230. https://doi.org/10.1093/erae/jbp021.",
        "Cabinet of Ministers of Ukraine. 2024. “Strategy of Food Security of Ukraine for the Period until 2027.” Order No. 684-r, July 23. https://zakon.rada.gov.ua/laws/show/684-2024-%D1%80.",
        "Chevalier, Judith A., Anil K. Kashyap, and Peter E. Rossi. 2003. “Why Don’t Prices Rise during Periods of Peak Demand? Evidence from Scanner Data.” American Economic Review 93 (1): 15-37.",
        "Chow, Gregory C., and An-loh Lin. 1971. “Best Linear Unbiased Interpolation, Distribution, and Extrapolation of Time Series by Related Series.” Review of Economics and Statistics 53 (4): 372-375.",
        "CME Group. n.d. “Class III Milk Futures.” Accessed May 8, 2026. https://www.cmegroup.com/markets/agriculture/dairy/class-iii-milk.html.",
        "Denton, Frank T. 1971. “Adjustment of Monthly or Quarterly Series to Annual Totals: An Approach Based on Quadratic Minimization.” Journal of the American Statistical Association 66 (333): 99-102.",
        "Enders, Walter, and Pierre L. Siklos. 2001. “Cointegration and Threshold Adjustment.” Journal of Business & Economic Statistics 19 (2): 166-176. https://doi.org/10.1198/073500101316970395.",
        "Engle, Robert F., and Clive W. J. Granger. 1987. “Co-integration and Error Correction: Representation, Estimation, and Testing.” Econometrica 55 (2): 251-276. https://doi.org/10.2307/1913236.",
        "European Commission Directorate-General for Agriculture and Rural Development. 2026. “Milk and Dairy Products Market Observatory.” Accessed May 8, 2026. https://agriculture.ec.europa.eu/data-and-analysis/markets/overviews/market-observatories/milk_en.",
        "Food and Agriculture Organization of the United Nations. 2013. Milk and Dairy Products Supply Chain in Ukraine. Rome: Food and Agriculture Organization of the United Nations. https://www.fao.org/4/i3446e/i3446e.pdf.",
        "International Trade Centre. 2024. Trade Map Data Workbooks for Ukraine, Harmonized System chapter 04 (HS 04) Dairy Produce and Related Animal-Origin Products. Project files share_dairy_exp.xlsx, share_dairy_imp.xlsx, and biggest_trade.xlsx.",
        "Johansen, Søren. 1988. “Statistical Analysis of Cointegration Vectors.” Journal of Economic Dynamics and Control 12: 231-254. https://doi.org/10.1016/0165-1889(88)90041-3.",
        "Kinnucan, Henry W., and Olan D. Forker. 1987. “Asymmetry in Farm-Retail Price Transmission for Major Dairy Products.” American Journal of Agricultural Economics 69 (2): 285-292. https://doi.org/10.2307/1242278.",
        "Kinnucan, Henry W., and Daowei Zhang. 2015. “Notes on Farm-Retail Price Transmission and Marketing Margin Behavior.” Agricultural Economics 46 (6): 729-737. https://doi.org/10.1111/agec.12194.",
        "Koester, Ulrich, and Stephan von Cramon-Taubadel, eds. 2023. Agricultural Price Formation in Theory and Reality. Newcastle upon Tyne: Cambridge Scholars Publishing.",
        "Litvinov, Valentyn. 2025a. Eurointegration 2.0: The Ukrainian Dairy Sector. Kyiv: Kyiv School of Economics, Center for Food and Land Use Research. https://kse.ua/wp-content/uploads/2025/10/Eurointegration-2.0.-The-Ukrainian-Dairy-Sector.pdf.",
        "Litvinov, Valentyn. 2025b. Eurointegration 2.0: The Ukrainian Dairy Sector. Volume 2: Dairy Supply Chain Bottlenecks. Kyiv: Kyiv School of Economics, Center for Food and Land Use Research. https://kse.ua/wp-content/uploads/2025/12/Eurointegration-2.0.-The-Ukrainian-Dairy-Sector-Volume-2-Dairy-Supply-Chain-Bottlenecks.pdf.",
        "Loy, Jens-Peter, Christoph R. Weiss, and Thomas Glauben. 2016. “Asymmetric Cost Pass-Through? Empirical Evidence on the Role of Market Power, Search and Menu Costs.” Journal of Economic Behavior & Organization 123: 184-192. https://doi.org/10.1016/j.jebo.2016.01.007.",
        "Loy, Jens-Peter. 2023. “Price Formation in Food Retailing.” In Agricultural Price Formation in Theory and Reality, edited by Ulrich Koester and Stephan von Cramon-Taubadel. Newcastle upon Tyne: Cambridge Scholars Publishing.",
        "Martinez, Steve W. 2002. Vertical Coordination of Marketing Systems: Lessons from the Poultry, Egg, and Pork Industries. Agricultural Economic Report No. 807. Washington, DC: USDA Economic Research Service. https://www.ers.usda.gov/webdocs/publications/42151/16919_aer807_1_.pdf.",
        "Meyer, Jochen, and Stephan von Cramon-Taubadel. 2004. “Asymmetric Price Transmission: A Survey.” Journal of Agricultural Economics 55 (3): 581-611. https://doi.org/10.1111/j.1477-9552.2004.tb00116.x.",
        "MilkUA.info. 2025. “In 2024, 10.4 Percent More Milk Was Received for Processing.” February 7. https://milkua.info/uk/post/u-2024-roci-na-pererobku-nadijslo-na-104-bilse-moloka.",
        "MilkUA.info. 2026. “Milk Prices under Pressure from Weak Demand, Blackouts and Frost.” January 22. https://milkua.info/en/post/milk-prices-under-pressure-from-weak-demand-blackouts-and-frost.",
        "Nakamura, Emi, and Jón Steinsson. 2008. “Five Facts about Prices: A Reevaluation of Menu Cost Models.” Quarterly Journal of Economics 123 (4): 1415-1464. https://doi.org/10.1162/qjec.2008.123.4.1415.",
        "Nivievskyi, Oleg, and Stephan von Cramon-Taubadel. 2008. “The Determinants of Dairy Farming Competitiveness in Ukraine.” Paper presented at the 12th EAAE Congress, Ghent, August 26-29. https://ageconsearch.umn.edu/record/44059/files/329.pdf.",
        "Nivievskyi, Oleg, and Stephan von Cramon-Taubadel. 2015. Dairy Supply Chain in Belarus: Bottlenecks and Scope for Improvements. https://www.researchgate.net/publication/282117525_Dairy_Supply_Chain_in_Belarus_Bottlenecks_and_Scope_for_Improvements.",
        "Nivievskyi, Oleg. 2009. “Price Support, Efficiency and Technology Change of Ukrainian Dairy Farms: Spatial Dependence in the Components of Productivity Growth.” Paper presented at the International Association of Agricultural Economists Conference, Beijing, August 16-22. https://ageconsearch.umn.edu/record/51403.",
        "Nivievskyi, Oleg. 2012. Increasing the Competitiveness of the Dairy Supply Chain in Ukraine: Role of the Government. German-Ukrainian Agricultural Policy Dialogue Policy Paper Series APD/PP/03/2012. Kyiv. https://www.ier.com.ua/files/publications/Policy_papers/Agriculture_dialogue/2012/APD_PP_2012_3_Dairy_Ukraine_en.pdf.",
        "Novus. n.d. “Online Supermarket.” Accessed May 8, 2026. https://novus.ua/.",
        "Opendatabot. 2025. “Retail Companies in the 2025 Opendatabot Index Saw Their Revenue Grow by 17%.” March 24. https://opendatabot.ua/en/analytics/index-retail-2025.",
        "Pesaran, M. Hashem, Yongcheol Shin, and Richard J. Smith. 2001. “Bounds Testing Approaches to the Analysis of Level Relationships.” Journal of Applied Econometrics 16 (3): 289-326. https://doi.org/10.1002/jae.616.",
        "ProZorro. n.d. “Public Procurement System and Business Intelligence Data Portal.” Accessed May 8, 2026. https://bi.prozorro.org/hub/.",
        "Sexton, Richard J. 2013. “Market Power, Misconceptions, and Modern Agricultural Markets.” American Journal of Agricultural Economics 95 (2): 209-219. https://doi.org/10.1093/ajae/aas102.",
        "Shin, Yongcheol, Byungchul Yu, and Matthew Greenwood-Nimmo. 2014. “Modelling Asymmetric Cointegration and Dynamic Multipliers in a Nonlinear ARDL Framework.” In Festschrift in Honor of Peter Schmidt: Econometric Methods and Applications, edited by William C. Horrace and Robin C. Sickles, 281-314. New York: Springer.",
        "Silpo. n.d. “Online Supermarket.” Accessed May 8, 2026. https://silpo.ua/.",
        "State Statistics Service of Ukraine. n.d. “Changes in Prices and Tariffs for Consumer Goods and Services.” Data.gov.ua dataset. Accessed May 8, 2026. https://data.gov.ua/dataset/c1cd8f39-c2ee-4f20-8c8a-6cdde6e6f343.",
        "State Statistics Service of Ukraine. n.d. “Changes in Producer Prices of Industrial Products.” Data.gov.ua dataset. Accessed May 8, 2026. https://data.gov.ua/dataset/332cd85e-5cab-47bf-8e62-4157cd8586c9.",
        "State Statistics Service of Ukraine. n.d. “Realization of Agricultural Products by Enterprises and Households.” Data.gov.ua dataset. Accessed May 8, 2026. https://data.gov.ua/dataset/0ecfe49a-49ff-4535-b3a4-537714329488.",
        "State Statistics Service of Ukraine. 2025. Agricultural Production Cost Index Workbook, Livestock Products. Project file індекс витрат сільгосп.xlsx.",
        "Tarassevych, Oleksandr. 2024. Dairy and Products Annual: Ukraine. USDA Foreign Agricultural Service, Report UP2024-0021. October 23.",
        "USDA Economic Research Service. 1999. Vertical Coordination and Consumer Welfare. Agricultural Economic Report No. 753. Washington, DC: USDA Economic Research Service. https://ers.usda.gov/sites/default/files/_laserfiche/publications/79929/AER-753.pdf.",
        "Vavra, Peter, and Barry K. Goodwin. 2005. Analysis of Price Transmission along the Food Chain. OECD Food, Agriculture and Fisheries Working Papers, No. 3. Paris: OECD Publishing. https://www.oecd.org/agriculture/price-transmission-in-the-food-chain.htm.",
    ]
    seen_refs = set()
    for w in works:
        key = re.sub(r"\s+", " ", w).strip().lower()
        if not key or key in seen_refs:
            continue
        seen_refs.add(key)
        add_p(doc, w)

    # Appendix.
    add_page_title(doc, "Appendix")
    add_p(doc, "Appendix A. Reliability Screen", style="Heading 2")
    add_table(doc, counts, "Table 10. Reliability screen by source block", "Source: author’s integrated evidence register.", max_rows=25)
    add_p(doc, "Appendix B. Abbreviation and Variable Discipline", style="Heading 2")
    add_table(
        doc,
        pd.DataFrame(
            [
                {"term": "farm-gate raw milk price", "variable meaning": "upstream price received by agricultural enterprises for milk sold", "raw source name": "Середня ціна продукції сільського господарства, реалізованої підприємствами"},
                {"term": "processor-level price", "variable meaning": "industrial producer price of processed dairy output", "raw source name": "Середні ціни виробників промислової продукції"},
                {"term": "official consumer price", "variable meaning": "SSSU average consumer price for dairy product rows", "raw source name": "Середні споживчі ціни на товари (послуги)"},
                {"term": "ProZorro procurement price", "variable meaning": "institutional procurement lot price after unit conversion", "raw source name": "Ціна за одиницю in ProzorroM(full)"},
                {"term": "retail SKU-day price", "variable meaning": "Silpo or Novus product-day observation with product and discount fields", "raw source name": "product_title, product_name, price_current, unit_price"},
            ]
        ),
        "Table 11. Abbreviation discipline and variable definitions",
        "Source: author’s source-name reconciliation.",
        max_rows=10,
    )
    add_p(doc, "Appendix C. Notes on Excluded or Down-Ranked Models", style="Heading 2")
    add_many(
        doc,
        [
            "Models are excluded from the main text when they fail at least one of the following checks: weak product mapping, too-short overlap, coefficient sign inconsistent with the economic object, peak coefficient flag, weak old-new validation, serial-correlation problems, or dependence on EU-based reconstruction without adequate validation.",
            "The appendix-only label is not a failure to disclose evidence. It is a discipline rule. A master’s thesis should show that weak results were considered and then placed where they belong. This prevents the empirical chapter from being overloaded and makes the final conclusion more credible.",
        ],
    )

    doc.save(DOCX_OUT)
    return qa(DOCX_OUT)


def qa(path: Path) -> dict:
    info = {
        "docx": str(path),
        "exists": path.exists(),
        "size_bytes": path.stat().st_size if path.exists() else 0,
    }
    with zipfile.ZipFile(path) as z:
        bad = z.testzip()
        names = z.namelist()
        imgs = [n for n in names if n.startswith("word/media/")]
        info["zip_test"] = bad
        info["embedded_images"] = len(imgs)
        info["has_styles"] = "word/styles.xml" in names
        info["has_settings"] = "word/settings.xml" in names
        info["has_update_fields"] = "w:updateFields" in z.read("word/settings.xml").decode("utf-8", errors="ignore")
    d = Document(path)
    text = "\n".join(p.text for p in d.paragraphs)
    info["paragraphs"] = len(d.paragraphs)
    info["tables"] = len(d.tables)
    info["word_count"] = sum(len(p.text.split()) for p in d.paragraphs)
    info["visible_toc_placeholder"] = "Open in Word and update" in text
    info["contains_main_oldmodel_warning"] = "old-model sheets" in text and "supporting or appendix evidence" in text
    info["contains_data_audit"] = "Processor_price is reported in hryvnia per tonne" in text
    info["contains_source_names_ua"] = "Середні ціни виробників промислової продукції" in text
    # Check images copied into the final output folder.
    checked = []
    for img in FIG_OUT.glob("*.png"):
        with Image.open(img) as im:
            checked.append({"file": img.name, "size": im.size})
    info["figure_files_checked"] = checked
    (REPORT_OUT / "final_v10_docx_qa.json").write_text(json.dumps(info, indent=2, ensure_ascii=False), encoding="utf-8")
    report = [
        "# Final V10 Build Report",
        "",
        "## What changed",
        "",
        "- Built from `draft/Maksym_Charniuk_MSc_thesis_draft_2.docx` to preserve Draft 2 styles.",
        "- Preserved the Draft 2 literature-review chapter structure and content, with one framing sentence.",
        "- Rewrote Introduction, Market Analysis, Methodology, Data, Results, Conclusion, Works Cited, and Appendix around market power and two hypotheses.",
        "- Used `newmodel.xlsx` as the main dataset and explicitly named raw SSSU source variables in Ukrainian.",
        "- Added mandatory audit coverage for farm-gate tonne magnitudes, processor tonne units, ProZorro non-breaking-space numbers, and retail product reclassification.",
        "- Demoted old reconstructed and retail short-window evidence unless validation and overlap justify supporting use.",
        "- Cleaned the Works Cited list into a curated Chicago author-date bibliography and repaired preserved literature-review citation punctuation.",
        "- Added a Loy-style descriptive-statistics table for the cleaned Chapter 5 data and clearer β/λ result-table notation.",
        "- Added a short OLS/IV/Lerner Index identification boundary, rather than estimating an unreliable direct markup model from unavailable margin and cost data.",
        "- Added concise 2024 sector facts from AgroTimes, MilkUA.info, and Opendatabot to strengthen the market-structure motivation.",
        "- Avoided LibreOffice/soffice validation because the local application crashes; used DOCX package QA instead.",
        "",
        "## QA",
        "",
        "```json",
        json.dumps(info, indent=2, ensure_ascii=False),
        "```",
    ]
    (REPORT_OUT / "final_v10_build_report.md").write_text("\n".join(report), encoding="utf-8")
    return info


if __name__ == "__main__":
    print(json.dumps(build_doc(), indent=2, ensure_ascii=False))
