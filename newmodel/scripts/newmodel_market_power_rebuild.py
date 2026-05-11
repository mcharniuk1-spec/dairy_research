#!/usr/bin/env python3
"""Rebuild the Ukrainian dairy market-power thesis evidence package.

The runner is intentionally explicit: it audits the new raw workbooks, builds
clean comparable UAH/kg datasets, estimates conservative price-transmission
models, exports figures/tables, and creates a DOCX draft based on the Draft 2
style shell. It does not mutate source workbooks or source drafts.
"""

from __future__ import annotations

import json
import math
import os
import re
import shutil
import textwrap
import warnings
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd

warnings.filterwarnings("ignore")

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.diagnostic import acorr_ljungbox, het_breuschpagan
from statsmodels.tsa.stattools import adfuller, coint, kpss

try:
    from statsmodels.tsa.vector_ar.vecm import VECM, select_coint_rank, select_order
except Exception:  # pragma: no cover - optional statsmodels component
    VECM = None
    select_coint_rank = None
    select_order = None

try:
    from statsmodels.tsa.ardl import ardl_select_order
except Exception:  # pragma: no cover - optional statsmodels component
    ardl_select_order = None

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
PACKAGE_ROOT = SCRIPT_DIR.parent
REPO_ROOT = PACKAGE_ROOT.parent
LEGACY_ROOT = Path("/Users/getapple/Documents/KSE/Master Thesis")


def resolve_root() -> Path:
    """Prefer the portable GitHub package; allow THESIS_ROOT for local reruns."""
    env_root = os.environ.get("THESIS_ROOT")
    if env_root:
        return Path(env_root).expanduser().resolve()
    if (PACKAGE_ROOT / "data" / "Newmodel_data" / "newmodel.xlsx").exists():
        return PACKAGE_ROOT
    return LEGACY_ROOT


ROOT = resolve_root()
OUT = ROOT / "outputs" / "newmodel_rebuild"
DATA_OUT = OUT / "clean_data"
FIG_OUT = OUT / "figures"
TABLE_OUT = OUT / "tables"
DOC_OUT = ROOT / "doc"

DATA_ROOT = ROOT / "data" / "Newmodel_data" if (ROOT / "data" / "Newmodel_data").exists() else ROOT / "Newmodel_data"
SOURCE_DOCS = ROOT / "doc" / "source" if (ROOT / "doc" / "source").exists() else ROOT
NEW = DATA_ROOT / "newmodel.xlsx"
PARTS = DATA_ROOT / "parts of main dataset"
ADD = DATA_ROOT / "additioanl materials"
FINAL_RESEARCH = REPO_ROOT if (REPO_ROOT / "outputs").exists() else ROOT / "FINAL_RESEARCH"
DRAFT2 = SOURCE_DOCS / "Maksym_Charniuk_MSc_thesis_draft_2.docx"
DRAFT3 = SOURCE_DOCS / "Maksym_Charniuk_MSc_thesis_draft_3.docx"
COMMENTED = SOURCE_DOCS / "Commented_draft2.docx"
TRANSCRIPT = SOURCE_DOCS / "Nivievskyi_5_05_transcript.docx"
LOY_PDF = ROOT / "references" / "loy2016.pdf"


SOURCE_LINKS = {
    "Food Security Strategy": "https://zakon.rada.gov.ua/laws/show/684-2024-%D1%80",
    "SSSU SDG 2.1.2": "https://sdg.ukrstat.gov.ua/uk/2-1-2/",
    "SSSU farm-gate agricultural sales": "https://data.gov.ua/dataset/0ecfe49a-49ff-4535-b3a4-537714329488",
    "SSSU producer prices": "https://data.gov.ua/dataset/332cd85e-5cab-47bf-8e62-4157cd8586c9",
    "SSSU consumer prices": "https://data.gov.ua/dataset/c1cd8f39-c2ee-4f20-8c8a-6cdde6e6f343",
    "SSSU economic accounts for agriculture": "https://data.gov.ua/dataset/91d1041e-5668-49fc-88a4-1c8b8ba08043",
    "ProZorro BI": "https://bi.prozorro.org/hub/",
    "Loy et al. 2016": "https://doi.org/10.1016/j.jebo.2016.01.007",
    "AMCU retailer memorandum": "https://amcu.gov.ua/news/amku-nadav-svij-visnovok-shchodo-dij-najbilshih-torgovelnih-merezh-v-ramkah-vikonannya-polozhen-ukladenogo-mizh-nimi-memorandumu",
    "AMCU AIC dialogue": "https://amcu.gov.ua/news/amku-prodovzhue-dialog-z-predstavnikami-rinkiv-agropromislovogo-kompleksu",
    "AMCU Yagotynsky case": "https://northmtv.amcu.gov.ua/news/pat-yagotinskiy-maslozavod-oshtrafovano-za-zlovzhivannya-monopolnim-stanovishchem-pid-chas-zakupivli-moloka-u-naselennya-2",
    "AMCU Rozhyshche case": "https://westernmtv.amcu.gov.ua/news/zakupilya-moloka-za-zanizhenimi-cinami-zahidne-mtv-oshtrafuvalo-sirzavod-na-volini",
}


PRODUCT_ORDER = [
    "raw_milk",
    "drinking_milk",
    "sour_cream",
    "kefir",
    "butter",
    "hard_cheese",
    "soft_cheese",
    "cottage_cheese",
    "skim_milk_powder",
    "cream",
    "condensed_milk",
    "yogurt",
    "other_dairy",
    "exclude_non_dairy",
]


PRODUCT_LABELS = {
    "raw_milk": "Raw milk",
    "drinking_milk": "Drinking milk",
    "sour_cream": "Sour cream",
    "kefir": "Kefir",
    "butter": "Butter",
    "hard_cheese": "Hard cheese",
    "soft_cheese": "Soft cheese",
    "cottage_cheese": "Cottage cheese",
    "skim_milk_powder": "Skim milk powder",
    "cream": "Cream",
    "condensed_milk": "Condensed milk",
    "yogurt": "Yogurt",
    "other_dairy": "Other dairy",
    "exclude_non_dairy": "Excluded non-dairy",
}


def ensure_dirs() -> None:
    for p in [OUT, DATA_OUT, FIG_OUT, TABLE_OUT, DOC_OUT]:
        p.mkdir(parents=True, exist_ok=True)


def clean_text(x) -> str:
    if pd.isna(x):
        return ""
    s = str(x)
    s = s.replace("&amp;", "&").replace("\xa0", " ").replace("\u202f", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def parse_num(x):
    if pd.isna(x):
        return np.nan
    if isinstance(x, (int, float, np.integer, np.floating)):
        return float(x)
    s = str(x)
    s = s.replace("\xa0", "").replace("\u202f", "").replace(" ", "")
    s = s.replace(",", ".")
    s = re.sub(r"[^0-9.\-]", "", s)
    if not s or s in {"-", ".", "-."}:
        return np.nan
    try:
        return float(s)
    except ValueError:
        return np.nan


def period_cols(df: pd.DataFrame, freq: str = "M") -> list[str]:
    if freq == "M":
        pattern = r"^\d{4}-M\d{2}$"
    elif freq == "MQ":
        pattern = r"^\d{4}-[MQ]\d{1,2}$"
    else:
        pattern = r"^\d{4}-[MQ]\d{1,2}$"
    return [c for c in df.columns if isinstance(c, str) and re.match(pattern, c)]


def period_to_date(p: str):
    m = re.match(r"^(\d{4})-M(\d{2})$", str(p))
    if m:
        return pd.Timestamp(year=int(m.group(1)), month=int(m.group(2)), day=1)
    q = re.match(r"^(\d{4})-Q(\d{1})$", str(p))
    if q:
        month = (int(q.group(2)) - 1) * 3 + 1
        return pd.Timestamp(year=int(q.group(1)), month=month, day=1)
    return pd.NaT


def to_markdown(df: pd.DataFrame, max_rows: int = 30, floatfmt: str = ".4g") -> str:
    if df is None or df.empty:
        return "_No rows._"
    d = df.copy().head(max_rows)
    for col in d.columns:
        if pd.api.types.is_float_dtype(d[col]):
            d[col] = d[col].map(lambda x: "" if pd.isna(x) else format(float(x), floatfmt))
        else:
            d[col] = d[col].map(lambda x: "" if pd.isna(x) else str(x))
    cols = [str(c) for c in d.columns]
    lines = ["| " + " | ".join(cols) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
    for _, row in d.iterrows():
        vals = [str(row[c]).replace("\n", " ") for c in d.columns]
        lines.append("| " + " | ".join(vals) + " |")
    if len(df) > max_rows:
        lines.append(f"\n_Showing {max_rows} of {len(df)} rows._")
    return "\n".join(lines)


def save_csv(df: pd.DataFrame, name: str) -> Path:
    path = DATA_OUT / name
    df.to_csv(path, index=False)
    return path


def normalize_key(s: str) -> str:
    s = clean_text(s).lower()
    s = s.replace("’", "'").replace("`", "'")
    s = re.sub(r"[^a-zа-яіїєґ0-9%., ]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def classify_product(*texts) -> str:
    txt = " ".join(normalize_key(t) for t in texts if clean_text(t))
    if not txt:
        return "exclude_non_dairy"

    plant_terms = [
        "соєв",
        "soja",
        "soya",
        "soy",
        "alpro",
        "sojasun",
        "кокос",
        "coconut",
        "мигдал",
        "almond",
        "рисов",
        "вівсян",
        "oat",
        "рослин",
    ]
    non_dairy_terms = [
        "яйц",
        "egg",
        "honey",
        "мед",
        "тістеч",
        "шоколад",
        "батончик",
        "печиво",
        "майонез",
        "маргарин",
        "морозиво",
        "ice cream",
    ]
    if any(t in txt for t in plant_terms):
        return "exclude_non_dairy"
    if any(t in txt for t in non_dairy_terms):
        return "exclude_non_dairy"

    if "молоко сух" in txt or "сухе молоко" in txt or "milk powder" in txt or "знежирен" in txt and "сух" in txt:
        return "skim_milk_powder"
    if "згущ" in txt or "condensed" in txt:
        return "condensed_milk"
    if "масло" in txt or "butter" in txt:
        if "верш" in txt or "солодковерш" in txt or "butter" in txt:
            return "butter"
    if "сметан" in txt or "sour cream" in txt:
        return "sour_cream"
    if "кефір" in txt or "kefir" in txt or "айран" in txt:
        return "kefir"
    if "йогурт" in txt or "yogurt" in txt or "yoghurt" in txt:
        return "yogurt"
    if "ряжан" in txt or "закваск" in txt:
        return "other_dairy"
    if ("вершки" in txt or "cream" in txt) and not ("молоко коров" in txt and ("пастер" in txt or "ультра" in txt)):
        return "cream"
    if "молоко" in txt or "milk" in txt:
        return "drinking_milk"
    if "кисломолоч" in txt or "творог" in txt or "cottage" in txt:
        return "cottage_cheese"
    if "сири м" in txt or "м'як" in txt and "сир" in txt or "soft cheese" in txt:
        return "soft_cheese"
    if "сир" in txt or "cheese" in txt or "бринз" in txt or "сулугун" in txt or "моцар" in txt:
        if "плавлен" in txt:
            return "other_dairy"
        if "кисломолоч" in txt or "творог" in txt:
            return "cottage_cheese"
        return "hard_cheese"

    return "exclude_non_dairy"


def parse_package_kg(*texts) -> float:
    txt = " ".join(normalize_key(t) for t in texts if clean_text(t))
    # Prefer explicit package units, not fat percentages.
    patterns = [
        (r"(\d+(?:[.,]\d+)?)\s*(?:кг|kg)\b", 1.0),
        (r"(\d+(?:[.,]\d+)?)\s*(?:г|g)\b", 0.001),
        (r"(\d+(?:[.,]\d+)?)\s*(?:л|l)\b", 1.0),
        (r"(\d+(?:[.,]\d+)?)\s*(?:мл|ml)\b", 0.001),
    ]
    candidates = []
    for pat, factor in patterns:
        for m in re.finditer(pat, txt):
            val = parse_num(m.group(1))
            if pd.notna(val):
                kg = val * factor
                if 0.03 <= kg <= 20:
                    candidates.append(kg)
    if candidates:
        return float(candidates[-1])
    return np.nan


def product_label(product: str) -> str:
    return PRODUCT_LABELS.get(product, product.replace("_", " ").title())


def load_doc_text(path: Path, limit: int | None = None) -> list[str]:
    doc = Document(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return texts[:limit] if limit else texts


def extract_doc_outline(path: Path) -> pd.DataFrame:
    doc = Document(path)
    rows = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        style = p.style.name if p.style is not None else ""
        if style.startswith("Heading") or style in {"Title", "Author"} or t.lower().startswith(("hypothesis", "figure ", "table ")):
            rows.append({"paragraph": i, "style": style, "text": t[:220]})
    return pd.DataFrame(rows)


def clean_farmgate() -> pd.DataFrame:
    df = pd.read_excel(PARTS / "Farm_milk_2015.xlsx")
    dates = period_cols(df, "M")
    id_cols = [c for c in df.columns if c not in dates]
    long = df.melt(id_vars=id_cols, value_vars=dates, var_name="period", value_name="price_raw")
    long["date"] = long["period"].map(period_to_date)
    long["price_raw"] = long["price_raw"].map(parse_num)
    long = long[long["Показник"].astype(str).str.contains("Середня ціна продукції", na=False)].copy()
    long["farm_type"] = np.where(
        long["Показник"].astype(str).str.contains("підприємствами", case=False, na=False),
        "enterprises",
        np.where(long["Показник"].astype(str).str.contains("населення", case=False, na=False), "households", "unknown"),
    )
    long["territory"] = long["Територіальний розріз"].map(clean_text)
    long["stage"] = "farm_gate"
    long["source"] = "SSSU agricultural sales raw workbook"
    long["product"] = "raw_milk"
    long["unit_raw"] = long["Одиниця виміру"].map(clean_text)
    long["unit_clean"] = "UAH/kg"
    long["price_uah_kg"] = long["price_raw"] / 1000.0
    long["quality_flag"] = np.select(
        [long["price_raw"].isna(), long["price_uah_kg"].le(0), long["price_uah_kg"].gt(100)],
        ["missing", "non_positive", "implausible_high"],
        default="ok",
    )
    keep = [
        "date",
        "period",
        "territory",
        "farm_type",
        "source",
        "stage",
        "product",
        "price_raw",
        "price_uah_kg",
        "unit_raw",
        "unit_clean",
        "quality_flag",
    ]
    return long[keep].sort_values(["territory", "farm_type", "date"])


def clean_processor() -> pd.DataFrame:
    df = pd.read_excel(PARTS / "Processor_2013_2026.xlsx")
    dates = period_cols(df, "M")
    id_cols = [c for c in df.columns if c not in dates]
    long = df.melt(id_vars=id_cols, value_vars=dates, var_name="period", value_name="price_raw")
    long["date"] = long["period"].map(period_to_date)
    long["price_raw"] = long["price_raw"].map(parse_num)
    long["territory"] = long["Територіальний розріз"].map(clean_text)
    long["stage"] = "processor"
    long["source"] = "SSSU industrial producer prices raw workbook"
    long["unit_raw"] = long["Одиниця виміру"].map(clean_text)
    long["unit_clean"] = "UAH/kg"
    long["product"] = long["Розріз"].map(lambda x: classify_product(x))
    # The official unit is hryvnia per tonne.
    long["price_uah_kg"] = long["price_raw"] / 1000.0
    long["quality_flag"] = np.select(
        [long["price_raw"].isna(), long["price_uah_kg"].le(0), long["price_uah_kg"].gt(1000)],
        ["missing", "non_positive", "implausible_high"],
        default="ok",
    )
    keep = [
        "date",
        "period",
        "territory",
        "source",
        "stage",
        "product",
        "Розріз",
        "price_raw",
        "price_uah_kg",
        "unit_raw",
        "unit_clean",
        "quality_flag",
    ]
    return long[keep].sort_values(["product", "date"])


def clean_consumer() -> pd.DataFrame:
    df = pd.read_excel(PARTS / "Consumer_2017_2026.xlsx")
    dates = period_cols(df, "M")
    id_cols = [c for c in df.columns if c not in dates]
    long = df.melt(id_vars=id_cols, value_vars=dates, var_name="period", value_name="price_raw")
    long["date"] = long["period"].map(period_to_date)
    long["price_raw"] = long["price_raw"].map(parse_num)
    long["territory"] = long["Територіальний розріз"].map(clean_text)
    long["stage"] = "official_consumer"
    long["source"] = "SSSU consumer prices raw workbook"
    long["unit_raw"] = long["Одиниця виміру"].map(clean_text)
    long["unit_clean"] = "UAH/kg"
    long["product"] = long["Тип товарів і послуг"].map(lambda x: classify_product(x))
    # Units are 1000 g or kilogram. Both are UAH/kg.
    long["price_uah_kg"] = long["price_raw"]
    long["quality_flag"] = np.select(
        [long["price_raw"].isna(), long["price_uah_kg"].le(0), long["price_uah_kg"].gt(1000)],
        ["missing", "non_positive", "implausible_high"],
        default="ok",
    )
    keep = [
        "date",
        "period",
        "territory",
        "source",
        "stage",
        "product",
        "Тип товарів і послуг",
        "price_raw",
        "price_uah_kg",
        "unit_raw",
        "unit_clean",
        "quality_flag",
    ]
    return long[keep].sort_values(["territory", "product", "date"])


def clean_farm_volumes() -> pd.DataFrame:
    path = PARTS / "farm_volumes.xlsx"
    df = pd.read_excel(path)
    dates = period_cols(df, "M")
    id_cols = [c for c in df.columns if c not in dates]
    long = df.melt(id_vars=id_cols, value_vars=dates, var_name="period", value_name="value_raw")
    long["date"] = long["period"].map(period_to_date)
    long["value"] = long["value_raw"].map(parse_num)
    long["territory"] = long["Територіальний розріз"].map(clean_text)
    long["indicator"] = long["Показник"].map(clean_text)
    long["flow"] = long["Розріз"].map(clean_text)
    long["milk_product_raw"] = long["Молоко та молочна продукція"].map(clean_text)
    long["unit_raw"] = long["Одиниця виміру"].map(clean_text)
    long["product"] = long["milk_product_raw"].map(lambda x: "raw_milk" if "Молоко коров'яче сире" in x else classify_product(x))
    long["value_uah_kg_or_tonne"] = np.where(
        long["indicator"].str.contains("Середня ціна купівлі", case=False, na=False),
        long["value"] / 1000.0,
        long["value"],
    )
    return long[
        [
            "date",
            "period",
            "territory",
            "indicator",
            "flow",
            "milk_product_raw",
            "product",
            "unit_raw",
            "value",
            "value_uah_kg_or_tonne",
        ]
    ].sort_values(["indicator", "territory", "product", "date"])


def clean_prozorro() -> pd.DataFrame:
    df = pd.read_excel(NEW, sheet_name="ProzorroM(full)")
    df = df[df["Ідентифікатор процедури"].map(clean_text).ne("Всього")].copy()
    df["date"] = pd.to_datetime(df["Дата публікації процедури"], errors="coerce").dt.normalize()
    text_cols = ["Профіль", "Товар", "Категорія", "CPV категорії"]
    for c in text_cols + ["Одиниця виміру", "Регіон організатора", "Організатор", "Переможець"]:
        if c in df.columns:
            df[c] = df[c].map(clean_text)
    df["quantity"] = df["Кількість"].map(parse_num)
    df["unit_price_raw"] = df["Ціна за одиницю"].map(parse_num)
    df["expected_value"] = df["Очікувана вартість"].map(parse_num)
    df["contract_value_initial"] = df["Сума договорів початкова"].map(parse_num)
    df["contract_value_current"] = df["Сума договорів поточна"].map(parse_num)
    # Lot/profile text has priority. CPV is only a fallback because broad CPV
    # strings such as "milk and cream" can otherwise misclassify milk as cream.
    primary_product = df.apply(lambda r: classify_product(r.get("Профіль"), r.get("Товар"), r.get("Категорія")), axis=1)
    fallback_product = df.apply(lambda r: classify_product(r.get("CPV категорії")), axis=1)
    df["product"] = np.where(
        primary_product.eq("exclude_non_dairy") & ~fallback_product.eq("exclude_non_dairy"),
        fallback_product,
        primary_product,
    )
    df["unit_raw"] = df["Одиниця виміру"]
    df["unit_clean"] = "UAH/kg"
    package_kg = df.apply(lambda r: parse_package_kg(r.get("Профіль"), r.get("Товар")), axis=1)
    kg_unit = df["unit_raw"].str.contains("кілограм", case=False, na=False)
    piece_unit = df["unit_raw"].str.contains("штука", case=False, na=False)
    df["package_kg_parsed"] = package_kg
    df["price_uah_kg"] = np.where(
        kg_unit,
        df["unit_price_raw"],
        np.where(piece_unit & package_kg.notna() & package_kg.gt(0), df["unit_price_raw"] / package_kg, np.nan),
    )
    df["stage"] = "prozorro_procurement"
    df["source"] = "ProZorro Market / tender export from newmodel workbook"
    flags = []
    for _, r in df.iterrows():
        if pd.isna(r["date"]):
            flags.append("missing_date")
        elif r["product"] == "exclude_non_dairy":
            flags.append("exclude_non_dairy")
        elif pd.isna(r["price_uah_kg"]):
            flags.append("no_reliable_uah_kg")
        elif r["price_uah_kg"] <= 0:
            flags.append("non_positive")
        elif r["price_uah_kg"] > 3000:
            flags.append("implausible_high")
        else:
            flags.append("ok")
    df["quality_flag"] = flags
    keep = [
        "date",
        "Ідентифікатор процедури",
        "Профіль",
        "Товар",
        "Категорія",
        "CPV категорії",
        "territory",
        "source",
        "stage",
        "product",
        "quantity",
        "unit_price_raw",
        "price_uah_kg",
        "unit_raw",
        "unit_clean",
        "package_kg_parsed",
        "expected_value",
        "contract_value_initial",
        "contract_value_current",
        "Організатор",
        "Переможець",
        "quality_flag",
    ]
    df["territory"] = df["Регіон організатора"].replace({"-": np.nan})
    return df[keep].sort_values(["product", "date"])


def clean_retail_sheet(sheet: str) -> pd.DataFrame:
    df = pd.read_excel(NEW, sheet_name=sheet)
    retailer = sheet.lower()
    df["date"] = pd.to_datetime(df["date"], errors="coerce").dt.normalize()
    for c in ["product_title", "product_name", "brand", "entity", "broader_category", "product_ua", "pack_unit_final", "unit_en"]:
        if c in df.columns:
            df[c] = df[c].map(clean_text)
    # Title/name are the authoritative classification layer. The legacy
    # product_ua/broader_category fields are used only if title/name are not
    # informative, because the old retail classification contains visible
    # mistakes.
    primary_product = df.apply(lambda r: classify_product(r.get("product_title"), r.get("product_name")), axis=1)
    fallback_product = df.apply(lambda r: classify_product(r.get("product_ua"), r.get("broader_category")), axis=1)
    df["product"] = np.where(
        primary_product.eq("exclude_non_dairy") & ~fallback_product.eq("exclude_non_dairy"),
        fallback_product,
        primary_product,
    )
    df["price_raw"] = df["price_current"].map(parse_num)
    df["price_package_uah"] = df["price_raw"]
    df["discount_value"] = df.get("discount_value", pd.Series(np.nan, index=df.index)).map(parse_num)
    if "discount_%" in df.columns:
        df["discount_pct"] = df["discount_%"].map(parse_num)
    else:
        df["discount_pct"] = np.nan
    df["unit_price_existing"] = df.get("unit_price", pd.Series(np.nan, index=df.index)).map(parse_num)
    df["qty_std_clean"] = df.get("qty_std", pd.Series(np.nan, index=df.index)).map(parse_num)
    title_package = df.apply(lambda r: parse_package_kg(r.get("product_title"), r.get("product_name")), axis=1)
    df["package_kg_parsed"] = title_package
    unit_en = df.get("unit_en", pd.Series("", index=df.index)).astype(str).str.lower()
    existing_valid = df["unit_price_existing"].between(1, 3000) & unit_en.isin(["kg", "liter"])
    computed = df["price_raw"] / df["package_kg_parsed"]
    computed_valid = computed.between(1, 3000)
    df["price_uah_kg"] = np.where(existing_valid, df["unit_price_existing"], np.where(computed_valid, computed, np.nan))
    baseline_package_price = df["price_raw"] + df["discount_value"].fillna(0)
    baseline_by_ratio = np.where(df["price_raw"].gt(0), df["price_uah_kg"] * baseline_package_price / df["price_raw"], np.nan)
    df["price_baseline_uah_kg"] = np.where(df["discount_value"].notna() & df["price_uah_kg"].notna(), baseline_by_ratio, df["price_uah_kg"])
    df["discount_incidence"] = np.where(df["discount_value"].fillna(0).gt(0) | df["discount_pct"].fillna(0).gt(0), 1, 0)
    sku_base = (
        retailer
        + "|"
        + df["product_title"].map(normalize_key)
        + "|"
        + df["brand"].map(normalize_key)
        + "|"
        + df["package_kg_parsed"].round(4).astype(str)
    )
    df["sku"] = sku_base
    df["territory"] = "Ukraine_observed_online"
    df["source"] = f"{sheet} retail web price sheet from newmodel workbook"
    df["stage"] = "retail"
    df["unit_raw"] = df.get("unit_en", pd.Series("", index=df.index))
    df["unit_clean"] = "UAH/kg"
    df["retailer"] = retailer
    flags = []
    unit_flags = []
    for _, r in df.iterrows():
        if pd.isna(r["date"]):
            flags.append("missing_date")
            unit_flags.append("missing_date")
        elif r["product"] == "exclude_non_dairy":
            flags.append("exclude_non_dairy")
            unit_flags.append("exclude_non_dairy")
        elif pd.isna(r["price_raw"]) or r["price_raw"] <= 0:
            flags.append("no_reliable_package_price")
            unit_flags.append("no_reliable_package_price")
        elif r["price_raw"] > 5000:
            flags.append("implausible_package_price")
            unit_flags.append("implausible_package_price")
        elif pd.isna(r["price_uah_kg"]):
            flags.append("ok")
            unit_flags.append("no_reliable_uah_kg")
        elif r["price_uah_kg"] <= 0:
            flags.append("ok")
            unit_flags.append("non_positive_uah_kg")
        elif r["price_uah_kg"] > 3000:
            flags.append("ok")
            unit_flags.append("implausible_high_uah_kg")
        else:
            flags.append("ok")
            unit_flags.append("ok_uah_kg")
    df["quality_flag"] = flags
    df["unit_quality_flag"] = unit_flags
    keep = [
        "date",
        "territory",
        "source",
        "stage",
        "retailer",
        "sku",
        "product",
        "product_title",
        "product_name",
        "brand",
        "price_raw",
        "price_package_uah",
        "price_uah_kg",
        "price_baseline_uah_kg",
        "discount_value",
        "discount_pct",
        "discount_incidence",
        "unit_raw",
        "unit_clean",
        "package_kg_parsed",
        "quality_flag",
        "unit_quality_flag",
    ]
    return df[keep].sort_values(["retailer", "product", "sku", "date"])


def build_prozorro_aggregates(proz: pd.DataFrame) -> dict[str, pd.DataFrame]:
    good = proz[(proz["quality_flag"] == "ok") & (proz["product"].isin(PRODUCT_ORDER)) & (proz["product"] != "exclude_non_dairy")].copy()
    good = good[good["price_uah_kg"].between(1, 3000)]

    def wavg(g):
        q = pd.to_numeric(g["quantity"], errors="coerce")
        p = pd.to_numeric(g["price_uah_kg"], errors="coerce")
        ok = q.notna() & q.gt(0) & p.notna()
        if ok.sum() == 0:
            return np.nan
        return np.average(p[ok], weights=np.minimum(q[ok], q[ok].quantile(0.95)))

    aggs = {}
    for freq, colname in [("D", "date"), ("W-MON", "week"), ("MS", "month")]:
        x = good.copy()
        if freq == "D":
            x[colname] = x["date"]
        elif freq == "W-MON":
            x[colname] = x["date"].dt.to_period("W-MON").dt.start_time
        else:
            x[colname] = x["date"].values.astype("datetime64[M]")
        grouped = (
            x.groupby(["product", colname], dropna=False)
            .apply(
                lambda g: pd.Series(
                    {
                        "price_median_uah_kg": g["price_uah_kg"].median(),
                        "price_weighted_uah_kg": wavg(g),
                        "n_lots": len(g),
                        "n_regions": g["territory"].nunique(dropna=True),
                        "expected_value_sum": g["expected_value"].sum(min_count=1),
                        "contract_current_sum": g["contract_value_current"].sum(min_count=1),
                    }
                )
            )
            .reset_index()
        )
        grouped = grouped.rename(columns={colname: "date"})
        aggs[freq] = grouped.sort_values(["product", "date"])
    return aggs


def aggregate_retail(retail: pd.DataFrame) -> dict[str, pd.DataFrame]:
    good = retail[(retail["quality_flag"] == "ok") & (retail["product"] != "exclude_non_dairy")].copy()
    unit_good = good[good["unit_quality_flag"].eq("ok_uah_kg")].copy()
    daily = (
        unit_good.groupby(["retailer", "product", "date"])
        .agg(
            price_observed_median_uah_kg=("price_uah_kg", "median"),
            price_baseline_median_uah_kg=("price_baseline_uah_kg", "median"),
            discount_share=("discount_incidence", "mean"),
            discount_depth_median_pct=("discount_pct", "median"),
            n_skus=("sku", "nunique"),
            n_obs=("sku", "size"),
        )
        .reset_index()
    )
    pooled = (
        unit_good.groupby(["product", "date"])
        .agg(
            price_observed_median_uah_kg=("price_uah_kg", "median"),
            price_baseline_median_uah_kg=("price_baseline_uah_kg", "median"),
            discount_share=("discount_incidence", "mean"),
            discount_depth_median_pct=("discount_pct", "median"),
            n_skus=("sku", "nunique"),
            n_obs=("sku", "size"),
        )
        .reset_index()
    )
    package_daily = (
        good.groupby(["retailer", "product", "date"])
        .agg(
            price_package_median_uah=("price_package_uah", "median"),
            discount_share=("discount_incidence", "mean"),
            discount_depth_median_pct=("discount_pct", "median"),
            n_skus=("sku", "nunique"),
            n_obs=("sku", "size"),
        )
        .reset_index()
    )
    package_pooled_daily = (
        good.groupby(["product", "date"])
        .agg(
            price_package_median_uah=("price_package_uah", "median"),
            discount_share=("discount_incidence", "mean"),
            discount_depth_median_pct=("discount_pct", "median"),
            n_skus=("sku", "nunique"),
            n_obs=("sku", "size"),
        )
        .reset_index()
    )
    weekly = daily.copy()
    weekly["date"] = weekly["date"].dt.to_period("W-MON").dt.start_time
    weekly = (
        weekly.groupby(["retailer", "product", "date"])
        .agg(
            price_observed_median_uah_kg=("price_observed_median_uah_kg", "median"),
            price_baseline_median_uah_kg=("price_baseline_median_uah_kg", "median"),
            discount_share=("discount_share", "mean"),
            discount_depth_median_pct=("discount_depth_median_pct", "median"),
            n_skus=("n_skus", "max"),
            n_obs=("n_obs", "sum"),
        )
        .reset_index()
    )
    pooled_weekly = pooled.copy()
    pooled_weekly["date"] = pooled_weekly["date"].dt.to_period("W-MON").dt.start_time
    pooled_weekly = (
        pooled_weekly.groupby(["product", "date"])
        .agg(
            price_observed_median_uah_kg=("price_observed_median_uah_kg", "median"),
            price_baseline_median_uah_kg=("price_baseline_median_uah_kg", "median"),
            discount_share=("discount_share", "mean"),
            discount_depth_median_pct=("discount_depth_median_pct", "median"),
            n_skus=("n_skus", "max"),
            n_obs=("n_obs", "sum"),
        )
        .reset_index()
    )
    package_weekly = package_daily.copy()
    package_weekly["date"] = package_weekly["date"].dt.to_period("W-MON").dt.start_time
    package_weekly = (
        package_weekly.groupby(["retailer", "product", "date"])
        .agg(
            price_package_median_uah=("price_package_median_uah", "median"),
            discount_share=("discount_share", "mean"),
            discount_depth_median_pct=("discount_depth_median_pct", "median"),
            n_skus=("n_skus", "max"),
            n_obs=("n_obs", "sum"),
        )
        .reset_index()
    )
    package_pooled_weekly = package_pooled_daily.copy()
    package_pooled_weekly["date"] = package_pooled_weekly["date"].dt.to_period("W-MON").dt.start_time
    package_pooled_weekly = (
        package_pooled_weekly.groupby(["product", "date"])
        .agg(
            price_package_median_uah=("price_package_median_uah", "median"),
            discount_share=("discount_share", "mean"),
            discount_depth_median_pct=("discount_depth_median_pct", "median"),
            n_skus=("n_skus", "max"),
            n_obs=("n_obs", "sum"),
        )
        .reset_index()
    )
    return {
        "daily": daily,
        "pooled_daily": pooled,
        "weekly": weekly,
        "pooled_weekly": pooled_weekly,
        "package_daily": package_daily,
        "package_pooled_daily": package_pooled_daily,
        "package_weekly": package_weekly,
        "package_pooled_weekly": package_pooled_weekly,
    }


def load_cost_index() -> pd.DataFrame:
    df = pd.read_excel(ADD / "індекс витрат сільгосп.xlsx")
    dates = period_cols(df, "MQ")
    id_cols = [c for c in df.columns if c not in dates]
    long = df.melt(id_vars=id_cols, value_vars=dates, var_name="period", value_name="index_value")
    long["date"] = long["period"].map(period_to_date)
    long["index_value"] = long["index_value"].map(parse_num)
    long = long[long["index_value"].notna()].copy()
    long["basis"] = long["Основа"].map(clean_text)
    long["territory"] = long["Територіальний розріз"].map(clean_text)
    long["product"] = "livestock_products"
    return long[["date", "period", "basis", "territory", "product", "index_value"]].sort_values("date")


def load_trade_tables() -> dict[str, pd.DataFrame]:
    out = {}
    for key, path in {
        "exports": ADD / "share_dairy_exp.xlsx",
        "imports": ADD / "share_dairy_imp.xlsx",
    }.items():
        df = pd.read_excel(path, header=13)
        df = df.dropna(subset=["Code"]).copy()
        df["code_clean"] = df["Code"].astype(str).str.replace("'", "", regex=False).str.strip()
        year_cols = [c for c in df.columns if isinstance(c, (int, float)) or str(c).replace(".0", "").isdigit()]
        rows = []
        for _, r in df.iterrows():
            for y in year_cols:
                try:
                    year = int(float(y))
                except Exception:
                    continue
                rows.append(
                    {
                        "flow": key,
                        "code": r["code_clean"],
                        "product_label": clean_text(r.get("Product label")),
                        "year": year,
                        "value_eur_mln": parse_num(r[y]),
                    }
                )
        long = pd.DataFrame(rows)
        long["is_dairy_only_0401_0406"] = long["code"].str.match(r"^040[1-6]$")
        out[key] = long
    biggest = pd.read_excel(ADD / "biggest_trade.xlsx", sheet_name="Sheet1", header=None)
    out["biggest_trade_raw"] = biggest
    return out


def stationarity_tests(series: pd.Series) -> dict:
    s = pd.to_numeric(series, errors="coerce").dropna()
    res = {"n": int(len(s)), "adf_p": np.nan, "kpss_p": np.nan, "integration": "insufficient"}
    if len(s) < 12:
        return res
    try:
        res["adf_p"] = float(adfuller(s, autolag="AIC")[1])
    except Exception:
        pass
    try:
        res["kpss_p"] = float(kpss(s, regression="c", nlags="auto")[1])
    except Exception:
        pass
    # Heuristic classification.
    if pd.notna(res["adf_p"]) and pd.notna(res["kpss_p"]):
        if res["adf_p"] < 0.10 and res["kpss_p"] > 0.05:
            res["integration"] = "I(0)"
        elif res["adf_p"] >= 0.10 and res["kpss_p"] <= 0.10:
            res["integration"] = "I(1)-likely"
        else:
            res["integration"] = "ambiguous"
    return res


def prepare_pair(y_df: pd.DataFrame, x_df: pd.DataFrame, y_col="price_uah_kg", x_col="price_uah_kg") -> pd.DataFrame:
    y = y_df[["date", y_col]].rename(columns={y_col: "y"}).copy()
    x = x_df[["date", x_col]].rename(columns={x_col: "x"}).copy()
    z = pd.merge(y, x, on="date", how="inner").dropna()
    z = z[(z["y"] > 0) & (z["x"] > 0)].sort_values("date")
    z["ln_y"] = np.log(z["y"])
    z["ln_x"] = np.log(z["x"])
    return z


def fit_ecm(z: pd.DataFrame, model_id: str, hypothesis: str, y_name: str, x_name: str, frequency: str, note: str = "") -> dict:
    out = {
        "model_id": model_id,
        "hypothesis": hypothesis,
        "method": "ECM",
        "y": y_name,
        "x": x_name,
        "frequency": frequency,
        "n": int(len(z)),
        "period_start": z["date"].min().date().isoformat() if len(z) else "",
        "period_end": z["date"].max().date().isoformat() if len(z) else "",
        "long_run_beta": np.nan,
        "coint_p": np.nan,
        "ect": np.nan,
        "ect_p": np.nan,
        "short_run_dx": np.nan,
        "short_run_dx_p": np.nan,
        "ljungbox_p": np.nan,
        "bp_p": np.nan,
        "reliability": "unreliable / appendix only",
        "interpretation": "",
        "note": note,
    }
    if len(z) < 24:
        out["interpretation"] = "Too few observations for a reliable equilibrium model."
        return out
    try:
        lr = sm.OLS(z["ln_y"], sm.add_constant(z["ln_x"])).fit()
        out["long_run_beta"] = float(lr.params.get("ln_x", np.nan))
        try:
            out["coint_p"] = float(coint(z["ln_y"], z["ln_x"])[1])
        except Exception:
            out["coint_p"] = np.nan
        e = lr.resid
        m = pd.DataFrame(
            {
                "date": z["date"],
                "dy": z["ln_y"].diff(),
                "dx": z["ln_x"].diff(),
                "dy_lag1": z["ln_y"].diff().shift(1),
                "dx_lag1": z["ln_x"].diff().shift(1),
                "ect_lag1": e.shift(1),
            }
        ).dropna()
        if len(m) < 20:
            out["interpretation"] = "Too few observations after lag construction."
            return out
        X = sm.add_constant(m[["dx", "dx_lag1", "dy_lag1", "ect_lag1"]])
        maxlags = max(1, min(4, int(len(m) ** 0.25)))
        er = sm.OLS(m["dy"], X).fit(cov_type="HAC", cov_kwds={"maxlags": maxlags})
        out["ect"] = float(er.params.get("ect_lag1", np.nan))
        out["ect_p"] = float(er.pvalues.get("ect_lag1", np.nan))
        out["short_run_dx"] = float(er.params.get("dx", np.nan))
        out["short_run_dx_p"] = float(er.pvalues.get("dx", np.nan))
        try:
            out["ljungbox_p"] = float(acorr_ljungbox(er.resid, lags=[min(6, max(1, len(er.resid) // 4))], return_df=True)["lb_pvalue"].iloc[0])
        except Exception:
            pass
        try:
            out["bp_p"] = float(het_breuschpagan(er.resid, X)[1])
        except Exception:
            pass
        reliable = (
            len(z) >= 48
            and pd.notna(out["coint_p"])
            and out["coint_p"] < 0.10
            and pd.notna(out["ect"])
            and out["ect"] < 0
            and pd.notna(out["ect_p"])
            and out["ect_p"] < 0.10
            and (pd.isna(out["ljungbox_p"]) or out["ljungbox_p"] >= 0.05)
        )
        probable = pd.notna(out["ect"]) and out["ect"] < 0 and pd.notna(out["ect_p"]) and out["ect_p"] < 0.20 and len(z) >= 30
        if reliable:
            out["reliability"] = "reliable"
        elif probable:
            out["reliability"] = "probable / needs validation"
        else:
            out["reliability"] = "unreliable / appendix only"
        beta = out["long_run_beta"]
        ect = out["ect"]
        if out["reliability"] == "reliable":
            out["interpretation"] = (
                f"Long-run pass-through is {beta:.2f}; the error-correction term is {ect:.2f}, "
                "so deviations from the long-run relation are corrected over time."
            )
        elif out["reliability"].startswith("probable"):
            out["interpretation"] = (
                f"The coefficient pattern is economically usable but diagnostically weaker: beta {beta:.2f}, ECT {ect:.2f}."
            )
        else:
            out["interpretation"] = "The specification is retained only as a diagnostic or appendix result."
    except Exception as exc:
        out["interpretation"] = f"Model failed: {exc}"
    return out


def fit_nardl(z: pd.DataFrame, model_id: str, hypothesis: str, y_name: str, x_name: str, frequency: str, note: str = "") -> dict:
    out = {
        "model_id": model_id,
        "hypothesis": hypothesis,
        "method": "NARDL-lite",
        "y": y_name,
        "x": x_name,
        "frequency": frequency,
        "n": int(len(z)),
        "period_start": z["date"].min().date().isoformat() if len(z) else "",
        "period_end": z["date"].max().date().isoformat() if len(z) else "",
        "pos_coef": np.nan,
        "pos_p": np.nan,
        "neg_coef": np.nan,
        "neg_p": np.nan,
        "asym_p": np.nan,
        "ect": np.nan,
        "ect_p": np.nan,
        "reliability": "unreliable / appendix only",
        "interpretation": "",
        "note": note,
    }
    if len(z) < 36:
        out["interpretation"] = "Too few observations for asymmetry testing."
        return out
    try:
        dx = z["ln_x"].diff()
        xpos = dx.clip(lower=0).fillna(0).cumsum()
        xneg = dx.clip(upper=0).fillna(0).cumsum()
        lrdat = pd.DataFrame({"ln_y": z["ln_y"], "xpos": xpos, "xneg": xneg}).dropna()
        lr = sm.OLS(lrdat["ln_y"], sm.add_constant(lrdat[["xpos", "xneg"]])).fit()
        resid = lr.resid.reindex(lrdat.index)
        m = pd.DataFrame(
            {
                "dy": z["ln_y"].diff(),
                "dpos": xpos.diff(),
                "dneg": xneg.diff(),
                "dy_lag1": z["ln_y"].diff().shift(1),
                "ect_lag1": resid.shift(1),
            }
        ).dropna()
        X = sm.add_constant(m[["dpos", "dneg", "dy_lag1", "ect_lag1"]])
        maxlags = max(1, min(4, int(len(m) ** 0.25)))
        res = sm.OLS(m["dy"], X).fit(cov_type="HAC", cov_kwds={"maxlags": maxlags})
        out["pos_coef"] = float(res.params.get("dpos", np.nan))
        out["pos_p"] = float(res.pvalues.get("dpos", np.nan))
        out["neg_coef"] = float(res.params.get("dneg", np.nan))
        out["neg_p"] = float(res.pvalues.get("dneg", np.nan))
        out["ect"] = float(res.params.get("ect_lag1", np.nan))
        out["ect_p"] = float(res.pvalues.get("ect_lag1", np.nan))
        try:
            out["asym_p"] = float(res.f_test("dpos = dneg").pvalue)
        except Exception:
            pass
        if len(z) >= 48 and out["ect"] < 0 and out["ect_p"] < 0.10 and pd.notna(out["asym_p"]) and out["asym_p"] < 0.10:
            out["reliability"] = "reliable"
        elif out["ect"] < 0 and out["ect_p"] < 0.20 and len(z) >= 36:
            out["reliability"] = "probable / needs validation"
        else:
            out["reliability"] = "unreliable / appendix only"
        if out["reliability"] == "reliable":
            out["interpretation"] = (
                f"Asymmetry is statistically meaningful: positive shocks {out['pos_coef']:.2f}, "
                f"negative shocks {out['neg_coef']:.2f}."
            )
        elif out["reliability"].startswith("probable"):
            out["interpretation"] = "Asymmetry is suggestive but should not be a headline claim."
        else:
            out["interpretation"] = "Asymmetry evidence is too weak for the main text."
    except Exception as exc:
        out["interpretation"] = f"Model failed: {exc}"
    return out


def fit_short_panel(df: pd.DataFrame, model_id: str, y_col: str, x_col: str, note: str) -> dict:
    out = {
        "model_id": model_id,
        "hypothesis": "H2",
        "method": "Short panel OLS-HAC",
        "y": y_col,
        "x": x_col,
        "frequency": "weekly/daily retail mechanism",
        "n": int(len(df)),
        "period_start": "",
        "period_end": "",
        "coef": np.nan,
        "pvalue": np.nan,
        "reliability": "unreliable / appendix only",
        "interpretation": "",
        "note": note,
    }
    if df.empty or len(df) < 20:
        out["interpretation"] = "Too few observations for a reliable retail mechanism regression."
        return out
    try:
        z = df.copy().sort_values(["product", "date"])
        z = z[z[y_col].gt(0) & z[x_col].gt(0)].copy()
        z["ln_y"] = np.log(z[y_col])
        z["ln_x"] = np.log(z[x_col])
        z["dln_y"] = z.groupby("product")["ln_y"].diff()
        z["dln_x"] = z.groupby("product")["ln_x"].diff()
        z["dln_y_lag"] = z.groupby("product")["dln_y"].shift(1)
        z = z.dropna(subset=["dln_y", "dln_x"])
        if len(z) < 20:
            out["interpretation"] = "Too few differenced observations after matching."
            return out
        controls = " + dln_y_lag" if z["dln_y_lag"].notna().sum() > 10 else ""
        if "discount_share" in z.columns and z["discount_share"].notna().sum() > 10 and z["discount_share"].std(skipna=True) > 0:
            controls += " + discount_share"
        formula = f"dln_y ~ dln_x{controls} + C(product)"
        res = smf.ols(formula, data=z).fit(cov_type="HAC", cov_kwds={"maxlags": 2})
        out["n"] = int(res.nobs)
        out["period_start"] = z["date"].min().date().isoformat()
        out["period_end"] = z["date"].max().date().isoformat()
        out["coef"] = float(res.params.get("dln_x", np.nan))
        out["pvalue"] = float(res.pvalues.get("dln_x", np.nan))
        if out["n"] >= 80 and pd.notna(out["pvalue"]) and out["pvalue"] < 0.10:
            out["reliability"] = "probable / needs validation"
        else:
            out["reliability"] = "unreliable / appendix only"
        out["interpretation"] = (
            "Short-window retail mechanism evidence; use for H2 timing and promotion discussion, not as a long-run market-power proof."
        )
    except Exception as exc:
        out["interpretation"] = f"Model failed: {exc}"
    return out


def fit_ardl_order(z: pd.DataFrame) -> str:
    if ardl_select_order is None or len(z) < 36:
        return ""
    try:
        y = z.set_index("date")["ln_y"].asfreq("MS")
        x = z.set_index("date")[["ln_x"]].asfreq("MS")
        both = pd.concat([y, x], axis=1).dropna()
        sel = ardl_select_order(both["ln_y"], maxlag=4, exog=both[["ln_x"]], maxorder=4, ic="aic", trend="c")
        return str(sel.model.ardl_order)
    except Exception as exc:
        return f"ARDL selection failed: {exc}"


def run_models(farm: pd.DataFrame, processor: pd.DataFrame, consumer: pd.DataFrame, proz_agg: dict, retail_aggs: dict) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    models = []
    stationarity_rows = []
    ardl_rows = []

    farm_main = farm[(farm["territory"] == "Україна") & (farm["farm_type"] == "enterprises") & (farm["quality_flag"] == "ok")]
    proc_main = processor[(processor["territory"] == "Україна") & (processor["quality_flag"] == "ok")]
    cons_main = consumer[(consumer["territory"] == "Україна") & (consumer["quality_flag"] == "ok")]

    # H1: farm-gate raw milk to processor-level prices.
    for product in sorted(proc_main["product"].dropna().unique()):
        if product in {"exclude_non_dairy", "other_dairy"}:
            continue
        y = proc_main[proc_main["product"] == product]
        z = prepare_pair(y, farm_main)
        for var, name in [(z.get("ln_y", pd.Series(dtype=float)), f"processor_{product}"), (z.get("ln_x", pd.Series(dtype=float)), "farmgate_raw_milk")]:
            st = stationarity_tests(var)
            st.update({"series": name, "block": "H1", "product": product})
            stationarity_rows.append(st)
        model_id = f"H1_farmgate_to_processor_{product}"
        ecm = fit_ecm(
            z,
            model_id,
            "H1",
            f"Processor {product_label(product)}",
            "Farm-gate raw milk",
            "monthly",
            "Observed SSSU monthly data; processor and farm-gate prices converted to UAH/kg.",
        )
        ecm["ardl_order"] = fit_ardl_order(z)
        models.append(ecm)
        models.append(
            fit_nardl(
                z,
                model_id + "_asymmetry",
                "H1",
                f"Processor {product_label(product)}",
                "Farm-gate raw milk",
                "monthly",
                "Asymmetry is a mechanism under H1, not a separate hypothesis.",
            )
        )
        ardl_rows.append({"model_id": model_id, "ardl_order": ecm.get("ardl_order", "")})

    # H2 official monthly: processor to official consumer.
    for product in ["drinking_milk", "sour_cream", "soft_cheese"]:
        cons_product = consumer[consumer["product"] == product]
        if product == "soft_cheese":
            proc_product = processor[processor["product"] == "hard_cheese"]
            note = "Soft consumer cheese compared with hard processor cheese; keep as appendix if diagnostics are weak."
            x_name = "Processor hard cheese"
        else:
            proc_product = processor[processor["product"] == product]
            note = "Observed SSSU monthly processor-to-consumer bridge."
            x_name = f"Processor {product_label(product)}"
        z = prepare_pair(cons_product[cons_product["territory"] == "Україна"], proc_product[proc_product["territory"] == "Україна"])
        model_id = f"H2_processor_to_consumer_{product}"
        ecm = fit_ecm(
            z,
            model_id,
            "H2",
            f"Official consumer {product_label(product)}",
            x_name,
            "monthly",
            note,
        )
        ecm["ardl_order"] = fit_ardl_order(z)
        models.append(ecm)
        models.append(
            fit_nardl(
                z,
                model_id + "_asymmetry",
                "H2",
                f"Official consumer {product_label(product)}",
                x_name,
                "monthly",
                note,
            )
        )

    # H1/H2 bridge: ProZorro monthly to farmgate/processor.
    proz_m = proz_agg["MS"].rename(columns={"price_median_uah_kg": "prozorro_price"})
    for product in sorted(proz_m["product"].dropna().unique()):
        if product not in {"drinking_milk", "sour_cream", "butter", "hard_cheese", "cottage_cheese", "cream", "condensed_milk", "skim_milk_powder"}:
            continue
        pz = proz_m[proz_m["product"] == product][["date", "prozorro_price"]].rename(columns={"prozorro_price": "price_uah_kg"})
        zf = prepare_pair(pz, farm_main)
        models.append(
            fit_ecm(
                zf,
                f"H1_farmgate_to_prozorro_{product}",
                "H1",
                f"ProZorro {product_label(product)}",
                "Farm-gate raw milk",
                "monthly",
                "Short ProZorro bridge; use as mechanism evidence only unless diagnostics are strong.",
            )
        )
        proc_product = proc_main[proc_main["product"] == product]
        if not proc_product.empty:
            zp = prepare_pair(pz, proc_product)
            models.append(
                fit_ecm(
                    zp,
                    f"H2_processor_to_prozorro_{product}",
                    "H2",
                    f"ProZorro {product_label(product)}",
                    f"Processor {product_label(product)}",
                    "monthly",
                    "Institutional procurement bridge, not retail.",
                )
            )

    # H2 short retail mechanism: ProZorro to retail.
    proz_w = proz_agg["W-MON"].rename(columns={"price_median_uah_kg": "prozorro_price"})
    pooled_w = retail_aggs["pooled_weekly"].rename(columns={"price_observed_median_uah_kg": "retail_observed", "price_baseline_median_uah_kg": "retail_baseline"})
    matched = pd.merge(pooled_w, proz_w[["product", "date", "prozorro_price", "n_lots"]], on=["product", "date"], how="inner")
    if len(matched) >= 10:
        models.append(
            fit_short_panel(
                matched,
                "H2_prozorro_to_pooled_retail_observed_weekly",
                "retail_observed",
                "prozorro_price",
                "Direct same-week ProZorro-retail product aggregate; short overlap.",
            )
        )
        models.append(
            fit_short_panel(
                matched,
                "H2_prozorro_to_pooled_retail_baseline_weekly",
                "retail_baseline",
                "prozorro_price",
                "Baseline non-discount retail price comparison; short overlap.",
            )
        )

    package_pooled_w = retail_aggs["package_pooled_weekly"]
    matched_pkg = pd.merge(package_pooled_w, proz_w[["product", "date", "prozorro_price", "n_lots"]], on=["product", "date"], how="inner")
    if len(matched_pkg) >= 10:
        models.append(
            fit_short_panel(
                matched_pkg,
                "H2_prozorro_to_pooled_retail_package_weekly",
                "price_package_median_uah",
                "prozorro_price",
                "SKU/package-price index response. Levels are not comparable to UAH/kg, but within-product log changes are useful as a short retail mechanism check.",
            )
        )

    # As-of daily retail mechanism: latest ProZorro quote within 30 days.
    proz_d = proz_agg["D"].rename(columns={"price_median_uah_kg": "prozorro_price"})
    pooled_d = retail_aggs["pooled_daily"].rename(columns={"price_observed_median_uah_kg": "retail_observed", "price_baseline_median_uah_kg": "retail_baseline"})
    asof_rows = []
    for product, gr in pooled_d.groupby("product"):
        pg = proz_d[proz_d["product"] == product].sort_values("date")
        if pg.empty:
            continue
        rg = gr.sort_values("date")
        merged = pd.merge_asof(rg, pg[["date", "prozorro_price"]], on="date", direction="backward", tolerance=pd.Timedelta(days=30))
        asof_rows.append(merged.assign(product=product))
    if asof_rows:
        asof = pd.concat(asof_rows, ignore_index=True).dropna(subset=["prozorro_price"])
        if len(asof) >= 20:
            models.append(
                fit_short_panel(
                    asof,
                    "H2_prozorro_to_pooled_retail_observed_daily_asof",
                    "retail_observed",
                    "prozorro_price",
                    "Daily retail matched to latest ProZorro price within 30 days; mechanism evidence only.",
                )
            )

    package_pooled_d = retail_aggs["package_pooled_daily"]
    asof_pkg_rows = []
    for product, gr in package_pooled_d.groupby("product"):
        pg = proz_d[proz_d["product"] == product].sort_values("date")
        if pg.empty:
            continue
        rg = gr.sort_values("date")
        merged = pd.merge_asof(rg, pg[["date", "prozorro_price"]], on="date", direction="backward", tolerance=pd.Timedelta(days=30))
        asof_pkg_rows.append(merged.assign(product=product))
    if asof_pkg_rows:
        asof_pkg = pd.concat(asof_pkg_rows, ignore_index=True).dropna(subset=["prozorro_price"])
        if len(asof_pkg) >= 20:
            models.append(
                fit_short_panel(
                    asof_pkg,
                    "H2_prozorro_to_pooled_retail_package_daily_asof",
                    "price_package_median_uah",
                    "prozorro_price",
                    "Daily SKU/package-price index matched to latest ProZorro price within 30 days. Use only as mechanism evidence.",
                )
            )

    return pd.DataFrame(models), pd.DataFrame(stationarity_rows), pd.DataFrame(ardl_rows)


def run_vecm_blocks(farm: pd.DataFrame, processor: pd.DataFrame, consumer: pd.DataFrame) -> pd.DataFrame:
    rows = []
    if VECM is None:
        return pd.DataFrame([{"system": "all", "status": "VECM unavailable"}])
    farm_main = farm[(farm["territory"] == "Україна") & (farm["farm_type"] == "enterprises") & (farm["quality_flag"] == "ok")]
    proc_main = processor[(processor["territory"] == "Україна") & (processor["quality_flag"] == "ok")]
    cons_main = consumer[(consumer["territory"] == "Україна") & (consumer["quality_flag"] == "ok")]
    for product in ["drinking_milk", "sour_cream"]:
        p = proc_main[proc_main["product"] == product][["date", "price_uah_kg"]].rename(columns={"price_uah_kg": "processor"})
        c = cons_main[cons_main["product"] == product][["date", "price_uah_kg"]].rename(columns={"price_uah_kg": "consumer"})
        f = farm_main[["date", "price_uah_kg"]].rename(columns={"price_uah_kg": "farmgate"})
        z = f.merge(p, on="date").merge(c, on="date").dropna()
        z = z[(z[["farmgate", "processor", "consumer"]] > 0).all(axis=1)]
        status = "not_run"
        rank = np.nan
        alpha_farm = alpha_proc = alpha_cons = np.nan
        if len(z) >= 48:
            try:
                lz = np.log(z.set_index("date")[["farmgate", "processor", "consumer"]])
                order = select_order(lz, maxlags=4, deterministic="ci")
                k_ar_diff = max(1, int(order.aic) if order.aic is not None and order.aic >= 1 else 1)
                cr = select_coint_rank(lz, det_order=0, k_ar_diff=k_ar_diff, method="trace", signif=0.10)
                rank = int(cr.rank)
                if rank > 0:
                    res = VECM(lz, k_ar_diff=k_ar_diff, coint_rank=rank, deterministic="ci").fit()
                    alpha = res.alpha[:, 0]
                    alpha_farm, alpha_proc, alpha_cons = [float(a) for a in alpha]
                    status = "estimated"
                else:
                    status = "no_cointegration_rank"
            except Exception as exc:
                status = f"failed: {exc}"
        else:
            status = "too_short"
        rows.append(
            {
                "system": f"farmgate_processor_consumer_{product}",
                "product": product,
                "n": len(z),
                "rank": rank,
                "alpha_farmgate": alpha_farm,
                "alpha_processor": alpha_proc,
                "alpha_consumer": alpha_cons,
                "status": status,
                "reliability": "probable / needs validation" if status == "estimated" and len(z) >= 60 else "unreliable / appendix only",
            }
        )
    return pd.DataFrame(rows)


def build_validation_tables(farm, processor, consumer) -> tuple[pd.DataFrame, pd.DataFrame]:
    # Audit the known corrupted integrated Consumer_price sheet against the raw part.
    integrated_cons = pd.read_excel(NEW, sheet_name="Consumer_price")
    dates = period_cols(integrated_cons, "M")
    integ_long = integrated_cons.melt(
        id_vars=["Територіальний розріз", "Тип товарів і послуг", "Одиниця виміру"],
        value_vars=dates,
        var_name="period",
        value_name="value",
    )
    integ_long["value"] = integ_long["value"].map(parse_num)
    corrupt = integ_long[integ_long["value"].gt(1000)].copy()
    corruption_summary = pd.DataFrame(
        [
            {
                "check": "Integrated Consumer_price implausible values above 1000 UAH/kg",
                "affected_rows": len(corrupt),
                "affected_ukraine_rows": int((corrupt["Територіальний розріз"] == "Україна").sum()),
                "decision": "Use raw Consumer_2017_2026.xlsx, not the combined sheet, for official consumer models.",
            }
        ]
    )

    comparisons = []
    if NEW.exists():
        # Old-model FarmGate comparison.
        try:
            old = pd.read_excel(NEW, sheet_name="FarmGate(oldmodel)")
            old["date"] = pd.to_datetime(old["date"], errors="coerce")
            old["month"] = old["date"].values.astype("datetime64[M]")
            old["territory"] = old.get("region", old.get("region_en", "")).map(clean_text)
            old["farm_type"] = old.get("farm_type", old.get("farm_type_en", "")).map(clean_text).str.lower()
            old["price"] = old["price_linear"].map(parse_num)
            o = old[(old["region"].eq("Україна")) | (old.get("region_en", pd.Series("", index=old.index)).eq("Ukraine"))].copy()
            o = o.groupby("month")["price"].median().reset_index().rename(columns={"month": "date", "price": "old_price"})
            n = farm[(farm["territory"] == "Україна") & (farm["farm_type"] == "enterprises")][["date", "price_uah_kg"]].rename(columns={"price_uah_kg": "new_price"})
            m = o.merge(n, on="date").dropna()
            comparisons.append(calc_validation("FarmGate(oldmodel) vs observed Farm_milk_2015", m))
        except Exception as exc:
            comparisons.append({"comparison": "FarmGate(oldmodel)", "error": str(exc)})
        try:
            old = pd.read_excel(NEW, sheet_name="Producer(oldmodel)")
            old["date"] = pd.to_datetime(old["date"], errors="coerce")
            old["date"] = old["date"].values.astype("datetime64[M]")
            old["product"] = old["product"].map(lambda x: classify_product(x))
            old["old_price"] = old["price_linear"].map(parse_num)
            o = old.groupby(["product", "date"])["old_price"].median().reset_index()
            n = processor.groupby(["product", "date"])["price_uah_kg"].median().reset_index().rename(columns={"price_uah_kg": "new_price"})
            for product, g in o.merge(n, on=["product", "date"]).dropna().groupby("product"):
                comparisons.append(calc_validation(f"Producer(oldmodel) vs Processor_price: {product}", g))
        except Exception as exc:
            comparisons.append({"comparison": "Producer(oldmodel)", "error": str(exc)})
        try:
            old = pd.read_excel(NEW, sheet_name="Consumer(oldmodel)")
            old["date"] = pd.to_datetime(old["date"], errors="coerce")
            old["date"] = old["date"].values.astype("datetime64[M]")
            old["product"] = old["product"].map(lambda x: classify_product(x))
            old["old_price"] = old["price_linear"].map(parse_num)
            o = old.groupby(["product", "date"])["old_price"].median().reset_index()
            n = consumer[consumer["territory"] == "Україна"].groupby(["product", "date"])["price_uah_kg"].median().reset_index().rename(columns={"price_uah_kg": "new_price"})
            for product, g in o.merge(n, on=["product", "date"]).dropna().groupby("product"):
                comparisons.append(calc_validation(f"Consumer(oldmodel) vs Consumer_2017_2026: {product}", g))
        except Exception as exc:
            comparisons.append({"comparison": "Consumer(oldmodel)", "error": str(exc)})
    return corruption_summary, pd.DataFrame(comparisons)


def calc_validation(name: str, m: pd.DataFrame) -> dict:
    if len(m) < 3:
        return {"comparison": name, "n": len(m), "decision": "Too few overlap observations."}
    x = pd.to_numeric(m["old_price"], errors="coerce")
    y = pd.to_numeric(m["new_price"], errors="coerce")
    ok = x.notna() & y.notna()
    x, y = x[ok], y[ok]
    if len(x) < 3:
        return {"comparison": name, "n": len(x), "decision": "Too few valid overlap observations."}
    corr = float(np.corrcoef(x, y)[0, 1]) if x.std() > 0 and y.std() > 0 else np.nan
    err = x - y
    mae = float(np.mean(np.abs(err)))
    rmse = float(np.sqrt(np.mean(err**2)))
    mape = float(np.mean(np.abs(err / y.replace(0, np.nan))) * 100)
    sign = float((np.sign(x.diff().dropna()) == np.sign(y.diff().dropna())).mean()) if len(x) > 3 else np.nan
    if pd.notna(corr) and corr >= 0.85 and mape < 15:
        decision = "Old reconstruction validates well; usable as robustness."
    elif pd.notna(corr) and corr >= 0.60:
        decision = "Partial validation; appendix/support only."
    else:
        decision = "Weak validation; do not use for main inference."
    return {
        "comparison": name,
        "n": int(len(x)),
        "correlation": corr,
        "MAE": mae,
        "RMSE": rmse,
        "MAPE_pct": mape,
        "sign_agreement": sign,
        "decision": decision,
    }


def regional_extension_tables(farm, consumer, prozorro, farm_volumes) -> dict[str, pd.DataFrame]:
    f = farm[(farm["farm_type"] == "enterprises") & (farm["quality_flag"] == "ok") & (farm["territory"] != "Україна")].copy()
    farm_disp = (
        f.groupby("date")
        .agg(
            region_count=("territory", "nunique"),
            farmgate_mean=("price_uah_kg", "mean"),
            farmgate_median=("price_uah_kg", "median"),
            farmgate_min=("price_uah_kg", "min"),
            farmgate_max=("price_uah_kg", "max"),
            farmgate_std=("price_uah_kg", "std"),
        )
        .reset_index()
    )
    farm_disp["farmgate_cv"] = farm_disp["farmgate_std"] / farm_disp["farmgate_mean"]

    c = consumer[(consumer["territory"] != "Україна") & (consumer["quality_flag"] == "ok")].copy()
    cons_disp = (
        c.groupby(["date", "product"])
        .agg(
            region_count=("territory", "nunique"),
            consumer_mean=("price_uah_kg", "mean"),
            consumer_median=("price_uah_kg", "median"),
            consumer_min=("price_uah_kg", "min"),
            consumer_max=("price_uah_kg", "max"),
            consumer_std=("price_uah_kg", "std"),
        )
        .reset_index()
    )
    cons_disp["consumer_cv"] = cons_disp["consumer_std"] / cons_disp["consumer_mean"]

    p = prozorro[(prozorro["quality_flag"] == "ok") & prozorro["territory"].notna()].copy()
    p["month"] = p["date"].values.astype("datetime64[M]")
    proz_reg = (
        p.groupby(["month", "product", "territory"])
        .agg(price_median_uah_kg=("price_uah_kg", "median"), n_lots=("price_uah_kg", "size"))
        .reset_index()
        .rename(columns={"month": "date"})
    )

    fv = farm_volumes.copy()
    fv_price = fv[fv["indicator"].str.contains("Середня ціна купівлі", case=False, na=False)].copy()
    return {
        "regional_farmgate_dispersion": farm_disp,
        "regional_consumer_dispersion": cons_disp,
        "regional_prozorro_medians": proz_reg,
        "farm_volumes_procurement_prices": fv_price,
    }


def make_figures(farm, processor, consumer, proz_aggs, retail_aggs, models, regional, trade, cost_index) -> list[dict]:
    figures = []

    def savefig(name, title, source, relevance):
        path = FIG_OUT / name
        plt.tight_layout()
        plt.savefig(path, dpi=200, bbox_inches="tight")
        plt.close()
        figures.append({"file": str(path), "title": title, "source": source, "relevance": relevance})
        return path

    # Figure 1: value-chain architecture.
    plt.figure(figsize=(11, 3.4))
    ax = plt.gca()
    ax.axis("off")
    boxes = [
        (0.02, 0.45, "Farm-gate\nraw milk\n(SSSU)"),
        (0.28, 0.45, "Processor-level\nindustrial producer prices\n(SSSU)"),
        (0.55, 0.45, "Institutional procurement\nProZorro lots"),
        (0.78, 0.45, "Retail / consumer\nSSSU + Silpo/Novus"),
    ]
    for x, y, txt in boxes:
        ax.add_patch(plt.Rectangle((x, y), 0.18, 0.28, fill=False, linewidth=1.6))
        ax.text(x + 0.09, y + 0.14, txt, ha="center", va="center", fontsize=10)
    for x1, x2 in [(0.20, 0.28), (0.46, 0.55), (0.73, 0.78)]:
        ax.annotate("", xy=(x2, 0.59), xytext=(x1, 0.59), arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.text(0.02, 0.18, "H1", fontsize=12, weight="bold")
    ax.annotate("", xy=(0.46, 0.25), xytext=(0.08, 0.25), arrowprops=dict(arrowstyle="<->", lw=1.2))
    ax.text(0.55, 0.18, "H2", fontsize=12, weight="bold")
    ax.annotate("", xy=(0.94, 0.25), xytext=(0.58, 0.25), arrowprops=dict(arrowstyle="<->", lw=1.2))
    ax.set_title("Dairy Value Chain and Empirical Hypotheses", fontsize=13, weight="bold")
    savefig("fig_01_value_chain.png", "Dairy value chain and empirical hypotheses", "Author design based on thesis data architecture.", "Core conceptual figure for Chapters 1-4.")

    # Figure 2: farmgate and processor monthly selected series.
    f = farm[(farm["territory"] == "Україна") & (farm["farm_type"] == "enterprises") & (farm["quality_flag"] == "ok")]
    p = processor[(processor["territory"] == "Україна") & (processor["quality_flag"] == "ok")]
    plt.figure(figsize=(11, 6))
    plt.plot(f["date"], f["price_uah_kg"], label="Farm-gate raw milk", lw=2.2, color="black")
    for prod in ["drinking_milk", "sour_cream", "butter", "hard_cheese"]:
        g = p[p["product"] == prod]
        if not g.empty:
            plt.plot(g["date"], g["price_uah_kg"], label=f"Processor {product_label(prod)}", alpha=0.85)
    plt.title("Observed Monthly Farm-Gate and Processor-Level Prices", weight="bold")
    plt.ylabel("UAH/kg")
    plt.legend(ncol=2, fontsize=8)
    plt.grid(alpha=0.25)
    savefig("fig_02_h1_monthly_prices.png", "Observed monthly farm-gate and processor-level prices", "SSSU raw workbooks; author conversion to UAH/kg.", "H1 data foundation.")

    # Figure 3: processor to consumer official series.
    c = consumer[(consumer["territory"] == "Україна") & (consumer["quality_flag"] == "ok")]
    plt.figure(figsize=(11, 5.5))
    for prod, color in [("drinking_milk", "#1f77b4"), ("sour_cream", "#2ca02c")]:
        gp = p[p["product"] == prod]
        gc = c[c["product"] == prod]
        plt.plot(gp["date"], gp["price_uah_kg"], ls="--", color=color, label=f"Processor {product_label(prod)}")
        plt.plot(gc["date"], gc["price_uah_kg"], color=color, label=f"Consumer {product_label(prod)}")
    plt.title("Processor-Level and Official Consumer Prices", weight="bold")
    plt.ylabel("UAH/kg")
    plt.legend(ncol=2, fontsize=8)
    plt.grid(alpha=0.25)
    savefig("fig_03_h2_official_bridge.png", "Processor-level and official consumer prices", "SSSU producer and consumer price workbooks.", "H2 official benchmark.")

    # Figure 4: model coefficient summary.
    rel = models[(models["method"] == "ECM") & models["model_id"].str.startswith(("H1_farmgate_to_processor", "H2_processor_to_consumer"))].copy()
    rel = rel.sort_values(["hypothesis", "reliability", "model_id"]).head(16)
    if not rel.empty:
        plt.figure(figsize=(11, 5.8))
        labels = rel["model_id"].str.replace("H1_farmgate_to_processor_", "H1 ", regex=False).str.replace("H2_processor_to_consumer_", "H2 ", regex=False)
        colors = rel["reliability"].map({"reliable": "#2166ac", "probable / needs validation": "#fdae61"}).fillna("#bdbdbd")
        plt.barh(labels, rel["long_run_beta"], color=colors)
        plt.axvline(0, color="black", lw=0.8)
        plt.title("Long-Run Pass-Through Estimates in Main ECM Blocks", weight="bold")
        plt.xlabel("Long-run elasticity / pass-through coefficient")
        plt.grid(axis="x", alpha=0.25)
        savefig("fig_04_main_ecm_coefficients.png", "Long-run pass-through estimates in main ECM blocks", "Author estimates from cleaned monthly data.", "Compact Chapter 5/6 result figure.")

    # Figure 5: ProZorro weekly medians.
    pw = proz_aggs["W-MON"].copy()
    keep = ["drinking_milk", "sour_cream", "butter", "hard_cheese", "cottage_cheese"]
    plt.figure(figsize=(11, 5.5))
    for prod in keep:
        g = pw[pw["product"] == prod]
        if len(g) > 2:
            plt.plot(g["date"], g["price_median_uah_kg"], label=product_label(prod))
    plt.title("ProZorro Dairy Procurement Prices, Weekly Median", weight="bold")
    plt.ylabel("UAH/kg")
    plt.legend(ncol=3, fontsize=8)
    plt.grid(alpha=0.25)
    savefig("fig_05_prozorro_weekly.png", "ProZorro dairy procurement prices, weekly median", "ProZorroM(full), author cleaning and product mapping.", "H1/H2 procurement bridge.")

    # Figure 6: retail daily product prices.
    pooled = retail_aggs["pooled_daily"]
    plt.figure(figsize=(11, 5.5))
    for prod in ["drinking_milk", "sour_cream", "butter", "hard_cheese", "cottage_cheese", "yogurt"]:
        g = pooled[pooled["product"] == prod]
        if len(g) > 2:
            plt.plot(g["date"], g["price_observed_median_uah_kg"], label=product_label(prod))
    plt.title("Silpo and Novus Retail Prices, Pooled Daily Median", weight="bold")
    plt.ylabel("UAH/kg")
    plt.legend(ncol=3, fontsize=8)
    plt.grid(alpha=0.25)
    savefig("fig_06_retail_pooled_daily.png", "Silpo and Novus retail prices, pooled daily median", "Silpo and Novus sheets, author SKU cleaning.", "H2 retail mechanism data.")

    pkg = retail_aggs["package_pooled_daily"]
    plt.figure(figsize=(11, 5.5))
    for prod in ["drinking_milk", "sour_cream", "butter", "hard_cheese", "cottage_cheese", "yogurt", "kefir"]:
        g = pkg[pkg["product"] == prod]
        if len(g) > 2:
            plt.plot(g["date"], g["price_package_median_uah"], label=product_label(prod))
    plt.title("Silpo and Novus Retail SKU Package Prices, Pooled Daily Median", weight="bold")
    plt.ylabel("UAH per observed SKU package")
    plt.legend(ncol=3, fontsize=8)
    plt.grid(alpha=0.25)
    savefig("fig_06b_retail_package_daily.png", "Silpo and Novus retail SKU package prices, pooled daily median", "Silpo and Novus sheets, author title/name product mapping.", "H2 retail SKU mechanism; levels are package-specific.")

    # Figure 7: Silpo discount.
    rd = retail_aggs["daily"]
    silpo = rd[rd["retailer"] == "silpo"]
    if not silpo.empty:
        plt.figure(figsize=(11, 5.5))
        for prod in ["drinking_milk", "sour_cream", "butter", "hard_cheese", "yogurt"]:
            g = silpo[silpo["product"] == prod]
            if len(g) > 2:
                plt.plot(g["date"], g["discount_share"], label=product_label(prod))
        plt.title("Silpo Discount Incidence by Product", weight="bold")
        plt.ylabel("Share of observed items with discount")
        plt.ylim(0, 1)
        plt.legend(ncol=3, fontsize=8)
        plt.grid(alpha=0.25)
        savefig("fig_07_silpo_discount_incidence.png", "Silpo discount incidence by product", "Silpo sheet, author SKU cleaning.", "Promotion mechanism under H2.")

    # Figure 8: regional farmgate dispersion.
    r = regional["regional_farmgate_dispersion"].dropna(subset=["farmgate_cv"])
    if not r.empty:
        plt.figure(figsize=(11, 4.8))
        plt.plot(r["date"], r["farmgate_cv"], color="#7b3294", lw=2)
        plt.title("Regional Dispersion of Enterprise Farm-Gate Raw Milk Prices", weight="bold")
        plt.ylabel("Coefficient of variation")
        plt.grid(alpha=0.25)
        savefig("fig_08_regional_farmgate_dispersion.png", "Regional dispersion of enterprise farm-gate raw milk prices", "SSSU agricultural sales raw workbook.", "Short regional extension for H1.")

    # Figure 9: dairy-only trade.
    exp = trade["exports"]
    imp = trade["imports"]
    e = exp[exp["is_dairy_only_0401_0406"]].groupby("year")["value_eur_mln"].sum()
    i = imp[imp["is_dairy_only_0401_0406"]].groupby("year")["value_eur_mln"].sum()
    tr = pd.concat([e.rename("exports"), i.rename("imports")], axis=1).reset_index()
    if not tr.empty:
        plt.figure(figsize=(10, 4.8))
        plt.plot(tr["year"], tr["exports"], marker="o", label="Dairy exports HS 0401-0406")
        plt.plot(tr["year"], tr["imports"], marker="o", label="Dairy imports HS 0401-0406")
        plt.title("Ukraine Dairy Trade, HS 0401-0406", weight="bold")
        plt.ylabel("EUR million")
        plt.legend(fontsize=8)
        plt.grid(alpha=0.25)
        savefig("fig_09_dairy_trade_hs0401_0406.png", "Ukraine dairy trade, HS 0401-0406", "ITC Trade Map based on SSSU and UN COMTRADE statistics.", "Chapter 2 trade context.")

    # Figure 10: livestock cost index.
    ci = cost_index[cost_index["basis"].str.contains("попереднього місяця", case=False, na=False)]
    if not ci.empty:
        plt.figure(figsize=(10, 4.8))
        plt.plot(ci["date"], ci["index_value"], color="#b2182b", lw=2)
        plt.axhline(100, color="black", lw=0.8)
        plt.title("Input-Cost Pressure in Livestock Production", weight="bold")
        plt.ylabel("Index, previous month = 100")
        plt.grid(alpha=0.25)
        savefig("fig_10_livestock_cost_index.png", "Input-cost pressure in livestock production", "SSSU agricultural production cost index workbook.", "Minor Chapter 1-2 cost-pressure context.")

    return figures


def write_markdown_outputs(
    farm,
    processor,
    consumer,
    prozorro,
    retail,
    farm_volumes,
    cost_index,
    trade,
    corruption_summary,
    validation,
    models,
    stationarity,
    vecm,
    regional,
    figures,
):
    row_counts = []
    for name, df in [
        ("clean_farmgate_monthly", farm),
        ("clean_processor_monthly", processor),
        ("clean_consumer_monthly", consumer),
        ("clean_prozorro_lot_level", prozorro),
        ("clean_retail_sku_day", retail),
        ("clean_farm_volumes", farm_volumes),
        ("clean_cost_index", cost_index),
    ]:
        row_counts.append(
            {
                "dataset": name,
                "rows": len(df),
                "date_min": df["date"].min().date().isoformat() if "date" in df and df["date"].notna().any() else "",
                "date_max": df["date"].max().date().isoformat() if "date" in df and df["date"].notna().any() else "",
                "products": ", ".join(sorted(map(str, df["product"].dropna().unique()))) if "product" in df else "",
                "quality_ok_share": round(float((df.get("quality_flag", pd.Series("ok", index=df.index)) == "ok").mean()), 3),
            }
        )
    row_counts = pd.DataFrame(row_counts)
    row_counts.to_csv(TABLE_OUT / "table_dataset_row_counts.csv", index=False)
    validation.to_csv(TABLE_OUT / "table_old_vs_new_validation.csv", index=False)
    models.to_csv(TABLE_OUT / "table_model_results_all.csv", index=False)
    stationarity.to_csv(TABLE_OUT / "table_stationarity.csv", index=False)
    vecm.to_csv(TABLE_OUT / "table_vecm_systems.csv", index=False)

    product_dictionary = pd.DataFrame(
        [
            {"product": p, "label": product_label(p), "role": product_role(p), "main_use": product_main_use(p)}
            for p in PRODUCT_ORDER
        ]
    )
    product_dictionary.to_csv(TABLE_OUT / "table_product_dictionary.csv", index=False)

    audit_md = f"""# New Model Dataset Audit

## Mandatory Audit Finding

The new modelling workflow must rebuild clean data from the raw files before any econometric model is estimated.

Known risks explicitly handled:

- `Processor_price` is reported in hryvnia per tonne and is converted to hryvnia per kilogram (UAH/kg) by dividing by `1000`.
- `Farm_price` / `Farm_milk_2015.xlsx` carries the short unit label `Гривня`, but the magnitude is economically a tonne price. It is converted to UAH/kg by dividing by `1000`.
- `ProzorroM(full)` stores numeric fields with non-breaking spaces and text formatting. Numeric columns are parsed before aggregation.
- Silpo and Novus have visible retail classification errors. Product labels are rebuilt from `product_title` and `product_name`.
- Every cleaned table contains a controlled `product` column.

## Cleaned Dataset Coverage

{to_markdown(row_counts)}

## Integrated Consumer Sheet Corruption Check

{to_markdown(corruption_summary)}

## Old-Model vs New Observed Validation

{to_markdown(validation, max_rows=40)}

## Product Dictionary

{to_markdown(product_dictionary)}

## Decision

Observed State Statistics Service of Ukraine (SSSU) monthly data are the main empirical base. Old reconstructed sheets are retained only as robustness or appendix evidence unless their validation is strong. ProZorro and retail data are useful for institutional and promotional mechanisms, but their shorter overlap makes them weaker than the official monthly core.
"""
    (OUT / "newmodel_dataset_audit.md").write_text(audit_md, encoding="utf-8")

    data_dict_rows = [
        {
            "dataset": "Farm-gate monthly",
            "source": "SSSU agricultural sales; Farm_milk_2015.xlsx",
            "frequency": "Monthly",
            "layer": "Farm-gate",
            "unit": "Converted to UAH/kg",
            "role": "Main H1 upstream price",
            "limitation": "Farm price label is abbreviated; conversion must be documented.",
        },
        {
            "dataset": "Processor monthly",
            "source": "SSSU producer prices; Processor_2013_2026.xlsx",
            "frequency": "Monthly",
            "layer": "Processor",
            "unit": "UAH/tonne converted to UAH/kg",
            "role": "Main H1 downstream and H2 upstream price",
            "limitation": "Ukraine-only; no regional processor series.",
        },
        {
            "dataset": "Consumer monthly",
            "source": "SSSU consumer prices; Consumer_2017_2026.xlsx",
            "frequency": "Monthly",
            "layer": "Official consumer",
            "unit": "UAH/kg",
            "role": "Main H2 official benchmark",
            "limitation": "Only selected products: milk, sour cream, soft cheese.",
        },
        {
            "dataset": "ProZorro lot-level",
            "source": "ProzorroM(full)",
            "frequency": "Event / lot",
            "layer": "Institutional procurement",
            "unit": "UAH/kg after parsing",
            "role": "H1/H2 institutional mechanism",
            "limitation": "Irregular events; piece units need package parsing.",
        },
        {
            "dataset": "Silpo/Novus SKU-day",
            "source": "Retail sheets in newmodel workbook",
            "frequency": "Daily observed online prices",
            "layer": "Retail",
            "unit": "UAH/kg where package parsing is reliable",
            "role": "H2 retail mechanism and discount analysis",
            "limitation": "Short period; classification and package parsing required.",
        },
    ]
    data_dict = pd.DataFrame(data_dict_rows)
    (OUT / "clean_newmodel_data_dictionary.md").write_text(
        "# Clean New Model Data Dictionary\n\n" + to_markdown(data_dict) + "\n\n" + "# Product Dictionary\n\n" + to_markdown(product_dictionary),
        encoding="utf-8",
    )

    rel_models = models.sort_values(["hypothesis", "reliability", "method", "model_id"])
    reliable = rel_models[rel_models["reliability"].eq("reliable")]
    probable = rel_models[rel_models["reliability"].eq("probable / needs validation")]
    model_md = f"""# New Model Results Summary

## Evidence Hierarchy

1. Main evidence: observed monthly SSSU farm-gate, processor-level, and official consumer prices.
2. Mechanism evidence: ProZorro procurement and Silpo/Novus retail prices.
3. Extension evidence: regional dispersion and old-model validation.
4. Appendix-only evidence: weak old reconstructed or short-overlap models.

## Reliable Models

{to_markdown(reliable[[c for c in ['model_id','hypothesis','method','n','period_start','period_end','long_run_beta','ect','ect_p','reliability','interpretation'] if c in reliable.columns]], max_rows=60)}

## Probable / Needs Validation Models

{to_markdown(probable[[c for c in ['model_id','hypothesis','method','n','period_start','period_end','long_run_beta','ect','ect_p','reliability','interpretation'] if c in probable.columns]], max_rows=60)}

## VECM Systems

{to_markdown(vecm)}

## Interpretation Rules

- A negative and significant error-correction term supports a long-run price relation and shows how quickly deviations are corrected.
- In this thesis, slow or incomplete correction is interpreted as evidence consistent with bargaining power, not as automatic legal proof of abuse.
- Nonlinear Autoregressive Distributed Lag (NARDL) asymmetry is treated as a mechanism under H1 or H2, not as a separate hypothesis.
- ProZorro and retail results are interpreted as institutional and promotional mechanisms because the overlap is short.
"""
    (OUT / "newmodel_results_summary.md").write_text(model_md, encoding="utf-8")

    protocol_md = """# Model Execution Protocol for H1 and H2

## H1

H1 tests whether market power exists between farm-gate raw milk producers and processors. The main models use observed monthly Ukrainian national data. The dependent variables are processor-level prices. The upstream variable is farm-gate raw milk price.

Execution order:

1. Harmonize all prices to UAH/kg.
2. Keep Ukraine national data for headline models.
3. Test stationarity with Augmented Dickey-Fuller (ADF) and Kwiatkowski-Phillips-Schmidt-Shin (KPSS).
4. Estimate long-run relation and Error Correction Model (ECM).
5. Use Autoregressive Distributed Lag (ARDL) lag selection as support.
6. Estimate Nonlinear Autoregressive Distributed Lag (NARDL) only as a mechanism test.
7. Use Vector Error Correction Model (VECM) only for three-layer monthly systems with enough observations.

## H2

H2 tests whether market power exists between processors/procurement and downstream retail actors. The main official benchmark is processor-level to official consumer price. ProZorro and retail SKU/day models explain institutional and promotional mechanisms.

Execution order:

1. Estimate processor to official consumer monthly ECM/NARDL models.
2. Aggregate ProZorro by product and date/week/month.
3. Build retail Stock Keeping Unit (SKU) panel after product reclassification.
4. Estimate short retail mechanism models only after overlap checks.
5. Treat discount incidence and discount depth as H2 mechanisms.

## Reliability

Main-text models must be economically interpretable, have enough observations, and pass core diagnostics. Weak models are retained only in appendix notes.
"""
    (OUT / "model_execution_protocol_H1_H2.md").write_text(protocol_md, encoding="utf-8")

    fig_df = pd.DataFrame(figures)
    fig_df.to_csv(TABLE_OUT / "table_figure_inventory.csv", index=False)
    fig_md = f"""# Figure and Table Rebuild Plan

## Figure Inventory

{to_markdown(fig_df, max_rows=50)}

## Keep in Main Text

- Value-chain architecture.
- Observed monthly farm-gate and processor-level prices.
- Processor-to-consumer official bridge.
- Main ECM coefficient summary.
- ProZorro weekly procurement prices.
- Retail pooled daily prices.
- Silpo discount incidence.
- Regional farm-gate dispersion as a short extension.
- HS 0401-0406 dairy trade if Chapter 2 needs trade context.

## Remove or Appendix

- Old EU reconstruction charts unless validation is strong.
- Dense coefficient plots without a clear hypothesis link.
- ProZorro or retail graphs with too few observations.
- Any graph with Ukrainian-only labels or excessive numeric labels.
"""
    (OUT / "figure_table_rebuild_plan.md").write_text(fig_md, encoding="utf-8")

    source_rows = [
        {
            "source": k,
            "link": v,
            "use": source_use(k),
            "reliability": source_reliability(k),
        }
        for k, v in SOURCE_LINKS.items()
    ]
    source_df = pd.DataFrame(source_rows)
    source_md = f"""# Source Reliability and Legal Cases

## Source Table

{to_markdown(source_df, max_rows=50)}

## Legal Interpretation Discipline

Legal and institutional sources motivate the research question. They do not prove unlawful behaviour in the thesis sample. The thesis should write that observed price dynamics are consistent with market-power mechanisms where they show delayed, incomplete, asymmetric, or strategic transmission.

## Value-Added and Dairy-Share Discipline

Use official State Statistics Service of Ukraine sources for value added where direct rows are available. If milk-specific value added is not directly available, present only sector-level or allocated estimates with a clear caveat. The Food Security Strategy fact that milk and dairy products accounted for 13.7 percent of household monetary food expenditure in 2021 can be used in the Introduction as a concise relevance fact after checking the citation in the final reference list.
"""
    (OUT / "source_reliability_and_legal_cases.md").write_text(source_md, encoding="utf-8")

    comment_summary = build_comments_transcript_summary()
    (OUT / "critical_comments_transcript_loy_summary.md").write_text(comment_summary, encoding="utf-8")

    rebuild_plan = f"""# Thesis Rebuild Plan: Market Power in the Dairy Value Chain in Ukraine

## Core Reframing

The thesis must no longer be framed as retail-to-farm price transmission. The research object is market power in the Ukrainian dairy value chain. Price transmission is the empirical method.

## Hypotheses

- H1: Market power exists between farm-gate raw milk producers and processors.
- H2: Market power exists between processors/procurement and downstream retail actors.

## Main Evidence Now Built

{to_markdown(row_counts)}

## Main Model Result Inventory

{to_markdown(rel_models[['model_id','hypothesis','method','n','reliability','interpretation']].head(80), max_rows=80)}

## Writing Instruction

Chapters 5 and 6 must be written from the cleaned model evidence. Paragraphs should state the empirical result, interpret it economically, connect it to market power, and return to H1 or H2.
"""
    (OUT / "thesis_rebuild_plan_market_power.md").write_text(rebuild_plan, encoding="utf-8")


def product_role(p: str) -> str:
    if p == "raw_milk":
        return "Farm-gate raw material"
    if p == "exclude_non_dairy":
        return "Excluded"
    if p == "other_dairy":
        return "Dairy but weakly comparable"
    return "Comparable dairy product group"


def product_main_use(p: str) -> str:
    if p == "raw_milk":
        return "H1 upstream"
    if p in {"drinking_milk", "sour_cream", "butter", "hard_cheese", "kefir", "skim_milk_powder"}:
        return "H1/H2 where matching is available"
    if p in {"cottage_cheese", "cream", "condensed_milk", "yogurt"}:
        return "H2 / retail-ProZorro mechanism if support is sufficient"
    return "Audit or exclusion"


def source_use(name: str) -> str:
    if "Food Security" in name:
        return "Introduction relevance and wartime food-price context"
    if "SDG" in name:
        return "Official milk and dairy consumption context"
    if "farm-gate" in name:
        return "Raw milk farm-gate price source"
    if "producer" in name:
        return "Processor-level price source"
    if "consumer" in name:
        return "Official consumer price source"
    if "ProZorro" in name:
        return "Institutional procurement context and data"
    if "Loy" in name:
        return "Methodological frame for market power, asymmetry, search/menu costs"
    if "AMCU" in name:
        return "Legal/institutional motivation"
    return "Source support"


def source_reliability(name: str) -> str:
    if "AMCU" in name or "SSSU" in name or "Food Security" in name or "ProZorro" in name:
        return "reliable official/institutional source"
    if "Loy" in name:
        return "peer-reviewed journal source"
    return "usable with verification"


def build_comments_transcript_summary() -> str:
    commented = load_doc_text(COMMENTED)
    transcript = load_doc_text(TRANSCRIPT)
    d2_outline = extract_doc_outline(DRAFT2)
    d3_outline = extract_doc_outline(DRAFT3)

    # Pull the most relevant comment headings from the diagnostic document.
    comment_headings = [t for t in commented if re.match(r"C\d+\s+-", t) or "Executive Summary" in t or "Model and Results" in t]
    transcript_hits = [
        t
        for t in transcript
        if any(
            key.lower() in t.lower()
            for key in [
                "market power",
                "price transmission",
                "hypothesis",
                "producer",
                "processor",
                "farm",
                "retail",
                "ProZorro",
                "Loy",
                "EU",
            ]
        )
    ][:80]
    return f"""# Critical Comments, Transcript, and Loy Summary

## Commented Draft 2: Main Revision Logic

{to_markdown(pd.DataFrame({'comment_or_section': comment_headings[:60]}), max_rows=60)}

## Draft 2 Structural Risk

Draft 2 still contains the old four-hypothesis and retail-to-farm framing. Its page setup and style system are valuable, but the argument must be rebuilt.

{to_markdown(d2_outline.head(40), max_rows=40)}

## Draft 3 Use

Draft 3 is useful as a low-priority idea text. It already moves toward two hypotheses, but its H1/H2 ordering differs from the final requested structure.

{to_markdown(d3_outline.head(30), max_rows=30)}

## Nivievskyi Transcript: Operational Reading

The transcript requires a narrow market-power thesis. Price transmission must be the method. The clean value chain is farm-gate raw milk, processor-level sale, institutional procurement, and retail/consumer prices. Direct farm-gate-to-retail models are not the main identification route. Observed Ukrainian official data must dominate the empirical strategy. EU reconstruction should be appendix/support only unless validated.

## Transcript Anchors

{to_markdown(pd.DataFrame({'transcript_anchor': transcript_hits}), max_rows=60)}

## Loy et al. 2016

Loy et al. study asymmetric cost pass-through using weekly store-level milk data and threshold error-correction models. Their central lesson for this thesis is that asymmetric transmission should not be mechanically equated with market power. Market power, menu costs, search costs, promotion practices, and perishability can all shape adjustment. The thesis should therefore interpret asymmetry as one market-power-consistent mechanism only when supported by the broader chain evidence.
"""


def add_heading(doc, text, level=1):
    if level == 0:
        p = doc.add_paragraph(text, style="Title")
    else:
        p = doc.add_heading(text, level=level)
    return p


def add_body(doc, text, style=None):
    for para in textwrap.dedent(text).strip().split("\n\n"):
        if not para.strip():
            continue
        p = doc.add_paragraph(para.strip(), style=style or "Normal")
        p.paragraph_format.space_after = Pt(6)
    return doc


def add_df_table(doc, df: pd.DataFrame, title: str, max_rows=12):
    doc.add_paragraph(title, style="Table" if "Table" in [s.name for s in doc.styles] else "Normal")
    if df.empty:
        doc.add_paragraph("No rows.")
        return
    d = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(d.columns))
    table_style_names = [s.name for s in doc.styles if s.type == 3]
    for candidate in ["Table Grid", "Grid Table 1 Light", "Normal Table"]:
        if candidate in table_style_names:
            try:
                table.style = candidate
                break
            except Exception:
                continue
    for j, col in enumerate(d.columns):
        table.rows[0].cells[j].text = str(col)
    for _, row in d.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(d.columns):
            val = row[col]
            if isinstance(val, float):
                cells[j].text = "" if pd.isna(val) else f"{val:.3g}"
            else:
                cells[j].text = "" if pd.isna(val) else str(val)
    if len(df) > max_rows:
        doc.add_paragraph(f"Note: table shows {max_rows} of {len(df)} rows.")


def add_figure(doc, path: str, caption: str, source: str):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    try:
        r.add_picture(path, width=Inches(6.2))
    except Exception:
        return
    cap = doc.add_paragraph(caption, style="figures" if "figures" in [s.name for s in doc.styles] else "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src = doc.add_paragraph(f"Source: {source}", style="source" if "source" in [s.name for s in doc.styles] else "Normal")
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER


def clear_doc_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def build_docx(models: pd.DataFrame, figures: list[dict], validation: pd.DataFrame, vecm: pd.DataFrame, regional: dict, trade: dict, cost_index: pd.DataFrame) -> Path:
    doc = Document(DRAFT2)
    clear_doc_body(doc)
    # Preserve draft 2 section geometry and styles; rebuild content.
    sec = doc.sections[0]
    sec.top_margin = Inches(1.5)
    sec.bottom_margin = Inches(1.33)
    sec.left_margin = Inches(2.0)
    sec.right_margin = Inches(1.37)

    title = doc.add_paragraph("MARKET POWER IN THE DAIRY VALUE CHAIN IN UKRAINE", style="Author" if "Author" in [s.name for s in doc.styles] else "Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("New Model Draft Based on Cleaned Observed Data", style="Title")
    doc.add_paragraph("Maksym Charniuk", style="Author" if "Author" in [s.name for s in doc.styles] else "Normal")
    doc.add_page_break()

    add_heading(doc, "LIST OF ABBREVIATIONS", level=0)
    abbr = [
        ("State Statistics Service of Ukraine (SSSU)", "Official Ukrainian statistical authority."),
        ("Antimonopoly Committee of Ukraine (AMCU)", "Ukrainian competition authority."),
        ("Stock Keeping Unit (SKU)", "Retail item identifier based on title, brand, package, and product group."),
        ("Harmonized System (HS)", "Trade classification system."),
        ("Common Procurement Vocabulary (CPV)", "Public-procurement product classification."),
        ("Autoregressive Distributed Lag (ARDL)", "Dynamic time-series model."),
        ("Error Correction Model (ECM)", "Model for short-run correction toward a long-run relation."),
        ("Vector Error Correction Model (VECM)", "System version of ECM."),
        ("Nonlinear Autoregressive Distributed Lag (NARDL)", "ARDL model allowing positive and negative shocks to differ."),
        ("Heteroskedasticity and Autocorrelation Consistent (HAC)", "Robust standard-error correction."),
        ("Gross Value Added (GVA)", "Value-added measure."),
        ("European Union (EU)", "External benchmark region."),
    ]
    add_df_table(doc, pd.DataFrame(abbr, columns=["Abbreviation", "Meaning"]), "Table A. Abbreviation discipline for the new draft", max_rows=20)
    doc.add_page_break()

    add_heading(doc, "INTRODUCTION", level=0)
    add_body(
        doc,
        """
        This thesis studies market power in the Ukrainian dairy value chain. The empirical object is not price transmission for its own sake. Price transmission is used as the method for detecting whether one layer of the chain can delay, absorb, or selectively pass on price changes.

        The Ukrainian dairy chain is economically relevant because milk and dairy products are a visible part of household food expenditure and because the chain connects many farmers, a narrower processing layer, public procurement, and large retail networks. The Food Security Strategy reports milk and dairy products at 13.7 percent of household monetary food expenditure in 2021. This makes dairy a useful case for studying food-price adjustment and bargaining power.

        The core research question is whether price movements along the dairy value chain reveal market power between upstream raw-milk producers, processors, institutional procurement, and retail actors.

        Hypothesis 1: Market power exists between farm-gate raw milk producers and processors.

        Hypothesis 2: Market power exists between processors/procurement and downstream retail actors.

        The thesis interprets market power cautiously. Delayed, incomplete, asymmetric, or promotion-mediated pass-through is evidence consistent with market power. It is not treated as proof of unlawful conduct.
        """,
    )

    add_heading(doc, "MARKET ANALYSIS AND INSTITUTIONAL BACKGROUND", level=0)
    add_body(
        doc,
        """
        The dairy value chain is treated as a three-layer economic structure with an institutional procurement bridge. The first layer is farm-gate raw milk. The second layer is processor-level output prices, measured through prices of industrial producers. The third layer is consumer and retail prices, measured through official consumer prices and observed Silpo and Novus retail prices. ProZorro is interpreted as an institutional procurement channel and a proxy for processor/procurement realization prices.

        The new model uses observed Ukrainian data first. Old reconstructed daily and weekly series are retained only as robustness evidence. This responds directly to the adviser comment that the thesis should not rely too heavily on EU-based reconstruction when Ukrainian observed monthly data are available.
        """,
    )
    add_figure(doc, figures[0]["file"], "Figure 2.1. Dairy value chain and empirical hypotheses.", figures[0]["source"])
    for fig in figures:
        if "trade" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 2.2. Ukraine dairy trade, HS 0401-0406.", fig["source"])
            break
    for fig in figures:
        if "cost_index" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 2.3. Input-cost pressure in livestock production.", fig["source"])
            break

    add_heading(doc, "LITERATURE REVIEW", level=0)
    add_body(
        doc,
        """
        The literature review should be organized around buyer power, asymmetric transmission, retail promotions, perishability, and adjustment costs. Loy et al. (2016) are especially important because they show that asymmetric cost pass-through cannot be mechanically reduced to market power. Market power, search costs, menu costs, mistakes, and perishability can all shape adjustment.

        This thesis follows that logic. It treats asymmetric adjustment as one mechanism under the two market-power hypotheses. Discounts, procurement-scale effects, and short-run retail reactions are not separate hypotheses. They explain how market power may appear in the Ukrainian dairy chain.
        """,
    )

    add_heading(doc, "METHODOLOGY", level=0)
    add_body(
        doc,
        """
        The empirical strategy starts with a data audit. All prices are harmonized to hryvnia per kilogram (UAH/kg). Processor prices reported in hryvnia per tonne are divided by 1000. Farm-gate prices with a short hryvnia unit label are treated as tonne prices after magnitude validation and are also divided by 1000. ProZorro numeric fields are cleaned from non-breaking spaces. Retail products are reclassified from product title and product name.

        The main H1 models use observed monthly national State Statistics Service of Ukraine data. The dependent variables are processor-level prices. The upstream price is farm-gate raw milk. The main H2 official models connect processor-level prices with official consumer prices. ProZorro and retail Stock Keeping Unit (SKU) models are used as mechanism evidence because their overlap is shorter.

        The main econometric families are Autoregressive Distributed Lag (ARDL), Error Correction Model (ECM), Nonlinear Autoregressive Distributed Lag (NARDL), and Vector Error Correction Model (VECM). The Error Correction Model is the main reporting format when the variables support a long-run relation.
        """,
    )
    add_df_table(doc, validation[["comparison", "n", "correlation", "MAPE_pct", "decision"]].head(12), "Table 4.1. Old-model reconstruction validation", max_rows=12)

    add_heading(doc, "EMPIRICAL RESULTS", level=0)
    add_heading(doc, "Hypothesis 1: Farm-Gate Producers and Processors", level=2)
    add_body(
        doc,
        """
        The H1 block is the strongest part of the new empirical design because it uses observed monthly Ukrainian data over a long period. The economic question is whether processor-level prices move with farm-gate raw milk prices completely and quickly, or whether correction is slow, incomplete, or asymmetric.
        """,
    )
    for fig in figures:
        if "h1_monthly" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.1. Observed monthly farm-gate and processor-level prices.", fig["source"])
    h1 = models[(models["hypothesis"] == "H1") & (models["method"] == "ECM")].copy()
    h1_main = h1[h1["model_id"].str.contains("farmgate_to_processor")].copy()
    add_df_table(
        doc,
        h1_main[["model_id", "n", "long_run_beta", "ect", "ect_p", "reliability", "interpretation"]].head(10),
        "Table 5.1. H1 monthly ECM results",
        max_rows=10,
    )
    add_body(
        doc,
        """
        The interpretation should not mechanically list coefficients. A negative and significant error-correction term means that deviations from the long-run farm-gate-to-processor relation are corrected over time. If correction is slow or incomplete, this is consistent with bargaining frictions at the processor level.
        """,
    )

    add_heading(doc, "Hypothesis 2: Processor/Procurement and Retail Actors", level=2)
    add_body(
        doc,
        """
        The H2 block has two layers. The official monthly layer links processor-level prices to official consumer prices. The retail mechanism layer links ProZorro procurement prices to Silpo and Novus observed retail prices, while separating observed prices from baseline non-discount prices.
        """,
    )
    for fig in figures:
        if "official_bridge" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.2. Processor-level and official consumer prices.", fig["source"])
    h2 = models[(models["hypothesis"] == "H2") & (models["method"] == "ECM")].copy()
    add_df_table(
        doc,
        h2[["model_id", "n", "long_run_beta", "ect", "ect_p", "reliability", "interpretation"]].head(12),
        "Table 5.2. H2 official and procurement ECM results",
        max_rows=12,
    )
    for fig in figures:
        if "prozorro_weekly" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.3. ProZorro dairy procurement prices.", fig["source"])
        if "retail_pooled_daily" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.4. Silpo and Novus pooled retail prices.", fig["source"])
        if "retail_package_daily" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.5. Silpo and Novus pooled retail SKU package prices.", fig["source"])
        if "discount_incidence" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.6. Silpo discount incidence by product.", fig["source"])

    add_heading(doc, "Regional Extension", level=2)
    add_body(
        doc,
        """
        Regional evidence is used only as a short extension. The main models remain national because processor-level official prices are Ukraine-only. Regional farm-gate dispersion and ProZorro regional medians are useful for showing that local procurement conditions differ, but they should not dominate the thesis.
        """,
    )
    for fig in figures:
        if "regional_farmgate" in fig["file"]:
            add_figure(doc, fig["file"], "Figure 5.7. Regional dispersion of enterprise farm-gate raw milk prices.", fig["source"])
    add_df_table(doc, vecm, "Table 5.3. VECM system feasibility", max_rows=8)

    add_heading(doc, "CONCLUSION AND POLICY DISCUSSION", level=0)
    add_body(
        doc,
        """
        The new empirical design makes the thesis narrower and more defensible. H1 is evaluated mainly with observed monthly farm-gate and processor-level prices. H2 is evaluated through official processor-to-consumer models and through ProZorro-retail mechanism evidence.

        The policy meaning is not that the models prove legal abuse. The result should be written more carefully: some parts of the chain show price-adjustment patterns consistent with bargaining asymmetry and market-power mechanisms. Processors may have room to delay or smooth the transmission of raw-milk shocks. Retailers and procurement channels may reshape upstream price signals through pricing, discounts, and institutional contracts.

        Future research should add confidential contract data, retailer settlement terms, processor-level regional data, and longer retail panels. These data would allow stronger identification of whether the observed price dynamics reflect market power, cost smoothing, perishability, or promotion strategy.
        """,
    )

    add_heading(doc, "REFERENCES AND SOURCE ANCHORS", level=0)
    for name, link in SOURCE_LINKS.items():
        doc.add_paragraph(f"{name}. {link}")

    out_path = DOC_OUT / "Maksym_Charniuk_MSc_thesis_market_power_newmodel_draft.docx"
    doc.save(out_path)
    return out_path


def save_all_clean_data(farm, processor, consumer, prozorro, proz_aggs, retail, retail_aggs, farm_volumes, regional, cost_index, trade):
    save_csv(farm, "clean_farmgate_monthly_ua_region.csv")
    save_csv(processor, "clean_processor_monthly_ua.csv")
    save_csv(consumer, "clean_consumer_monthly_ua_region.csv")
    save_csv(prozorro, "clean_prozorro_lot_level.csv")
    for key, df in proz_aggs.items():
        safe = {"D": "daily", "W-MON": "weekly", "MS": "monthly"}[key]
        save_csv(df, f"clean_prozorro_{safe}_product.csv")
    save_csv(retail, "clean_retail_sku_day.csv")
    for key, df in retail_aggs.items():
        save_csv(df, f"clean_retail_{key}.csv")
    save_csv(farm_volumes, "clean_farm_volumes_long.csv")
    for key, df in regional.items():
        save_csv(df, f"clean_{key}.csv")
    save_csv(cost_index, "clean_livestock_cost_index.csv")
    for key, df in trade.items():
        save_csv(df, f"clean_trade_{key}.csv")


def render_check_docx(path: Path) -> dict:
    info = {"docx": str(path), "exists": path.exists(), "paragraphs": None, "tables": None, "render_status": "not_attempted"}
    if not path.exists():
        return info
    try:
        d = Document(path)
        info["paragraphs"] = len(d.paragraphs)
        info["tables"] = len(d.tables)
    except Exception as exc:
        info["render_status"] = f"docx_read_failed: {exc}"
        return info
    soffice = shutil.which("soffice")
    if not soffice:
        info["render_status"] = "soffice_unavailable_docx_structure_checked"
    else:
        info["render_status"] = "soffice_available_manual_render_not_run_in_script"
    return info


def main():
    ensure_dirs()

    farm = clean_farmgate()
    processor = clean_processor()
    consumer = clean_consumer()
    farm_volumes = clean_farm_volumes()
    prozorro = clean_prozorro()
    proz_aggs = build_prozorro_aggregates(prozorro)
    retail = pd.concat([clean_retail_sheet("Silpo"), clean_retail_sheet("Novus")], ignore_index=True)
    retail_aggs = aggregate_retail(retail)
    cost_index = load_cost_index()
    trade = load_trade_tables()

    corruption_summary, validation = build_validation_tables(farm, processor, consumer)
    models, stationarity, ardl_orders = run_models(farm, processor, consumer, proz_aggs, retail_aggs)
    vecm = run_vecm_blocks(farm, processor, consumer)
    regional = regional_extension_tables(farm, consumer, prozorro, farm_volumes)

    save_all_clean_data(farm, processor, consumer, prozorro, proz_aggs, retail, retail_aggs, farm_volumes, regional, cost_index, trade)
    stationarity.to_csv(TABLE_OUT / "table_stationarity.csv", index=False)
    ardl_orders.to_csv(TABLE_OUT / "table_ardl_orders.csv", index=False)
    vecm.to_csv(TABLE_OUT / "table_vecm_systems.csv", index=False)
    models.to_csv(TABLE_OUT / "table_model_results_all.csv", index=False)
    corruption_summary.to_csv(TABLE_OUT / "table_integrated_consumer_corruption.csv", index=False)
    validation.to_csv(TABLE_OUT / "table_old_vs_new_validation.csv", index=False)

    figures = make_figures(farm, processor, consumer, proz_aggs, retail_aggs, models, regional, trade, cost_index)
    write_markdown_outputs(
        farm,
        processor,
        consumer,
        prozorro,
        retail,
        farm_volumes,
        cost_index,
        trade,
        corruption_summary,
        validation,
        models,
        stationarity,
        vecm,
        regional,
        figures,
    )
    docx_path = build_docx(models, figures, validation, vecm, regional, trade, cost_index)
    render_info = render_check_docx(docx_path)
    (OUT / "run_manifest.json").write_text(
        json.dumps(
            {
                "outputs": {
                    "root": str(OUT),
                    "docx": str(docx_path),
                    "clean_data": str(DATA_OUT),
                    "figures": str(FIG_OUT),
                    "tables": str(TABLE_OUT),
                },
                "render_check": render_info,
                "model_counts": models["reliability"].value_counts(dropna=False).to_dict() if not models.empty else {},
                "figure_count": len(figures),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps({"status": "ok", "out": str(OUT), "docx": str(docx_path), "figures": len(figures)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
